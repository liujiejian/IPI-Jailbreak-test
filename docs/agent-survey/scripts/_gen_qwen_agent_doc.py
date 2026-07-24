# -*- coding: utf-8 -*-
"""Generate Qwen-Agent research document — accuracy-focused with 5-round audit."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "QwenAgent.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def add_code(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def cite(text: str, url: str) -> None:
        doc.add_paragraph(f"{text}（来源：{url}）")

    doc.add_heading("Qwen-Agent", level=0)
    doc.add_paragraph(
        "本文档研究阿里巴巴 Qwen 团队开源 Qwen-Agent 框架机制，术语严格对齐官方资料。"
        "用户指定入口：https://github.com/QwenLM/Qwen-Agent；"
        "技术细节以 README、qwenlm.github.io/Qwen-Agent 文档站与源码为准。"
        "README 注明：框架现为 Qwen Chat（https://chat.qwen.ai/）后端。"
    )

    # ── Scope ──
    doc.add_heading("文档范围与术语分层（必读）", level=1)
    doc.add_paragraph("研究须区分以下层次，不可混用：")
    bullet(
        "框架原子层：BaseChatModel（LLM + Function Calling）与 BaseTool（@register_tool 工具注册表）。"
    )
    bullet(
        "Agent 运行时（class Agent 基类）：run(messages) 流式生成器；"
        "子类实现 _run() 工作流。"
    )
    bullet(
        "FnCallAgent / Assistant：默认 function calling 循环（LLM → tool → FUNCTION 消息 → 再 LLM），"
        "MAX_LLM_CALL_PER_RUN 默认 20。"
    )
    bullet(
        "ReActChat：文本 ReAct 模板（Thought/Action/Observation），"
        "非 JSON tool_calls API 路径。"
    )
    bullet(
        "Memory 类（qwen_agent.memory.Memory）：专用于文件/RAG 检索的 Agent，"
        "≠ 聊天会话历史；Assistant 内嵌 self.mem 做知识库检索。"
    )
    bullet(
        "应用层：BrowserQwen（Chrome 扩展 + qwen_server 三进程）、"
        "examples 演示脚本、Qwen Chat 产品后端。"
    )
    bullet(
        "多 Agent：Router / GroupChat / VirtualMemoryAgent 等，"
        "均须 MultiAgentHub._agents 且 name 唯一。"
    )
    cite(
        "README；agent.py；agents/__init__.py",
        "https://github.com/QwenLM/Qwen-Agent",
    )
    cite(
        "官方文档站 Overview",
        "https://qwenlm.github.io/Qwen-Agent/en/guide/",
    )

    # ── 1. LLM 请求采集 ──
    doc.add_heading("研究背景与 LLM 请求采集", level=1)
    doc.add_paragraph(
        "采集目标：每次 LLM 调用前的 messages、functions schema、"
        "function_call 参数、FUNCTION 角色工具返回、"
        "RAG 知识注入片段与多 Agent 路由轨迹。"
    )

    doc.add_heading("1. logging 与 QWEN_AGENT_DEBUG", level=2)
    bullet(
        "qwen_agent.log：标准 logging.StreamHandler；"
        "QWEN_AGENT_DEBUG=1 时 DEBUG 级别。"
    )
    bullet(
        "格式：时间 - 文件名 - 行号 - level - message；"
        "无内置 OpenTelemetry / LangSmith trace。"
    )
    cite(
        "qwen_agent/log.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/log.py",
    )

    doc.add_heading("2. agent.run() 流式输出", level=2)
    bullet(
        "Agent.run(messages)：默认流式 yield 增量 response 列表；"
        "run_nonstream() 取最后一次完整响应。"
    )
    bullet(
        "_call_llm(..., stream=True)：所有 Agent 默认 LLM 流式；"
        "可用 typewriter_print 做终端增量渲染。"
    )
    bullet(
        "返回类型：输入全为 dict 时 yield dict 列表，否则 yield Message 列表。"
    )
    cite(
        "qwen_agent/agent.py > run / _call_llm",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agent.py",
    )

    doc.add_heading("3. FnCallAgent 逐步可观测链", level=2)
    bullet(
        "循环：_call_llm(functions=...) → _detect_tool(message.function_call) "
        "→ _call_tool → Message(role=FUNCTION, name=..., content=result)。"
    )
    bullet(
        "工具返回字符串或 List[ContentItem]（多模态）；"
        "错误时 _call_tool 捕获异常并返回 traceback 文本。"
    )
    bullet(
        "研究侧可在每次 yield 后读取 response 列表末条 FUNCTION 消息。"
    )
    cite(
        "qwen_agent/agents/fncall_agent.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agents/fncall_agent.py",
    )

    doc.add_heading("4. BaseChatModel.chat 与 fncall 模板", level=2)
    bullet(
        "llm.chat(messages, functions, stream, extra_generate_cfg)："
        "preprocess/postprocess fncall 消息。"
    )
    bullet(
        "默认 NousFnCallPrompt：工具 schema 写入 system；"
        "调用格式 <tool_call>{\"name\":...,\"arguments\":...}</tool_call>；"
        "工具结果包裹 <tool_response>。"
    )
    bullet(
        "generate_cfg：fncall_prompt_type（默认 nous）、thought_in_content、"
        "max_input_tokens、use_raw_api（Qwen3-Coder 原生 API tool call）。"
    )
    bullet(
        "reasoning_content 字段：与 content 分离的推理链（Qwen3/QwQ 等）。"
    )
    bullet(
        "可选 diskcache 缓存 LLM 响应（cache_dir / generate_cfg）。"
    )
    cite(
        "llm/fncall_prompts/nous_fncall_prompt.py；README FAQ",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/llm/fncall_prompts/nous_fncall_prompt.py",
    )

    doc.add_heading("5. Assistant RAG 注入可观测", level=2)
    bullet(
        "Assistant._prepend_knowledge_prompt：调用 self.mem.run(messages) 检索，"
        "将片段格式化为 # Knowledge Base 块追加到 system prompt。"
    )
    bullet(
        "logger.debug 记录 retrieved/formatted knowledge 类型与内容。"
    )
    cite(
        "qwen_agent/agents/assistant.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agents/assistant.py",
    )

    doc.add_heading("6. 无内置 Trace 框架", level=2)
    doc.add_paragraph(
        "所查 README 与核心模块均无 SmolagentsInstrumentor / LangSmith 式 tracing；"
        "观测依赖 logging、messages 列表导出与应用侧 hook。"
    )
    cite(
        "README；log.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/README.md",
    )

    # ── 2. Prompt ──
    doc.add_heading("Prompt 设计组成", level=1)

    doc.add_heading("1. Agent.system_message", level=2)
    bullet(
        "Agent 初始化 system_message（默认空）；"
        "run() 时若无首条 SYSTEM 则 insert(0)，"
        "否则 prepend 到已有 system content。"
    )
    bullet(
        "lang 参数：根据消息是否含中文自动选 zh/en（可覆盖）。"
    )
    cite(
        "qwen_agent/agent.py > run",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agent.py",
    )

    doc.add_heading("2. Nous Function Call 模板", level=2)
    bullet(
        "FN_CALL_TEMPLATE：# Tools + JSON function signatures + "
        "<tool_call> XML 调用格式；原生支持并行工具调用（Parallel Function Calls）。"
    )
    bullet(
        "SPECIAL_CODE_MODE=true 时 code_interpreter 使用 "
        "FN_CALL_TEMPLATE_WITH_CI（code 放 <code> 块）。"
    )
    bullet(
        "FUNCTION 消息在 preprocess 时转为 USER 侧 <tool_response> 文本块。"
    )
    cite(
        "nous_fncall_prompt.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/llm/fncall_prompts/nous_fncall_prompt.py",
    )

    doc.add_heading("3. ReActChat 文本模板", level=2)
    bullet(
        "PROMPT_REACT：Question/Thought/Action/Action Input/Observation/Final Answer；"
        "stop=['Observation:', 'Observation:\\n']。"
    )
    bullet(
        "_detect_tool 解析 \\nAction: / \\nAction Input: 文本，非 function_call 字段。"
    )
    cite(
        "qwen_agent/agents/react_chat.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agents/react_chat.py",
    )

    doc.add_heading("4. Router 多 Agent 提示", level=2)
    bullet(
        "ROUTER_PROMPT：列出 agent_descs；"
        "模型输出 Call: {name} 或 Reply: ...；"
        "stop=['Reply:', 'Reply:\\n']。"
    )
    bullet(
        "解析 Call: 后委派给对应 Agent.run()，覆盖路由消息。"
    )
    cite(
        "qwen_agent/agents/router.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agents/router.py",
    )

    doc.add_heading("5. Assistant 知识库注入模板", level=2)
    bullet(
        "KNOWLEDGE_TEMPLATE_zh/en：# 知识库 / # Knowledge Base + "
        "KNOWLEDGE_SNIPPET 按 source 格式化检索片段。"
    )
    bullet(
        "VirtualMemoryAgent：检索结果写入 system 后，"
        "FUNCTION 返回固定文案 \"The relevant content has already been retrieved...\"。"
    )
    cite(
        "assistant.py；virtual_memory_agent.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agents/virtual_memory_agent.py",
    )

    doc.add_heading("6. 典型 FnCallAgent 执行轮次示意", level=2)
    add_code(
        """# 研究采集点（非固定 JSON schema）：
# - system（含 knowledge 注入 + tool schema）
# - user / assistant（含 function_call）
# - function（tool 返回，preprocess 为 <tool_response>）
# - Router: Call: sub_agent → 子 Agent 完整 run"""
    )
    cite(
        "FnCallAgent._run；schema.py Message roles",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/llm/schema.py",
    )

    # ── 3. Memory ──
    doc.add_heading("Memory 与持久化机制", level=1)
    doc.add_paragraph(
        "Qwen-Agent 中「Memory」一词有两层含义，研究时须严格区分："
        "框架 Memory 类 = RAG 文件检索 Agent；"
        "会话历史 = 应用侧 messages 列表；"
        "BrowserQwen = qwen_server 本地 jsonl/目录持久化。"
    )

    doc.add_heading("1. Memory 类（RAG 文件 Agent）", level=2)
    bullet(
        "继承 Agent；内置 retrieval + doc_parser 工具；"
        "处理 pdf/docx/pptx/txt/csv/tsv/xlsx/xls/html。"
    )
    bullet(
        "get_rag_files：system_files + 消息中提取的会话文件 URL。"
    )
    bullet(
        "rag_keygen_strategy：GenKeyword / SplitQueryThenGenKeyword 等 "
        "（可用 LLM 生成检索关键词）。"
    )
    bullet(
        "rag_searchers 默认 ['keyword_search', 'front_page_search']；"
        "max_ref_token 默认 20000（环境变量可覆盖）。"
    )
    bullet(
        "FnCallAgent 默认 self.mem = Memory(llm=..., files=files)；"
        "Qwen3/QwQ + dashscope 时用 qwen-turbo 作 mem LLM 以省 token。"
    )
    cite(
        "qwen_agent/memory/memory.py；settings.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/memory/memory.py",
    )

    doc.add_heading("2. storage 工具（键值文件存储）", level=2)
    bullet(
        "register_tool('storage')：operate=put/get/delete/scan；"
        "根目录默认 workspace/tools/storage。"
    )
    bullet(
        "应用层可模拟跨轮状态，但非框架自动会话记忆。"
    )
    cite(
        "qwen_agent/tools/storage.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/storage.py",
    )

    doc.add_heading("3. 会话历史模型", level=2)
    bullet(
        "README 示例：messages = [] 循环 append user + extend bot.run 返回；"
        "框架不内置 thread_id / checkpointer。"
    )
    bullet(
        "同 Agent 实例多次 run(messages) 依赖调用方传入累积 messages。"
    )
    cite(
        "README > Developing Your Own Agent",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/README.md",
    )

    doc.add_heading("4. BrowserQwen / qwen_server 持久化", level=2)
    bullet(
        "run_server.py 启动三子进程：database_server、workstation_server、assistant_server。"
    )
    bullet(
        "database_server：FastAPI；Memory 缓存浏览页；"
        "meta_data.jsonl 记录浏览元数据；history/ 目录存对话历史；"
        "download_root 存抓取页面文本。"
    )
    bullet(
        "cache_page 任务：扩展 POST 网页 content → 本地文件 → mem.run 解析标题。"
    )
    bullet(
        "此为 BrowserQwen 应用层持久化，≠ 通用 Agent SDK 内置能力。"
    )
    cite(
        "run_server.py；database_server.py；browser_qwen.md",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/browser_qwen.md",
    )

    doc.add_heading("5. 外部 MCP server-memory", level=2)
    bullet(
        "README MCP 示例含 @modelcontextprotocol/server-memory；"
        "属 MCP 远程工具，非框架内置跨会话语义记忆。"
    )
    cite(
        "README FAQ > MCP",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/README.md",
    )

    # ── 4. Session ──
    doc.add_heading("Session 与会话模型", level=1)
    bullet(
        "框架层无 thread_id / session key；"
        "会话 = 调用方维护的 messages 列表 + 可选 files 附件。"
    )
    bullet(
        "Message 支持多模态 ContentItem：text/image/file/audio/video。"
    )
    bullet(
        "BrowserQwen：按 URL 维度 history_dir 存对话；"
        "workstation Gradio 端口默认 7864。"
    )
    bullet(
        "GroupChat / UserAgent：多角色对话模拟，"
        "各 Agent 须有唯一 name。"
    )
    cite(
        "llm/schema.py；multi_agent_hub.py；browser_qwen.md",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/llm/schema.py",
    )

    # ── 5. Automation ──
    doc.add_heading("Automation 与定时任务", level=1)
    bullet(
        "框架层无 OpenClaw 式 Heartbeat、内置 cron 或定时调度。"
    )
    bullet(
        "run_server.py 常驻三服务供 BrowserQwen 扩展调用，"
        "非 agent 自主周期任务。"
    )
    bullet(
        "Qwen Chat 产品后端使用本框架，但定时/推送属产品层，"
        "开源仓库未记录等价机制。"
    )
    cite(
        "README；run_server.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/README.md",
    )

    # ── 6. External Web ──
    doc.add_heading("外部 Web 访问", level=1)
    doc.add_paragraph(
        "Qwen-Agent 无统一 web_fetch 或 EXTERNAL_UNTRUSTED_CONTENT；"
        "外部内容经可选 Tool、MCP 或 BrowserQwen 注入 agent。"
    )

    doc.add_heading("1. web_search（Serper API）", level=2)
    bullet(
        "register_tool('web_search')：POST google.serper.dev/search；"
        "须环境变量 SERPER_API_KEY。"
    )
    bullet(
        "返回 markdown 代码块包裹的 title+snippet 列表。"
    )
    cite(
        "qwen_agent/tools/web_search.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/web_search.py",
    )

    doc.add_heading("2. web_extractor", level=2)
    bullet(
        "register_tool('web_extractor')：参数 url；"
        "调用 SimpleDocParser 抓取并解析网页为文本。"
    )
    cite(
        "qwen_agent/tools/web_extractor.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/web_extractor.py",
    )

    doc.add_heading("3. MCP fetch / 其他 Web 工具", level=2)
    bullet(
        "官方 Overview 示例：mcp-server-fetch（uvx mcp-server-fetch）。"
    )
    bullet(
        "MCPManager：stdio / SSE / streamable-http 连接；"
        "动态注册 server_name-tool_name 为 BaseTool；"
        "支持 list_resources / read_resource。"
    )
    bullet(
        "BrowserQwen：Chrome 扩展将当前页 content POST 至 database_server cache_page。"
    )
    cite(
        "qwenlm.github.io guide；mcp_manager.py；browser_qwen.md",
        "https://qwenlm.github.io/Qwen-Agent/en/guide/",
    )

    doc.add_heading("4. 与 OpenClaw web_fetch 差异（研究对照）", level=2)
    bullet("OpenClaw：内置 web_fetch + EXTERNAL_UNTRUSTED_CONTENT 包裹。")
    bullet(
        "Qwen-Agent：web_extractor / web_search / MCP fetch 返回原文，"
        "经 FUNCTION 消息或 RAG system 注入 LLM；无框架级 untrusted 边界。"
    )

    # ── 7. Security ──
    doc.add_heading("Security 机制", level=1)

    doc.add_heading("1. code_interpreter（Docker 沙箱）", level=2)
    bullet(
        "register_tool('code_interpreter')：Docker 容器内 Jupyter kernel；"
        "仅挂载指定 work_dir（M6_CODE_INTERPRETER_WORK_DIR）。"
    )
    bullet(
        "README 与 Disclaimer：基础沙盒隔离，生产环境仍须谨慎。"
    )
    bullet(
        "须 Docker 安装运行；首次构建 code-interpreter:latest 镜像。"
    )
    cite(
        "qwen_agent/tools/code_interpreter.py；README FAQ",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/code_interpreter.py",
    )

    doc.add_heading("2. python_executor（非沙箱）", level=2)
    bullet(
        "PythonExecutor：description 明确 Not sandboxed；"
        "默认不 register_tool（# Do not register by default）。"
    )
    bullet(
        "tir_math.py / Qwen2.5-Math Demo 使用；"
        "README：仅本地测试，不可生产。"
    )
    bullet(
        "GenericRuntime.exec_code 阻止 input() 与 os.system()，"
        "但仍为本地 exec，非隔离边界。"
    )
    cite(
        "qwen_agent/tools/python_executor.py；README_CN",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/python_executor.py",
    )

    doc.add_heading("3. MCP 与 Hub 远程工具", level=2)
    bullet(
        "MCP 工具返回原文进入 agent；"
        "恶意 MCP server 可经 tool description 或返回内容注入。"
    )
    bullet(
        "MCPManager 单例管理子进程；atexit/SIGINT 清理。"
    )
    cite(
        "mcp_manager.py；README FAQ",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/mcp_manager.py",
    )

    doc.add_heading("4. IPI 研究含义", level=2)
    doc.add_paragraph(
        "注入面：web_extractor / web_search 返回、"
        "BrowserQwen 缓存页文本、MCP fetch/read_resource、"
        "RAG 检索片段（Memory → system prompt）、"
        "Router 子 Agent 返回。"
        "内容经 <tool_response> 或 knowledge 块进入后续 LLM 轮次，"
        "无 EXTERNAL_UNTRUSTED_CONTENT 自动消毒。"
    )

    doc.add_heading("5. 无框架级 Guardrail / PII 脱敏", level=2)
    doc.add_paragraph(
        "所查 agent/llm/tools 均无 HallucinationGuardrail、"
        "PII 中间件或 tool output 统一消毒层。"
    )
    cite(
        "qwen_agent/agent.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agent.py",
    )

    # ── 8. Architecture ──
    doc.add_heading("架构总览", level=1)
    bullet(
        "包结构：qwen_agent（框架核心）、qwen_server（BrowserQwen 服务）、"
        "browser_qwen（Chrome 扩展）、examples（演示）。"
    )
    bullet(
        "内置 Agent：Assistant、FnCallAgent、ReActChat、Router、GroupChat、"
        "DocQAAgent、ParallelDocQA、VirtualMemoryAgent、TIRMathAgent 等。"
    )
    bullet(
        "LLM 后端：qwen_dashscope、openai 兼容（vLLM/Ollama）、"
        "transformers 等（get_chat_model 注册表）。"
    )
    bullet(
        "工具注册表 TOOL_REGISTRY：code_interpreter、web_search、web_extractor、"
        "storage、image_gen、retrieval、image_search 等。"
    )
    bullet(
        "Gradio 5 WebUI（qwen_agent.gui）：快速部署 Demo；"
        "Python ≥3.10（GUI 要求）。"
    )
    cite(
        "agents/__init__.py；tools/__init__.py；README",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agents/__init__.py",
    )

    doc.add_heading(
        "与 OpenClaw / OpenHands / Hermes / LangGraph / CrewAI / AutoGen / Smolagents 的对照（研究映射）",
        level=2,
    )
    doc.add_paragraph("非 Qwen-Agent 官方术语：")
    bullet(
        "FnCallAgent 循环 ≈ OpenManus ToolCallAgent / AutoGen tool 循环；"
        "ReActChat ≈ 文本 ReAct 路径。"
    )
    bullet(
        "Memory 类（RAG）≈ CrewAI Knowledge / LangGraph store 检索注入；"
        "≠ 跨会话语义 Memory。"
    )
    bullet(
        "messages 列表 ≈ Smolagents AgentMemory 外置化；"
        "无 LangGraph checkpointer。"
    )
    bullet(
        "code_interpreter Docker ≈ Qwen-Agent 特色沙箱；"
        "对比 Smolagents e2b/docker executor。"
    )
    bullet(
        "BrowserQwen ≈ Smolagents webagent / OpenManus BrowserUseTool 应用层；"
        "web_extractor ≈ VisitWebpageTool 轻量 HTTP 抓取。"
    )
    bullet(
        "Nous <tool_call> 模板 ≈ 自研 fncall prompt；"
        "use_raw_api ≈ Qwen3-Coder 原生 OpenAI tools API。"
    )

    # ── 9. Research ──
    doc.add_heading("研究建议", level=1)
    bullet("采集：QWEN_AGENT_DEBUG=1 + 每轮导出 messages；Assistant 关注 system 中 knowledge 块。")
    bullet("对比 FnCallAgent vs ReActChat 同任务 IPI 易感面与步数。")
    bullet("IPI：web_extractor、web_search、MCP fetch、BrowserQwen 缓存页、RAG 检索注入。")
    bullet("安全：code_interpreter Docker vs python_executor 本地；验证 SERPER/MCP 依赖。")
    bullet("多 Agent：Router Call: 委派链；VirtualMemoryAgent 检索改写 system 时机。")
    bullet("长文档：assistant_rag.py vs parallel_doc_qa.py 两条 RAG 路径。")

    # ── 10. Five-round audit ──
    doc.add_heading("自查：潜在幻觉与核实结论（第五轮）", level=1)
    doc.add_paragraph(
        "2026-07-12 第五轮复核：对照 GitHub README/README_CN、"
        "qwenlm.github.io guide、agent/fncall_agent/assistant/memory、"
        "web_search/web_extractor/code_interpreter/python_executor、"
        "mcp_manager、run_server/database_server、browser_qwen.md 全部关键断言。"
        "结论：第四轮修正项仍成立；本轮未发现新的机制性幻觉。"
    )
    checks = [
        ("✓ 框架基于 Qwen Function Calling；FnCallAgent + Assistant 为主路径", "https://github.com/QwenLM/Qwen-Agent/blob/main/README.md"),
        ("✓ ReActChat 为文本 ReAct 模板，与 FnCallAgent function_call 路径并存", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agents/react_chat.py"),
        ("✓ Memory 类 = RAG 文件检索 Agent；≠ 聊天会话历史", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/memory/memory.py"),
        ("✓ 会话历史由应用侧 messages 列表维护；无框架 checkpointer", "https://github.com/QwenLM/Qwen-Agent/blob/main/README.md"),
        ("✓ 默认 NousFnCallPrompt + <tool_call>/<tool_response>；支持并行工具调用", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/llm/fncall_prompts/nous_fncall_prompt.py"),
        ("✓ code_interpreter 基于 Docker 容器 + Jupyter kernel", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/code_interpreter.py"),
        ("✓ python_executor 明确 Not sandboxed；默认不注册", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/python_executor.py"),
        ("✓ web_search 依赖 SERPER_API_KEY（Serper API）", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/web_search.py"),
        ("✓ web_extractor 经 SimpleDocParser 抓取 URL", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/web_extractor.py"),
        ("✓ MCP：mcpServers 配置 → MCPManager 动态注册工具", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/mcp_manager.py"),
        ("✓ BrowserQwen：run_server 三进程 + Chrome 扩展 + 本地 jsonl 持久化", "https://github.com/QwenLM/Qwen-Agent/blob/main/browser_qwen.md"),
        ("✓ MAX_LLM_CALL_PER_RUN 默认 20", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/settings.py"),
        ("✓ Qwen Chat 后端使用本框架（README 原文）", "https://github.com/QwenLM/Qwen-Agent/blob/main/README.md"),
        ("△ Memory 命名易混淆 — 框架 Memory=RAG Agent；MCP server-memory=外部工具（第四轮已分述）", "https://github.com/QwenLM/Qwen-Agent/blob/main/README.md"),
        ("△ use_raw_api — Qwen3-Coder 推荐开启；QwQ/Qwen3 vLLM 建议不加 hermes parser（README FAQ）", "https://github.com/QwenLM/Qwen-Agent/blob/main/README.md"),
        ("△ BrowserQwen 持久化 — 应用层 qwen_server，非通用 Agent.run 内置", "https://github.com/QwenLM/Qwen-Agent/blob/main/run_server.py"),
        ("✗ 内置 web_fetch / EXTERNAL_UNTRUSTED_CONTENT — 所查代码均无", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/web_extractor.py"),
        ("✗ OpenClaw 式 Heartbeat / 内置 cron — 框架无", "https://github.com/QwenLM/Qwen-Agent/blob/main/README.md"),
        ("✗ 框架级跨会话语义 Memory / OpenTelemetry trace — 所查代码无", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/log.py"),
        ("✗ python_executor 可作生产安全沙箱 — 源码与 README 明确否定", "https://github.com/QwenLM/Qwen-Agent/blob/main/README_CN.md"),
        ("✗ 所有 Web 工具零配置可用 — web_search 须 SERPER_API_KEY", "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/web_search.py"),
        ("— 第五轮：上述条目全部复检通过，无新增 ✗/△", "https://github.com/QwenLM/Qwen-Agent"),
    ]
    for text, url in checks:
        cite(text, url)

    doc.add_heading("防幻觉机制说明（五轮复核流程）", level=2)
    bullet("第一轮：确认 GitHub README 与 qwenlm.github.io 文档站权威来源，建立 Agent/Tool/应用分层。")
    bullet("第二轮：核对 Memory 类（RAG）≠ 会话历史 ≠ MCP server-memory。")
    bullet("第三轮：核对 Web 路径（web_search/web_extractor/MCP/BrowserQwen）及 code_interpreter vs python_executor。")
    bullet("第四轮：核对 Nous fncall 模板 vs use_raw_api；BrowserQwen 持久化为应用层。")
    bullet("第五轮：全量重读核心源码路径，确认无 OpenClaw 式机制误标。")

    doc.add_heading("主要参考链接", level=1)
    for link in [
        "https://github.com/QwenLM/Qwen-Agent",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/README.md",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/README_CN.md",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/browser_qwen.md",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/run_server.py",
        "https://qwenlm.github.io/Qwen-Agent/en/guide/",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agent.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agents/fncall_agent.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agents/assistant.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/agents/react_chat.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/memory/memory.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/llm/fncall_prompts/nous_fncall_prompt.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/web_search.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/web_extractor.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/code_interpreter.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_agent/tools/mcp_manager.py",
        "https://github.com/QwenLM/Qwen-Agent/blob/main/qwen_server/database_server.py",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
