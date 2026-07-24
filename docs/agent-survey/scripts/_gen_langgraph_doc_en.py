# -*- coding: utf-8 -*-
"""Generate LangGraph research notes (English).

Same structure as Chinese _gen_langgraph_doc.py:
Session / task modes / web_fetch(security notice) / tools,
plus Prompt composition & memory/persistence. Five-round anti-hallucination vs
docs.langchain.com (oss/python/langgraph).
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "en" / "LangGraph.docx"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label} (Source: {url})")

    # ── Title ──
    h("LangGraph", 0)
    p(
        "This document follows the same structure as OpenHands / Hermes / Claude-Code / Dify / AutoGPT / "
        "Cline / Codex-CLI / CrewAI series: "
        "Session, task modes, web_fetch (and Security Notice comparison), Tools; "
        "plus Prompt composition, memory and persistence mechanisms, and five-round "
        "anti-hallucination verification at the end. "
        "Subject: LangGraph (low-level orchestration runtime: durable execution, streaming, HITL, persistence)."
    )
    p("Official documentation index: https://docs.langchain.com/llms.txt")
    p(
        "Product layering (do not conflate): "
        "LangGraph = orchestration and persistence; "
        "LangChain = models/tools/`create_agent` and higher layers; "
        "Deep Agents = optional harness; "
        "LangSmith Deployment / Agent Server = deployment and platform capabilities (including Cron)."
    )
    cite("LangGraph overview", "https://docs.langchain.com/oss/python/langgraph/overview")
    p(
        "Execution chain (research view): "
        "invoke/stream(config.thread_id) → StateGraph nodes → "
        "(optional) bind_tools / create_agent tool loop → "
        "checkpointer persisted → interrupt/resume or next turn same thread."
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "Official session unit is **thread**: selected via `config[\"configurable\"][\"thread_id\"]` "
        "as checkpointer namespace. Same thread_id continues same conversation/workflow state; "
        "new thread_id = empty state new cursor."
    )
    cite("Persistence", "https://docs.langchain.com/oss/python/langgraph/persistence")
    cite("Interrupts", "https://docs.langchain.com/oss/python/langgraph/interrupts")

    h("1. thread_id and Checkpointer", 2)
    bullet(
        "After `graph.compile(checkpointer=...)`, invoke/stream must carry thread_id "
        "to continue state across calls."
    )
    bullet(
        "Checkpoint = graph state snapshot at a moment; thread = ordered set of checkpoints. "
        "Optional checkpoint_id locates a historical snapshot (time travel)."
    )
    bullet(
        "Implementations: InMemorySaver (in-process, lost on restart); "
        "SqliteSaver / PostgresSaver / MongoDB / Redis and other production backends."
    )
    bullet("PostgresSaver recommends thread_id < 255 characters.")
    cite(
        "Add memory / checkpointers",
        "https://docs.langchain.com/oss/python/langgraph/add-memory",
    )

    h("2. Interrupt / Resume (same thread)", 2)
    bullet(
        "`interrupt(payload)` pauses and persists; caller uses `Command(resume=...)` "
        "to continue on same thread_id; resume value becomes interrupt() return value."
    )
    bullet(
        "stream_events(..., version=\"v3\") can expose stream.interrupted / stream.interrupts."
    )

    h("3. Agent Server (platform side)", 2)
    bullet(
        "LangSmith Agent Server can auto-manage checkpointer/store infrastructure; "
        "thread model extended in Deployment API (including Thread Cron)."
    )

    h("4. Research mapping (non-official terminology)", 2)
    bullet("Codex/Cline session ≈ LangGraph thread_id.")
    bullet("CrewAI conversational session_id ≈ thread_id (semantically close, different API).")
    bullet("Checkpoint lineage (CrewAI) ≈ LangGraph checkpoint chain + optional fork/time-travel.")

    # ══════════════════════════════════════════════════════════════
    # Task modes
    # ══════════════════════════════════════════════════════════════
    h("Task Modes", 1)

    h("1. Workflows vs Agents", 2)
    bullet(
        "Workflow: predetermined code path (node/edge order); "
        "Agent: dynamically chooses process and uses tools."
    )
    bullet(
        "LangGraph provides orchestration infrastructure; common agent loop via LangChain `create_agent` "
        "(compiled to LangGraph) or hand-written StateGraph."
    )
    cite(
        "Workflows and agents",
        "https://docs.langchain.com/oss/python/langgraph/workflows-agents",
    )

    h("2. Invocation forms", 2)
    bullet("`invoke` / `ainvoke`, `stream` / `stream_events`.")
    bullet("Functional API and Graph API as two programming surfaces (officially coexist).")
    bullet("Subgraphs have independent/configurable checkpoint namespace.")

    h("3. Human-in-the-loop", 2)
    bullet("Dynamic interrupt (inside any node); static breakpoint (before/after nodes) also available.")
    bullet("Requires checkpointer; without persistence cannot reliably resume cross-process.")

    h("4. Periodic task boundary (aligned with survey table \"framework none*\")", 2)
    bullet(
        "Open-source LangGraph library itself provides no built-in cron / Heartbeat."
    )
    bullet(
        "LangSmith Deployment \"Use cron jobs\": cron creates new thread + fixed input runs assistant — "
        "deployment platform, not framework core API; survey note \"framework none*\" aligns."
    )
    cite("Use cron jobs", "https://docs.langchain.com/langsmith/cron-jobs")

    h("5. Local deployment", 2)
    bullet("`pip install -U langgraph` (or uv add); pure library embedded in application process.")
    bullet("Production can deploy to LangSmith Deployment / self-hosted Agent Server.")

    h("Task Modes · verification summary", 2)
    bullet(
        "Keep: workflow/agent; thread+checkpointer; interrupt; "
        "periodic tasks platform Cron only; pip library local deployment."
    )
    bullet("Drop: writing LangSmith Cron as langgraph package built-in; OpenClaw Heartbeat.")

    # ══════════════════════════════════════════════════════════════
    # web_fetch
    # ══════════════════════════════════════════════════════════════
    h("web_fetch (and Security Notice)", 1)
    p(
        "Conclusion (aligned with survey table \"must build yourself\"): LangGraph has **no** built-in "
        "`web_fetch` / browser tool, "
        "nor documented OpenClaw-style SECURITY NOTICE wrapping. "
        "External web capability from: user-defined `@tool` / LangChain integration tools / "
        "model-side server-side tools (if provider supports) — "
        "security appearance wrapping after injection into messages is developer's choice."
    )

    h("1. Tools must be built or integrated", 2)
    bullet(
        "Official emphasis on familiarity with LangChain models & tools; "
        "examples often use `llm.bind_tools([...])` or `create_agent(..., tools=[...])`."
    )
    bullet(
        "Integration catalog may include Tavily Search and other web search tools (LangChain integrations), "
        "not LangGraph core built-in."
    )
    bullet(
        "Some chat models have server-side web search / code interpreter — "
        "model provider capability, not graph runtime built-in scrape."
    )
    cite("Tools (LangChain)", "https://docs.langchain.com/oss/python/langchain/tools")
    cite(
        "Tavily Search (integrations)",
        "https://docs.langchain.com/oss/python/integrations/tools/tavily_search",
    )

    h("2. Security Notice comparison", 2)
    bullet(
        "LangGraph/LangChain tool doc examples return plain strings or structured objects; "
        "no mandatory EXTERNAL_UNTRUSTED_CONTENT / SECURITY NOTICE."
    )
    bullet(
        "Research should treat web as untrusted at application layer (can reference Codex \"treat as untrusted\" engineering practice), "
        "but do not write as LangGraph default behavior."
    )

    h("Threat model mapping (research, non-official)", 2)
    bullet(
        "Custom scrape/search tool → ToolMessage → subsequent node/agent: "
        "typical indirect injection surface; framework does not write notice."
    )
    bullet(
        "If Store persists web-derived facts: cross-thread pollution surface expands."
    )
    bullet(
        "Survey table row: CrewAI/Cline have ready scrape/fetch tool names; "
        "LangGraph row \"must build yourself\" is accurate."
    )

    code(
        """# Research sketch: custom external web tool attached to graph/agent
@tool
def web_search(query: str) -> str:
    \"\"\"Search the web for information.\"\"\"
    return fetch_somehow(query)  # return value has no framework-mandated SECURITY NOTICE

# llm.bind_tools([web_search]) or create_agent(model, tools=[web_search])
# Official docs do not mandate untrusted wrapper
"""
    )

    h("web_fetch · verification summary", 2)
    bullet("Keep: no built-in web_fetch; tools custom/integrated; no SECURITY NOTICE.")
    bullet("Drop: writing Tavily/OpenAI web_search as LangGraph built-in same-name tool.")

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "Orchestration layer \"natively\" supports tool-calling loop and graph node side effects, "
        "but **tool implementations must be defined** (table: \"native, tools must be custom\")."
    )

    h("1. Attachment methods", 2)
    bullet("LangChain: `@tool`, `bind_tools`, `create_agent(tools=...)`")
    bullet("StateGraph nodes manually execute tool_calls and write back to MessagesState.")
    bullet("ToolRuntime can access state / context / store / tool_call_id.")

    h("2. Store / Command", 2)
    bullet("Tools can update state via Command; watch reducer when parallel tool updates fields.")
    bullet("Tools can read/write BaseStore (long-term memory).")

    h("Tools · verification summary", 2)
    bullet("Keep: native tool-calling orchestration + custom/integration tools.")
    bullet("Drop: fabricating built-in fixed tool whitelist (WebFetch/Bash…).")

    # ══════════════════════════════════════════════════════════════
    # Five-round anti-hallucination
    # ══════════════════════════════════════════════════════════════
    h("Five-Round Anti-Hallucination Verification", 1)

    h("Round 1 — Session", 2)
    bullet("Checked: persistence; add-memory; interrupts.")
    bullet("Keep: thread_id; checkpointer; interrupt/Command(resume).")
    bullet("Drop: Claude JSONL path; CrewAI kickoff_id conflation.")

    h("Round 2 — Task modes", 2)
    bullet("Checked: overview; workflows-agents; cron-jobs; deploy.")
    bullet("Keep: workflow/agent; pip; HITL; Cron Deployment only.")
    bullet("Drop: framework built-in scheduler.")

    h("Round 3 — web_fetch / security", 2)
    bullet("Checked: overview (requires models+tools); langchain tools; integrations.")
    bullet("Keep: must build yourself; no SECURITY NOTICE.")
    bullet("Drop: built-in web_fetch.")

    h("Round 4 — Tools", 2)
    bullet("Checked: langchain tools; workflows-agents bind_tools.")
    bullet("Keep: native orchestration + custom tools.")
    bullet("Drop: counting Deep Agents filesystem tools as LangGraph core default.")

    h("Round 5 — Prompt / Memory cross-check", 2)
    bullet("Checked: add-memory; stores; persistence table.")
    bullet(
        "Freeze: "
        "(1) Session ≡ thread_id + checkpoints; "
        "(2) Tasks = graph invoke/stream ± HITL; periodic = platform Cron; "
        "(3) External web = custom/integration tools; "
        "(4) Prompt = application state messages + developer node logic; "
        "(5) Memory = checkpointer (short) + store (long)."
    )
    bullet(
        "Open questions: whether specific integration tool returns include security prefix; "
        "Agent Server default checkpointer backend brand; "
        "subgraph checkpoint sharing best practices need per-version verification."
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt composition
    # ══════════════════════════════════════════════════════════════
    h("Prompt Composition", 1)
    p(
        "LangGraph **does not abstract prompts or architecture** (overview): "
        "prompts assembled by you in nodes/`create_agent` system layer. "
        "Common carriers: `MessagesState[\"messages\"]`, system messages, structured output schema."
    )

    h("1. Messages state", 2)
    bullet("Multi-turn: messages persisted with thread via checkpointer.")
    bullet("Tool results usually written back as ToolMessage in message list.")

    h("2. create_agent / hand-written nodes", 2)
    bullet("High-level: LangChain agents encapsulate tool loop and middleware.")
    bullet("Low-level: `model.invoke(state[\"messages\"])` fully controlled inside nodes.")

    h("3. No default AGENTS.md", 2)
    bullet("Framework does not load repo AGENTS.md; can load via application code or Deep Agents etc.")

    h("Prompt Composition · verification summary", 2)
    bullet("Keep: developer owns prompt; messages-centric.")
    bullet("Drop: claiming LangGraph built-in fixed system template / SECURITY NOTICE.")

    # ══════════════════════════════════════════════════════════════
    # Memory and persistence
    # ══════════════════════════════════════════════════════════════
    h("Memory and Persistence", 1)
    p(
        "Official dual track (aligned with survey table Memory=checkpointer/store):"
    )

    h("1. Checkpointer (short / thread-scoped)", 2)
    bullet("Persists graph state snapshots: conversation continuity, HITL, time travel, fault tolerance.")
    bullet("Access: thread_id (+ optional checkpoint_id).")
    bullet("In-memory implementation dev only; production uses DB backend.")

    h("2. Store (long / cross-thread)", 2)
    bullet(
        "Application-defined KV: user preferences, facts, shared knowledge; "
        "`compile(..., store=)`; nodes/tools access via Runtime."
    )
    bullet("InMemoryStore vs PostgresStore / MongoDBStore / RedisStore, etc.")
    bullet("Can enable semantic search (semantic search over store items).")
    cite("Stores", "https://docs.langchain.com/oss/python/langgraph/stores")

    h("3. Relation to \"session\"", 2)
    bullet("Multi-turn chat: usually checkpointer + MessagesState suffices.")
    bullet("Cross-session memory: Store (or external DB).")
    bullet("Often compile both together.")

    code(
        """from langgraph.checkpoint.memory import InMemorySaver
from langgraph.store.memory import InMemoryStore

graph = builder.compile(
    checkpointer=InMemorySaver(),
    store=InMemoryStore(),
)
graph.invoke(inputs, {"configurable": {"thread_id": "thread-1"}})
"""
    )

    h("Memory and Persistence · verification summary", 2)
    bullet("Keep: checkpointer≠store; thread short memory / cross-thread long memory.")
    bullet("Drop: Hermes MEMORY.md; CrewAI unified Memory class as LangGraph API.")

    h("Primary Reference Links", 1)
    for link in [
        "https://docs.langchain.com/llms.txt",
        "https://docs.langchain.com/oss/python/langgraph/overview",
        "https://docs.langchain.com/oss/python/langgraph/persistence",
        "https://docs.langchain.com/oss/python/langgraph/add-memory",
        "https://docs.langchain.com/oss/python/langgraph/checkpointers",
        "https://docs.langchain.com/oss/python/langgraph/stores",
        "https://docs.langchain.com/oss/python/langgraph/interrupts",
        "https://docs.langchain.com/oss/python/langgraph/streaming",
        "https://docs.langchain.com/oss/python/langgraph/workflows-agents",
        "https://docs.langchain.com/oss/python/langchain/tools",
        "https://docs.langchain.com/oss/python/langchain/agents",
        "https://docs.langchain.com/oss/python/integrations/tools/tavily_search",
        "https://docs.langchain.com/langsmith/cron-jobs",
        "https://docs.langchain.com/langsmith/agent-server",
        "https://docs.langchain.com/oss/python/concepts/memory",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
