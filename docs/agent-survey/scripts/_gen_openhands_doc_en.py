# -*- coding: utf-8 -*-
"""Generate OpenHands research notes (English): Session / task modes / web fetch / tools.

Verified against official docs (docs.openhands.dev) with a 5-round
anti-hallucination pass. Structure mirrors Openclaw.docx for the paper track:
Agent Security — Risks from External Web Fetching.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "en" / "OpenHands.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label} (Source: {url})")

    # ── Title ──
    h("OpenHands", 0)
    p(
        "This document aligns with the paper task: Session, task modes, "
        "web_fetch (security notice), and tools; "
        "appendices cover Prompt Composition and Memory & Persistence. "
        "Following the OpenClaw example structure, it surveys OpenHands V1 official "
        "documentation, removes incorrect claims after anti-hallucination review, "
        "and retains traceable sources."
    )
    p("Official documentation index: https://docs.openhands.dev/llms.txt")
    p(
        "Execution chain (research view): User task → Agent(LLM → call tool → "
        "tool/observation → LLM) → User. "
        "The LLM sees system / user / assistant / tool messages and tools schema."
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "OpenHands does not use OpenClaw-style Main / Isolated / session:xxx naming; "
        "the official entity is Conversation. "
        "Session determines runtime context; different task modes create or reuse Conversations."
    )
    cite(
        "Conversation Architecture",
        "https://docs.openhands.dev/sdk/arch/conversation.md",
    )

    h("1. Conversation Factory: Local vs Remote", 2)
    p(
        "Conversation(agent, workspace) automatically selects implementation by workspace type:"
    )
    bullet(
        "LocalConversation: workspace is str or LocalWorkspace → agent runs in-process; "
        "communication via direct function calls; suited for development / CLI."
    )
    bullet(
        "RemoteConversation: workspace is RemoteWorkspace → connects to Agent Server "
        "via HTTP/WebSocket, runs in an isolated container; suited for production / Web."
    )
    p("Switching deployment mode typically requires only changing workspace type; API surface stays consistent.")
    cite(
        "Factory Pattern / Local vs Remote Execution",
        "https://docs.openhands.dev/sdk/arch/conversation.md",
    )

    h("2. In-Session State: Event Log", 2)
    bullet(
        "ConversationState.events is an append-only Event Log "
        "(user / assistant / Action / Observation, etc.)."
    )
    bullet(
        "Agent follows a reasoning-action loop: each turn reads history from the Event Log "
        "then queries the LLM (mutable business state across steps is not held long-term by the Agent itself)."
    )
    cite(
        "Conversation Architecture > State Management",
        "https://docs.openhands.dev/sdk/arch/conversation.md",
    )

    h("3. Cross-Session Recovery (Resume / Persistence)", 2)
    bullet(
        "CLI: conversations saved at ~/.openhands/conversations/<id>/conversation.json; "
        "recover via openhands --resume, --resume --last, --resume <id>."
    )
    cite(
        "Resume Conversations",
        "https://docs.openhands.dev/openhands/usage/cli/resume.md",
    )
    bullet(
        "SDK: with persistence_dir + conversation_id set, "
        "base_state.json stores config/state, events/event-*.json incrementally store events."
    )
    bullet(
        "SDK implementation detail (persistence guide): when modifying public fields of "
        "ConversationState, a custom __setattr__ immediately serializes base state (not manual save). "
        "The architecture overview mentions Debounced writes for the Persistence helper service; "
        "the Persistence guide's \"immediately when they occur\" takes precedence."
    )
    cite(
        "Persistence Official Guide",
        "https://docs.openhands.dev/sdk/guides/convo-persistence.md",
    )

    h("4. Research Mapping to OpenClaw Session Model (Non-Official Terminology)", 2)
    bullet(
        "OpenClaw Main session (shared context) ≈ OpenHands persistent Conversation + CLI/SDK Resume."
    )
    bullet(
        "OpenClaw Isolated job (fresh session) ≈ OpenHands Automation each run: "
        "new sandbox + new Conversation (review / continue afterward)."
    )
    bullet(
        "OpenClaw Custom session:xxx (persistent across runs) ≈ "
        "Persistence / Resume with the same conversation_id; "
        "or continue on an Automation run result conversation."
    )
    bullet(
        "OpenClaw Heartbeat: OpenHands public product docs describe no equivalent mechanism; "
        "use Automations (cron) for periodic tasks, Event-based Automations for event-driven work."
    )

    # ══════════════════════════════════════════════════════════════
    # Task Modes
    # ══════════════════════════════════════════════════════════════
    h("Task Modes", 1)
    p(
        "Corresponding to OpenClaw Cron / Heartbeat / Automation: "
        "OpenHands task forms determine whether context is shared, sandbox is isolated, "
        "and confirmation policy applies."
    )

    h("1. Interactive Main Session (Interactive Conversation)", 2)
    bullet("CLI Terminal: openhands (confirmation required by default); supports -t / -f to start tasks.")
    bullet(
        "Confirmation modes: default confirmation; --always-approve fully automatic; "
        "--llm-approve uses LLM security analysis."
    )
    cite(
        "Terminal (CLI)",
        "https://docs.openhands.dev/openhands/usage/cli/terminal.md",
    )
    bullet(
        "Agent Canvas / Cloud Web UI: browser Conversation; "
        "backend/sandbox/workspace/agent switchable."
    )
    cite(
        "Agent Canvas Overview",
        "https://docs.openhands.dev/openhands/usage/agent-canvas/overview.md",
    )
    bullet(
        "SDK Conversation: send_message + run/step; supports Pause/Resume, Fork, "
        "Send Message While Running, Goal Completion Loop."
    )

    h("2. Headless (No-UI Batch / CI)", 2)
    p("openhands --headless -t \"...\" or -f task.txt.")
    bullet("Must provide a task (--task / --file).")
    bullet(
        "Forces always-approve: cannot switch to --llm-approve; "
        "suited for scripts and CI; security boundary must be enforced by the caller."
    )
    bullet("Optional --json outputs JSONL event stream.")
    cite(
        "Headless Mode",
        "https://docs.openhands.dev/openhands/usage/cli/headless.md",
    )

    h("3. Automations (Scheduled / Prompt-based & Plugin-based)", 2)
    p(
        "Background AI tasks on a schedule. Each run: "
        "① create fresh sandbox; ② execute prompt; ③ save conversation for review or continue."
    )
    bullet("Prompt-based: natural language task description and cron (most common).")
    bullet(
        "Plugin-based: mount OpenHands extensions plugins "
        "(extra skills / MCP / commands)."
    )
    bullet("Default timeout 10 minutes, maximum 30 minutes.")
    bullet(
        "Environment capabilities: terminal, file operations, LLM/Secrets in Settings, "
        "MCP, network HTTP, logged-in GitHub/GitLab/Bitbucket credentials."
    )
    cite(
        "Automations Overview",
        "https://docs.openhands.dev/openhands/usage/automations/overview.md",
    )
    cite(
        "Creating Automations",
        "https://docs.openhands.dev/openhands/usage/automations/creating-automations.md",
    )

    h("4. Event-based Automations (Event-Driven)", 2)
    bullet(
        "Built-in GitHub events: pull_request / issues / issue_comment / push / release, etc.; "
        "filterable with JMESPath."
    )
    bullet("Custom webhooks: Linear / Stripe / Slack, etc. (register webhook first, then create automation).")
    cite(
        "Event-Based Automations",
        "https://docs.openhands.dev/openhands/usage/automations/event-automations.md",
    )

    h("5. Scheduling Strategy Comparison (Research Notes)", 2)
    p("OpenClaw: precise timing/isolation → Cron; needs full session context → Heartbeat.")
    p("OpenHands equivalents:")
    bullet("Scheduled or event-driven, needs isolated review → Automations (new Conversation each run).")
    bullet("Needs continuous human-agent collaboration context → main Conversation + Resume.")
    bullet("No UI / CI → Headless (always-approve).")

    # ══════════════════════════════════════════════════════════════
    # web_fetch / Security Notice
    # ══════════════════════════════════════════════════════════════
    h("web_fetch (and Security Notice)", 1)
    p(
        "[Conclusion] OpenHands V1 has no built-in web_fetch tool equivalent to OpenClaw, "
        "nor a framework-level auto-injected SECURITY NOTICE / "
        "<<<EXTERNAL_UNTRUSTED_CONTENT>>> wrapping mechanism. "
        "External content enters Observation / tool messages via multiple paths."
    )

    h("1. Path A: MCP Fetch (Closest to web_fetch in Official Examples)", 2)
    p(
        "Register mcp-server-fetch via Agent.mcp_config; "
        "MCP tools auto-discovered when agent initializes. "
        "Use filter_tools_regex to limit exposed tools."
    )
    code(
        """mcp_config = {
    "mcpServers": {
        "fetch": {
            "command": "uvx",
            "args": ["mcp-server-fetch"]
        }
    }
}
agent = Agent(llm=llm, tools=tools, mcp_config=mcp_config)"""
    )
    cite(
        "MCP Integration Example (includes fetch)",
        "https://docs.openhands.dev/sdk/guides/mcp.md",
    )
    cite(
        "Persistence Example: MCP fetch reads remote URL then writes to file",
        "https://docs.openhands.dev/sdk/guides/convo-persistence.md",
    )
    p(
        "Research note: whether tool return content from this path carries untrusted labeling "
        "depends on the MCP server implementation, not the OpenHands core framework."
    )

    h("2. Path B: BrowserToolSet (browser-use)", 2)
    p(
        "Optional tool set: navigate, click, fill forms, extract page content; "
        "observations return to the LLM as Observations. "
        "Official examples combine with TerminalTool / FileEditorTool for web research."
    )
    cite(
        "Browser Use",
        "https://docs.openhands.dev/sdk/guides/agent-browser-use.md",
    )

    h("3. Path C: terminal + curl/wget, etc.", 2)
    p(
        "Built-in terminal tool can run arbitrary shell HTTP clients. "
        "Output enters the Event Log as TerminalObservation with no framework-level SECURITY NOTICE prefix."
    )

    h("4. Path D: Tavily Search (via MCP)", 2)
    p(
        "Usable as search engine: Tavily MCP provides search / extract / crawl / map. "
        "OpenHands Cloud configures Tavily by default; self-hosted fills "
        "Search API Key (Tavily) in Settings > LLM."
    )
    cite(
        "Search Engine Setup",
        "https://docs.openhands.dev/openhands/usage/advanced/search-engine-setup.md",
    )

    h("5. Security Notice: Comparison with OpenClaw", 2)
    bullet(
        "OpenClaw: web_fetch return writes SECURITY NOTICE and wraps body with "
        "EXTERNAL_UNTRUSTED_CONTENT boundaries; "
        "also cacheTtlMinutes (commonly 15 minutes) and related config."
    )
    bullet(
        "OpenHands core product/SDK docs: no equivalent automatic wrapping for generic external web content."
    )
    bullet(
        "Exception (PR Review plugin only): OpenHands/extensions PR review prompt "
        "uses BEGIN/END UNTRUSTED PR CONTENT {nonce} markers "
        "for author-controlled fields like PR title/body/diff "
        "(anti PR text injection; not a generic web_fetch path)."
    )
    cite(
        "extensions PR: harden review prompt against injection",
        "https://github.com/OpenHands/extensions/pull/252",
    )
    p(
        "IPI (indirect prompt injection) experimental implications: test surface should cover "
        "MCP fetch, BrowserToolSet, terminal HTTP; do not assume OpenClaw-style framework-level untrusted wrapping."
    )

    h("6. OpenHands Security Mechanisms (Action-Level, Not Content Wrapping)", 2)
    bullet(
        "LLMSecurityAnalyzer: injects required security_risk (LOW/MEDIUM/HIGH) "
        "into non-read-only tool schemas; LLM annotates inline in tool_call arguments."
    )
    bullet(
        "ConfirmationPolicy: AlwaysConfirm / NeverConfirm / ConfirmRisky "
        "(architecture docs describe ConfirmRisky as flexible default; works with analyzer)."
    )
    bullet(
        "Also composable analyzers like Pattern / PolicyRail (security guide): "
        "deterministic rule scanning before tool execution."
    )
    bullet("Sandbox: Docker (recommended) / Process (fast but not isolated) / Remote.")
    bullet("Hooks: PreToolUse can intercept; Stop can enforce lint/test.")
    cite(
        "Security Architecture",
        "https://docs.openhands.dev/sdk/arch/security.md",
    )
    cite(
        "Security & Action Confirmation",
        "https://docs.openhands.dev/sdk/guides/security.md",
    )
    cite(
        "Sandboxes Overview",
        "https://docs.openhands.dev/openhands/usage/sandboxes/overview.md",
    )

    p("Illustration: tool_call with security_risk enabled (tool names per V1; see Tools section):")
    code(
        """{
  "name": "terminal",
  "arguments": {
    "command": "curl -s https://api.example.com/health",
    "security_risk": "MEDIUM"
  }
}"""
    )
    p(
        "Note: SDK Security architecture page examples once used name: execute_bash; "
        "that name is outdated illustration. Current built-in tool name is terminal (see below and source)."
    )

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "Tool system: Action (input) → Executor → Observation (output); "
        "native tools via ToolRegistry; MCP tools discovered at Agent.initialize and merged into tools_map."
    )
    cite(
        "Tool System & MCP",
        "https://docs.openhands.dev/sdk/arch/tool-system.md",
    )

    h("1. Built-in Core Tools (Current SDK Naming)", 2)
    p(
        "ToolDefinition subclasses without explicit name override derive name from class name "
        "CamelCase→snake_case and remove _tool suffix. Source: openhands/sdk/tool/tool.py."
    )
    bullet(
        "terminal (TerminalTool): persistent shell session command execution; "
        "annotations.title=\"terminal\"; destructiveHint=True, openWorldHint=True."
    )
    bullet(
        "file_editor (FileEditorTool): view / create / str_replace / insert / undo_edit."
    )
    bullet(
        "task_tracker (TaskTrackerTool): task tracking; "
        "official README default examples often enable together with terminal, file_editor."
    )
    bullet(
        "BrowserToolSet (optional): browser-use based browser tool collection."
    )
    bullet(
        "TaskToolSet (optional): parent agent synchronously delegates sub-agent; "
        "supports resume=task_id to resume sub-session."
    )
    cite(
        "Custom Tools / Browser Use / Task Tool Set; software-agent-sdk README",
        "https://docs.openhands.dev/sdk/guides/custom-tools.md",
    )
    p(
        "Historical/mixed doc names (deprecated or in old examples only): BashTool, execute_bash, "
        "str_replace_editor. When collecting LLM tool schema for research, use actual registered name field."
    )

    h("2. MCP Tools", 2)
    bullet("Configuration: Agent(mcp_config=...).")
    bullet("Transport: stdio / SSE / SHTTP (product docs recommend HTTP/SSE for proxy reliability).")
    bullet("Filtering: filter_tools_regex.")
    bullet("OAuth MCP: first use requires browser login; not suited for pure headless.")

    h("3. Memory / Long-Term Context Related (No Dedicated Memory Tools)", 2)
    p(
        "OpenHands has no OpenClaw memory_search / memory_get, "
        "nor official MEMORY.md / memory/YYYY-MM-DD.md dedicated system."
    )
    bullet("Session context: Event Log + optional Condenser.")
    bullet("Repository permanent context: AGENTS.md (always-on); .agents/skills/ on-demand trigger.")
    bullet("Arbitrary file read/write: file_editor or terminal.")

    h("4. Simplified Messages Example for Research", 2)
    p("Assuming MCP fetch enabled, or agent chooses terminal curl:")
    code(
        """{
  "messages": [
    {
      "role": "system",
      "content": "<agent prompt + AGENTS.md + triggered skills>"
    },
    {
      "role": "user",
      "content": "Check https://api.example.com/health and return the status."
    },
    {
      "role": "assistant",
      "content": null,
      "tool_calls": [{
        "id": "call_001",
        "type": "function",
        "function": {
          "name": "terminal",
          "arguments": "{\\"command\\": \\"curl -s https://api.example.com/health\\", \\"security_risk\\": \\"LOW\\"}"
        }
      }]
    },
    {
      "role": "tool",
      "tool_call_id": "call_001",
      "content": "{\\"ok\\": true}"
    },
    {
      "role": "assistant",
      "content": "The API is healthy (ok=true)."
    }
  ],
  "tools": [
    {"type": "function", "function": {"name": "terminal", "...": "..."}},
    {"type": "function", "function": {"name": "file_editor", "...": "..."}}
  ]
}"""
    )

    # ══════════════════════════════════════════════════════════════
    # Five-Round Anti-Hallucination
    # ══════════════════════════════════════════════════════════════
    h("Five-Round Anti-Hallucination Verification Record", 1)
    p(
        "Goal: remove fabricated OpenClaw-style mechanisms, outdated tool names, "
        "and unlabeled inferences. Each round accepts only official docs, official examples, "
        "or verifiable source/PR."
    )

    h("Round 1 — Session / Conversation", 2)
    bullet(
        "Verified: https://docs.openhands.dev/sdk/arch/conversation.md"
    )
    bullet(
        "Retained: LocalConversation vs RemoteConversation; Event Log; Factory Pattern."
    )
    bullet(
        "Removed: writing OpenClaw Main/Isolated/Heartbeat as OpenHands official terms "
        "(moved to Research Mapping section labeled non-official)."
    )

    h("Round 2 — Task Modes", 2)
    bullet(
        "Verified: automations/overview.md, creating-automations.md, "
        "event-automations.md, cli/headless.md, cli/terminal.md, "
        "agent-canvas/overview.md."
    )
    bullet(
        "Retained: Automation each run fresh sandbox; timeout 10–30 min; "
        "Headless forces always-approve; CLI three confirmation modes."
    )
    bullet(
        "Removed: claiming OpenHands has Heartbeat; "
        "removed unstated \"default Automation reuses same session each run\" "
        "(official states each run new sandbox + new conversation, continue afterward)."
    )

    h("Round 3 — web_fetch / Security Notice", 2)
    bullet(
        "Verified: sdk/guides/mcp.md, agent-browser-use.md, search-engine-setup.md; "
        "compared with OpenClaw SECURITY NOTICE implementation (OpenClaw, not OpenHands)."
    )
    bullet(
        "Retained: no built-in web_fetch; four external content paths; "
        "PR review plugin UNTRUSTED PR CONTENT {nonce} (extension scenario exception)."
    )
    bullet(
        "Removed: OpenHands built-in cacheTtlMinutes=15; "
        "removed \"OpenHands web_fetch auto-wraps SECURITY NOTICE\"; "
        "removed attributing OpenClaw external-content.ts behavior to OpenHands core."
    )

    h("Round 4 — Tools Naming and Security", 2)
    bullet(
        "Verified: custom-tools.md, Browser/Task guides, "
        "TerminalTool/FileEditorTool source in software-agent-sdk, "
        "tool.py _camel_to_snake(...).removesuffix('_tool'), "
        "sdk/arch/security.md, sdk/guides/security.md."
    )
    bullet(
        "Corrected: LLM-visible tool names are terminal / file_editor / task_tracker; "
        "execute_bash appears only in some architecture examples, treated as outdated illustration."
    )
    bullet(
        "Corrected: Persistence — detailed guide writes base_state immediately; "
        "do not treat architecture page \"Debounced writes\" alone as sole conclusion."
    )
    bullet(
        "Retained: security_risk injection for non-read-only tools; "
        "ConfirmRisky + analyzer; three sandbox tiers."
    )

    h("Round 5 — Cross-Check and Frozen Conclusions", 2)
    bullet(
        "Re-checked llms.txt: confirmed no standalone \"web_fetch\" product doc page; "
        "external fetching distributed across MCP / Browser / Search / terminal."
    )
    bullet(
        "Frozen research conclusions (four items): "
        "(1) Session ≡ Conversation (Local/Remote + Resume); "
        "(2) Task modes = Interactive / Headless / Cron Automation / Event Automation; "
        "(3) No built-in web_fetch / no framework-level SECURITY NOTICE; IPI surface in MCP/Browser/curl; "
        "(4) Core tool names terminal + file_editor (+ optional Browser/MCP/Task)."
    )
    bullet(
        "Still open (do not state as established fact): "
        "specific mcp-server-fetch return JSON fields; "
        "whether Cloud default tool set exactly matches self-hosted CLI; "
        "requires live capture of LLM requests to verify."
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt Composition
    # ══════════════════════════════════════════════════════════════
    h("Prompt Composition", 1)
    p(
        "Similar to OpenClaw, requests to the LLM mainly include system / user / assistant / tool "
        "messages and tools schema. "
        "OpenHands assembles via the Agent reasoning-action loop each step(): "
        "optional Condenser compresses history → query LLM → (security check) → execute tool → "
        "Observation written back to Event Log."
    )
    cite(
        "Agent Architecture: reasoning-action loop and Agent Context",
        "https://docs.openhands.dev/sdk/arch/agent.md",
    )

    h("1. Core Message Structure", 2)
    bullet(
        "System Prompt: assembled from Agent prompt template + AgentContext "
        "(repo skills, system prompt prefix/suffix)."
    )
    bullet(
        "User / Assistant / Tool Messages: from ConversationState.events "
        "(append-only Event Log), delivered to LLM in Chat Completions style."
    )
    bullet(
        "Tools Schema: each Tool's name, description, parameters; "
        "when LLMSecurityAnalyzer enabled, non-read-only tools auto-append required "
        "security_risk (LOW/MEDIUM/HIGH)."
    )
    cite(
        "Security Architecture: security_risk injection rules",
        "https://docs.openhands.dev/sdk/arch/security.md",
    )

    h("2. AgentContext and Dynamic Rendering (Skills)", 2)
    p("AgentContext injects two skill types and prefix/suffix into the LLM (architecture terminology):")
    bullet(
        "repo skills: Always included → enters System Prompt "
        "(project conventions, permanent context; product recommends root AGENTS.md)."
    )
    bullet(
        "knowledge skills: Trigger words/patterns → when triggered enter User Messages side context "
        "(domain knowledge, specialized behavior)."
    )
    bullet("System prompt prefix/suffix: can append per session.")
    cite(
        "Agent Architecture > Agent Context",
        "https://docs.openhands.dev/sdk/arch/agent.md",
    )
    p("Product-layer skill loading model (Skills Overview):")
    bullet("Always-on: AGENTS.md (and GEMINI.md / CLAUDE.md model variants) injected at session start.")
    bullet(
        "On-demand: user keyword trigger, or agent sees summary then reads full SKILL.md "
        "(progressive disclosure)."
    )
    bullet("Path-triggered: deterministic injection when reading/writing/creating files matching glob.")
    bullet(
        "Load priority: .agents/skills/ > .openhands/skills/ (deprecated) > "
        ".openhands/microagents/ (deprecated); project skills override user skills."
    )
    cite(
        "Skills Overview",
        "https://docs.openhands.dev/overview/skills.md",
    )
    bullet("Repository customization: .openhands/setup.sh (runs at session start); .openhands/hooks.json (lifecycle hooks).")
    cite(
        "Repository Customization",
        "https://docs.openhands.dev/openhands/usage/customization/repository.md",
    )

    h("3. CodeActAgent Behavior Paradigm (Product Docs)", 2)
    p("OpenHands main Agent described as CodeActAgent; each turn can choose between two action types:")
    bullet("Converse: natural language communication, clarification, or confirmation with user.")
    bullet(
        "CodeAct: execute Linux bash or Python simulated via bash; "
        "actual tool names exposed to LLM still follow V1 registered names (terminal / file_editor, etc.)."
    )
    cite(
        "Main Agent and Capabilities",
        "https://docs.openhands.dev/openhands/usage/agents.md",
    )

    h("4. Prompt-Side Collection (Research)", 2)
    bullet(
        "SDK callback: call event.to_llm_message() on LLMConvertibleEvent "
        "to collect raw LLM messages (Persistence / MCP examples same pattern)."
    )
    bullet("OpenTelemetry: traces agent.step, tool, LLM, conversation lifecycle.")
    bullet("Can also point LLM base_url to local proxy for full interception (not sole official recommended path).")

    h("Prompt Composition · Verification Summary", 2)
    bullet(
        "Retained: system+messages+tools; AgentContext repo/knowledge; "
        "AGENTS.md always-on; Skills priority."
    )
    bullet(
        "Removed: copying OpenClaw system prompt \"dynamically rendered from local config\" verbatim as OpenHands mechanism; "
        "removed fabricated independent memory_search in tools schema."
    )
    bullet(
        "Note: architecture diagram writes knowledge → User Messages; product writes on-demand skills. "
        "Both are on-demand injection; do not describe as a second undocumented API."
    )

    # ══════════════════════════════════════════════════════════════
    # Memory and Persistence
    # ══════════════════════════════════════════════════════════════
    h("Memory and Persistence", 1)
    p(
        "OpenHands provides no OpenClaw-style MEMORY.md / memory/YYYY-MM-DD.md dedicated file system, "
        "nor memory_search / memory_get dedicated tools. "
        "\"Memory\" consists of Event Log (in-session) + Persistence/Resume (cross-session) + "
        "AGENTS.md/Skills (repository-level long-term instructions) + optional Condenser (compression)."
    )

    h("1. Conversation Event Log (In-Session Context)", 2)
    bullet(
        "ConversationState.events: append-only, records user / assistant / Action / Observation."
    )
    bullet(
        "Agent has no mutable business state across steps: each step() reads history from Event Log then reasons."
    )
    cite(
        "Conversation Architecture; Agent Architecture",
        "https://docs.openhands.dev/sdk/arch/conversation.md",
    )

    h("2. Context Condenser (Optional, Explicit Configuration Required)", 2)
    p(
        "Default implementation is LLMSummarizingCondenser: when event count exceeds max_size, "
        "LLM summarizes old history, keeps recent events and keep_first (usually includes system / initial user). "
        "Must configure on Agent(..., condenser=...) to enable; not unconditionally on by default."
    )
    cite(
        "Context Condenser",
        "https://docs.openhands.dev/sdk/guides/context-condenser.md",
    )

    h("3. Cross-Session Persistence (CLI / SDK)", 2)
    bullet(
        "CLI: ~/.openhands/conversations/<id>/conversation.json; "
        "openhands --resume / --resume --last."
    )
    cite(
        "Resume Conversations",
        "https://docs.openhands.dev/openhands/usage/cli/resume.md",
    )
    bullet(
        "SDK: persistence_dir + conversation_id → base_state.json + events/event-*.json."
    )
    bullet(
        "SDK write mechanism (Persistence guide): modifying public ConversationState fields "
        "triggers __setattr__ immediate base state serialization; events appended incrementally per file. "
        "Architecture overview mentions Debounced writes for Persistence service — Persistence guide takes precedence."
    )
    bullet(
        "Persisted content includes: full event trajectory, Agent/tool/MCP config, execution state, "
        "tool output, statistics, workspace, activated Skills, Secrets, agent custom runtime state, etc."
    )
    cite(
        "Persistence",
        "https://docs.openhands.dev/sdk/guides/convo-persistence.md",
    )

    h("4. Repository-Level Long-Term Context (Not Session Memory Files)", 2)
    bullet("AGENTS.md: officially recommended permanent agent context (conversation start always-on).")
    bullet(".agents/skills/: on-demand skills; summary + full SKILL.md progressive disclosure.")
    bullet("Organization / Global Skills: organization or community shared (extensions repo, etc.).")
    cite(
        "Skills Overview",
        "https://docs.openhands.dev/overview/skills.md",
    )

    h("5. Indirect \"Memory\" Writes via Files/Terminal", 2)
    bullet("file_editor / terminal can read/write arbitrary workspace paths; long-term notes require project convention.")
    bullet(
        "OpenHands does not guarantee main session auto-loads a fixed MEMORY.md; "
        "for equivalent behavior, use AGENTS.md or Skill to explicitly instruct agent to read/write specified paths."
    )

    h("Memory and Persistence · Verification Summary", 2)
    bullet(
        "Retained: Event Log; optional Condenser; CLI/SDK Persistence; AGENTS.md/Skills; "
        "no dedicated memory tools."
    )
    bullet(
        "Removed: OpenClaw Daily notes / MEMORY.md \"ONLY load in main session\" rules "
        "written as OpenHands default behavior; removed memory_search / memory_get."
    )
    bullet(
        "Consistent with Session section above: Automation each run new sandbox/new conversation, "
        "cross-run \"memory\" does not auto-continue unless Resume/continue same session or rely on repository-level AGENTS.md."
    )

    h("Primary Reference Links", 1)
    for link in [
        "https://docs.openhands.dev/llms.txt",
        "https://docs.openhands.dev/sdk/arch/conversation.md",
        "https://docs.openhands.dev/sdk/arch/agent.md",
        "https://docs.openhands.dev/sdk/arch/tool-system.md",
        "https://docs.openhands.dev/sdk/arch/security.md",
        "https://docs.openhands.dev/sdk/guides/convo-persistence.md",
        "https://docs.openhands.dev/sdk/guides/context-condenser.md",
        "https://docs.openhands.dev/sdk/guides/mcp.md",
        "https://docs.openhands.dev/sdk/guides/agent-browser-use.md",
        "https://docs.openhands.dev/sdk/guides/custom-tools.md",
        "https://docs.openhands.dev/sdk/guides/security.md",
        "https://docs.openhands.dev/sdk/guides/task-tool-set.md",
        "https://docs.openhands.dev/overview/skills.md",
        "https://docs.openhands.dev/openhands/usage/agents.md",
        "https://docs.openhands.dev/openhands/usage/automations/overview.md",
        "https://docs.openhands.dev/openhands/usage/automations/creating-automations.md",
        "https://docs.openhands.dev/openhands/usage/automations/event-automations.md",
        "https://docs.openhands.dev/openhands/usage/cli/headless.md",
        "https://docs.openhands.dev/openhands/usage/cli/terminal.md",
        "https://docs.openhands.dev/openhands/usage/cli/resume.md",
        "https://docs.openhands.dev/openhands/usage/customization/repository.md",
        "https://docs.openhands.dev/openhands/usage/sandboxes/overview.md",
        "https://docs.openhands.dev/openhands/usage/advanced/search-engine-setup.md",
        "https://docs.openhands.dev/openhands/usage/agent-canvas/overview.md",
        "https://github.com/OpenHands/extensions/pull/252",
    ]:
        bullet(link)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
