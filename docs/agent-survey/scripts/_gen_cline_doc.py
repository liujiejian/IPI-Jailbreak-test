# -*- coding: utf-8 -*-
"""Generate Cline research notes (same structure as OpenHands / Hermes / Dify / AutoGPT).

Paper track: Session / 任务模式 / web_fetch(security notice) / tools,
plus Prompt 组成 & 记忆持久化. 5-round anti-hallucination vs docs.cline.bot.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "Cline.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label}（来源：{url}）")

    # ── Title ──
    h("Cline", 0)
    p(
        "本文档结构对齐 OpenHands.docx / Hermes.docx / Claude-Code.docx / Dify.docx / AutoGPT.docx："
        "Session、任务模式、web_fetch（及 Security Notice 对照）、Tools；"
        "文末附 Prompt 组成、记忆与持久化机制，以及五轮防幻觉核对。"
        "对象为开源 coding agent Cline（VS Code / JetBrains 扩展、CLI、Kanban、SDK 共享 agent core）。"
    )
    p("官方文档索引：https://docs.cline.bot/llms.txt")
    p(
        "定位：编辑器与终端中的 AI coding agent — 读写文件、跑命令、取网页/浏览器、"
        "MCP 扩展；默认每步需用户批准（Auto Approve / YOLO 可放宽）。"
    )
    cite("Cline Overview", "https://docs.cline.bot/cline-overview")
    p(
        "执行链路（研究视角）：User → agentic loop（Plan/Act）→ "
        "tools（read_files / editor / bash / fetch_web / MCP…）→ "
        "结果回灌会话 →（可选）Checkpoints 快照工作区。"
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "ClineCore / Hub 体系下，Session 是带完整对话历史、tool call 记录与元数据的持久实体，"
        "有 sessionId；可跨客户端 attach，断开后仍可由 Spoke 继续执行。"
    )
    cite(
        "Hub-Spoke Architecture · Session Persistence",
        "https://docs.cline.bot/sdk/architecture/hub-spoke",
    )
    cite("ClineCore", "https://docs.cline.bot/sdk/clinecore")

    h("1. 存储位置", 2)
    bullet(
        "默认：~/.cline/data/sessions/ — sessions.db（SQLite 索引）+ "
        "[session-id].json（权威会话快照）。"
    )
    bullet(
        "配置根：~/.cline/（全局）与项目 .cline/；"
        "可用 --config / CLINE_DATA_DIR / --data-dir 隔离。"
    )
    bullet(
        "ClineCore.start() 返回 session.sessionId、manifestPath、messagesPath 等。"
    )
    cite("Config", "https://docs.cline.bot/getting-started/config")

    h("2. Hub-Spoke 与多客户端", 2)
    bullet(
        "Hub：本机单例 daemon（默认 127.0.0.1:25463），协调会话与审批路由，不跑 agent loop。"
    )
    bullet("Spoke：工人进程跑 agent loop 与工具；Client（CLI/IDE）可来去。")
    bullet(
        "backendMode：auto | hub | remote | local；"
        "local = 进程内 + 本地存储，无共享会话。"
    )
    bullet(
        "参与者角色：creator / participant / observer；"
        "CLI 可用 --id <session-id> 恢复；cline history 管理历史。"
    )
    cite("CLI Reference", "https://docs.cline.bot/cli/cli-reference")

    h("3. Task 与 Session 的关系（产品语义）", 2)
    bullet(
        "用户界面常称「Task」；slash `/newtask` 会在蒸馏上下文后开启新任务"
        "（新干净上下文窗口）。"
    )
    bullet(
        "研究映射：同一 sessionId 下的多轮消息 ≈ Claude Code 项目会话；"
        "Schedule/Kanban 每次触发更接近独立或挂接新 session（见任务模式）。"
    )

    h("4. 与其它 Agent 对照（非官方术语）", 2)
    bullet("Claude Code ~/.claude/projects/*.jsonl ≈ Cline ~/.cline/data/sessions/*.json。")
    bullet("Hermes SessionDB ≈ Cline sessions.db + JSON 快照。")
    bullet("Dify conversation_id API 字段 ≠ Cline sessionId（协议不同）。")

    # ══════════════════════════════════════════════════════════════
    # 任务模式
    # ══════════════════════════════════════════════════════════════
    h("任务模式", 1)

    h("1. Plan & Act（双模式）", 2)
    bullet(
        "Plan：可读代码库、搜索、讨论策略；不能改文件、不能执行命令。"
    )
    bullet(
        "Act：保留完整 Plan 对话上下文后执行改文件/跑命令；可多次来回切换。"
    )
    bullet("可分别为 Plan/Act 配置不同模型；CLI `-p/--plan` 启动 Plan。")
    bullet("`/deep-planning`：更深规划 → implementation_plan.md → 新实现任务。")
    cite("Plan & Act Mode", "https://docs.cline.bot/core-workflows/plan-and-act")

    h("2. 产品入口", 2)
    bullet("IDE：VS Code / JetBrains 扩展（及跨编辑器 ACP：Cursor、Windsurf、Zed、Neovim 等）。")
    bullet("CLI 交互：`cline` / TUI（`-i`）；`--zen` 后台 hub 会话。")
    bullet(
        "Headless：`--json`、stdin pipe、stdout 重定向 → CI/脚本；"
        "`--auto-approve`、`--timeout`。"
    )
    bullet("Kanban：多智能体任务板 + 每卡 worktree。")
    cite("CLI Overview", "https://docs.cline.bot/usage/cli-overview")

    h("3. Auto Approve / YOLO", 2)
    bullet(
        "按工具类别审批：读/写项目或全部文件、安全/全部命令、"
        "Use the browser、Use MCP、通知等。"
    )
    bullet(
        "YOLO：自动批准一切（含 Plan→Act 切换）；官方标危险，建议隔离环境。"
    )
    cite(
        "Auto Approve & YOLO Mode",
        "https://docs.cline.bot/features/auto-approve",
    )

    h("4. Scheduling / Teams / Subagents", 2)
    bullet(
        "Schedule：cron 经 hub；CLI/SDK/Kanban 可用；"
        "文档标明暂不适用于 VS Code / JetBrains 扩展。"
    )
    bullet("Agent Teams：多 agent 协作；Subagents：并行调研不填满主上下文。")
    cite("Scheduling", "https://docs.cline.bot/cli/scheduling")
    cite("Subagents", "https://docs.cline.bot/features/subagents")

    h("任务模式 · 核对摘要", 2)
    bullet(
        "保留：Plan/Act；IDE/CLI/Headless/Kanban；Auto Approve/YOLO；"
        "Schedule（面受限）；Teams/Subagents。"
    )
    bullet("剔除：OpenClaw Heartbeat；把 Schedule 写成扩展内建能力。")

    # ══════════════════════════════════════════════════════════════
    # web_fetch
    # ══════════════════════════════════════════════════════════════
    h("web_fetch（及 Security Notice）", 1)
    p(
        "结论：官方工具参考将外部检索内置工具命名为 fetch_web"
        "（HTTP + HTML→markdown），归类 External retrieval。"
        "公开文档未记载 OpenClaw 式 SECURITY NOTICE / EXTERNAL_UNTRUSTED_CONTENT 包装，"
        "也未声明 Claude Code WebFetch 的「孤立小模型抽取 + ~15min cache」行为。"
        "另有独立的 Browser 能力面（Auto Approve「Use the browser」）；"
        "Overview 写「use a browser」。"
    )

    h("1. 文档中的内置名（注意命名漂移）", 2)
    bullet(
        "tools-reference（对齐 ClineCore）："
        "`fetch_web` — HTTP requests with HTML-to-markdown conversion。"
    )
    bullet(
        "sdk/tools 表另列 `fetch_web_content`（Fetch web content），"
        "与上表命名不完全一致——研究冻结时并列记录，勿只认一个为绝对真理。"
    )
    bullet(
        "旧扩展/示例中的 XML 风格名（read_file / execute_command 等）"
        "与当前 runtime 名不同；文档要求对齐新名。"
    )
    cite(
        "Tools Reference · Built-In",
        "https://docs.cline.bot/tools-reference/all-cline-tools",
    )
    cite("SDK Tools", "https://docs.cline.bot/sdk/tools")

    h("2. Browser vs fetch_web", 2)
    bullet(
        "Auto Approve 单独一项「Use the browser」："
        "Browser tool for web fetching and searching——"
        "产品语义上与通用文件/终端审批并列。"
    )
    bullet(
        "公开用户文档未给出 browser 工具的固定内部函数名列表"
        "（如历史社区讨论中的 browser_action）；"
        "论文勿捏造未在 docs.cline.bot 出现的参数协议。"
    )
    bullet("MCP 可再叠加浏览器/抓取类第三方工具。")

    h("3. 审批与启用条件", 2)
    bullet("默认设计：工具调用前需用户批准（除非 Auto Approve / YOLO）。")
    bullet(
        "SDK toolPolicies / requestToolApproval 可细粒度控制各工具是否 autoApprove。"
    )
    bullet(
        "扩展实现侧（非 Mintlify 主文档条款）：历史源码中有面向 Cline 账号的 "
        "web_fetch API 通路与 clineWebToolsEnabled 开关——"
        "标记为「实现细节 / 待官方文档固化」，不写成 llms.txt 正式契约。"
    )

    h("4. Security Notice 对照", 2)
    bullet(
        "核对 docs：未找到 SECURITY NOTICE / untrusted_tool_result / "
        "EXTERNAL_UNTRUSTED_CONTENT 默认包装说明。"
    )
    bullet(
        "安全相关公开面："
        "审批门禁、.clineignore、CLINE_COMMAND_PERMISSIONS、"
        "规则/hooks/plugins「仅信任来源」、Enterprise MCP allowlist / YOLO 管控。"
    )
    bullet(
        "企业 Prompt Storage 可把对话备份到 S3/R2（合规），"
        "不是对网页内容的 untrusted 标记。"
    )

    h("威胁模型映射（研究，非官方）", 2)
    bullet(
        "外网页 markdown/HTML→模型上下文→后续 editor/bash/MCP："
        "Indirect Prompt Injection / Goal Hijacking 论文面。"
    )
    bullet(
        "Browser 交互面可能更大（导航/表单）；与单纯 fetch_web 应分 ablation。"
    )
    bullet(
        "对比：OpenClaw web_fetch+NOTICE；Claude Code WebFetch+隔离抽取；"
        "Hermes web_extract+untrusted 标签；"
        "Cline ≈ fetch_web(+browser) + 审批，文档层无 NOTICE。"
    )

    code(
        """研究示意（非官方消息绑定）：
User: Summarize https://example.com/docs
→ tool: fetch_web { url: "https://example.com/docs" }
← tool result: (HTML→markdown 正文，文档未写 SECURITY NOTICE 外壳)
→ 模型可能继续 editor/bash/MCP
说明：具体 tool schema 以实现/SDK 为准；勿写成 OpenClaw EXTERNAL_UNTRUSTED_CONTENT。"""
    )

    h("web_fetch · 核对摘要", 2)
    bullet(
        "保留：fetch_web（及文档中的 fetch_web_content 别名表）；"
        "browser 审批类别；无公开 SECURITY NOTICE；HTML→markdown。"
    )
    bullet(
        "剔除：默认 15min cache+孤立抽取；"
        "剔除把 OpenClaw NOTICE 文案写成 Cline 内置。"
    )

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "ClineCore 内置工具 + MCP（.cline/mcp.json）+ "
        "自定义工具/插件（文档：SDK/CLI/Kanban；扩展侧暂不全面适用）。"
    )

    h("1. ClineCore 内置（tools-reference）", 2)
    bullet("`bash` — 执行 shell。")
    bullet("`editor` — 查看与编辑文件。")
    bullet("`read_files` — 批量读文件。")
    bullet("`apply_patch` — 统一 diff 补丁。")
    bullet("`search` — ripgrep 代码搜索。")
    bullet("`fetch_web` — 外网 HTTP + HTML→markdown。")
    bullet("`ask_question` — 向用户提问。")
    p(
        "SDK Tools 页另列：search_codebase、run_commands、fetch_web_content、"
        "skills、submit_and_exit 等——同一 harness 文档存在命名演进，引用时注明来源页。"
    )

    h("2. MCP / Plugins / Skills", 2)
    bullet("MCP：与内置工具一并加载；Enterprise 可 allowlist。")
    bullet("Plugins：注册工具、hooks、改行为（SDK/CLI/Kanban）。")
    bullet("Skills：模块化 SKILL.md；可用 slash `/skill-name` 触发。")
    cite("MCP Overview", "https://docs.cline.bot/mcp/mcp-overview")
    cite("Skills", "https://docs.cline.bot/customization/skills")

    h("3. 访问控制", 2)
    bullet(".clineignore：限制可访问路径。")
    bullet("CLINE_COMMAND_PERMISSIONS：allow/deny 命令模式与重定向。")
    bullet("沙箱：CLINE_SANDBOX / CLINE_SANDBOX_DATA_DIR。")

    h("Tools · 核对摘要", 2)
    bullet("保留：ClineCore 表 + MCP + 审批策略；命名以官方表为准并标注漂移。")
    bullet("剔除：伪造固定 PascalCase WebFetch 为唯一名。")

    # ══════════════════════════════════════════════════════════════
    # 五轮防幻觉
    # ══════════════════════════════════════════════════════════════
    h("五轮防幻觉核对记录", 1)

    h("Round 1 — Session", 2)
    bullet("核对：hub-spoke；clinecore；config；cli-reference --id/history。")
    bullet(
        "保留：sessions.db + [id].json；sessionId；hub 多客户端；local vs hub。"
    )
    bullet("剔除：Claude Code JSONL 路径；Dify conversation_id 字段名。")

    h("Round 2 — 任务模式", 2)
    bullet("核对：plan-and-act；cli-overview；auto-approve；scheduling；kanban。")
    bullet("保留：Plan/Act；Headless；YOLO；Schedule 面限制。")
    bullet("剔除：扩展内可用 schedule（与文档 Warning 矛盾）。")

    h("Round 3 — web_fetch / 安全", 2)
    bullet("核对：all-cline-tools；sdk/tools；auto-approve browser；overview。")
    bullet(
        "保留：fetch_web；HTML→markdown；browser 类别；无 SECURITY NOTICE 文档。"
    )
    bullet(
        "剔除：OpenClaw NOTICE；Claude Code 15min cache 声明；"
        "未文档化的 browser_action 参数表。"
    )

    h("Round 4 — Tools", 2)
    bullet("核对：all-cline-tools；sdk/tools；mcp；plugins/skills；clineignore。")
    bullet("保留：内置七件套（reference 表）+ MCP/插件；旧名废弃说明。")
    bullet("剔除：仅认 sdk 页或仅认 reference 页而不写漂移。")

    h("Round 5 — Prompt / Memory 交叉复检", 2)
    bullet("核对：cline-rules；memory-bank；using-commands；checkpoints；clinecore systemPrompt。")
    bullet(
        "冻结："
        "(1) Session = sessionId + ~/.cline/data/sessions；"
        "(2) 模式 = Plan/Act + 入口（IDE/CLI/Headless/Kanban）+ 审批级别；"
        "(3) 外网 = fetch_web(+browser/MCP)，无文档 SECURITY NOTICE；"
        "(4) Prompt = systemPrompt + Rules + Skills；"
        "(5) 跨会话项目记忆 = Memory Bank 方法论（规则驱动 markdown），"
        "非 Hermes 自动 MEMORY.md 装载。"
    )
    bullet(
        "开放问题：fetch_web 与 fetch_web_content 运行时是否同一实现；"
        "browser 工具官方 schema；扩展 web_fetch API 是否仍为主路径。"
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt 组成
    # ══════════════════════════════════════════════════════════════
    h("Prompt 组成", 1)
    p(
        "运行时由 harness 组装：可配置 systemPrompt + 自动装载的 Rules/Skills 指令 "
        "+ 用户消息/@文件上下文 + tool 结果。完整系统模板未完全公开。"
    )

    h("1. systemPrompt（SDK/ClineCore）", 2)
    bullet("start/config 可传 systemPrompt（示例如 helpful coding assistant）。")
    bullet("CLI 侧亦有覆盖系统提示相关能力（见 CLI Reference 更新）。")

    h("2. Rules（持久指令）", 2)
    bullet(
        "工作区 .clinerules/、全局 Rules 目录、以及自动检测 "
        ".cursorrules / .windsurfrules / AGENTS.md。"
    )
    bullet("工作区与全局合并；冲突时工作区优先。")
    cite("Rules", "https://docs.cline.bot/customization/cline-rules")

    h("3. Skills 与 slash", 2)
    bullet("Skills 模块说明进入能力指令；`/skill` 触发。")
    bullet(
        "`/deep-planning`、`/newtask`、`/smol`（`/compact`）、`/newrule` "
        "管理规划与上下文。"
    )
    cite("Using Commands", "https://docs.cline.bot/core-workflows/using-commands")

    h("4. 用户上下文注入", 2)
    bullet("@ 提及与拖拽文件进对话（Adding Context）。")
    bullet("CLI：`@./image.png`、管道 stdin 作为任务上下文。")

    h("Prompt 组成 · 核对摘要", 2)
    bullet("保留：systemPrompt + Rules + Skills + @context + tool results。")
    bullet("剔除：CLAUDE.md 自动体系或 Hermes SOUL.md 作为 Cline 默认文件名。")

    # ══════════════════════════════════════════════════════════════
    # 记忆与持久化
    # ══════════════════════════════════════════════════════════════
    h("记忆与持久化机制", 1)
    p(
        "分层："
        "(A) 会话级 JSON/SQLite 持久化；(B) 上下文压缩 slash；"
        "(C) Memory Bank 项目文档方法论；(D) Checkpoints 工作区影子 Git；"
        "(E) 可选企业 Prompt Storage。"
    )

    h("1. Session 消息持久化", 2)
    bullet("完整 history + tool records 写入 session JSON；list/get/readMessages API。")
    bullet("Hub 下跨进程/客户端恢复；与 MEMORY.md 文件体系无关。")

    h("2. 上下文窗口管理", 2)
    bullet("`/smol`：同任务内压缩对话。")
    bullet("`/newtask`：蒸馏后新任务、净上下文。")
    bullet("CLI compaction 相关标志见 CLI Reference（agentic/basic/off 等，以当页为准）。")

    h("3. Memory Bank（最佳实践，非内置 DB）", 2)
    bullet(
        "通过 Rules 注入的方法论：项目 memory-bank/*.md"
        "（projectbrief、productContext、activeContext、systemPatterns、"
        "techContext、progress）。"
    )
    bullet(
        "指令强调「会话间记忆重置，须靠 Memory Bank 文件」——"
        "需用户 initialize/update；不是自动 MEMORY.md loader。"
    )
    cite("Memory Bank", "https://docs.cline.bot/best-practices/memory-bank")

    h("4. Checkpoints（代码状态，非语义记忆）", 2)
    bullet("影子 Git：每次工具改文件/命令后快照；可 Restore Files / Task / Both。")
    bullet("对话可保留而代码回滚——服务实验，不是跨会话知识库。")
    cite("Checkpoints", "https://docs.cline.bot/core-workflows/checkpoints")

    h("5. 其它", 2)
    bullet("Skills/Rules/Hooks/cron 规格文件本身即持久配置记忆。")
    bullet("Enterprise Prompt Storage：对话备份至 S3/R2。")

    h("记忆与持久化 · 核对摘要", 2)
    bullet(
        "保留：session JSON；/smol·/newtask；Memory Bank 方法论；Checkpoints；"
        "企业备份可选。"
    )
    bullet(
        "剔除：把 Memory Bank 写成开箱即自动的内置向量记忆；"
        "剔除 Hermes ~/.hermes/memories/MEMORY.md 路径。"
    )

    h("主要参考链接", 1)
    for link in [
        "https://docs.cline.bot/llms.txt",
        "https://docs.cline.bot/cline-overview",
        "https://docs.cline.bot/sdk/architecture/hub-spoke",
        "https://docs.cline.bot/sdk/clinecore",
        "https://docs.cline.bot/getting-started/config",
        "https://docs.cline.bot/core-workflows/plan-and-act",
        "https://docs.cline.bot/core-workflows/checkpoints",
        "https://docs.cline.bot/core-workflows/using-commands",
        "https://docs.cline.bot/features/auto-approve",
        "https://docs.cline.bot/tools-reference/all-cline-tools",
        "https://docs.cline.bot/sdk/tools",
        "https://docs.cline.bot/mcp/mcp-overview",
        "https://docs.cline.bot/customization/cline-rules",
        "https://docs.cline.bot/customization/skills",
        "https://docs.cline.bot/best-practices/memory-bank",
        "https://docs.cline.bot/usage/cli-overview",
        "https://docs.cline.bot/cli/cli-reference",
        "https://docs.cline.bot/cli/scheduling",
        "https://docs.cline.bot/customization/clineignore",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
