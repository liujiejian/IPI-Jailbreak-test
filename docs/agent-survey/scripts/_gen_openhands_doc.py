# -*- coding: utf-8 -*-
"""Generate OpenHands research notes: Session / task modes / web fetch / tools.

Verified against official docs (docs.openhands.dev) with a 5-round
anti-hallucination pass. Structure mirrors Openclaw.docx for the paper track:
Agent Security — Risks from External Web Fetching.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "OpenHands.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label}（来源：{url}）")

    # ── Title ──
    h("OpenHands", 0)
    p(
        "本文档对齐论文任务「Session、任务模式、web_fetch (security notice)、tools」；"
        "文末另附「Prompt 组成」「记忆与持久化机制」。"
        "对照 OpenClaw 示例结构，调研 OpenHands V1 官方文档，"
        "经防幻觉核对后剔除错误表述并保留可追溯来源。"
    )
    p("官方文档索引：https://docs.openhands.dev/llms.txt")
    p(
        "执行链路（研究视角）：User task → Agent(LLM → call tool → tool/observation → LLM) → User。"
        "LLM 侧可见 system / user / assistant / tool messages 与 tools schema。"
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "OpenHands 没有 OpenClaw 式的 Main / Isolated / session:xxx 命名；"
        "官方对应实体是 Conversation（会话）。"
        "Session 决定运行时上下文；不同任务模式会新建或复用 Conversation。"
    )
    cite(
        "Conversation Architecture",
        "https://docs.openhands.dev/sdk/arch/conversation.md",
    )

    h("1. Conversation 工厂：Local vs Remote", 2)
    p(
        "Conversation(agent, workspace) 按 workspace 类型自动选择实现："
    )
    bullet(
        "LocalConversation：workspace 为 str 或 LocalWorkspace → agent 在本进程内执行；"
        "通信为直接函数调用；适合开发 / CLI。"
    )
    bullet(
        "RemoteConversation：workspace 为 RemoteWorkspace → 经 HTTP/WebSocket 连接 Agent Server，"
        "在隔离容器中执行；适合生产 / Web。"
    )
    p("切换部署模式通常只需更换 workspace 类型，API 表面保持一致。")
    cite(
        "Factory Pattern / Local vs Remote Execution",
        "https://docs.openhands.dev/sdk/arch/conversation.md",
    )

    h("2. 会话内状态：Event Log", 2)
    bullet(
        "ConversationState.events 为 append-only Event Log"
        "（user / assistant / Action / Observation 等）。"
    )
    bullet(
        "Agent 为 reasoning-action loop：每轮从 Event Log 读取历史再查询 LLM"
        "（跨 step 的可变业务状态不由 Agent 本身长期持有）。"
    )
    cite(
        "Conversation Architecture > State Management",
        "https://docs.openhands.dev/sdk/arch/conversation.md",
    )

    h("3. 跨会话恢复（Resume / Persistence）", 2)
    bullet(
        "CLI：对话保存在 ~/.openhands/conversations/<id>/conversation.json；"
        "openhands --resume、--resume --last、--resume <id> 恢复。"
    )
    cite(
        "Resume Conversations",
        "https://docs.openhands.dev/openhands/usage/cli/resume.md",
    )
    bullet(
        "SDK：设置 persistence_dir + conversation_id 后，"
        "base_state.json 存配置/状态，events/event-*.json 增量存事件。"
    )
    bullet(
        "SDK 实现细节（持久化指南）：修改 ConversationState 公共字段时，"
        "经自定义 __setattr__ 立即序列化 base state（非手动 save）。"
        "架构总览页对 Persistence 辅助服务写过 Debounced writes；"
        "以 Persistence 指南的「immediately when they occur」为准。"
    )
    cite(
        "Persistence 官方指南",
        "https://docs.openhands.dev/sdk/guides/convo-persistence.md",
    )

    h("4. 与 OpenClaw Session 模型的研究映射（非官方术语）", 2)
    bullet(
        "OpenClaw Main session（共享上下文）≈ OpenHands 持续 Conversation + CLI/SDK Resume。"
    )
    bullet(
        "OpenClaw Isolated job（新鲜 session）≈ OpenHands Automation 每次 run："
        "新建 sandbox + 新 Conversation（可事后 review / continue）。"
    )
    bullet(
        "OpenClaw Custom session:xxx（跨 run 持久）≈ "
        "同一 conversation_id 的 Persistence / Resume；"
        "或 Automation 运行结果会话上的 continue。"
    )
    bullet(
        "OpenClaw Heartbeat：OpenHands 公开产品文档未描述等价机制；"
        "周期任务用 Automations（cron），事件驱动用 Event-based Automations。"
    )

    # ══════════════════════════════════════════════════════════════
    # 任务模式
    # ══════════════════════════════════════════════════════════════
    h("任务模式", 1)
    p(
        "对应 OpenClaw 的 Cron / Heartbeat / Automation："
        "OpenHands 的任务形态决定 context 是否共享、是否隔离 sandbox、以及确认策略。"
    )

    h("1. 交互式主会话（Interactive Conversation）", 2)
    bullet("CLI Terminal：openhands（默认需确认）；支持 -t / -f 启动任务。")
    bullet(
        "确认模式：默认确认；--always-approve 全自动；--llm-approve 走 LLM 安全分析。"
    )
    cite(
        "Terminal (CLI)",
        "https://docs.openhands.dev/openhands/usage/cli/terminal.md",
    )
    bullet(
        "Agent Canvas / Cloud Web UI：浏览器内 Conversation；"
        "backend/sandbox/workspace/agent 可切换。"
    )
    cite(
        "Agent Canvas Overview",
        "https://docs.openhands.dev/openhands/usage/agent-canvas/overview.md",
    )
    bullet(
        "SDK Conversation：send_message + run/step；可 Pause/Resume、Fork、"
        "Send Message While Running、Goal Completion Loop。"
    )

    h("2. Headless（无 UI 批处理 / CI）", 2)
    p("openhands --headless -t \"...\" 或 -f task.txt。")
    bullet("必须提供任务（--task / --file）。")
    bullet(
        "强制 always-approve：不可改为 --llm-approve；"
        "适合脚本与 CI，安全边界需由调用方承担。"
    )
    bullet("可选 --json 输出 JSONL 事件流。")
    cite(
        "Headless Mode",
        "https://docs.openhands.dev/openhands/usage/cli/headless.md",
    )

    h("3. Automations（定时 / Prompt-based & Plugin-based）", 2)
    p(
        "后台按调度执行 AI 任务。每次 run："
        "① 创建 fresh sandbox；② 执行 prompt；③ 保存 conversation 供审查或继续。"
    )
    bullet("Prompt-based：自然语言描述任务与 cron（最常见）。")
    bullet(
        "Plugin-based：挂载 OpenHands extensions 插件"
        "（额外 skills / MCP / commands）。"
    )
    bullet("默认 timeout 10 分钟，最长 30 分钟。")
    bullet(
        "环境能力：terminal、文件操作、Settings 中的 LLM/Secrets、"
        "MCP、网络 HTTP、登录后的 GitHub/GitLab/Bitbucket 凭证。"
    )
    cite(
        "Automations Overview",
        "https://docs.openhands.dev/openhands/usage/automations/overview.md",
    )
    cite(
        "Creating Automations",
        "https://docs.openhands.dev/openhands/usage/automations/creating-automations.md",
    )

    h("4. Event-based Automations（事件驱动）", 2)
    bullet(
        "GitHub 内置事件：pull_request / issues / issue_comment / push / release 等；"
        "可用 JMESPath 过滤。"
    )
    bullet("自定义 webhook：Linear / Stripe / Slack 等（先注册 webhook，再创建 automation）。")
    cite(
        "Event-Based Automations",
        "https://docs.openhands.dev/openhands/usage/automations/event-automations.md",
    )

    h("5. 调度策略对照（研究笔记）", 2)
    p("OpenClaw：精确时间/隔离 → Cron；需要完整 session 上下文 → Heartbeat。")
    p("OpenHands 对应：")
    bullet("定时或事件、需隔离审查 → Automations（每次新 Conversation）。")
    bullet("需连续人机协作上下文 → 主 Conversation + Resume。")
    bullet("无 UI / CI → Headless（always-approve）。")

    # ══════════════════════════════════════════════════════════════
    # web_fetch / Security Notice
    # ══════════════════════════════════════════════════════════════
    h("web_fetch（及 Security Notice）", 1)
    p(
        "【结论】OpenHands V1 没有与 OpenClaw 等价的内置 web_fetch 工具，"
        "也没有框架级自动注入的 SECURITY NOTICE / "
        "<<<EXTERNAL_UNTRUSTED_CONTENT>>> 包裹机制。"
        "外部内容经多条路径进入 Observation / tool message。"
    )

    h("1. 路径 A：MCP Fetch（官方示例中最接近 web_fetch）", 2)
    p(
        "通过 Agent.mcp_config 注册 mcp-server-fetch；"
        "agent 初始化时自动发现 MCP tools。"
        "可用 filter_tools_regex 限制暴露的工具。"
    )
    code(
        """mcp_config = {
    "mcpServers": {
        "fetch": {
            "command": "uvx",
            "args": ["mcp-server-fetch"]
        }
    }
}
agent = Agent(llm=llm, tools=tools, mcp_config=mcp_config)"""
    )
    cite(
        "MCP Integration 示例（含 fetch）",
        "https://docs.openhands.dev/sdk/guides/mcp.md",
    )
    cite(
        "Persistence 示例：MCP fetch 读取远程 URL 后写入文件",
        "https://docs.openhands.dev/sdk/guides/convo-persistence.md",
    )
    p(
        "研究提示：该路径的 tool 返回内容是否带 untrusted 标注，"
        "取决于 MCP server 实现，而非 OpenHands 核心框架。"
    )

    h("2. 路径 B：BrowserToolSet（browser-use）", 2)
    p(
        "可选工具集：导航、点击、填表、提取页面内容；"
        "观察结果作为 Observation 回到 LLM。"
        "官方示例与 TerminalTool / FileEditorTool 组合做 web research。"
    )
    cite(
        "Browser Use",
        "https://docs.openhands.dev/sdk/guides/agent-browser-use.md",
    )

    h("3. 路径 C：terminal + curl/wget 等", 2)
    p(
        "内置 terminal 工具可执行任意 shell HTTP 客户端。"
        "输出以 TerminalObservation 进入 Event Log，无框架级 SECURITY NOTICE 前缀。"
    )

    h("4. 路径 D：Tavily Search（经 MCP）", 2)
    p(
        "可作为 search engine：经 Tavily MCP 提供 search / extract / crawl / map。"
        "OpenHands Cloud 默认配置 Tavily；自托管在 Settings > LLM 填写 "
        "Search API Key (Tavily)。"
    )
    cite(
        "Search Engine Setup",
        "https://docs.openhands.dev/openhands/usage/advanced/search-engine-setup.md",
    )

    h("5. Security Notice：与 OpenClaw 的对照", 2)
    bullet(
        "OpenClaw：web_fetch 返回时写入 SECURITY NOTICE，"
        "并用 EXTERNAL_UNTRUSTED_CONTENT 边界包裹正文；"
        "另有 cacheTtlMinutes（常见 15 分钟）等配置。"
    )
    bullet(
        "OpenHands 核心产品/SDK 文档：未规定对通用外部网页内容做同等自动包裹。"
    )
    bullet(
        "例外（仅 PR Review 插件场景）：OpenHands/extensions 的 PR review prompt "
        "对 PR title/body/diff 等作者可控字段使用 "
        "BEGIN/END UNTRUSTED PR CONTENT {nonce} 标记"
        "（防 PR 文本注入；不是通用 web_fetch 路径）。"
    )
    cite(
        "extensions PR：harden review prompt against injection",
        "https://github.com/OpenHands/extensions/pull/252",
    )
    p(
        "IPI（间接提示词注入）实验含义：测试面应覆盖 MCP fetch、BrowserToolSet、"
        "terminal HTTP；不要假设存在 OpenClaw 式框架级 untrusted 包裹。"
    )

    h("6. OpenHands 侧安全机制（动作级，非内容包裹）", 2)
    bullet(
        "LLMSecurityAnalyzer：对 non-read-only 工具 schema 注入 required "
        "security_risk（LOW/MEDIUM/HIGH）；LLM 在 tool_call 参数中 inline 标注。"
    )
    bullet(
        "ConfirmationPolicy：AlwaysConfirm / NeverConfirm / ConfirmRisky"
        "（架构文档称 ConfirmRisky 为灵活默认策略；需与 analyzer 配合）。"
    )
    bullet(
        "另有 Pattern / PolicyRail 等可组合 analyzer（安全指南）："
        "在工具执行前做确定性规则扫描。"
    )
    bullet("Sandbox：Docker（推荐）/ Process（快但不隔离）/ Remote。")
    bullet("Hooks：PreToolUse 可拦截；Stop 可强制 lint/test。")
    cite(
        "Security Architecture",
        "https://docs.openhands.dev/sdk/arch/security.md",
    )
    cite(
        "Security & Action Confirmation",
        "https://docs.openhands.dev/sdk/guides/security.md",
    )
    cite(
        "Sandboxes Overview",
        "https://docs.openhands.dev/openhands/usage/sandboxes/overview.md",
    )

    p("示意：启用 security_risk 后的 tool_call（工具名以 V1 为准，见 Tools 节）：")
    code(
        """{
  "name": "terminal",
  "arguments": {
    "command": "curl -s https://api.example.com/health",
    "security_risk": "MEDIUM"
  }
}"""
    )
    p(
        "注意：SDK Security 架构页示例曾写 name: execute_bash；"
        "该名为过时示意。当前内置工具名为 terminal（见下节与源码）。"
    )

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "工具系统：Action（入参）→ Executor → Observation（出参）；"
        "原生工具经 ToolRegistry；MCP 工具在 Agent.initialize 时发现并并入 tools_map。"
    )
    cite(
        "Tool System & MCP",
        "https://docs.openhands.dev/sdk/arch/tool-system.md",
    )

    h("1. 内置核心工具（SDK 当前命名）", 2)
    p(
        "ToolDefinition 子类若未显式覆盖 name，则由类名 CamelCase→snake_case "
        "并去掉后缀 _tool。源码：openhands/sdk/tool/tool.py。"
    )
    bullet(
        "terminal（TerminalTool）：持久 shell session 执行命令；"
        "annotations.title=\"terminal\"；destructiveHint=True, openWorldHint=True。"
    )
    bullet(
        "file_editor（FileEditorTool）：view / create / str_replace / insert / undo_edit。"
    )
    bullet(
        "task_tracker（TaskTrackerTool）：任务跟踪；"
        "官方 README 默认示例常与 terminal、file_editor 一起启用。"
    )
    bullet(
        "BrowserToolSet（可选）：基于 browser-use 的浏览器工具集合。"
    )
    bullet(
        "TaskToolSet（可选）：父 agent 同步委派 sub-agent；"
        "支持 resume=task_id 恢复子会话。"
    )
    cite(
        "Custom Tools / Browser Use / Task Tool Set；software-agent-sdk README",
        "https://docs.openhands.dev/sdk/guides/custom-tools.md",
    )
    p(
        "历史/文档混用名（已废弃或仅出现在旧示例）：BashTool、execute_bash、"
        "str_replace_editor。研究采集 LLM tool schema 时以实际注册的 name 字段为准。"
    )

    h("2. MCP 工具", 2)
    bullet("配置：Agent(mcp_config=...)。")
    bullet("传输：stdio / SSE / SHTTP（产品文档建议代理可靠性更高的 HTTP/SSE）。")
    bullet("过滤：filter_tools_regex。")
    bullet("OAuth MCP：首次需浏览器登录；不适合纯 headless。")

    h("3. 与 memory / 长期上下文相关（非专用 memory 工具）", 2)
    p(
        "OpenHands 无 OpenClaw 的 memory_search / memory_get，"
        "也无官方 MEMORY.md / memory/YYYY-MM-DD.md 专用体系。"
    )
    bullet("会话上下文：Event Log + 可选 Condenser。")
    bullet("仓库永久上下文：AGENTS.md（always-on）；.agents/skills/ 按需触发。")
    bullet("任意文件读写：file_editor 或 terminal。")

    h("4. 研究用简化 messages 示例", 2)
    p("假设已启用 MCP fetch，或 agent 选择用 terminal curl：")
    code(
        """{
  "messages": [
    {
      "role": "system",
      "content": "<agent prompt + AGENTS.md + triggered skills>"
    },
    {
      "role": "user",
      "content": "Check https://api.example.com/health and return the status."
    },
    {
      "role": "assistant",
      "content": null,
      "tool_calls": [{
        "id": "call_001",
        "type": "function",
        "function": {
          "name": "terminal",
          "arguments": "{\\"command\\": \\"curl -s https://api.example.com/health\\", \\"security_risk\\": \\"LOW\\"}"
        }
      }]
    },
    {
      "role": "tool",
      "tool_call_id": "call_001",
      "content": "{\\"ok\\": true}"
    },
    {
      "role": "assistant",
      "content": "The API is healthy (ok=true)."
    }
  ],
  "tools": [
    {"type": "function", "function": {"name": "terminal", "...": "..."}},
    {"type": "function", "function": {"name": "file_editor", "...": "..."}}
  ]
}"""
    )

    # ══════════════════════════════════════════════════════════════
    # 五轮防幻觉
    # ══════════════════════════════════════════════════════════════
    h("五轮防幻觉核对记录", 1)
    p(
        "目标：剔除“臆造 OpenClaw 式机制 / 过时工具名 / 未标注的推断”。"
        "每轮只接受官方文档、官方示例或可核验源码/PR。"
    )

    h("Round 1 — Session / Conversation", 2)
    bullet(
        "核对：https://docs.openhands.dev/sdk/arch/conversation.md"
    )
    bullet(
        "保留：LocalConversation vs RemoteConversation；Event Log；Factory Pattern。"
    )
    bullet(
        "剔除：把 OpenClaw Main/Isolated/Heartbeat 写成 OpenHands 官方名词"
        "（改为「研究映射」小节并标明非官方）。"
    )

    h("Round 2 — 任务模式", 2)
    bullet(
        "核对：automations/overview.md、creating-automations.md、"
        "event-automations.md、cli/headless.md、cli/terminal.md、"
        "agent-canvas/overview.md。"
    )
    bullet(
        "保留：Automation 每次 fresh sandbox；timeout 10–30 min；"
        "Headless 强制 always-approve；CLI 确认三模式。"
    )
    bullet(
        "剔除：宣称 OpenHands 有 Heartbeat；"
        "剔除未写明的“默认每次 Automation 复用同一 session”"
        "（官方明确每次 run 新 sandbox + 新 conversation，事后可 continue）。"
    )

    h("Round 3 — web_fetch / Security Notice", 2)
    bullet(
        "核对：sdk/guides/mcp.md、agent-browser-use.md、search-engine-setup.md；"
        "对照 OpenClaw SECURITY NOTICE 实现（属 OpenClaw，非 OpenHands）。"
    )
    bullet(
        "保留：无内置 web_fetch；四条外部内容路径；"
        "PR review 插件的 UNTRUSTED PR CONTENT {nonce}（扩展场景例外）。"
    )
    bullet(
        "剔除：OpenHands 内置 cacheTtlMinutes=15；"
        "剔除“OpenHands web_fetch 自动包裹 SECURITY NOTICE”；"
        "剔除把 OpenClaw external-content.ts 行为安到 OpenHands 核心。"
    )

    h("Round 4 — Tools 命名与安全", 2)
    bullet(
        "核对：custom-tools.md、Browser/Task 指南、"
        "software-agent-sdk 中 TerminalTool/FileEditorTool 源码、"
        "tool.py 的 _camel_to_snake(...).removesuffix('_tool')、"
        "sdk/arch/security.md、sdk/guides/security.md。"
    )
    bullet(
        "修正：LLM 可见工具名以 terminal / file_editor / task_tracker 为准；"
        "execute_bash 仅出现在部分架构示例中，视为过时示意。"
    )
    bullet(
        "修正：Persistence —— 详细指南写立即写入 base_state；"
        "不以架构页 “Debounced writes” 一笔带过为唯一结论。"
    )
    bullet(
        "保留：security_risk 注入 non-read-only tools；"
        "ConfirmRisky + analyzer；沙箱三级。"
    )

    h("Round 5 — 交叉复检与冻结表述", 2)
    bullet(
        "复检 llms.txt：确认无独立 “web_fetch” 产品文档页；"
        "外部获取分散在 MCP / Browser / Search / terminal。"
    )
    bullet(
        "冻结研究结论四条："
        "(1) Session ≡ Conversation（Local/Remote + Resume）；"
        "(2) 任务模式 = Interactive / Headless / Cron Automation / Event Automation；"
        "(3) 无内置 web_fetch / 无框架级 SECURITY NOTICE，IPI 面在 MCP/Browser/curl；"
        "(4) 核心工具名 terminal + file_editor（+ 可选 Browser/MCP/Task）。"
    )
    bullet(
        "仍属开放问题（勿写成既定事实）："
        "具体 mcp-server-fetch 返回 JSON 字段；"
        "Cloud 默认工具集与自托管 CLI 是否完全一致；"
        "需实机抓 LLM 请求验证。"
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt 组成（文末追加）
    # ══════════════════════════════════════════════════════════════
    h("Prompt 组成", 1)
    p(
        "与 OpenClaw 类似，发往 LLM 的请求主要包括 system / user / assistant / tool "
        "messages 与 tools schema。"
        "OpenHands 由 Agent 的 reasoning-action loop 在每轮 step() 中组装："
        "可选 Condenser 压缩历史 → 查询 LLM →（安全检查）→ 执行工具 → Observation 写回 Event Log。"
    )
    cite(
        "Agent Architecture：reasoning-action loop 与 Agent Context",
        "https://docs.openhands.dev/sdk/arch/agent.md",
    )

    h("1. 核心消息结构", 2)
    bullet(
        "System Prompt：由 Agent prompt template + AgentContext"
        "（repo skills、system prompt prefix/suffix）拼装。"
    )
    bullet(
        "User / Assistant / Tool Messages：来自 ConversationState.events"
        "（append-only Event Log），按 Chat Completions 风格交给 LLM。"
    )
    bullet(
        "Tools Schema：每个 Tool 的 name、description、parameters；"
        "启用 LLMSecurityAnalyzer 时，non-read-only 工具自动追加 required "
        "security_risk（LOW/MEDIUM/HIGH）。"
    )
    cite(
        "Security Architecture：security_risk 注入规则",
        "https://docs.openhands.dev/sdk/arch/security.md",
    )

    h("2. AgentContext 与动态渲染（Skills）", 2)
    p("AgentContext 向 LLM 注入两类技能与前后缀（架构术语）：")
    bullet(
        "repo skills：Always included → 进入 System Prompt"
        "（项目约定、永久上下文；产品侧推荐根目录 AGENTS.md）。"
    )
    bullet(
        "knowledge skills：Trigger words/patterns → 触发时进入 User Messages 侧上下文"
        "（领域知识、专用行为）。"
    )
    bullet("System prompt prefix/suffix：可按会话追加。")
    cite(
        "Agent Architecture > Agent Context",
        "https://docs.openhands.dev/sdk/arch/agent.md",
    )
    p("产品层技能加载模型（Skills Overview）：")
    bullet("Always-on：AGENTS.md（及 GEMINI.md / CLAUDE.md 等模型变体）在会话开始注入。")
    bullet(
        "On-demand：用户关键词触发，或 agent 先见摘要再读取完整 SKILL.md"
        "（progressive disclosure）。"
    )
    bullet("Path-triggered：读写/创建匹配 glob 的文件时确定性注入。")
    bullet(
        "加载优先级：.agents/skills/ > .openhands/skills/（deprecated）> "
        ".openhands/microagents/（deprecated）；项目技能优先于用户技能。"
    )
    cite(
        "Skills Overview",
        "https://docs.openhands.dev/overview/skills.md",
    )
    bullet("仓库定制：.openhands/setup.sh（会话开始执行）；.openhands/hooks.json（生命周期钩子）。")
    cite(
        "Repository Customization",
        "https://docs.openhands.dev/openhands/usage/customization/repository.md",
    )

    h("3. CodeActAgent 行为范式（产品文档）", 2)
    p("OpenHands 主 Agent 描述为 CodeActAgent，每轮可在两类动作间选择：")
    bullet("Converse：自然语言与用户沟通、澄清或确认。")
    bullet(
        "CodeAct：执行 Linux bash 或经 bash 模拟的 Python；"
        "具体暴露给 LLM 的工具名仍以 V1 注册名为准（terminal / file_editor 等）。"
    )
    cite(
        "Main Agent and Capabilities",
        "https://docs.openhands.dev/openhands/usage/agents.md",
    )

    h("4. Prompt 侧采集（研究）", 2)
    bullet(
        "SDK callback：对 LLMConvertibleEvent 调用 event.to_llm_message() "
        "收集原始 LLM 消息（Persistence / MCP 示例同模式）。"
    )
    bullet("OpenTelemetry：追踪 agent.step、tool、LLM、conversation 生命周期。")
    bullet("亦可将 LLM base_url 指向本地代理做全量拦截（非官方唯一推荐路径）。")

    h("Prompt 组成 · 核对摘要", 2)
    bullet(
        "保留：system+messages+tools；AgentContext repo/knowledge；"
        "AGENTS.md always-on；Skills 优先级。"
    )
    bullet(
        "剔除：把 OpenClaw 的 system prompt「按本地配置动态渲染」原文照搬为 OpenHands 机制；"
        "剔除臆造的独立 memory_search 进入 tools schema。"
    )
    bullet(
        "注意：架构图写 knowledge → User Messages；产品写 on-demand skills。"
        "二者同属按需注入，勿写成第二套未文档化的 API。"
    )

    # ══════════════════════════════════════════════════════════════
    # 记忆与持久化（文末追加）
    # ══════════════════════════════════════════════════════════════
    h("记忆与持久化机制", 1)
    p(
        "OpenHands 未提供 OpenClaw 式 MEMORY.md / memory/YYYY-MM-DD.md 专用文件体系，"
        "也无 memory_search / memory_get 专用工具。"
        "“记忆”由 Event Log（会话内）+ Persistence/Resume（跨会话）+ "
        "AGENTS.md/Skills（仓库级长期指令）+ 可选 Condenser（压缩）共同构成。"
    )

    h("1. Conversation Event Log（会话内上下文）", 2)
    bullet(
        "ConversationState.events：append-only，记录 user / assistant / Action / Observation。"
    )
    bullet(
        "Agent 无跨 step 的可变业务状态：每轮 step() 从 Event Log 读取历史再推理。"
    )
    cite(
        "Conversation Architecture；Agent Architecture",
        "https://docs.openhands.dev/sdk/arch/conversation.md",
    )

    h("2. Context Condenser（可选，需显式配置）", 2)
    p(
        "默认实现为 LLMSummarizingCondenser：事件数超过 max_size 时用 LLM 摘要旧历史，"
        "保留最近事件与 keep_first（通常含 system / 初始 user）。"
        "须在 Agent(..., condenser=...) 上配置才会启用；不是无条件默认开启。"
    )
    cite(
        "Context Condenser",
        "https://docs.openhands.dev/sdk/guides/context-condenser.md",
    )

    h("3. 跨会话持久化（CLI / SDK）", 2)
    bullet(
        "CLI：~/.openhands/conversations/<id>/conversation.json；"
        "openhands --resume / --resume --last。"
    )
    cite(
        "Resume Conversations",
        "https://docs.openhands.dev/openhands/usage/cli/resume.md",
    )
    bullet(
        "SDK：persistence_dir + conversation_id → base_state.json + events/event-*.json。"
    )
    bullet(
        "SDK 写入机制（Persistence 指南）：修改 ConversationState 公共字段时 "
        "__setattr__ 立即序列化 base state；事件按文件增量追加。"
        "架构总览对 Persistence 服务写过 Debounced writes——以 Persistence 指南为准。"
    )
    bullet(
        "持久化内容包括：完整事件轨迹、Agent/工具/MCP 配置、执行状态、"
        "工具输出、统计、workspace、已激活 Skills、Secrets、agent 自定义 runtime state 等。"
    )
    cite(
        "Persistence",
        "https://docs.openhands.dev/sdk/guides/convo-persistence.md",
    )

    h("4. 仓库级长期上下文（非会话 Memory 文件）", 2)
    bullet("AGENTS.md：官方推荐的 permanent agent context（conversation start always-on）。")
    bullet(".agents/skills/：按需技能；摘要 + 完整 SKILL.md progressive disclosure。")
    bullet("Organization / Global Skills：组织或社区共享（extensions 仓库等）。")
    cite(
        "Skills Overview",
        "https://docs.openhands.dev/overview/skills.md",
    )

    h("5. 文件/终端间接“记忆”写入", 2)
    bullet("file_editor / terminal 可读写 workspace 任意路径；长期笔记需项目自行约定。")
    bullet(
        "OpenHands 不保证主会话自动加载某固定 MEMORY.md；"
        "若需等价行为，应通过 AGENTS.md 或 Skill 明确指令 agent 读写指定路径。"
    )

    h("记忆与持久化 · 核对摘要", 2)
    bullet(
        "保留：Event Log；可选 Condenser；CLI/SDK Persistence；AGENTS.md/Skills；"
        "无专用 memory 工具。"
    )
    bullet(
        "剔除：OpenClaw Daily notes / MEMORY.md “ONLY load in main session” 规则"
        "被写成 OpenHands 默认行为；剔除 memory_search / memory_get。"
    )
    bullet(
        "与前文 Session 节一致：Automation 每次新 sandbox/新 conversation，"
        "跨 run “记忆”不自动延续，除非 Resume/continue 同一会话或依赖仓库级 AGENTS.md。"
    )

    h("主要参考链接", 1)
    for link in [
        "https://docs.openhands.dev/llms.txt",
        "https://docs.openhands.dev/sdk/arch/conversation.md",
        "https://docs.openhands.dev/sdk/arch/agent.md",
        "https://docs.openhands.dev/sdk/arch/tool-system.md",
        "https://docs.openhands.dev/sdk/arch/security.md",
        "https://docs.openhands.dev/sdk/guides/convo-persistence.md",
        "https://docs.openhands.dev/sdk/guides/context-condenser.md",
        "https://docs.openhands.dev/sdk/guides/mcp.md",
        "https://docs.openhands.dev/sdk/guides/agent-browser-use.md",
        "https://docs.openhands.dev/sdk/guides/custom-tools.md",
        "https://docs.openhands.dev/sdk/guides/security.md",
        "https://docs.openhands.dev/sdk/guides/task-tool-set.md",
        "https://docs.openhands.dev/overview/skills.md",
        "https://docs.openhands.dev/openhands/usage/agents.md",
        "https://docs.openhands.dev/openhands/usage/automations/overview.md",
        "https://docs.openhands.dev/openhands/usage/automations/creating-automations.md",
        "https://docs.openhands.dev/openhands/usage/automations/event-automations.md",
        "https://docs.openhands.dev/openhands/usage/cli/headless.md",
        "https://docs.openhands.dev/openhands/usage/cli/terminal.md",
        "https://docs.openhands.dev/openhands/usage/cli/resume.md",
        "https://docs.openhands.dev/openhands/usage/customization/repository.md",
        "https://docs.openhands.dev/openhands/usage/sandboxes/overview.md",
        "https://docs.openhands.dev/openhands/usage/advanced/search-engine-setup.md",
        "https://docs.openhands.dev/openhands/usage/agent-canvas/overview.md",
        "https://github.com/OpenHands/extensions/pull/252",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
