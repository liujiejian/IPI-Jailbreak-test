# -*- coding: utf-8 -*-
"""Generate CrewAI research notes (English).

Same structure as Chinese _gen_crewai_doc.py:
Session / task modes / web_fetch(security notice) / tools,
plus Prompt composition & memory/persistence. Five-round anti-hallucination vs docs.crewai.com (v1.15.2).
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "en" / "CrewAI.docx"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label} (Source: {url})")

    # ── Title ──
    h("CrewAI", 0)
    p(
        "This document follows the same structure as OpenHands.docx / Hermes.docx / Claude-Code.docx / "
        "Dify.docx / AutoGPT.docx / Cline.docx / Codex-CLI.docx: "
        "Session, task modes, web_fetch (and Security Notice comparison), Tools; "
        "plus Prompt composition, memory and persistence mechanisms, and five-round "
        "anti-hallucination verification at the end. "
        "Subject: open-source multi-agent framework CrewAI (Crews + Flows); "
        "CrewAI AMP enterprise layer for deployment/trigger comparison only."
    )
    p("Official documentation index: https://docs.crewai.com/llms.txt")
    p(
        "Positioning: Flows manage state and control flow; Crews (role-based Agent + Task) solve collaboratively. "
        "Official recommendation: start production apps with Flow; delegate complex autonomous tasks to Crew."
    )
    cite(
        "Introduction",
        "https://docs.crewai.com/v1.15.2/en/introduction",
    )
    p(
        "Execution chain (research view · Crew): "
        "kickoff → Agent(LLM) → Tool/MCP/App call → results fed back → Task complete → "
        "(optional) Memory remember / Checkpoint."
    )
    p(
        "Execution chain (research view · Conversational Flow): "
        "handle_turn(message, session_id) → same session message history + this Flow run → "
        "append_assistant_message."
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "CrewAI \"session\" semantics split by entry point: "
        "(A) Conversational Flow session_id/message history; "
        "(B) Crew/Flow kickoff runtime + Checkpoint lineage; "
        "(C) AMP API kickoff_id / resume. Do not treat these as the same field."
    )

    h("1. Conversational Flows (true multi-turn chat session)", 2)
    bullet(
        "Each user input = one new flow run, but shares same session_id."
    )
    bullet(
        "`handle_turn(message, session_id=...)` → internally "
        "`kickoff(inputs={\"id\": session_id})` → `state.id`; "
        "messages written to `state.messages`."
    )
    bullet(
        "`stream_turn` / `flow.chat()` REPL / `ChatSession` (SSE·WebSocket)."
    )
    bullet(
        "`Flow.kickoff()` does not accept `user_message=` / `session_id=` keywords; "
        "conversational surface uses handle_turn."
    )
    cite(
        "Conversational Flows",
        "https://docs.crewai.com/v1.15.2/en/guides/flows/conversational-flows",
    )

    h("2. Crew / Flow run + Checkpoint", 2)
    bullet(
        "Ordinary `crew.kickoff()` / Flow run is single (or batch) workflow, not chat session."
    )
    bullet(
        "Checkpoint: event-driven snapshot (default task_completed); "
        "stores full config, task progress, intermediate outputs, memory/knowledge, etc.; "
        "can resume / fork (new lineage)."
    )
    bullet(
        "Storage: JsonProvider (e.g. `./.checkpoints/`) or SqliteProvider; "
        "`max_checkpoints` can trim."
    )
    bullet(
        "Flow Persistence: persisted on `state.id`; fork can carry state with new id."
    )
    cite(
        "Checkpointing",
        "https://docs.crewai.com/v1.15.2/en/concepts/checkpointing",
    )

    h("3. AMP API (enterprise)", 2)
    bullet("POST /kickoff starts deployed crew; GET /status/{kickoff_id}; POST /resume (human feedback continuation).")
    bullet("Different namespace from local Conversational session_id.")
    cite(
        "AMP kickoff / resume",
        "https://docs.crewai.com/v1.15.2/en/api-reference/kickoff",
    )

    h("4. Research mapping (non-official terminology)", 2)
    bullet(
        "Codex/Cline interactive session ≈ CrewAI Conversational Flow session_id."
    )
    bullet(
        "OpenClaw Isolated / Automation single run ≈ crew.kickoff / AMP kickoff_id."
    )
    bullet(
        "Checkpoint resume ≈ failure recovery, not chat continuation API."
    )

    # ══════════════════════════════════════════════════════════════
    # Task modes
    # ══════════════════════════════════════════════════════════════
    h("Task Modes", 1)

    h("1. Flows (backbone)", 2)
    bullet("@start / @listen event-driven; conditional or/and, Router, loops.")
    bullet("Can embed Crew or direct Agent; built-in Flow memory methods (remember/recall/extract).")
    bullet("HITL: `@human_feedback`, ask(), etc. (step approval ≠ next chat utterance).")
    cite("Flows", "https://docs.crewai.com/v1.15.2/en/concepts/flows")

    h("2. Crews + Processes", 2)
    bullet("Crew = Agents + Tasks; `kickoff` / async / for_each / streaming.")
    bullet("Process.sequential: task list order, prior output as subsequent context.")
    bullet(
        "Process.hierarchical: manager Agent/LLM delegates and validates; requires manager_llm or manager_agent."
    )
    bullet("`planning=True` can enable planning capability.")
    cite("Crews", "https://docs.crewai.com/v1.15.2/en/concepts/crews")
    cite("Processes", "https://docs.crewai.com/v1.15.2/en/concepts/processes")

    h("3. Local CLI / AMP Automations", 2)
    bullet("crewai CLI: project scaffolding, run, replay, etc.")
    bullet(
        "AMP: Automations, Triggers (Gmail/Slack/Webhook…), "
        "Crew Studio, Traces, Webhook Streaming."
    )

    h("Task Modes · verification summary", 2)
    bullet("Keep: Flow as production backbone; Crew sequential/hierarchical; Conversational handle_turn; Checkpoint.")
    bullet("Drop: writing \"single Agent chat only\" as full framework capability.")

    # ══════════════════════════════════════════════════════════════
    # web_fetch
    # ══════════════════════════════════════════════════════════════
    h("web_fetch (and Security Notice)", 1)
    p(
        "Conclusion: CrewAI framework layer has no built-in tool named `web_fetch` / `WebFetch`, "
        "nor documented OpenClaw-style SECURITY NOTICE / EXTERNAL_UNTRUSTED_CONTENT wrapping. "
        "Web content enters context via installable crewai_tools / MCP / Apps capabilities; "
        "MCP docs separately warn about tool metadata prompt injection."
    )

    h("1. Typical web tools (no single built-in name)", 2)
    bullet(
        "`ScrapeWebsiteTool`: HTTP fetch page and parse HTML; "
        "can fix website_url or arbitrary URL at runtime."
    )
    bullet(
        "Firecrawl / Scrapfly / Selenium / Spider / Stagehand / Browserbase / "
        "Hyperbrowser / Bright Data / Oxylabs / You.com Contents and other scraping suites."
    )
    bullet(
        "Search side: `SerperDevTool`, Tavily, Exa, Brave, WebsiteSearchTool (website content RAG), etc."
    )
    bullet("ApifyActorsTool, MultiOnTool: more crawler/browser automation platform oriented.")
    cite(
        "Web Scraping Overview",
        "https://docs.crewai.com/v1.15.2/en/tools/web-scraping/overview",
    )
    cite(
        "ScrapeWebsiteTool",
        "https://docs.crewai.com/v1.15.2/en/tools/web-scraping/scrapewebsitetool",
    )
    cite(
        "SerperDevTool",
        "https://docs.crewai.com/v1.15.2/en/tools/search-research/serperdevtool",
    )

    h("2. Knowledge URL source ≠ live web_fetch", 2)
    bullet(
        "Knowledge: document/file/URL semantic retrieval (RAG) injects \"what to know\"; "
        "different threat timing from agent calling ScrapeWebsiteTool live scrape."
    )

    h("3. Security Notice comparison", 2)
    bullet(
        "Tool docs do not describe wrapping scrape/search results in SECURITY NOTICE."
    )
    bullet(
        "MCP Security: trusted MCP only; "
        "malicious tool metadata can pollute LLM at \"list tools\" stage; "
        "includes \"hijack reasoning / prompt injection\" wording."
    )
    bullet(
        "Memory Privacy Note: memory content sent to analysis LLM — "
        "sensitive data should use local LLM; not web untrusted wrapper."
    )
    cite(
        "MCP Security Considerations",
        "https://docs.crewai.com/v1.15.2/en/mcp/security",
    )

    h("Threat model mapping (research, non-official)", 2)
    bullet(
        "External scrape/search → tool result → Agent LLM → "
        "subsequent Tool/App/delegation: indirect injection surface."
    )
    bullet(
        "Malicious MCP metadata: injection without actual tool call — additional attack surface."
    )
    bullet(
        "If memory enabled and stores web-derived content: cross kickoff/session persistent pollution "
        "(analyze with scopes/source)."
    )
    bullet(
        "Compare: Codex web_search+treat untrusted; Cline fetch_web; "
        "OpenClaw web_fetch+NOTICE; CrewAI = optional tool ecosystem + MCP warning."
    )

    code(
        """Typical (official example pattern):
from crewai_tools import SerperDevTool, ScrapeWebsiteTool
agent = Agent(..., tools=[SerperDevTool(), ScrapeWebsiteTool()])
# Tool return body enters subsequent LLM turns; docs have no SECURITY NOTICE wrapper convention
"""
    )

    h("web_fetch · verification summary", 2)
    bullet(
        "Keep: no framework-level web_fetch; Scrape/Search/Browser tool family; "
        "no NOTICE; MCP metadata risk."
    )
    bullet("Drop: default built-in web_fetch; treating Knowledge URL ingest as live fetch.")

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "Five extension categories: **Tools / MCP / Apps / Skills / Knowledge**. "
        "First three resolve to BaseTool unified list at runtime; latter two modify prompt/context."
    )
    cite(
        "Agent Capabilities",
        "https://docs.crewai.com/v1.15.2/en/concepts/agent-capabilities",
    )

    h("1. Action: Tools · MCP · Apps", 2)
    bullet("Tools: `pip install 'crewai[tools]'`; custom BaseTool; optional cache/async/Pydantic output.")
    bullet("MCP: remote tool server; connect trusted parties only.")
    bullet("Apps: platform integrations (Gmail, etc.) via CrewAI platform token.")

    h("2. Context: Skills · Knowledge", 2)
    bullet("Skills: filesystem skill packs, inject how to think; not callable.")
    bullet("Knowledge: RAG facts; different from unified Memory storage.")

    h("3. Collaboration / other", 2)
    bullet("coworker delegation tool, code execution options (Agent params), reasoning, etc. see Agents docs.")

    h("Tools · verification summary", 2)
    bullet("Keep: five-capability layering; Tools≠Skills≠Knowledge.")
    bullet("Drop: single fixed web_fetch built-in name.")

    # ══════════════════════════════════════════════════════════════
    # Five-round anti-hallucination
    # ══════════════════════════════════════════════════════════════
    h("Five-Round Anti-Hallucination Verification", 1)

    h("Round 1 — Session", 2)
    bullet("Checked: conversational-flows; checkpointing; AMP kickoff/resume.")
    bullet(
        "Keep: session_id+handle_turn; kickoff runtime; checkpoint lineage; kickoff_id."
    )
    bullet("Drop: writing kickoff_id as chat session_id.")

    h("Round 2 — Task modes", 2)
    bullet("Checked: introduction; flows; crews; processes; planning.")
    bullet("Keep: Flow+Crew; sequential/hierarchical; Conversational; AMP triggers.")
    bullet("Drop: claiming no Flow, single Agent only.")

    h("Round 3 — web_fetch / security", 2)
    bullet("Checked: web-scraping overview; scrapewebsitetool; mcp/security.")
    bullet("Keep: tool ecosystem web fetch; no SECURITY NOTICE; MCP metadata PI.")
    bullet("Drop: OpenClaw NOTICE; copying Codex web_search default mode to CrewAI.")

    h("Round 4 — Tools", 2)
    bullet("Checked: agent-capabilities; tools; skills; knowledge.")
    bullet("Keep: five categories; Tools/MCP/Apps unified as BaseTool.")
    bullet("Drop: treating Skills as callable scrape tool.")

    h("Round 5 — Prompt / Memory cross-check", 2)
    bullet("Checked: agents; memory; knowledge; skills.")
    bullet(
        "Freeze: "
        "(1) Session = conversational session_id or kickoff/checkpoint; "
        "(2) Tasks = Flow / Crew(process) / AMP; "
        "(3) External web = optional scrape/search tools + MCP risk; "
        "(4) Prompt = role/goal/backstory + Task + Skills/Knowledge; "
        "(5) Memory = unified Memory class (not Checkpoint)."
    )
    bullet(
        "Open questions: whether specific tool result templates include undocumented security prefix; "
        "default memory storage path and cross-process isolation; "
        "whether AMP and local memory files share."
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt composition
    # ══════════════════════════════════════════════════════════════
    h("Prompt Composition", 1)
    p(
        "Agent: role / goal / backstory (and optional system template params) + "
        "Task description/expected_output + "
        "Skills instructions + Knowledge retrieval snippets + "
        "tool schema and tool returns."
    )

    h("1. Agent persona fields", 2)
    bullet("role, goal, backstory are critical params; affect tool selection and style.")
    bullet("Can define via JSONC/YAML/code; supports reasoning, dynamic date, and other advanced options.")
    cite("Agents", "https://docs.crewai.com/v1.15.2/en/concepts/agents")

    h("2. Task and collaboration context", 2)
    bullet("Task.description / expected_output; context= prior task outputs.")
    bullet("In Hierarchical mode, manager planning and delegation enter execution context.")

    h("3. Skills / Knowledge", 2)
    bullet("Skills: domain procedure injects \"how to think\".")
    bullet("Knowledge: semantic retrieval injects \"what to know\".")

    h("4. Context window", 2)
    bullet(
        "`respect_context_window` (default True) can auto-handle overly long context; "
        "large materials prefer RAG/Knowledge over stuffing prompt."
    )

    h("Prompt Composition · verification summary", 2)
    bullet("Keep: role/goal/backstory + task + skills/knowledge + tools.")
    bullet("Drop: AGENTS.md / CLAUDE.md as CrewAI default loading (unless user-defined convention).")

    # ══════════════════════════════════════════════════════════════
    # Memory and persistence
    # ══════════════════════════════════════════════════════════════
    h("Memory and Persistence", 1)
    p(
        "Three lines must be separated: "
        "① unified Memory class (semantic memory); "
        "② Checkpointing (execution snapshot); "
        "③ Conversational state.messages / Flow state persistence."
    )

    h("1. Unified Memory", 2)
    bullet(
        "Single `Memory` API replaces legacy short/long/entity/external split; "
        "on save LLM infers scope/categories/importance; "
        "recall combines semantic+recency+importance."
    )
    bullet(
        "Usage: standalone; `Crew(memory=True)` or pass Memory; "
        "Agent-level scoped view; within Flow `remember/recall/extract_memories`."
    )
    bullet(
        "Default embedder often OpenAI text-embedding-3-large (configurable); "
        "analysis LLM default gpt-4o-mini configurable (including Ollama)."
    )
    bullet("Scopes / Slices / private+source tags; RecallFlow deep retrieval.")
    cite("Memory", "https://docs.crewai.com/v1.15.2/en/concepts/memory")

    h("2. Checkpointing (not semantic memory)", 2)
    bullet("Restores task progress and intermediate artifacts; can rehydrate memory/knowledge config state.")
    bullet("Purpose: failure recovery / fork; not \"user preference knowledge base\" API.")

    h("3. Conversational / Flow state", 2)
    bullet("ConversationState.messages across turns with same session_id.")
    bullet("Flow Persistence snapshots business state by state.id.")

    h("4. Knowledge (RAG)", 2)
    bullet("Retrieval augmentation from document/URL sources; coexists with unified Memory class but different purpose.")

    h("Memory and Persistence · verification summary", 2)
    bullet(
        "Keep: Memory vs Checkpoint vs chat messages; scopes; analysis LLM privacy note."
    )
    bullet(
        "Drop: treating Checkpoint directory as MEMORY.md; "
        "default OpenClaw-style MEMORY.md."
    )

    h("Primary Reference Links", 1)
    for link in [
        "https://docs.crewai.com/llms.txt",
        "https://docs.crewai.com/v1.15.2/en/introduction",
        "https://docs.crewai.com/v1.15.2/en/concepts/flows",
        "https://docs.crewai.com/v1.15.2/en/concepts/crews",
        "https://docs.crewai.com/v1.15.2/en/concepts/processes",
        "https://docs.crewai.com/v1.15.2/en/concepts/agents",
        "https://docs.crewai.com/v1.15.2/en/concepts/tools",
        "https://docs.crewai.com/v1.15.2/en/concepts/agent-capabilities",
        "https://docs.crewai.com/v1.15.2/en/concepts/memory",
        "https://docs.crewai.com/v1.15.2/en/concepts/knowledge",
        "https://docs.crewai.com/v1.15.2/en/concepts/checkpointing",
        "https://docs.crewai.com/v1.15.2/en/concepts/skills",
        "https://docs.crewai.com/v1.15.2/en/guides/flows/conversational-flows",
        "https://docs.crewai.com/v1.15.2/en/tools/web-scraping/overview",
        "https://docs.crewai.com/v1.15.2/en/tools/web-scraping/scrapewebsitetool",
        "https://docs.crewai.com/v1.15.2/en/tools/search-research/serperdevtool",
        "https://docs.crewai.com/v1.15.2/en/mcp/security",
        "https://docs.crewai.com/v1.15.2/en/api-reference/kickoff",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
