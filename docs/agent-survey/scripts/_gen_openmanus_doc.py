# -*- coding: utf-8 -*-
"""Generate OpenManus research document — accuracy-focused with 5-round audit."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "OpenManus.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def add_code(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def cite(text: str, url: str) -> None:
        doc.add_paragraph(f"{text}（来源：{url}）")

    doc.add_heading("OpenManus", level=0)
    doc.add_paragraph(
        "本文档研究 OpenManus 开源 Agent 框架机制，术语严格对齐官方资料。"
        "用户指定入口：https://openmanus.github.io/；"
        "技术细节以 GitHub 仓库 FoundationAgents/OpenManus 的 README 与源码为准。"
    )

    # ── Scope ──
    doc.add_heading("文档范围与术语分层（必读）", level=1)
    doc.add_paragraph(
        "openmanus.github.io 仅为营销落地页（Agent Framework / Tool Integration / Open Source），"
        "不含 API 级技术文档。研究须以仓库源码与 README 为权威来源，并区分以下层次："
    )
    bullet(
        "单 Agent 运行时（main.py → Manus）：ReAct think-act 循环 + 本地工具 + 可选 MCP 远程工具。"
    )
    bullet(
        "MCP 专用模式（run_mcp.py → MCPAgent）：仅使用 MCP server 暴露的工具（stdio/SSE）。"
    )
    bullet(
        "多 Agent 规划流（run_flow.py → PlanningFlow）：README 标注为 unstable；"
        "FlowFactory 当前仅 FlowType.PLANNING。"
    )
    bullet(
        "内置 MCP Server（app/mcp/server.py）：将 bash/browser/editor/terminate 注册为 FastMCP 工具，"
        "供外部 MCP 客户端或 run_mcp 连接。"
    )
    bullet(
        "OpenManus-RL（github.com/OpenManus/OpenManus-RL）：独立 RL/GRPO 调优项目，"
        "≠ FoundationAgents/OpenManus 主框架。"
    )
    cite(
        "官网；README；main.py / run_mcp.py / run_flow.py",
        "https://openmanus.github.io/",
    )
    cite(
        "FoundationAgents/OpenManus 仓库",
        "https://github.com/FoundationAgents/OpenManus",
    )

    # ── 1. LLM 请求采集 ──
    doc.add_heading("研究背景与 LLM 请求采集", level=1)
    doc.add_paragraph(
        "采集目标：每次 LLM 调用前的 messages、tools schema，"
        "tool_calls 参数与返回、browser 截图 base64、token 计数日志。"
    )

    doc.add_heading("1. loguru 日志（app/logger.py）", level=2)
    bullet(
        "基于 loguru：stderr（默认 INFO）+ logs/{name}_{timestamp}.log（DEBUG）。"
    )
    bullet(
        "ToolCallAgent.think/act 记录：thoughts、选定工具名、参数、执行结果；"
        "token 超限记 TokenLimitExceeded。"
    )
    cite(
        "app/logger.py；app/agent/toolcall.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/logger.py",
    )

    doc.add_heading("2. Memory.messages（对话历史）", level=2)
    bullet(
        "Memory 类（app/schema.py）：messages 列表，默认 max_messages=100，"
        "超出时保留最近 100 条。"
    )
    bullet(
        "每条 Message 含 role/content/tool_calls/tool_call_id/base64_image；"
        "to_dict_list() 可导出为 LLM API 格式。"
    )
    bullet(
        "研究侧可直接读取 agent.memory.messages 或 agent.messages 属性采样完整上下文。"
    )
    cite(
        "app/schema.py > Memory / Message",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/schema.py",
    )

    doc.add_heading("3. LLM.ask_tool 调用链", level=2)
    bullet(
        "ToolCallAgent.think() 调用 llm.ask_tool(messages, system_msgs, tools, tool_choice)。"
    )
    bullet(
        "TokenCounter（tiktoken）在发送前估算 input tokens；"
        "超 max_input_tokens 抛 TokenLimitExceeded。"
    )
    bullet(
        "tenacity 重试 APIError/RateLimitError；TokenLimitExceeded 不重试。"
    )
    bullet(
        "支持 OpenAI / Azure / AWS Bedrock / Ollama / Jiekou 等（config.toml api_type）。"
    )
    cite(
        "app/llm.py；app/agent/toolcall.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/llm.py",
    )

    doc.add_heading("4. 工具执行可观测输出", level=2)
    bullet(
        "execute_tool 将结果格式化为 "
        "\"Observed output of cmd `{name}` executed:\\n{result}\" 写入 role=tool 的 Message。"
    )
    bullet(
        "BrowserUseTool 可附带 base64_image（截图）在 tool Message 中。"
    )
    bullet(
        "max_observe 可截断工具返回字符串长度（Manus 默认 10000）。"
    )
    cite(
        "app/agent/toolcall.py；app/tool/browser_use_tool.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/toolcall.py",
    )

    doc.add_heading("5. 无内置 Trace/Streaming 框架", level=2)
    doc.add_paragraph(
        "所查 README 与核心模块均无 OpenTelemetry/LangSmith 式 tracing，"
        "亦无 AgentChat 类 run_stream 事件流；观测依赖 loguru 与 memory 导出。"
    )
    cite(
        "README；app/logger.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/README.md",
    )

    # ── 2. Prompt ──
    doc.add_heading("Prompt 设计组成", level=1)

    doc.add_heading("1. Manus Agent 层", level=2)
    bullet(
        "system_prompt（app/prompt/manus.py）："
        "\"You are OpenManus...\" + workspace 目录 {directory}。"
    )
    bullet(
        "next_step_prompt：引导主动选工具、分步执行、说明结果；"
        "终止须调用 terminate。"
    )
    bullet(
        "使用浏览器时，think() 临时将 next_step_prompt 替换为 "
        "browser_context_helper.format_next_step_prompt()（含页面状态）。"
    )
    cite(
        "app/prompt/manus.py；app/agent/manus.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/prompt/manus.py",
    )

    doc.add_heading("2. ToolCallAgent / ReAct 层", level=2)
    bullet(
        "继承链：BaseAgent → ReActAgent → ToolCallAgent → Manus。"
    )
    bullet(
        "每 step：think()（LLM 决策）→ act()（执行 tool_calls）；"
        "next_step_prompt 在 think 前作为 user Message 追加。"
    )
    bullet(
        "tool_choices：none / auto / required（默认 auto）。"
    )
    bullet(
        "特殊工具 Terminate：执行后 state=FINISHED，结束 run 循环。"
    )
    cite(
        "app/agent/react.py；app/agent/toolcall.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/toolcall.py",
    )

    doc.add_heading("3. PlanningFlow 层（run_flow）", level=2)
    bullet(
        "独立 LLM 调用创建计划：system \"planning assistant\" + PlanningTool。"
    )
    bullet(
        "每步向 executor agent 发送 step_prompt（含 CURRENT PLAN STATUS + YOUR CURRENT TASK）。"
    )
    bullet(
        "可选 DataAnalysis agent（config runflow.use_data_analysis_agent=true）。"
    )
    cite(
        "app/flow/planning.py；config.example.toml [runflow]",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/flow/planning.py",
    )

    doc.add_heading("4. MCPAgent 层（run_mcp）", level=2)
    bullet(
        "system_prompt 追加 \"Available MCP tools: {tool_names}\"；"
        "工具 schema 每 5 步 refresh，变更时写入 system Message。"
    )
    cite(
        "app/agent/mcp.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/mcp.py",
    )

    doc.add_heading("5. 典型 LLM 请求示意（Manus think 轮次）", level=2)
    add_code(
        """# ask_tool 输入（简化）：
system_msgs: [Message(role=system, content=SYSTEM_PROMPT.format(directory=...))]
messages: [
  ...memory.messages...,  # user/assistant/tool 历史
  Message(role=user, content=NEXT_STEP_PROMPT),  # 每轮 think 前追加
]
tools: available_tools.to_params()  # python_execute, browser_use, str_replace_editor, ask_human, terminate, +MCP"""
    )
    cite(
        "app/agent/toolcall.py > think()",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/toolcall.py",
    )

    # ── 3. Memory ──
    doc.add_heading("Memory 与持久化机制", level=1)
    doc.add_paragraph(
        "OpenManus 的 Memory 是会话内消息缓冲区，非 CrewAI/LangGraph 式语义长期记忆或 checkpoint。"
    )

    doc.add_heading("1. 会话内 Memory（app/schema.py）", level=2)
    bullet("add_message / add_messages：追加 user/system/assistant/tool 消息。")
    bullet("get_recent_messages(n)：取最近 n 条。")
    bullet("clear()：清空；max_messages=100 滑动窗口。")
    bullet("update_memory()（BaseAgent）：封装 role→Message 工厂方法。")

    doc.add_heading("2. 无跨会话内置持久化", level=2)
    bullet(
        "所查代码无 save_state/load_state、无数据库/向量库 Memory；"
        "进程结束后对话历史丢失。"
    )
    bullet(
        "应用层可自行序列化 agent.memory.to_dict_list() 到 JSON 文件恢复。"
    )

    doc.add_heading("3. 规划状态（PlanningFlow 局部）", level=2)
    bullet(
        "PlanningTool 内存存储 plans dict（plan_id → steps/step_statuses）；"
        "仅 run_flow 会话有效，非全局持久化。"
    )
    cite(
        "app/schema.py；app/flow/planning.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/schema.py",
    )

    # ── 4. Session ──
    doc.add_heading("Session 与会话模型", level=1)
    bullet(
        "单进程单 Agent 实例：agent.run(request) 从 IDLE→RUNNING，"
        "request 作为 user Message 追加；循环至 FINISHED 或 max_steps。"
    )
    bullet(
        "Manus max_steps=20；ToolCallAgent 默认 max_steps=30；MCPAgent max_steps=20。"
    )
    bullet(
        "达 max_steps 后 current_step 归零、state=IDLE（BaseAgent.run 原文）。"
    )
    bullet(
        "run_flow 整体 timeout 3600s（asyncio.wait_for）。"
    )
    bullet(
        "SANDBOX_CLIENT.cleanup() 在 run 结束时调用（BaseAgent.run）。"
    )
    cite(
        "app/agent/base.py；run_flow.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/base.py",
    )

    # ── 5. Automation ──
    doc.add_heading("Automation 与定时任务", level=1)
    bullet("框架无内置 cron/Heartbeat/触发器；需外部调度器调用 python main.py --prompt \"...\"")
    bullet("README 仅提供终端交互：python main.py 后 input prompt。")
    bullet("run_mcp.py 支持 --interactive 循环或 --prompt 单次。")
    cite(
        "README > Quick Start；run_mcp.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/README.md",
    )

    # ── 6. External Web ──
    doc.add_heading("外部 Web 访问", level=1)
    doc.add_paragraph(
        "OpenManus 无统一 web_fetch 抽象；外部内容经浏览器工具、搜索工具或 MCP 进入 agent。"
    )

    doc.add_heading("1. BrowserUseTool（Manus 默认工具）", level=2)
    bullet(
        "基于 browser-use + Playwright Chromium；动作含 go_to_url、click、"
        "extract_content、web_search（内部调 WebSearch）等。"
    )
    bullet(
        "extract_content：将 page HTML markdownify 后送 LLM ask_tool 结构化提取。"
    )
    bullet(
        "返回 JSON 状态 + base64 截图；内容进入 tool Message，无 untrusted 包裹。"
    )
    bullet("默认 headless=False、disable_security=True（config.browser）。")
    cite(
        "app/tool/browser_use_tool.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/tool/browser_use_tool.py",
    )

    doc.add_heading("2. WebSearch 工具（库内存在，非 Manus 默认挂载）", level=2)
    bullet(
        "app/tool/web_search.py：Google/Baidu/DuckDuckGo/Bing 搜索引擎；"
        "可选 fetch_content 用 requests+BeautifulSoup 抓取结果页（截断 10000 字符）。"
    )
    bullet(
        "通过 BrowserUseTool.web_search 动作或 Crawl4aiTool 间接使用；"
        "Manus.available_tools 默认不含独立 WebSearch。"
    )
    cite(
        "app/tool/web_search.py；app/tool/__init__.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/tool/web_search.py",
    )

    doc.add_heading("3. MCP 远程抓取", level=2)
    bullet(
        "Manus 可连接 config/mcp.json 中 sse/stdio MCP server，"
        "动态将 MCPClientTool 加入 available_tools。"
    )
    bullet(
        "README 示例用 mcp-server-fetch（uvx）；属外部 MCP 实现，非 OpenManus 内置 API。"
    )
    cite(
        "app/agent/manus.py；config.example.toml [mcp]",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/manus.py",
    )

    doc.add_paragraph("与 OpenClaw web_fetch 差异（研究对照）：")
    bullet("OpenClaw：内置 web_fetch + EXTERNAL_UNTRUSTED_CONTENT。")
    bullet(
        "OpenManus：browser_use / web_search / MCP 自选；"
        "网页与搜索结果以 tool Message 原文进入 memory，无框架级信任边界标记。"
    )

    # ── 7. Security ──
    doc.add_heading("Security 机制", level=1)

    doc.add_heading("1. 代码执行", level=2)
    bullet(
        "PythonExecute：multiprocessing 子进程 + 默认 5s timeout；"
        "非 Docker 沙箱（与 config [sandbox] 可选 Docker 沙箱分离）。"
    )
    bullet(
        "SandboxSettings（config.toml）：use_sandbox 默认 false；"
        "可配 image/work_dir/network_enabled/timeout。"
    )
    bullet("Bash 工具在 MCP Server 中注册，具 shell 访问能力。")
    cite(
        "app/tool/python_execute.py；config.example.toml [sandbox]",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/tool/python_execute.py",
    )

    doc.add_heading("2. Human-in-the-Loop", level=2)
    bullet(
        "AskHuman 工具：execute 内调用 input() 阻塞等待终端用户输入。"
    )
    bullet(
        "Manus 将 AskHuman 列入默认 available_tools（极端情况求助人类）。"
    )
    cite(
        "app/tool/ask_human.py；app/agent/manus.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/tool/ask_human.py",
    )

    doc.add_heading("3. 浏览器与 MCP 风险", level=2)
    bullet(
        "BrowserUseTool 默认 disable_security=True；可访问任意 URL，"
        "页面内容直接进入 LLM 上下文（网页 IPI 面）。"
    )
    bullet(
        "MCP 工具 schema 由远程 server 提供；连接不可信 server 存在 metadata 注入风险。"
    )
    bullet(
        "is_stuck() 检测 assistant 重复回复，追加策略变更 next_step_prompt；"
        "非安全 guardrail。"
    )
    cite(
        "app/agent/base.py；app/tool/browser_use_tool.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/base.py",
    )

    doc.add_heading("4. 无框架级 Guardrail", level=2)
    doc.add_paragraph(
        "所查代码无 HallucinationGuardrail、PII 脱敏、"
        "EXTERNAL_UNTRUSTED_CONTENT 或 tool output 消毒中间件。"
    )

    doc.add_heading("5. IPI 研究含义", level=2)
    doc.add_paragraph(
        "注入面：browser extract_content / web_search raw_content、"
        "MCP tool 返回、StrReplaceEditor 读取的本地文件、"
        "PlanningFlow step_prompt 中的外部文本。"
        "tool 结果经 \"Observed output of cmd...\" 包装后仍保留攻击载荷。"
    )

    # ── 8. Architecture ──
    doc.add_heading("架构总览", level=1)
    bullet(
        "执行模型：BaseAgent.run() while 循环 → step() → think()+act()（ReAct）。"
    )
    bullet(
        "Manus 默认工具：PythonExecute、BrowserUseTool、StrReplaceEditor、"
        "AskHuman、Terminate + 动态 MCP 工具。"
    )
    bullet(
        "ToolCollection：工具注册、to_params() 生成 OpenAI function schema、execute() 分发。"
    )
    bullet(
        "配置：config/config.toml 单例 Config（LLM/Search/Browser/Sandbox/MCP/Runflow）。"
    )
    bullet(
        "工作区：workspace/ 目录（SYSTEM_PROMPT 注入 {directory}）。"
    )
    bullet(
        "致谢提及 MetaGPT、OpenHands、browser-use、crawl4ai 等上游项目。"
    )
    cite(
        "app/agent/manus.py；app/config.py；README",
        "https://github.com/FoundationAgents/OpenManus/blob/main/README.md",
    )

    doc.add_heading("与 OpenClaw / OpenHands / Hermes / LangGraph / CrewAI / AutoGen 的对照（研究映射）", level=2)
    doc.add_paragraph("非 OpenManus 官方术语：")
    bullet(
        "ReAct think-act ≈ AutoGen ToolCallAgent / CrewAI Agent 单轮 tool 循环；"
        "max_steps 类似 max_iter。"
    )
    bullet(
        "Memory.messages ≈ 会话 thread history；无 LangGraph checkpointer / CrewAI Checkpoint。"
    )
    bullet(
        "PlanningFlow ≈ CrewAI Flow 简化版 / AutoGen GraphFlow 规划子模块；"
        "README 标 unstable。"
    )
    bullet(
        "BrowserUseTool ≈ AutoGen MultimodalWebSurfer / OpenHands browser；"
        "无 OpenClaw 单一 web_fetch。"
    )
    bullet(
        "MetaGPT 同源团队；OpenManus 为轻量开源替代 Manus  invite-only 实现。"
    )

    # ── 9. Research ──
    doc.add_heading("研究建议", level=1)
    bullet("采集：开启 DEBUG 日志 + 每 step 导出 memory.to_dict_list()。")
    bullet("对比三条入口：main.py（Manus）/ run_mcp.py / run_flow.py 工具集差异。")
    bullet("IPI：browser extract_content、web_search fetch_content、MCP 返回、文件编辑器读入。")
    bullet("安全：PythonExecute 超时 vs sandbox；AskHuman 阻塞 vs 异步 HITL。")
    bullet("持久化：验证无内置跨会话恢复后，测试自研 JSON 序列化方案。")

    # ── 10. Five-round audit ──
    doc.add_heading("自查：潜在幻觉与核实结论（第五轮）", level=1)
    doc.add_paragraph(
        "2026-07-12 第五轮复核：对照 openmanus.github.io、README、config.example.toml、"
        "main/run_mcp/run_flow、base/toolcall/manus/mcp agents、schema/llm、"
        "browser_use/web_search、planning flow、config、mcp server 全部关键断言。"
        "结论：第四轮修正项仍成立；本轮未发现新的机制性幻觉。"
    )
    checks = [
        ("✓ 官网仅为特性概览；技术文档以 GitHub 仓库为准", "https://openmanus.github.io/"),
        ("✓ 主仓库为 FoundationAgents/OpenManus（非 OpenManus/OpenManus-RL）", "https://github.com/FoundationAgents/OpenManus"),
        ("✓ 三入口：main.py / run_mcp.py / run_flow.py", "https://github.com/FoundationAgents/OpenManus/blob/main/README.md"),
        ("✓ ReAct：think()→act()；llm.ask_tool 传 tools schema", "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/toolcall.py"),
        ("✓ Memory=会话消息列表 max 100；非语义长期记忆", "https://github.com/FoundationAgents/OpenManus/blob/main/app/schema.py"),
        ("✓ Manus 默认工具含 browser/python/editor/ask_human/terminate", "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/manus.py"),
        ("✓ WebSearch 在代码库中但不在 Manus 默认 available_tools", "https://github.com/FoundationAgents/OpenManus/blob/main/app/tool/web_search.py"),
        ("✓ BrowserUseTool 用 browser-use+Playwright；内容进 tool Message", "https://github.com/FoundationAgents/OpenManus/blob/main/app/tool/browser_use_tool.py"),
        ("✓ MCP：mcp.json + Manus.initialize_mcp_servers / MCPAgent", "https://github.com/FoundationAgents/OpenManus/blob/main/app/config.py"),
        ("✓ run_flow 标 unstable；PlanningFlow 唯一 FlowType", "https://github.com/FoundationAgents/OpenManus/blob/main/README.md"),
        ("✓ AskHuman 用阻塞 input()", "https://github.com/FoundationAgents/OpenManus/blob/main/app/tool/ask_human.py"),
        ("✓ loguru 写 logs/；无内置 OTel trace", "https://github.com/FoundationAgents/OpenManus/blob/main/app/logger.py"),
        ("△ openmanus.github.io 无 /docs API 文档（404）— 研究须读 GitHub 源码", "https://openmanus.github.io/"),
        ("△ PythonExecute 为 multiprocessing 非 Docker — 与 [sandbox] 可选容器不同路径", "https://github.com/FoundationAgents/OpenManus/blob/main/app/tool/python_execute.py"),
        ("△ MCP fetch 示例（mcp-server-fetch）为外部 uvx 包 — 非 OpenManus 内置函数", "https://github.com/FoundationAgents/OpenManus/blob/main/README.md"),
        ("✗ 内置 web_fetch / EXTERNAL_UNTRUSTED_CONTENT — 所查代码均无", "https://github.com/FoundationAgents/OpenManus/blob/main/app/tool/browser_use_tool.py"),
        ("✗ OpenClaw 式 Heartbeat / 定时任务 — 框架无", "https://github.com/FoundationAgents/OpenManus/blob/main/README.md"),
        ("✗ 跨会话 save_state/checkpoint — 框架无；仅 Memory 滑动窗口", "https://github.com/FoundationAgents/OpenManus/blob/main/app/schema.py"),
        ("✗ 框架级 HallucinationGuardrail / PII 脱敏 — 所查代码无", "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/toolcall.py"),
        ("✗ 官网「Documentation」链完整 API 参考 — 落地页无技术文档站", "https://openmanus.github.io/"),
        ("— 第五轮：上述条目全部复检通过，无新增 ✗/△", "https://github.com/FoundationAgents/OpenManus"),
    ]
    for text, url in checks:
        cite(text, url)

    doc.add_heading("防幻觉机制说明（五轮复核流程）", level=2)
    bullet("第一轮：确认官网 vs GitHub 权威来源，建立三入口分层。")
    bullet("第二轮：核对 Memory（消息缓冲）≠ 语义/RAG 长期记忆。")
    bullet("第三轮：核对 Web 路径（browser_use / web_search / MCP）及默认工具集。")
    bullet("第四轮：核对 OpenManus-RL 与主仓库分离；run_flow unstable 标注。")
    bullet("第五轮：全量重读核心源码路径，确认无 OpenClaw 式机制误标。")

    doc.add_heading("主要参考链接", level=1)
    for link in [
        "https://openmanus.github.io/",
        "https://github.com/FoundationAgents/OpenManus",
        "https://github.com/FoundationAgents/OpenManus/blob/main/README.md",
        "https://github.com/FoundationAgents/OpenManus/blob/main/config/config.example.toml",
        "https://github.com/FoundationAgents/OpenManus/blob/main/main.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/run_mcp.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/run_flow.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/base.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/toolcall.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/manus.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/agent/mcp.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/schema.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/llm.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/flow/planning.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/tool/browser_use_tool.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/tool/web_search.py",
        "https://github.com/FoundationAgents/OpenManus/blob/main/app/mcp/server.py",
        "https://github.com/OpenManus/OpenManus-RL",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
