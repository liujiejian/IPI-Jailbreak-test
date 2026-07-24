# -*- coding: utf-8 -*-
"""Generate LangGraph research notes (same structure as prior agent docs).

Paper track: Session / 任务模式 / web_fetch(security notice) / tools,
plus Prompt 组成 & 记忆持久化. 5-round anti-hallucination vs
docs.langchain.com (oss/python/langgraph).
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "LangGraph.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label}（来源：{url}）")

    # ── Title ──
    h("LangGraph", 0)
    p(
        "本文档结构对齐 OpenHands / Hermes / Claude-Code / Dify / AutoGPT / "
        "Cline / Codex-CLI / CrewAI 系列："
        "Session、任务模式、web_fetch（及 Security Notice 对照）、Tools；"
        "文末附 Prompt 组成、记忆与持久化机制，以及五轮防幻觉核对。"
        "对象为 LangGraph（低层编排 runtime：durable execution、streaming、HITL、persistence）。"
    )
    p("官方文档索引：https://docs.langchain.com/llms.txt")
    p(
        "产品分层（勿混）："
        "LangGraph = 编排与持久化；"
        "LangChain = models/tools/`create_agent` 等更高层；"
        "Deep Agents = 可选 harness；"
        "LangSmith Deployment / Agent Server = 部署与平台能力（含 Cron）。"
    )
    cite("LangGraph overview", "https://docs.langchain.com/oss/python/langgraph/overview")
    p(
        "执行链路（研究视角）："
        "invoke/stream(config.thread_id) → StateGraph 节点 → "
        "（可选）bind_tools / create_agent tool loop → "
        "checkpointer 落盘 → interrupt/resume 或下一轮同 thread。"
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "官方会话单元是 **thread**：通过 `config[\"configurable\"][\"thread_id\"]` "
        "选定 checkpointer 命名空间。同一 thread_id 续跑同一对话/工作流状态；"
        "新 thread_id = 空状态新游标。"
    )
    cite("Persistence", "https://docs.langchain.com/oss/python/langgraph/persistence")
    cite("Interrupts", "https://docs.langchain.com/oss/python/langgraph/interrupts")

    h("1. thread_id 与 Checkpointer", 2)
    bullet(
        "`graph.compile(checkpointer=...)` 后，invoke/stream 必须带 thread_id "
        "才能跨调用续状态。"
    )
    bullet(
        "Checkpoint = 某时刻图状态快照；thread = 一组有序 checkpoint。"
        "可选 checkpoint_id 定位某一历史快照（time travel）。"
    )
    bullet(
        "实现：InMemorySaver（进程内，重启丢失）；"
        "SqliteSaver / PostgresSaver / MongoDB / Redis 等生产后端。"
    )
    bullet("PostgresSaver 建议 thread_id < 255 字符。")
    cite(
        "Add memory / checkpointers",
        "https://docs.langchain.com/oss/python/langgraph/add-memory",
    )

    h("2. Interrupt / Resume（同 thread）", 2)
    bullet(
        "`interrupt(payload)` 暂停并持久化；调用方用 `Command(resume=...)` "
        "在同一 thread_id 上继续，resume 值成为 interrupt() 返回值。"
    )
    bullet(
        "stream_events(..., version=\"v3\") 可暴露 stream.interrupted / stream.interrupts。"
    )

    h("3. Agent Server（平台侧）", 2)
    bullet(
        "LangSmith Agent Server 可自动管理 checkpointer/store 基础设施；"
        "线程模型仍在 Deployment API 中扩展（含 Thread Cron）。"
    )

    h("4. 研究映射（非官方术语）", 2)
    bullet("Codex/Cline session ≈ LangGraph thread_id。")
    bullet("CrewAI conversational session_id ≈ thread_id（语义近，API 不同）。")
    bullet("Checkpoint lineage (CrewAI) ≈ LangGraph checkpoint 链 + optional fork/time-travel。")

    # ══════════════════════════════════════════════════════════════
    # 任务模式
    # ══════════════════════════════════════════════════════════════
    h("任务模式", 1)

    h("1. Workflows vs Agents", 2)
    bullet(
        "Workflow：预定代码路径（节点/边顺序）；"
        "Agent：动态自己选过程并使用 tools。"
    )
    bullet(
        "LangGraph 提供编排基建；常见 agent loop 可用 LangChain `create_agent` "
        "（编译为 LangGraph）或手写 StateGraph。"
    )
    cite(
        "Workflows and agents",
        "https://docs.langchain.com/oss/python/langgraph/workflows-agents",
    )

    h("2. 调用形态", 2)
    bullet("`invoke` / `ainvoke`、`stream` / `stream_events`。")
    bullet("Functional API 与 Graph API 两种编程面（官方并存）。")
    bullet("子图（subgraph）有独立/可配置的 checkpoint 命名空间。")

    h("3. Human-in-the-loop", 2)
    bullet("动态 interrupt（任意节点内）；静态 breakpoint（节点前后）也可用。")
    bullet("依赖 checkpointer；无持久化则无法可靠跨进程 resume。")

    h("4. 周期任务边界（对齐调研表「框架无*」）", 2)
    bullet(
        "开源 LangGraph 库本身不提供内置 cron / Heartbeat。"
    )
    bullet(
        "LangSmith Deployment「Use cron jobs」：按 cron 建新 thread + 固定 input 跑 assistant——"
        "属部署平台，非框架核心 API；表注「框架无*」与此一致。"
    )
    cite("Use cron jobs", "https://docs.langchain.com/langsmith/cron-jobs")

    h("5. 本地部署", 2)
    bullet("`pip install -U langgraph`（或 uv add）；纯库嵌入应用进程。")
    bullet("生产可部署到 LangSmith Deployment / 自建 Agent Server。")

    h("任务模式 · 核对摘要", 2)
    bullet(
        "保留：workflow/agent；thread+checkpointer；interrupt；"
        "周期任务仅平台 Cron；pip 库本地部署。"
    )
    bullet("剔除：把 LangSmith Cron 写成 langgraph 包内置；剔除 OpenClaw Heartbeat。")

    # ══════════════════════════════════════════════════════════════
    # web_fetch
    # ══════════════════════════════════════════════════════════════
    h("web_fetch（及 Security Notice）", 1)
    p(
        "结论（对齐表「须自建」）：LangGraph **没有**内置 `web_fetch` / 浏览器工具，"
        "也未文档化 OpenClaw 式 SECURITY NOTICE 包装。"
        "外网能力来自：用户自定义 `@tool` / LangChain 集成工具 / "
        "模型侧 server-side tools（若提供商支持）——"
        "回注入 messages 后的安全外观式包装由开发者自行决定。"
    )

    h("1. 工具需自建或接入集成", 2)
    bullet(
        "官方强调熟悉 LangChain models & tools；"
        "示例常用 `llm.bind_tools([...])` 或 `create_agent(..., tools=[...])`。"
    )
    bullet(
        "集成目录可含 Tavily Search 等网页搜索工具（LangChain integrations），"
        "不是 LangGraph core 内置。"
    )
    bullet(
        "部分 chat model 有 server-side web search / code interpreter——"
        "属模型提供方能力，非图运行时内置 scrape。"
    )
    cite("Tools (LangChain)", "https://docs.langchain.com/oss/python/langchain/tools")
    cite(
        "Tavily Search (integrations)",
        "https://docs.langchain.com/oss/python/integrations/tools/tavily_search",
    )

    h("2. Security Notice 对照", 2)
    bullet(
        "LangGraph/LangChain 工具文档示例返回普通字符串或结构化对象；"
        "无强制 EXTERNAL_UNTRUSTED_CONTENT / SECURITY NOTICE。"
    )
    bullet(
        "研论须在应用层把网页视为 untrusted（可参考 Codex「treat as untrusted」工程实践），"
        "但不要写成 LangGraph 默认行为。"
    )

    h("威胁模型映射（研究，非官方）", 2)
    bullet(
        "自建 scrape/search tool → ToolMessage → 后续节点/agent："
        "典型间接注入面；框架不代写 notice。"
    )
    bullet(
        "若用 Store 持久化网页派生事实：跨 thread 污染面扩大。"
    )
    bullet(
        "对比表行：CrewAI/Cline 有现成 scrape/fetch 工具名；"
        "LangGraph 行「须自建」准确。"
    )

    code(
        """# 研究示意：外网工具自建后挂到图/agent
@tool
def web_search(query: str) -> str:
    \"\"\"Search the web for information.\"\"\"
    return fetch_somehow(query)  # 返回值无框架规定的 SECURITY NOTICE

# llm.bind_tools([web_search]) 或 create_agent(model, tools=[web_search])
# 官方不强制 untrusted 外壳
"""
    )

    h("web_fetch · 核对摘要", 2)
    bullet("保留：无内置 web_fetch；工具自建/集成；无 SECURITY NOTICE。")
    bullet("剔除：把 Tavily/OpenAI web_search 写成 LangGraph 内置同名工具。")

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "编排层「原生」支持 tool-calling 循环与图节点副作用，"
        "但**工具实现本身需定义**（表：「原生, 工具需自定义」）。"
    )

    h("1. 挂载方式", 2)
    bullet("LangChain：`@tool`、`bind_tools`、`create_agent(tools=...)`。")
    bullet("StateGraph 节点内手动执行 tool_calls 并写回 MessagesState。")
    bullet("ToolRuntime 可访问 state / context / store / tool_call_id。")

    h("2. 与 Store / Command", 2)
    bullet("工具可通过 Command 更新状态；并行工具更新字段时注意 reducer。")
    bullet("工具可读写 BaseStore（长期记忆）。")

    h("Tools · 核对摘要", 2)
    bullet("保留：原生 tool-calling 编排 + 自定义/集成工具。")
    bullet("剔除：伪造内置固定工具表白名单（WebFetch/Bash…）。")

    # ══════════════════════════════════════════════════════════════
    # 五轮防幻觉
    # ══════════════════════════════════════════════════════════════
    h("五轮防幻觉核对记录", 1)

    h("Round 1 — Session", 2)
    bullet("核对：persistence；add-memory；interrupts。")
    bullet("保留：thread_id；checkpointer；interrupt/Command(resume)。")
    bullet("剔除：Claude JSONL 路径；CrewAI kickoff_id 混用。")

    h("Round 2 — 任务模式", 2)
    bullet("核对：overview；workflows-agents；cron-jobs；deploy。")
    bullet("保留：workflow/agent；pip；HITL；Cron 仅 Deployment。")
    bullet("剔除：框架内置调度器。")

    h("Round 3 — web_fetch / 安全", 2)
    bullet("核对：overview（须 models+tools）；langchain tools；integrations。")
    bullet("保留：须自建；无 SECURITY NOTICE。")
    bullet("剔除：内置 web_fetch。")

    h("Round 4 — Tools", 2)
    bullet("核对：langchain tools；workflows-agents bind_tools。")
    bullet("保留：原生编排 + 自定义工具。")
    bullet("剔除：把 Deep Agents 文件系统工具算作 LangGraph core 默认。")

    h("Round 5 — Prompt / Memory 交叉复检", 2)
    bullet("核对：add-memory；stores；persistence 表。")
    bullet(
        "冻结："
        "(1) Session ≡ thread_id + checkpoints；"
        "(2) 任务 = graph invoke/stream ± HITL；周期=平台 Cron；"
        "(3) 外网 = 自建/集成工具；"
        "(4) Prompt = 应用 state messages + 开发者节点逻辑；"
        "(5) 记忆 = checkpointer（短）+ store（长）。"
    )
    bullet(
        "开放问题：具体集成工具返回是否自带安全前缀；"
        "Agent Server 默认 checkpointer 后端品牌；"
        "子图 checkpoint 共享最佳实践需按版本核对。"
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt 组成
    # ══════════════════════════════════════════════════════════════
    h("Prompt 组成", 1)
    p(
        "LangGraph **不抽象 prompts or architecture**（overview）："
        "prompt 由你在节点/`create_agent` 系统层组装。"
        "常见载体：`MessagesState[\"messages\"]`、system 消息、structured output schema。"
    )

    h("1. Messages 状态", 2)
    bullet("多轮时 messages 经 checkpointer 随 thread 持久。")
    bullet("工具结果通常以 ToolMessage 写回消息列表。")

    h("2. create_agent / 手写节点", 2)
    bullet("高阶：LangChain agents 封装 tool loop 与中间件。")
    bullet("低阶：节点内 `model.invoke(state[\"messages\"])` 完全自控。")

    h("3. 无默认 AGENTS.md", 2)
    bullet("框架不装载仓库 AGENTS.md；可用应用代码或 Deep Agents 等另层加载。")

    h("Prompt 组成 · 核对摘要", 2)
    bullet("保留：开发者拥有 prompt；messages-centric。")
    bullet("剔除：宣称 LangGraph 内置固定 system template / SECURITY NOTICE。")

    # ══════════════════════════════════════════════════════════════
    # 记忆与持久化
    # ══════════════════════════════════════════════════════════════
    h("记忆与持久化机制", 1)
    p(
        "官方双轨（与调研表 Memory=checkpointer/store 一致）："
    )

    h("1. Checkpointer（短时 / thread-scoped）", 2)
    bullet("持久化图状态快照：对话连续性、HITL、time travel、容错。")
    bullet("访问：thread_id（+ 可选 checkpoint_id）。")
    bullet("内存实现仅适合开发；生产用 DB 后端。")

    h("2. Store（长时 / cross-thread）", 2)
    bullet(
        "应用自定义 KV：用户偏好、事实、共享知识；"
        "`compile(..., store=)`；节点/工具经 Runtime 访问。"
    )
    bullet("InMemoryStore vs PostgresStore / MongoDBStore / RedisStore 等。")
    bullet("可启用语义检索（semantic search over store items）。")
    cite("Stores", "https://docs.langchain.com/oss/python/langgraph/stores")

    h("3. 与「会话」关系", 2)
    bullet("多轮聊天：通常 checkpointer + MessagesState 即够。")
    bullet("跨多次会话记忆：Store（或外部 DB）。")
    bullet("二者常一起 compile。")

    code(
        """from langgraph.checkpoint.memory import InMemorySaver
from langgraph.store.memory import InMemoryStore

graph = builder.compile(
    checkpointer=InMemorySaver(),
    store=InMemoryStore(),
)
graph.invoke(inputs, {\"configurable\": {\"thread_id\": \"thread-1\"}})
"""
    )

    h("记忆与持久化 · 核对摘要", 2)
    bullet("保留：checkpointer≠store；thread 短记忆 / 跨 thread 长记忆。")
    bullet("剔除：Hermes MEMORY.md；CrewAI unified Memory 类当成 LangGraph API。")

    h("主要参考链接", 1)
    for link in [
        "https://docs.langchain.com/llms.txt",
        "https://docs.langchain.com/oss/python/langgraph/overview",
        "https://docs.langchain.com/oss/python/langgraph/persistence",
        "https://docs.langchain.com/oss/python/langgraph/add-memory",
        "https://docs.langchain.com/oss/python/langgraph/checkpointers",
        "https://docs.langchain.com/oss/python/langgraph/stores",
        "https://docs.langchain.com/oss/python/langgraph/interrupts",
        "https://docs.langchain.com/oss/python/langgraph/streaming",
        "https://docs.langchain.com/oss/python/langgraph/workflows-agents",
        "https://docs.langchain.com/oss/python/langchain/tools",
        "https://docs.langchain.com/oss/python/langchain/agents",
        "https://docs.langchain.com/oss/python/integrations/tools/tavily_search",
        "https://docs.langchain.com/langsmith/cron-jobs",
        "https://docs.langchain.com/langsmith/agent-server",
        "https://docs.langchain.com/oss/python/concepts/memory",
    ]:
        bullet(link)

    doc.save(str(OUT))
    # keep legacy filename used earlier in the folder
    alt = Path(__file__).resolve().parent / "LangGraph Agent.docx"
    doc.save(str(alt))
    print(f"Saved: {OUT}")
    print(f"Also: {alt}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
