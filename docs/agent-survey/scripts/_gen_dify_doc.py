# -*- coding: utf-8 -*-
"""Generate Dify research notes (same structure as OpenHands / Hermes / Claude-Code).

Paper track: Session / 任务模式 / web_fetch(security notice) / tools,
plus Prompt 组成 & 记忆持久化. 5-round anti-hallucination vs docs.dify.ai.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "Dify.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label}（来源：{url}）")

    # ── Title ──
    h("Dify", 0)
    p(
        "本文档结构对齐 OpenHands.docx / Hermes.docx / Claude-Code.docx："
        "Session、任务模式、web_fetch（及 Security Notice 对照）、Tools；"
        "文末附 Prompt 组成、记忆与持久化机制，以及五轮防幻觉核对。"
        "对象为开源 AI 应用平台 Dify（Studio 编排 Workflow/Chatflow/Agent 等应用）。"
    )
    p("官方文档索引：https://docs.dify.ai/llms.txt")
    p(
        "定位：在 Studio 中用拖拽编排 agentic 流程，发布为 Web App、API 或 MCP Server。"
        "底层统一为 Workflow 引擎；Chatflow / Agent 等是其上的产品形态。"
    )
    cite(
        "Key Concepts",
        "https://docs.dify.ai/en/use-dify/getting-started/key-concepts",
    )
    p(
        "执行链路（研究视角，Agent/Chatflow）："
        "User → chat-messages（conversation）→ LLM/Agent 节点 → "
        "Tool / HTTP Request / Knowledge Retrieval → 结果回注入 prompt → User。"
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "Dify 会话实体是 Conversation，由系统变量 sys.conversation_id 标识。"
        "同一 conversation 归并多轮消息，使 LLM 在同一话题与上下文上续聊。"
        "Chatflow 每轮对话触发一次 workflow 运行。"
    )
    cite(
        "Key Concepts > Chatflow variables",
        "https://docs.dify.ai/en/use-dify/getting-started/key-concepts",
    )

    h("1. conversation_id 与 API 续聊", 2)
    bullet(
        "聊天类 API：POST /v1/chat-messages；"
        "新会话留空 conversation_id，响应返回新 ID；后续请求带上以续聊。"
    )
    bullet(
        "已有 conversation_id 时，新的 inputs 会被忽略，仅处理 query"
        "（动态状态改用 conversation variables）。"
    )
    bullet(
        "Service API 与 WebApp 会话不共享："
        "相同 user 在两种渠道下历史分离。"
    )
    bullet("sys.dialogue_count：对话轮次计数，可配合 If-Else 分支。")
    cite(
        "Developing with APIs / Chat messages",
        "https://docs.dify.ai/en/self-host/use-dify/publish/developing-with-apis",
    )

    h("2. Workflow run vs Conversation", 2)
    bullet(
        "Workflow 应用：单轮任务；可用 sys.workflow_run_id 追踪每次执行；"
        "无 Chatflow 式多轮 conversation（除非外层自行管理）。"
    )
    bullet(
        "Chatflow：每轮 = 一次 Chatflow 运行 + 同一 conversation_id 下的历史/变量。"
    )
    bullet("日志/Traces：可在 Monitor > Logs 查看触发源与运行记录。")

    h("3. 与 OpenClaw / Hermes / Claude Code 的研究映射（非官方术语）", 2)
    bullet(
        "OpenClaw Main session / Claude Code 交互会话 ≈ "
        "Dify Chatflow/Chatbot 的同一 conversation_id。"
    )
    bullet(
        "OpenClaw Isolated / Automation 新会话 ≈ "
        "Workflow 单次 run，或 Schedule/Webhook Trigger 每次触发的独立 run。"
    )
    bullet(
        "跨「会话」知识：Knowledge Base / 环境变量；"
        "非 MEMORY.md 文件体系（见记忆节）。"
    )

    # ══════════════════════════════════════════════════════════════
    # 任务模式
    # ══════════════════════════════════════════════════════════════
    h("任务模式", 1)
    p(
        "应用类型与启动方式决定是否多轮、是否持久 conversation、以及是否无人值守。"
    )

    h("1. Chatflow（推荐多轮）", 2)
    bullet("每轮用户输入触发完整 Chatflow；支持 Memory、Conversation Variables、Answer 流式输出。")
    bullet("入口固定为 User Input（非 Trigger）。")
    cite(
        "Workflow & Chatflow",
        "https://docs.dify.ai/en/cloud/use-dify/build/workflow-chatflow.md",
    )

    h("2. Workflow（单轮 / 批处理 / 触发器）", 2)
    bullet("适合单次任务与批量；WebApp/API 易批跑。")
    bullet(
        "Start：User Input（人/API 触发）或 Trigger（自动）。"
    )

    h("3. Trigger（仅 Workflow）", 2)
    bullet("Schedule Trigger：cron / 可视化周期；每 workflow 最多一个 schedule。")
    bullet("Integration Trigger：第三方事件订阅（如 Slack）。")
    bullet("Webhook Trigger：外部 HTTP 回调；可从 query/header/body 抽变量。")
    bullet("可多 Trigger；Quick Settings 可启停已发布 Trigger。")
    cite(
        "Trigger Overview；Schedule Trigger",
        "https://docs.dify.ai/en/use-dify/nodes/trigger/overview",
    )

    h("4. 基础应用形态（遗留简化 UI）", 2)
    bullet("Chatbot：模型 + prompt 的简单对话。")
    bullet("Agent（含 New Agent）：自主工具调用的聊天应用。")
    bullet("Text Generator：单轮补全。")
    p("官方说明：底层仍是同一 workflow 引擎，界面为简化遗留形态。")

    h("5. 发布与调用面", 2)
    bullet("Web App / Embed。")
    bullet("Service API（Bearer Key；Cloud 基址 api.dify.ai/v1）。")
    bullet("MCP Server：把 Dify App 暴露给 Claude Desktop / Cursor 等。")
    bullet("difyctl：终端/CI/coding agent 调用。")

    h("6. 调度策略对照（研究笔记）", 2)
    p("OpenClaw：精确隔离 → Cron；完整 session → Heartbeat。")
    p("Dify 对应：")
    bullet("定时无人值守 → Workflow Schedule Trigger（新 run，非 Chatflow 续聊）。")
    bullet("事件驱动 → Webhook / Integration Trigger。")
    bullet("需对话记忆 → Chatflow + conversation_id（非 Trigger）。")

    # ══════════════════════════════════════════════════════════════
    # web_fetch / Security Notice
    # ══════════════════════════════════════════════════════════════
    h("web_fetch（及 Security Notice）", 1)
    p(
        "【结论】Dify 没有名为 web_fetch 的内置 Agent 工具，"
        "也没有 OpenClaw 式框架级 SECURITY NOTICE / EXTERNAL_UNTRUSTED_CONTENT 包裹。"
        "外部网页/API 内容主要通过 HTTP Request 节点、Tool/Agent 工具、"
        "以及 Knowledge「从网站导入」（Jina/Firecrawl）进入系统。"
    )

    h("1. 路径 A：HTTP Request 节点（工作流级 fetch）", 2)
    bullet(
        "支持 GET/HEAD/POST/PUT/PATCH/DELETE；"
        "URL/Headers/Body/Auth 可插入 {{variable}}。"
    )
    bullet(
        "输出拆为 Response Body、Status Code、Headers、Files、Size — "
        "下游以结构化变量引用，非自动加 SECURITY NOTICE。"
    )
    bullet("可配超时、重试（最多约 10 次）、失败分支、ssl_verify。")
    cite(
        "HTTP Request Node",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/http-request",
    )

    h("2. 路径 B：Tool / Agent 工具调用", 2)
    bullet(
        "Agent 节点或 Tool 节点调用 Workspace 工具插件、Swagger、MCP、Workflow-as-Tool。"
        "例如搜索类/爬取类插件（视安装与 Marketplace 而定）。"
    )
    bullet("工具返回作为 Observation / 节点输出进入后续 LLM，官方未规定自动 untrusted 边界文本。")
    cite(
        "Agent Node；Tool Node；Dify Tools",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/agent.md",
    )

    h("3. 路径 C：Knowledge 网站导入（延迟 / 索引面）", 2)
    p(
        "「Sync from website」用 Firecrawl 或 Jina Reader 将公开页解析为 Markdown 再入库。"
        "属于知识库构建时 ingest，不是同一轮 Agent 的实时 web_fetch；"
        "触发注入面是后续 Knowledge Retrieval → LLM Context。"
    )
    cite(
        "Import Data from Website",
        "https://docs.dify.ai/en/cloud/use-dify/knowledge/create-knowledge/import-text-data/sync-from-website.md",
    )

    h("4. Security Notice 对照与 SSRF（基础设施）", 2)
    bullet("OpenClaw：web_fetch 结果带 SECURITY NOTICE 包裹。")
    bullet("Hermes：<untrusted_tool_result>。")
    bullet("Claude Code：WebFetch 隔离上下文 + lossy 抽取 + 15min cache。")
    bullet(
        "Dify：产品节点文档不描述同等「内容标记」；"
        "自托管默认有 ssrf_proxy（Squid）拦截/过滤沙箱服务的出站请求，"
        "防 SSRF；域名未入 allowed_domains 会被拦。"
    )
    cite(
        "SSRF Proxy（Docker Issues / External KB）",
        "https://docs.dify.ai/en/self-host/deploy/troubleshooting/docker-issues",
    )
    p(
        "IPI 研究含义：假设 HTTP/Tool 响应正文可被模型当指令；"
        "应在编排层自加分隔/清洗，或依赖 SSRF ACL，不能假设 OpenClaw 式自动包裹。"
        "知识库网页路径是「索引后检索投毒」威胁，区别于实时 tool output。"
    )

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "工具在 Workspace Integrations > Tools 管理，可用于 "
        "Workflow/Chatflow 的 Tool 节点、Agent 节点，以及 Agent 类应用。"
    )
    cite(
        "Dify Tools",
        "https://docs.dify.ai/en/cloud/use-dify/workspace/tools.md",
    )

    h("1. 工具类型", 2)
    bullet("Tool Plugin：Marketplace / 内置（如 Current Time）；部分需授权。")
    bullet("Swagger/OpenAPI：粘贴或 URL 导入生成工具界面。")
    bullet("Workflow as Tool：仅 User Input 起步的 Workflow（Chatflow 不可作 Tool）。")
    bullet("MCP：仅 HTTP transport；可 Dynamic Client Registration / 自定义 Header/超时。")

    h("2. 编排中的关键节点", 2)
    bullet("Agent：Function Calling 或 ReAct；配置工具、Instructions、Max Iterations、Memory。")
    bullet("Tool：固化调用某工具 action（非自主选工具）。")
    bullet("HTTP Request：通用出站 HTTP。")
    bullet("Knowledge Retrieval：检索 KB → context 供 LLM。")
    bullet("Code：沙箱内 Python/JS（与 ssrf/sandbox 组件相关）。")
    bullet("Variable Assigner：更新 Conversation Variables。")

    h("3. Agent 策略", 2)
    bullet("Function Calling：走模型原生 tools 参数。")
    bullet("ReAct：Thought → Action → Observation 提示循环。")
    bullet("可从 Marketplace 安装更多 Agent Strategies。")

    h("4. 研究用简化 messages 示意（Chatflow + Agent）", 2)
    code(
        """{
  "messages": [
    {
      "role": "system",
      "content": "<Agent Instructions / LLM System prompt with {{vars}}>"
    },
    {
      "role": "user",
      "content": "Check https://api.example.com/health and return the status."
    },
    {
      "role": "assistant",
      "content": null,
      "tool_calls": [{
        "id": "call_001",
        "function": {
          "name": "<configured_http_or_plugin_tool>",
          "arguments": "{\\"url\\": \\"https://api.example.com/health\\"}"
        }
      }]
    },
    {
      "role": "tool",
      "tool_call_id": "call_001",
      "content": "{\\"status\\": 200, \\"body\\": \\"{\\\\\\"ok\\\\\\": true}\\"}"
    },
    {
      "role": "assistant",
      "content": "The API is healthy (status 200, ok=true)."
    }
  ]
}"""
    )
    p(
        "说明：实际走 Dify 工作流事件与 chat-messages API；"
        "工具名来自所装插件/Swagger，非固定 web_fetch。"
    )

    # ══════════════════════════════════════════════════════════════
    # 五轮防幻觉
    # ══════════════════════════════════════════════════════════════
    h("五轮防幻觉核对记录", 1)

    h("Round 1 — Session", 2)
    bullet("核对：key-concepts；chat-messages / developing-with-apis。")
    bullet("保留：conversation_id；API/WebApp 隔离；dialogue_count。")
    bullet("剔除：把 Claude Code JSONL 路径或 Hermes SessionDB 写成 Dify 存储格式。")

    h("Round 2 — 任务模式", 2)
    bullet("核对：trigger/overview；schedule-trigger；workflow-chatflow；build/agent。")
    bullet("保留：Chatflow vs Workflow；Trigger 仅 Workflow；Schedule 每应用最多一个。")
    bullet(
        "剔除：宣称 Chatflow 可用 Schedule Trigger；"
        "剔除 Dify Heartbeat API。"
    )

    h("Round 3 — web_fetch / 安全", 2)
    bullet(
        "核对：nodes/http-request；sync-from-website；ssrf_proxy Docker docs。"
    )
    bullet(
        "保留：无内置 web_fetch 名；HTTP Request 结构化输出；"
        "Jina/Firecrawl 为 KB ingest；SSRF 代理。"
    )
    bullet(
        "剔除：OpenClaw SECURITY NOTICE 默认存在；"
        "剔除「HTTP 响应必有 15min cache」（Claude Code/OpenClaw 特性，非 Dify HTTP 节点文档项）；"
        "剔除把网站导入当成实时 Agent web_fetch。"
    )

    h("Round 4 — Tools", 2)
    bullet("核对：workspace/tools；nodes/tools；nodes/agent。")
    bullet("保留：Plugin / Swagger / Workflow-as-Tool / MCP；Agent FC/ReAct。")
    bullet("剔除：把 Claude Code `WebFetch` PascalCase 工具当成 Dify 内置名。")

    h("Round 5 — Prompt / Memory 交叉复检", 2)
    bullet("核对：nodes/llm；key-concepts Conversation Variables；knowledge integrate。")
    bullet(
        "冻结："
        "(1) Session ≡ conversation_id；"
        "(2) 任务 = Chatflow / Workflow(+Triggers) / Agent·Chatbot；"
        "(3) 外部网 = HTTP Request + Tools + KB crawl；无框架 SECURITY NOTICE；"
        "(4) 记忆 = TokenBuffer Memory + Conversation Vars + Knowledge。"
    )
    bullet(
        "开放问题：Cloud 与自托管 SSRF 策略差异；"
        "具体 Marketplace「网页抓取」插件返回格式；"
        "需实机抓 LLM messages 验证 Memory 窗口注入模板。"
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt 组成
    # ══════════════════════════════════════════════════════════════
    h("Prompt 组成", 1)
    p(
        "LLM 节点按模型类型使用 Chat roles（System / User / Assistant）或 Completion 文本。"
        "提示中用 {{variable}} 引用工作流变量；可切换 Jinja2 做循环/条件。"
    )
    cite(
        "LLM Node",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/llm",
    )

    h("1. 系统 / 用户消息与变量", 2)
    bullet("System：角色与约束；User：{{user_input}} 等。")
    bullet("变量在送达模型前替换；支持深层路径 {{api_response.data.items[0].id}}。")
    bullet("Agent Instructions：自然语言 + Jinja2；Query 可来自上游节点。")

    h("2. Context / RAG", 2)
    bullet(
        "Knowledge Retrieval 输出接到 LLM 的 context；"
        "提示中引用如 {{knowledge_retrieval.result}}；Dify 可跟踪引用。"
    )

    h("3. Memory 注入（Chatflow）", 2)
    bullet(
        "开启 Memory：将先前轮次以 formatted user–assistant 形式并入后续 prompt；"
        "可编辑 USER 模板；TokenBufferMemory；节点级、不跨不同 conversation。"
    )
    bullet("Agent 节点同样有 Memory 窗口（消息条数/成本权衡）。")

    h("4. 采集研究提示", 2)
    bullet("Monitor Logs / 工作流运行详情查看节点 IO。")
    bullet("Service API streaming 事件观察文本块；无内置「全量 LLM mock」文档路径。")
    bullet("自托管可在 API 与模型 Provider 之间加代理抓包。")

    h("Prompt 组成 · 核对摘要", 2)
    bullet("保留：roles + {{var}} + Jinja2 + Knowledge context + Memory 开关。")
    bullet("剔除：Hermes SOUL.md / Claude Code CLAUDE.md 作为 Dify 默认装载文件。")

    # ══════════════════════════════════════════════════════════════
    # 记忆与持久化
    # ══════════════════════════════════════════════════════════════
    h("记忆与持久化机制", 1)
    p(
        "Dify 无 OpenClaw/Hermes 式 MEMORY.md 专用文件；"
        "也无 Claude Code auto memory 目录。"
        "跨轮与跨会话能力分散在：对话 Memory、Conversation Variables、Knowledge、Logs。"
    )

    h("1. TokenBuffer Memory（节点级对话窗）", 2)
    bullet("LLM / Agent 节点可选启用；缓冲近期用户–助手对。")
    bullet("窗越大上下文越全、token 越贵；不写入别的 conversation。")

    h("2. Conversation Variables（Chatflow）", 2)
    bullet("会话级可变状态（待办、累计成本等）；Variable Assigner 更新。")
    bullet("API 可有 conversation variables 的 list/update 端点。")
    cite(
        "Key Concepts > Conversation Variables",
        "https://docs.dify.ai/en/use-dify/getting-started/key-concepts",
    )

    h("3. Knowledge Bases（持久外部记忆 / RAG）", 2)
    bullet("文档/Notion/网站爬取入库；检索节点注入 LLM。")
    bullet("外部知识 API + SSRF 白名单（自托管）可接自建检索服务。")
    bullet("威胁面：索引时投毒 → 多会话 Retrieve（类似论文 RAG IPI，非实时 web_fetch）。")

    h("4. Environment Variables", 2)
    bullet("应用级密钥/常量，DSL 分享时可与机密分离；运行中不随便改。")

    h("5. 与其它 Agent 对照", 2)
    bullet(
        "Hermes MEMORY.md frozen snapshot ≈ Dify Conversation Vars + Memory 窗口"
        "（机制不同：无 md 文件快照层）。"
    )
    bullet(
        "Claude Code CLAUDE.md ≈ 部分可用「应用提示词 / DSL 固定文案」类比，"
        "但非仓库根文件自动加载体系。"
    )

    h("记忆与持久化 · 核对摘要", 2)
    bullet("保留：三种主记忆面（buffer / conv vars / KB）；API 续聊靠 conversation_id。")
    bullet(
        "剔除：默认存在 MEMORY.md；"
        "剔除 Trigger 运行自动继承某用户的 Chatflow 完整记忆"
        "（Trigger 开启新 Workflow run，无 Chatflow conversation）。"
    )

    h("主要参考链接", 1)
    for link in [
        "https://docs.dify.ai/llms.txt",
        "https://docs.dify.ai/en/use-dify/getting-started/key-concepts",
        "https://docs.dify.ai/en/cloud/use-dify/build/workflow-chatflow.md",
        "https://docs.dify.ai/en/use-dify/nodes/trigger/overview",
        "https://docs.dify.ai/en/use-dify/nodes/trigger/schedule-trigger",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/http-request",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/agent.md",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/llm",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/tools.md",
        "https://docs.dify.ai/en/cloud/use-dify/workspace/tools.md",
        "https://docs.dify.ai/en/cloud/use-dify/knowledge/create-knowledge/import-text-data/sync-from-website.md",
        "https://docs.dify.ai/en/self-host/use-dify/publish/developing-with-apis",
        "https://docs.dify.ai/en/self-host/deploy/troubleshooting/docker-issues",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
