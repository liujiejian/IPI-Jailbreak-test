# -*- coding: utf-8 -*-
"""Generate Claude Code research notes (English, same structure as OpenHands / Hermes).

Paper track: Session / task modes / web_fetch(security notice) / tools,
plus Prompt Composition & Memory/Persistence. 5-round anti-hallucination vs code.claude.com docs.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "en" / "Claude-Code.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label} (Source: {url})")

    # ── Title ──
    h("Claude Code", 0)
    p(
        "This document follows the same structure as OpenHands.docx / Hermes.docx: "
        "Session, task modes, web_fetch (and Security Notice comparison), Tools; "
        "appendices cover Prompt Composition, Memory & Persistence, and five-round anti-hallucination verification. "
        "Subject is Anthropic Claude Code (terminal / IDE / Desktop / Web agentic coding harness)."
    )
    p("Official documentation index: https://code.claude.com/docs/llms.txt")
    p(
        "Positioning: agentic harness on top of Claude models — provides tools, context management, "
        "and execution environment so the model can read/write codebases, run commands, search and fetch web pages."
    )
    cite(
        "How Claude Code works; Overview",
        "https://code.claude.com/docs/en/how-claude-code-works.md",
    )
    p(
        "Execution chain (research view): User → agentic loop (gather context → take action → verify) "
        "→ tools (Read/Edit/Bash/WebFetch…) → results fed back into context."
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "Session is a saved conversation bound to a project directory. "
        "CLI continuously writes local transcript (JSONL); "
        "Desktop / Web / VS Code extension each maintain independent session history; this page focuses on CLI."
    )
    cite(
        "Manage sessions",
        "https://code.claude.com/docs/en/sessions.md",
    )

    h("1. Storage and Recovery", 2)
    bullet(
        "Default path: ~/.claude/projects/<sanitized-cwd>/<session-id>.jsonl "
        "(directory name replaces non-alphanumeric chars in working path with -)."
    )
    bullet(
        "Recovery: claude --continue (most recent session); "
        "claude --resume / --resume <name|id>; /resume; --from-pr."
    )
    bullet(
        "claude -p / Agent SDK sessions do not appear in interactive picker, but can --resume with session ID "
        "(must look up within same project directory or its git worktree)."
    )
    bullet(
        "Configurable: CLAUDE_CONFIG_DIR; cleanupPeriodDays (default ~30 days); "
        "CLAUDE_CODE_SKIP_PROMPT_HISTORY; --no-session-persistence with -p."
    )
    p("Note: official states JSONL entry format is internal implementation, variable across versions; scripts should use /export or -p JSON, not hard-parse.")

    h("2. Branch / Fork / Context Management", 2)
    bullet("/branch or --fork-session: copy history to new session ID, original session preserved.")
    bullet("Session \"allow this time\" permissions do not carry to fork.")
    bullet("/clear clears current context (old session still resumable); /compact summary compression.")
    bullet("Checkpointing: file edit snapshots before changes, Esc Esc rewind (separate from git).")

    h("3. Research Mapping to OpenClaw / Hermes (Non-Official Terminology)", 2)
    bullet("OpenClaw Main session ≈ Claude Code interactive CLI/IDE main session + transcript.")
    bullet(
        "OpenClaw Isolated / Hermes Cron fresh ≈ "
        "Routines / Desktop scheduled / CI -p independent runs "
        "(or same-machine /loop but still attached to original session)."
    )
    bullet("Cross-session knowledge: CLAUDE.md + auto memory (not auto-merge transcript into new session).")

    # ══════════════════════════════════════════════════════════════
    # Task Modes
    # ══════════════════════════════════════════════════════════════
    h("Task Modes", 1)
    p("Determines whether session is shared, how permissions are approved, whether local machine must be online, and whether scheduling persists across restarts.")

    h("1. Interactive Session (Interactive)", 2)
    bullet("Shift+Tab cycles permission modes: Manual / Accept edits / Plan / Auto.")
    bullet(
        "Plan: explore and propose, no source edits; Accept edits: file edits and some filesystem commands without prompt."
    )
    bullet("Execution environment: Local / Cloud (Anthropic VM) / Remote Control (local execution, remote UI).")
    cite(
        "How Claude Code works; Permission modes",
        "https://code.claude.com/docs/en/permission-modes.md",
    )

    h("2. Non-Interactive / Headless (claude -p)", 2)
    bullet("claude -p \"...\": scripts and CI; configurable --allowedTools, --permission-mode, --output-format.")
    bullet(
        "--bare: skip auto-discovery of hooks/skills/plugins/MCP/auto memory/CLAUDE.md, "
        "only explicit flags apply; recommended for CI; may become -p default in future."
    )
    bullet("-p first trust confirmation not written to disk; Background Bash terminates ~5s after result returns.")
    cite(
        "Run Claude Code programmatically (headless)",
        "https://code.claude.com/docs/en/headless.md",
    )

    h("3. In-Session Scheduling: /loop and Cron* Tools", 2)
    bullet(
        "Tasks attached to current session: closing session stops them; --resume/--continue can restore unexpired tasks "
        "(loops ~7 days; one-shot before due time)."
    )
    bullet("Tools: CronCreate / CronList / CronDelete; max ~50 per session.")
    bullet("Also /loop, natural language one-shot reminder, Monitor (streaming background output).")
    bullet("Contrast: Routines (cloud) and Desktop scheduled do not depend on open interactive session.")
    cite(
        "Run prompts on a schedule",
        "https://code.claude.com/docs/en/scheduled-tasks.md",
    )

    h("4. Persistent Automation: Routines / Desktop / CI", 2)
    bullet(
        "Routines: Anthropic-hosted cloud, triggered by schedule / API / GitHub events; "
        "minimum interval ~1 hour; no local files (fresh clone); RemoteTrigger / /schedule."
    )
    bullet("Desktop scheduled tasks: run locally, can access local files; minimum ~1 minute.")
    bullet("GitHub Actions and other CI schedules.")
    cite(
        "Automate work with routines",
        "https://code.claude.com/docs/en/routines.md",
    )

    h("5. Scheduling Strategy Comparison (Research Notes)", 2)
    p("OpenClaw: precise isolation → Cron; full session → Heartbeat.")
    p("Claude Code equivalents:")
    bullet("In-session polling/reminders → /loop + Cron* (requires session alive).")
    bullet("Unattended, cross-machine reliable → Routines (cloud) or CI.")
    bullet("Needs local files → Desktop scheduled or locally kept-open session /loop.")

    # ══════════════════════════════════════════════════════════════
    # web_fetch / Security Notice
    # ══════════════════════════════════════════════════════════════
    h("web_fetch (and Security Notice)", 1)
    p(
        "Built-in tool name is WebFetch (PascalCase, same in permission rules); "
        "also WebSearch (returns titles/URLs only, does not fetch body)."
    )
    cite(
        "Tools reference > WebFetch / WebSearch",
        "https://code.claude.com/docs/en/tools-reference.md",
    )

    h("1. WebFetch Behavior", 2)
    bullet("Parameters: URL + prompt describing extraction target.")
    bullet("HTML → Markdown; small fast model extracts per prompt.")
    bullet(
        "[Key] For most requests, main-session Claude receives the small model's answer, not full raw page "
        "(lossy by design). For raw content use Bash curl, or more specific second prompt."
    )
    bullet("HTTP auto-upgrades to HTTPS; large pages truncated at fixed character limit.")
    bullet(
        "[Similar to OpenClaw] Response cached ~15 minutes: repeated fetch of same URL uses cache. "
        "IPI/testing should add timestamp or equivalent URL parameter to avoid cache."
    )
    bullet("User-Agent starts with Claude-User; Accept prefers Markdown.")

    h("2. Security Notice Comparison: Isolated Context (Not SECURITY NOTICE Verbatim)", 2)
    p("OpenClaw: tool results wrapped with SECURITY NOTICE + EXTERNAL_UNTRUSTED_CONTENT.")
    p("Hermes: <untrusted_tool_result source=\"web_extract\">.")
    p(
        "Claude Code official security page: "
        "\"Isolated context windows: Web fetch uses a separate context window "
        "to avoid injecting potentially malicious prompts\" — "
        "external pages first processed by extraction model in isolated context, then result returned to main agent, "
        "reducing risk of raw malicious instructions entering main session directly."
    )
    cite(
        "Security > Protect against prompt injection",
        "https://code.claude.com/docs/en/security.md",
    )
    bullet(
        "Permissions: under default/acceptEdits, first WebFetch to new domain prompts "
        "(few built-in doc domains pre-approved); rules like WebFetch(domain:example.com)."
    )
    bullet(
        "auto / bypassPermissions can skip domain prompts; "
        "explicit deny/ask/allow overrides pre-approved set."
    )
    bullet(
        "Network Bash (curl/wget) not auto-approved by default; "
        "deny WebFetch does not block Bash curl — needs permissions + optional sandbox network rules as two layers."
    )
    bullet("Sandbox allowedDomains/deniedDomains complements WebFetch permissions; do not conflate.")

    h("3. IPI Research Implications", 2)
    bullet(
        "Primary injection surface remains: WebFetch extraction result entering main session; "
        "WebSearch returned URL list; Bash network; MCP tool output; "
        "untrusted CLAUDE.md / repository files."
    )
    bullet(
        "Compared to \"raw full-page feed\": lossy extraction can reduce raw injection, "
        "but attacker may poison extraction model/prompt, or user may bypass isolation via curl."
    )
    bullet("Reproduction experiments should add URL cache-buster (15min).")

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "Built-in tool names are exact strings in permissions/hooks/SDK; "
        "extensions via MCP (separate names) and Skill (executed via Skill tool, no new tool table entry)."
    )
    cite(
        "Tools reference; How Claude Code works",
        "https://code.claude.com/docs/en/tools-reference.md",
    )

    h("1. Tools Most Relevant to the Paper", 2)
    bullet("Web: WebFetch, WebSearch.")
    bullet("Files: Read, Edit, Write, Glob, Grep, NotebookEdit, LSP.")
    bullet("Execution: Bash (+ optional PowerShell); Monitor.")
    bullet("Orchestration: Agent (subagent), Skill, Task*, Workflow, AskUserQuestion.")
    bullet("Session scheduling: CronCreate / CronList / CronDelete; ScheduleWakeup (/loop).")
    bullet("Cloud routines: RemoteTrigger (Routines, plan-limited).")

    h("2. Permissions and Sandbox (Tool Gates)", 2)
    bullet(
        "Permission rules: evaluated before tool execution, cover Bash/Read/Edit/WebFetch/MCP, etc."
    )
    bullet(
        "Sandbox: OS-level restrictions on Bash subprocess filesystem/network; "
        "not equivalent to permission mode; /sandbox configures boundaries."
    )
    bullet(
        "Whole-process isolation optional via sandbox-runtime / container / Cloud VM "
        "(MCP and hooks default on host unless fully enclosed in boundary)."
    )
    cite(
        "Permissions; Sandboxing; Security",
        "https://code.claude.com/docs/en/permissions.md",
    )

    h("3. Hooks (Deterministic Policy, Not Memory)", 2)
    bullet(
        "PreToolUse / PostToolUse / SessionStart / Stop, etc.: "
        "can intercept dangerous commands, inject context; complements CLAUDE.md \"soft guidance\"."
    )
    cite(
        "Hooks guide",
        "https://code.claude.com/docs/en/hooks-guide.md",
    )

    h("4. Simplified Messages Example for Research", 2)
    code(
        """{
  "messages": [
    {
      "role": "system",
      "content": "<Claude Code system + CLAUDE.md + auto MEMORY.md snippet + skills index>"
    },
    {
      "role": "user",
      "content": "Check https://api.example.com/health and return the status."
    },
    {
      "role": "assistant",
      "content": null,
      "tool_calls": [{
        "id": "call_001",
        "type": "function",
        "function": {
          "name": "WebFetch",
          "arguments": "{\\"url\\": \\"https://api.example.com/health\\", \\"prompt\\": \\"Return HTTP status and body summary.\\"}"
        }
      }]
    },
    {
      "role": "tool",
      "tool_call_id": "call_001",
      "content": "<extractor model answer, not necessarily raw body>"
    },
    {
      "role": "assistant",
      "content": "The endpoint looks healthy."
    }
  ],
  "tools": [
    {"name": "WebFetch", "...": "..."},
    {"name": "WebSearch", "...": "..."},
    {"name": "Bash", "...": "..."},
    {"name": "Read", "...": "..."},
    {"name": "Edit", "...": "..."}
  ]
}"""
    )
    p("Illustration only: real protocol is Anthropic Messages / Agent SDK event stream; tool name casing per Tools reference.")

    # ══════════════════════════════════════════════════════════════
    # Five-Round Anti-Hallucination
    # ══════════════════════════════════════════════════════════════
    h("Five-Round Anti-Hallucination Verification Record", 1)

    h("Round 1 — Session", 2)
    bullet("Verified: sessions.md, how-claude-code-works.md.")
    bullet("Retained: JSONL per project; --continue/--resume/--fork; /clear /compact.")
    bullet("Removed: writing Hermes SessionDB or OpenClaw session:xxx as Claude Code API.")

    h("Round 2 — Task Modes", 2)
    bullet("Verified: scheduled-tasks.md, routines.md, headless.md, permission-modes.")
    bullet(
        "Retained: /loop session scope + 7 days; Routines cloud persistent; Desktop; -p/--bare."
    )
    bullet(
        "Removed: claiming /loop equals Hermes gateway cron (cross-restart unrelated session still runs); "
        "removed Claude Code has no WebFetch (it does have same-named tool)."
    )

    h("Round 3 — WebFetch / Security", 2)
    bullet("Verified: tools-reference WebFetch, security.md, permissions.md, sandboxing.md.")
    bullet(
        "Retained: lossy extraction; 15min cache; isolated context; WebFetch(domain:); curl not auto-blocked by deny WebFetch."
    )
    bullet(
        "Removed: OpenClaw SECURITY NOTICE verbatim copy; "
        "removed \"main session always sees full HTML\"; "
        "removed \"no cache\" (official states 15 minutes)."
    )

    h("Round 4 — Tools", 2)
    bullet("Verified: tools-reference full tool name table.")
    bullet(
        "Retained: PascalCase built-in names (WebFetch, Bash, Read…); MCP/Skill extension model."
    )
    bullet(
        "Removed: OpenHands terminal/file_editor or Hermes web_extract as Claude Code built-in names."
    )

    h("Round 5 — Prompt / Memory Cross-Check", 2)
    bullet("Verified: memory.md, prompt-caching.md, how-claude-code-works context.")
    bullet(
        "Frozen: "
        "(1) Session = local JSONL project sessions; "
        "(2) Tasks = interactive / -p / session Cron / Routines / Desktop / CI; "
        "(3) WebFetch = isolated extraction context + 15min cache, not SECURITY NOTICE wrapping; "
        "(4) Memory = CLAUDE.md + auto MEMORY.md (soft context, hard policy via hooks/permissions)."
    )
    bullet(
        "Open questions: whether live tool_result contains any visible untrusted marker "
        "(official emphasizes isolated context and lossy answer, does not promise OpenClaw-style marker string); "
        "Cloud vs local CLI default permission differences; version changes in Cron/TodoWrite behavior."
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt Composition
    # ══════════════════════════════════════════════════════════════
    h("Prompt Composition", 1)
    p(
        "Each API request resends: system prompt + project context + history messages/tool results + new messages. "
        "Claude Code auto-manages prompt caching; cache matches by prefix exactly."
    )
    cite(
        "How Claude Code uses prompt caching",
        "https://code.claude.com/docs/en/prompt-caching.md",
    )

    h("1. Cache Layers (Conceptual)", 2)
    bullet("System prompt layer: tool definitions, etc.; changes easily invalidate entire prefix.")
    bullet(
        "Project context: CLAUDE.md, auto memory, unscoped rules — "
        "loaded at session start; rebuildable after /clear or /compact."
    )
    bullet("Conversation: messages and tool results; generally append without breaking prefix.")
    bullet(
        "Model and effort also cache key: /model, /effort switch causes full window cache miss."
    )

    h("2. CLAUDE.md and System Prompt Customization", 2)
    bullet(
        "CLAUDE.md / .claude/rules / managed / user / local multi-level loading "
        "(see Memory section); context not mandatory config."
    )
    bullet(
        "CLAUDE.md edits do not mid-session rewrite established cache prefix by default; "
        "new content often needs new session, /clear, or /compact to enter subsequent requests."
    )
    bullet(
        "-p / SDK: --system-prompt replaces; "
        "--append-system-prompt / --append-system-prompt-file appends."
    )
    bullet("Skills: descriptions visible at startup, full content loaded on demand; can defer MCP tools (ToolSearch).")

    h("3. Agent SDK / Collection", 2)
    bullet("Agent SDK can modify system prompt preset, permission callbacks, session resume.")
    bullet(
        "Research capture: ANTHROPIC_BASE_URL / LLM gateway; "
        "or -p --output-format stream-json; do not rely on private JSONL schema."
    )
    cite(
        "Modifying system prompts (Agent SDK)",
        "https://code.claude.com/docs/en/agent-sdk/modifying-system-prompts.md",
    )

    h("Prompt Composition · Verification Summary", 2)
    bullet("Retained: three-layer cache view; CLAUDE.md not immediately effective mid-session; lossy WebFetch not in system.")
    bullet("Removed: writing Hermes SOUL/MEMORY volatile tier names as Claude Code API.")

    # ══════════════════════════════════════════════════════════════
    # Memory and Persistence
    # ══════════════════════════════════════════════════════════════
    h("Memory and Persistence", 1)
    p(
        "Each session context window starts empty; cross-session relies on "
        "CLAUDE.md (you write) and auto memory (Claude writes). "
        "Both are context, not hard policy — hard gates use PreToolUse hooks / permissions."
    )
    cite(
        "How Claude remembers your project",
        "https://code.claude.com/docs/en/memory.md",
    )

    h("1. CLAUDE.md Hierarchy", 2)
    bullet("Managed: OS-level org policy path (e.g. /etc/claude-code/CLAUDE.md).")
    bullet("User: ~/.claude/CLAUDE.md.")
    bullet("Project: ./CLAUDE.md or ./.claude/CLAUDE.md.")
    bullet("Local: ./CLAUDE.local.md (gitignore personal preferences).")
    bullet("Parent directory CLAUDE.md fully loaded; subdirs on demand; .claude/rules can path-scope.")
    bullet("@imports can nest (max ~4 hops); /init can generate starter file.")

    h("2. Auto Memory", 2)
    bullet(
        "Directory: ~/.claude/projects/<project>/memory/ "
        "(shared across git repo worktrees); contains MEMORY.md index + topic files."
    )
    bullet(
        "Each session loads only first 200 lines or first 25KB of MEMORY.md (whichever first); "
        "detailed notes in topic files."
    )
    bullet("On by default; /memory to disable; CLAUDE_CODE_DISABLE_AUTO_MEMORY=1; configurable autoMemoryDirectory.")
    bullet("Machine-local, does not auto-sync to cloud environments.")

    h("3. Session Transcript Persistence", 2)
    bullet("JSONL saves full conversation and tool trajectory, supports resume/fork/export.")
    bullet("Separate from \"memory\": new session does not auto-load old transcript unless --resume.")

    h("4. Comparison with OpenClaw / Hermes", 2)
    bullet(
        "OpenClaw MEMORY.md + daily notes ≈ Claude Code CLAUDE.md (instructions) "
        "+ auto MEMORY.md (learning); no official daily notes file convention."
    )
    bullet(
        "Hermes frozen snapshot mid-session not updating prompt ≈ "
        "Claude Code CLAUDE.md edits often wait for compact/new session; "
        "but auto memory can read/write files mid-session (UI shows Writing/Recalling memory)."
    )

    h("Memory and Persistence · Verification Summary", 2)
    bullet("Retained: dual-track memory; MEMORY.md load limit; soft context vs hooks.")
    bullet(
        "Removed: conflating platform API Memory Tool (/memories) with Claude Code product auto memory "
        "(former is Messages API general tool, hosted on application side; product doc path is "
        "~/.claude/projects/.../memory/)."
    )

    h("Primary Reference Links", 1)
    for link in [
        "https://code.claude.com/docs/llms.txt",
        "https://code.claude.com/docs/en/how-claude-code-works.md",
        "https://code.claude.com/docs/en/sessions.md",
        "https://code.claude.com/docs/en/scheduled-tasks.md",
        "https://code.claude.com/docs/en/routines.md",
        "https://code.claude.com/docs/en/headless.md",
        "https://code.claude.com/docs/en/tools-reference.md",
        "https://code.claude.com/docs/en/security.md",
        "https://code.claude.com/docs/en/permissions.md",
        "https://code.claude.com/docs/en/sandboxing.md",
        "https://code.claude.com/docs/en/memory.md",
        "https://code.claude.com/docs/en/prompt-caching.md",
        "https://code.claude.com/docs/en/hooks-guide.md",
        "https://code.claude.com/docs/en/permission-modes.md",
    ]:
        bullet(link)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
