# -*- coding: utf-8 -*-
"""Generate Codex CLI research notes (same structure as prior agent docs).

Paper track: Session / 任务模式 / web_fetch(security notice) / tools,
plus Prompt 组成 & 记忆持久化. 5-round anti-hallucination vs
developers.openai.com/codex (learn.chatgpt.com mirrors).
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "Codex-CLI.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label}（来源：{url}）")

    # ── Title ──
    h("Codex CLI", 0)
    p(
        "本文档结构对齐 OpenHands.docx / Hermes.docx / Claude-Code.docx / "
        "Dify.docx / AutoGPT.docx / Cline.docx："
        "Session、任务模式、web_fetch（及 Security Notice 对照）、Tools；"
        "文末附 Prompt 组成、记忆与持久化机制，以及五轮防幻觉核对。"
        "主对象为本地终端代理 Codex CLI（与 IDE 扩展 / Desktop 共享配置层）；"
        "Codex Cloud 仅作对照，不作本文主威胁面。"
    )
    p("官方文档索引：https://developers.openai.com/codex/llms.txt")
    p(
        "定位：本机 coding agent — OS 沙箱内读改文件、执行命令；"
        "可选 web search / MCP / Skills / Subagents；"
        "默认网络关闭于 workspace-write 沙箱内命令。"
    )
    cite(
        "Agent approvals & security",
        "https://developers.openai.com/codex/agent-approvals-security",
    )
    p(
        "执行链路（研究视角 · CLI）：User → TUI/`codex exec` turn → "
        "tools（shell / web_search / MCP / skill scripts…）→ "
        "结果入 transcript →（可选）Memories / AGENTS.md 反馈。"
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "交互式 CLI 以可恢复的 session / conversation 为单位持久化；"
        "非交互 `codex exec` 的 JSON 流会出现 thread_id。"
        "本地状态根目录为 CODEX_HOME（默认 ~/.codex）。"
    )

    h("1. 存储与续聊", 2)
    bullet(
        "默认将本地会话 transcript 写入 CODEX_HOME"
        "（文档示例：~/.codex/history.jsonl）。"
    )
    bullet(
        "[history] persistence = \"none\" 可关闭；"
        "history.max_bytes 可截断最旧条目并 compact。"
    )
    bullet(
        "`codex resume`：按 ID 或最近会话续聊；"
        "`--last` 默认限定当前工作目录，`--all` 跨目录。"
    )
    bullet(
        "`codex archive` / `unarchive` / `delete`；"
        "TUI 有 `/resume` `/archive` `/delete` `/fork` `/rename`。"
    )
    bullet(
        "`codex exec --ephemeral`：不把 session rollout 文件持久化到磁盘。"
    )
    cite(
        "History persistence",
        "https://developers.openai.com/codex/config-file/config-advanced",
    )
    cite(
        "CLI reference · resume / exec",
        "https://developers.openai.com/codex/cli/reference",
    )

    h("2. Task vs Session（产品语义）", 2)
    bullet(
        "`/new`：同一 CLI session 内开启新 task（清任务上下文，不离开进程）。"
    )
    bullet(
        "`/fork`：分叉当前 conversation 为新 task。"
    )
    bullet(
        "`/compact`：压缩可见对话以腾上下文 window。"
    )
    bullet(
        "JSONL 事件：thread.started / turn.* / item.*（含 web searches、MCP、command_execution）。"
    )

    h("3. 与其它 Agent 对照（非官方术语）", 2)
    bullet("Claude Code 项目 JSONL ≈ Codex history.jsonl + resume。")
    bullet("Cline sessions.db/[id].json ≈ Codex session ID + history 持久化。")
    bullet("勿把 Dify conversation_id API 字段写成 Codex 契约。")

    # ══════════════════════════════════════════════════════════════
    # 任务模式
    # ══════════════════════════════════════════════════════════════
    h("任务模式", 1)

    h("1. Interactive TUI（`codex`）", 2)
    bullet("无子命令启动交互 UI；可挂图；slash 控制会话中途策略。")
    bullet(
        "Web search 默认 cached；`--search` 切 live。"
        "低摩擦本地开发常用：`--sandbox workspace-write --ask-for-approval on-request`。"
    )
    cite("CLI reference · interactive", "https://developers.openai.com/codex/cli/reference")

    h("2. Non-interactive（`codex exec`）", 2)
    bullet("脚本/CI：进度在 stderr，最终消息在 stdout；`--json` 得 JSONL 事件流。")
    bullet("默认 read-only 沙箱；需写盘时显式 `--sandbox workspace-write`。")
    bullet(
        "`--ignore-user-config` / `--ignore-rules` 用于受控自动化；"
        "`--full-auto` 已弃用兼容路径。"
    )
    cite(
        "Non-interactive mode",
        "https://developers.openai.com/codex/non-interactive-mode",
    )

    h("3. Sandbox × Approval 组合", 2)
    bullet(
        "两层：**Sandbox**（技术边界：写哪、能否联网）+"
        "**Approval policy**（何时停下来问人）。"
    )
    bullet(
        "常用：Auto（workspace-write + on-request）；"
        "read-only；untrusted 审批；"
        "`--yolo` / danger-full-access（无沙箱无审批，官方不推荐）。"
    )
    bullet(
        "`approvals_reviewer = auto_review`：合格审批请求可先过 reviewer agent。"
    )
    bullet(
        "版本目录推荐 Auto；非版本目录倾向 read-only；"
        "可 `/permissions` 中途切换。"
    )
    cite(
        "Agent approvals & security",
        "https://developers.openai.com/codex/agent-approvals-security",
    )
    cite("Sandbox", "https://developers.openai.com/codex/sandboxing")

    h("4. 自动化 / 多代理 / 云（边界）", 2)
    bullet("Scheduled tasks / Automations：文档以 Desktop 「Scheduled」为主；可配 skills。")
    bullet("Subagents / multi_agent feature：可开协作工具。")
    bullet(
        "Cloud：隔离容器 + 两阶段（setup 可联网装依赖，agent 阶段默认离线）——"
        "与本机 CLI 沙箱不同，论文勿混。"
    )

    h("任务模式 · 核对摘要", 2)
    bullet("保留：TUI / exec；sandbox×approval；yolo；resume；exec 默认只读。")
    bullet("剔除：把 Cloud 互联网默认写成 CLI workspace-write 默认网络开启。")

    # ══════════════════════════════════════════════════════════════
    # web_fetch
    # ══════════════════════════════════════════════════════════════
    h("web_fetch（及 Security Notice）", 1)
    p(
        "结论：Codex CLI 文档中的外网页入口是 first-party **web search tool**"
        "（配置键 `web_search`），不是名为 `web_fetch` / `WebFetch` 的工具。"
        "官方明确要求将全部 web results 视为 untrusted input，"
        "并讨论 cached 模式降低任意 live 内容带来的 prompt injection，"
        "但公开文档未给出 OpenClaw 式 SECURITY NOTICE / EXTERNAL_UNTRUSTED_CONTENT 包装文本。"
    )

    h("1. web_search 模式", 2)
    bullet("`cached`（默认）：OpenAI 维护的网页索引/缓存，非即时任意拉页。")
    bullet("`indexed`：外部访问经搜索索引 gating。")
    bullet("`live`（=`--search`）：取较新的 live 数据。")
    bullet("`disabled`：关闭该工具。")
    bullet(
        "使用 `--yolo` 或其它 full-access 沙箱时，web search 默认为 live。"
    )
    cite("Web search", "https://developers.openai.com/codex/web-search")
    cite(
        "Config basics · Web search mode",
        "https://developers.openai.com/codex/config-file/config-basic",
    )

    h("2. Untrusted 表述 vs SECURITY NOTICE", 2)
    bullet(
        "文档原意：「Treat all web results as untrusted input」；"
        "cached「降低但不消除」PI；"
        "启用网络或 live search 时需谨慎——「agent 可能 fetch and follow untrusted instructions」。"
    )
    bullet(
        "这与 Claude Code「孤立小模型抽取 + ~15min cache」描述不同；"
        "与 OpenClaw 固定 NOTICE 外壳不同——勿互相抄贴。"
    )
    bullet(
        "另：`memories.disable_on_external_context` 可将含 MCP / web search 的任务"
        "排除出 memory 生成（legacy alias：no_memories_if_mcp_or_web_search）。"
    )

    h("3. 与「命令出网」分离", 2)
    bullet(
        "默认 workspace-write：**spawned commands 无网络**"
        "（需 `[sandbox_workspace_write] network_access = true`）。"
    )
    bullet(
        "可在不开全量命令网络的情况下使用 web search tool"
        "（approvals 页明确对照 platform web_search guide）。"
    )
    bullet(
        "可选 features.network_proxy：对已开启的命令网络做域名 allowlist / 本地阻断 / DNS 检查。"
    )

    h("威胁模型映射（研究，非官方）", 2)
    bullet(
        "cached/live web 结果 → 模型上下文 → shell/MCP/改文件："
        "间接注入 / 目标劫持面；官方已点名 PI。"
    )
    bullet(
        "live + 命令 network_access=true 扩大可执行外联面；"
        "yolo 同时去掉沙箱与审批。"
    )
    bullet(
        "对比表：OpenClaw web_fetch+NOTICE；Claude WebFetch；"
        "Cline fetch_web；Codex = web_search(cached|live) + 「treat as untrusted」。"
    )

    code(
        """配置示例（官方文档）：
web_search = \"cached\"   # 默认
# web_search = \"live\"    # 同 --search
# web_search = \"disabled\"

# 命令出网（与 web_search 分离）
[sandbox_workspace_write]
network_access = true
"""
    )

    h("web_fetch · 核对摘要", 2)
    bullet(
        "保留：工具语义=web_search；四模式；默认 cached；untrusted 警告；"
        "无 SECURITY NOTICE 文档；命令网络默认关。"
    )
    bullet("剔除：内置工具名 web_fetch；OpenClaw NOTICE 默认存在。")

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "能力来自：沙箱内 shell/exec、web_search、MCP、Skills 脚本、"
        "（可选）multi_agent/subagent 协作工具；配置可关 namespace。"
    )

    h("1. 内置执行与检索", 2)
    bullet("`shell_tool` feature（默认 true）：默认 shell 工具。")
    bullet("`unified_exec`：统一 PTY-backed exec（Windows 默认策略不同）。")
    bullet("web_search feature / 顶层 `web_search = ...` 控制检索工具。")
    bullet("JSON item 类型含 command_execution、file changes、web searches、MCP tool calls、plan updates。")

    h("2. MCP / Skills / Rules", 2)
    bullet(
        "MCP：STDIO / Streamable HTTP；配置在 config.toml；"
        "Desktop/CLI/IDE 共享；`/mcp` 列出工具。"
    )
    bullet("Skills：SKILL.md + scripts/references；`/skills`；可审批 skill scripts。")
    bullet(
        "Rules（experimental）：`prefix_rule` 控制沙箱外命令 allow/prompt/forbidden。"
    )
    cite("MCP", "https://developers.openai.com/codex/extend/mcp")
    cite("Rules", "https://developers.openai.com/codex/agent-configuration/rules")

    h("3. 审批面延伸到非 shell", 2)
    bullet("带副作用的 app/connector/MCP 调用也会审批。")
    bullet("标注 destructive 的 MCP 调用在有注解时总是要审批。")
    bullet("granular approval 可分别控制 sandbox / rules / MCP / request_permissions / skill_approval。")

    h("Tools · 核对摘要", 2)
    bullet("保留：shell/exec + web_search + MCP + Skills + Rules。")
    bullet("剔除：伪造固定 PascalCase WebFetch 为唯一名。")

    # ══════════════════════════════════════════════════════════════
    # 五轮防幻觉
    # ══════════════════════════════════════════════════════════════
    h("五轮防幻觉核对记录", 1)

    h("Round 1 — Session", 2)
    bullet("核对：config-advanced History；cli-reference resume/archive；exec JSON thread_id。")
    bullet("保留：~/.codex/history.jsonl；resume；/new·/fork·/compact；ephemeral。")
    bullet("剔除：Claude ~/.claude/projects 路径；Cline sessions.db 结构。")

    h("Round 2 — 任务模式", 2)
    bullet("核对：cli-reference；non-interactive；approvals；sandboxing。")
    bullet("保留：TUI vs exec；sandbox×approval 表；yolo；Cloud 两阶段对照。")
    bullet("剔除：把 Desktop Scheduled 断言为 CLI 唯一调度引擎而不加备注。")

    h("Round 3 — web_fetch / 安全", 2)
    bullet("核对：web-search；config-basic web_search；approvals 网络段。")
    bullet(
        "保留：web_search 四模式；cached 默认；treat as untrusted；"
        "命令网络默认关；无 SECURITY NOTICE 包装文档。"
    )
    bullet("剔除：工具名 web_fetch；OpenClaw NOTICE；Claude 15min isolated extract 照搬。")

    h("Round 4 — Tools", 2)
    bullet("核对：features 表 shell_tool/web_search；mcp；rules；skills。")
    bullet("保留：shell + web_search + MCP + Skills；destructive MCP 审批。")
    bullet("剔除：只谈 curl 而无 web_search 官方工具。")

    h("Round 5 — Prompt / Memory 交叉复检", 2)
    bullet("核对：agents-md；memories；customization overview；/compact·/memories。")
    bullet(
        "冻结："
        "(1) Session = history + resume/session id；"
        "(2) 模式 = TUI/exec + sandbox/approval；"
        "(3) 外网检索 = web_search(cached|live|…) + untrusted 指引；"
        "(4) Prompt = AGENTS.md 链 + Skills +（可选）Memories；"
        "(5) Memories 默认关，路径 ~/.codex/memories/。"
    )
    bullet(
        "开放问题：web_search 结果进入模型的具体消息包装格式；"
        "live 是否等价任意 URL fetch；"
        "history.jsonl 与 per-session rollout 文件的精确布局需实机核对。"
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt 组成
    # ══════════════════════════════════════════════════════════════
    h("Prompt 组成", 1)
    p(
        "启动时装载指令链 + 用户消息/附件 + 工具结果；"
        "完整 system 模板未全部公开。可配置 compact prompt 覆盖等。"
    )

    h("1. AGENTS.md 指令链", 2)
    bullet(
        "全局：~/.codex/AGENTS.override.md 优先，否则 AGENTS.md。"
    )
    bullet(
        "项目：自 Git 根走到 cwd，每层至多一个"
        "AGENTS.override.md / AGENTS.md / fallback 文件名。"
    )
    bullet(
        "自上而下拼接；越靠近 cwd 越靠后、越覆盖；"
        "受 project_doc_max_bytes（默认 32KiB）限制。"
    )
    cite(
        "Custom instructions with AGENTS.md",
        "https://developers.openai.com/codex/agent-configuration/agents-md",
    )

    h("2. Skills / MCP instructions / 用户层", 2)
    bullet("Skills 元数据对 agent 可见；SKILL.md 在选用时加载。")
    bullet("MCP initialization 的 instructions 字段作 server-wide 指引。")
    bullet("personality、/personality；debug 可查 instruction discovery。")

    h("3. Memories 注入（可选）", 2)
    bullet("`memories.use_memories`：是否注入已有记忆到未来会话。")
    bullet("与「必须遵守」的团队规则分离：规则应放 AGENTS.md，勿只靠 memories。")

    h("4. 上下文管理 slash", 2)
    bullet("`/compact`；compact_prompt / experimental_compact_prompt_file 可覆盖摘要提示。")
    bullet("`/status`：模型、审批、可写根、上下文余量。")

    h("Prompt 组成 · 核对摘要", 2)
    bullet("保留：AGENTS 发现序；Skills；可选 Memories；compact。")
    bullet("剔除：CLAUDE.md 或 Hermes SOUL.md 作为 Codex 默认文件名。")

    # ══════════════════════════════════════════════════════════════
    # 记忆与持久化
    # ══════════════════════════════════════════════════════════════
    h("记忆与持久化机制", 1)

    h("1. Session transcripts（历史）", 2)
    bullet("~/.codex/history.jsonl（可关/可封顶）；resume/archive/delete/fork。")
    bullet("exec JSON 暴露 thread/turn 事件流。")

    h("2. Local Memories（实验特性，默认关）", 2)
    bullet("[features] memories = true；文件在 ~/.codex/memories/。")
    bullet(
        "后台从空闲足够久的合格任务生成；跳过活跃/过短会话；字段脱敏。"
    )
    bullet(
        "`/memories`：当前任务是否 use / generate；不等同改全局开关。"
    )
    bullet(
        "可用 disable_on_external_context 避免把 MCP/web_search 任务写入记忆——"
        "对 PI 持久化特别有研究意义。"
    )
    cite("Memories", "https://developers.openai.com/codex/customization/memories")

    h("3. AGENTS.md（持久规则层）", 2)
    bullet("跨会话、可入 Git 的团队指导；官方建议与 Memories 分工。")

    h("4. 其它本地状态", 2)
    bullet("config.toml / profiles / auth / logs / caches under CODEX_HOME。")
    bullet("受信任项目才加载 `.codex/` 项目层（config/hooks/rules）。")
    bullet("Chronicle：Desktop-only，从屏幕恢复近况以助记忆——非 CLI 核心路径。")

    h("记忆与持久化 · 核对摘要", 2)
    bullet(
        "保留：history；可选 memories 目录；AGENTS；"
        "external_context 排除项；信任边界加载 .codex/。"
    )
    bullet(
        "剔除：Memories 默认开启；"
        "剔除 Hermes 固定 MEMORY.md 路径等同 ~/.codex/memories。"
    )

    h("主要参考链接", 1)
    for link in [
        "https://developers.openai.com/codex/llms.txt",
        "https://developers.openai.com/codex/cli/reference",
        "https://developers.openai.com/codex/non-interactive-mode",
        "https://developers.openai.com/codex/agent-approvals-security",
        "https://developers.openai.com/codex/sandboxing",
        "https://developers.openai.com/codex/web-search",
        "https://developers.openai.com/codex/config-file/config-basic",
        "https://developers.openai.com/codex/config-file/config-advanced",
        "https://developers.openai.com/codex/agent-configuration/agents-md",
        "https://developers.openai.com/codex/agent-configuration/rules",
        "https://developers.openai.com/codex/extend/mcp",
        "https://developers.openai.com/codex/customization/overview",
        "https://developers.openai.com/codex/customization/memories",
        "https://developers.openai.com/codex/build-skills",
        "https://developers.openai.com/codex/cli/slash-commands",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
