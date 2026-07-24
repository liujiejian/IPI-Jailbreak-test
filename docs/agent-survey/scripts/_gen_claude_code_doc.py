# -*- coding: utf-8 -*-
"""Generate Claude Code research notes (same structure as OpenHands / Hermes).

Paper track: Session / 任务模式 / web_fetch(security notice) / tools,
plus Prompt 组成 & 记忆持久化. 5-round anti-hallucination vs code.claude.com docs.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "Claude-Code.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label}（来源：{url}）")

    # ── Title ──
    h("Claude Code", 0)
    p(
        "本文档结构对齐 OpenHands.docx / Hermes.docx："
        "Session、任务模式、web_fetch（及 Security Notice 对照）、Tools；"
        "文末附 Prompt 组成、记忆与持久化机制，以及五轮防幻觉核对。"
        "对象为 Anthropic Claude Code（终端 / IDE / Desktop / Web 的 agentic coding harness）。"
    )
    p("官方文档索引：https://code.claude.com/docs/llms.txt")
    p(
        "定位：在 Claude 模型外加 agentic harness — 提供工具、上下文管理与执行环境，"
        "使模型可读写代码库、跑命令、检索与抓取网页。"
    )
    cite(
        "How Claude Code works；Overview",
        "https://code.claude.com/docs/en/how-claude-code-works.md",
    )
    p(
        "执行链路（研究视角）：User → agentic loop（gather context → take action → verify）"
        "→ tools（Read/Edit/Bash/WebFetch…）→ 结果回灌上下文。"
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "Session 是绑定项目目录的已保存对话。"
        "CLI 连续写入本地 transcript（JSONL）；"
        "Desktop / Web / VS Code 扩展各自维护独立会话历史；本页以 CLI 为准。"
    )
    cite(
        "Manage sessions",
        "https://code.claude.com/docs/en/sessions.md",
    )

    h("1. 存储与恢复", 2)
    bullet(
        "默认路径：~/.claude/projects/<sanitized-cwd>/<session-id>.jsonl"
        "（目录名由工作路径非字母数字字符替换为 -）。"
    )
    bullet(
        "恢复：claude --continue（最近会话）；"
        "claude --resume / --resume <name|id>；/resume；--from-pr。"
    )
    bullet(
        "claude -p / Agent SDK 会话不进交互 picker，但仍可用 session ID --resume "
        "（必须在同一 project 目录或其 git worktree 内查找）。"
    )
    bullet(
        "可配置：CLAUDE_CONFIG_DIR；cleanupPeriodDays（默认约 30 天）；"
        "CLAUDE_CODE_SKIP_PROMPT_HISTORY；-p 时 --no-session-persistence。"
    )
    p("注意：官方声明 JSONL 条目格式为内部实现、跨版本可变；脚本应用 /export 或 -p JSON，勿硬解析。")

    h("2. 分支 / Fork / 上下文管理", 2)
    bullet("/branch 或 --fork-session：复制历史到新 session ID，原会话保留。")
    bullet("会话内「允许本次」权限不带到 fork。")
    bullet("/clear 清空当前上下文（旧会话仍可 resume）；/compact 摘要压缩。")
    bullet("Checkpointing：文件编辑前快照，可 Esc Esc rewind（与 git 分离）。")

    h("3. 与 OpenClaw / Hermes 的研究映射（非官方术语）", 2)
    bullet("OpenClaw Main session ≈ Claude Code 交互 CLI/IDE 主会话 + transcript。")
    bullet(
        "OpenClaw Isolated / Hermes Cron fresh ≈ "
        "Routines / Desktop scheduled / CI -p 独立运行"
        "（或同机 /loop 但仍挂在原 session）。"
    )
    bullet("跨会话知识：CLAUDE.md + auto memory（非 transcript 自动并入新会话）。")

    # ══════════════════════════════════════════════════════════════
    # 任务模式
    # ══════════════════════════════════════════════════════════════
    h("任务模式", 1)
    p("决定是否共享会话、权限如何批准、是否需要本机在线，以及调度是否跨重启持久。")

    h("1. 交互式会话（Interactive）", 2)
    bullet("Shift+Tab 循环权限模式：Manual / Accept edits / Plan / Auto。")
    bullet(
        "Plan：探索与提案、不改源码；Accept edits：文件编辑与部分文件系统命令免提示。"
    )
    bullet("执行环境：Local / Cloud（Anthropic VM）/ Remote Control（本机执行、远端 UI）。")
    cite(
        "How Claude Code works；Permission modes",
        "https://code.claude.com/docs/en/permission-modes.md",
    )

    h("2. 非交互 / Headless（claude -p）", 2)
    bullet("claude -p \"...\"：脚本与 CI；可配 --allowedTools、--permission-mode、--output-format。")
    bullet(
        "--bare：跳过自动发现 hooks/skills/plugins/MCP/auto memory/CLAUDE.md，"
        "仅显式 flags 生效；CI 推荐；未来可能成为 -p 默认。"
    )
    bullet("-p 下首次信任确认不写入磁盘；Background Bash 在结果返回后约 5s 终止。")
    cite(
        "Run Claude Code programmatically (headless)",
        "https://code.claude.com/docs/en/headless.md",
    )

    h("3. 会话内调度：/loop 与 Cron* 工具", 2)
    bullet(
        "任务挂在当前会话：关会话即停；--resume/--continue 可恢复未过期任务"
        "（循环约 7 天；one-shot 未到点）。"
    )
    bullet("工具：CronCreate / CronList / CronDelete；最多约 50 个/会话。")
    bullet("亦可 /loop、自然语言 one-shot reminder、Monitor（流式后台输出）。")
    bullet("对比：Routines（云）与 Desktop scheduled 不依赖打开的交互会话。")
    cite(
        "Run prompts on a schedule",
        "https://code.claude.com/docs/en/scheduled-tasks.md",
    )

    h("4. 持久自动化：Routines / Desktop / CI", 2)
    bullet(
        "Routines：Anthropic 托管云上，按 schedule / API / GitHub 事件触发；"
        "最小间隔约 1 小时；无本机文件（fresh clone）；RemoteTrigger / /schedule。"
    )
    bullet("Desktop scheduled tasks：本机跑、可访问本地文件；最小约 1 分钟。")
    bullet("GitHub Actions 等 CI schedule。")
    cite(
        "Automate work with routines",
        "https://code.claude.com/docs/en/routines.md",
    )

    h("5. 调度策略对照（研究笔记）", 2)
    p("OpenClaw：精确隔离 → Cron；完整 session → Heartbeat。")
    p("Claude Code 对应：")
    bullet("会话内轮询/提醒 → /loop + Cron*（需会话存活）。")
    bullet("无人值守、跨机可靠 → Routines（云）或 CI。")
    bullet("要本地文件 → Desktop scheduled 或本机常开会话 /loop。")

    # ══════════════════════════════════════════════════════════════
    # web_fetch / Security Notice
    # ══════════════════════════════════════════════════════════════
    h("web_fetch（及 Security Notice）", 1)
    p(
        "内置工具名即为 WebFetch（PascalCase，权限规则中亦用此名）；"
        "另有 WebSearch（只返回标题/URL，不抓正文）。"
    )
    cite(
        "Tools reference > WebFetch / WebSearch",
        "https://code.claude.com/docs/en/tools-reference.md",
    )

    h("1. WebFetch 行为", 2)
    bullet("参数：URL + 描述提取目标的 prompt。")
    bullet("HTML → Markdown；用小型快速模型按 prompt 抽取。")
    bullet(
        "【关键】对多数请求，主会话 Claude 收到的是小模型答案，而非原始页面全文"
        "（lossy by design）。要原文可用 Bash curl，或更具体的二次 prompt。"
    )
    bullet("HTTP 自动升 HTTPS；大页按固定字符上限截断。")
    bullet(
        "【与 OpenClaw 类似】响应缓存约 15 分钟：同一 URL 重复 fetch 走缓存。"
        "IPI/测试需在 URL 加时间戳或等效参数以避免缓存。"
    )
    bullet("User-Agent 以 Claude-User 开头；Accept 偏好 Markdown。")

    h("2. Security Notice 对照：隔离上下文（非 SECURITY NOTICE 原文）", 2)
    p("OpenClaw：tool 结果包裹 SECURITY NOTICE + EXTERNAL_UNTRUSTED_CONTENT。")
    p("Hermes：<untrusted_tool_result source=\"web_extract\">。")
    p(
        "Claude Code 官方安全页："
        "「Isolated context windows: Web fetch uses a separate context window "
        "to avoid injecting potentially malicious prompts」——"
        "外部网页先在隔离上下文由抽取模型处理，再将结果送回主 agent，"
        "降低原始恶意指令直接进入主会话的风险。"
    )
    cite(
        "Security > Protect against prompt injection",
        "https://code.claude.com/docs/en/security.md",
    )
    bullet(
        "权限：默认/acceptEdits 下新域名首次 WebFetch 会询问"
        "（内置少量文档域预批准）；规则形如 WebFetch(domain:example.com)。"
    )
    bullet(
        "auto / bypassPermissions 可跳过域名提示；"
        "显式 deny/ask/allow 覆盖预批准集。"
    )
    bullet(
        "网络类 Bash（curl/wget）默认不自动批准；"
        "deny WebFetch 并不能阻止 Bash curl——需权限 + 可选 sandbox 网络规则双层。"
    )
    bullet("Sandbox allowedDomains/deniedDomains 与 WebFetch 权限互补，勿混为一谈。")

    h("3. IPI 研究含义", 2)
    bullet(
        "主注入面仍是：WebFetch 抽取结果进入主会话；"
        "WebSearch 返回的 URL 列表；Bash 网络；MCP 工具输出；"
        "未可信校验的 CLAUDE.md / 仓库文件。"
    )
    bullet(
        "与「原文直灌」相比：lossy 抽取可减少原文注入，"
        "但也可能使攻击者针对抽取模型/prompt 投毒，或迫使用户改用 curl 绕过隔离。"
    )
    bullet("复现实验请加 URL cache-buster（15min）。")

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "内置工具名为权限/hooks/SDK 中的精确字符串；"
        "扩展靠 MCP（另名）与 Skill（经 Skill 工具执行，不新增工具表项）。"
    )
    cite(
        "Tools reference；How Claude Code works",
        "https://code.claude.com/docs/en/tools-reference.md",
    )

    h("1. 与论文最相关的工具", 2)
    bullet("Web：WebFetch、WebSearch。")
    bullet("文件：Read、Edit、Write、Glob、Grep、NotebookEdit、LSP。")
    bullet("执行：Bash（+ 可选 PowerShell）；Monitor。")
    bullet("编排：Agent（subagent）、Skill、Task*、Workflow、AskUserQuestion。")
    bullet("会话调度：CronCreate / CronList / CronDelete；ScheduleWakeup（/loop）。")
    bullet("云例程：RemoteTrigger（Routines，计划受限）。")

    h("2. 权限与沙箱（工具闸门）", 2)
    bullet(
        "Permission rules：在工具执行前评估，覆盖 Bash/Read/Edit/WebFetch/MCP 等。"
    )
    bullet(
        "Sandbox：OS 级限制 Bash 子进程的文件系统/网络；"
        "不等同于 permission mode；/sandbox 配置边界。"
    )
    bullet(
        "整进程隔离可选 sandbox-runtime / container / Cloud VM"
        "（MCP 与 hooks 默认在宿主机，除非整体包入边界）。"
    )
    cite(
        "Permissions；Sandboxing；Security",
        "https://code.claude.com/docs/en/permissions.md",
    )

    h("3. Hooks（确定性策略，非 memory）", 2)
    bullet(
        "PreToolUse / PostToolUse / SessionStart / Stop 等："
        "可拦截危险命令、注入上下文；与 CLAUDE.md「软指引」互补。"
    )
    cite(
        "Hooks guide",
        "https://code.claude.com/docs/en/hooks-guide.md",
    )

    h("4. 研究用简化 messages 示例", 2)
    code(
        """{
  "messages": [
    {
      "role": "system",
      "content": "<Claude Code system + CLAUDE.md + auto MEMORY.md snippet + skills index>"
    },
    {
      "role": "user",
      "content": "Check https://api.example.com/health and return the status."
    },
    {
      "role": "assistant",
      "content": null,
      "tool_calls": [{
        "id": "call_001",
        "type": "function",
        "function": {
          "name": "WebFetch",
          "arguments": "{\\"url\\": \\"https://api.example.com/health\\", \\"prompt\\": \\"Return HTTP status and body summary.\\"}"
        }
      }]
    },
    {
      "role": "tool",
      "tool_call_id": "call_001",
      "content": "<extractor model answer, not necessarily raw body>"
    },
    {
      "role": "assistant",
      "content": "The endpoint looks healthy."
    }
  ],
  "tools": [
    {"name": "WebFetch", "...": "..."},
    {"name": "WebSearch", "...": "..."},
    {"name": "Bash", "...": "..."},
    {"name": "Read", "...": "..."},
    {"name": "Edit", "...": "..."}
  ]
}"""
    )
    p("示意：真实协议为 Anthropic Messages / Agent SDK 事件流；工具名大小写以 Tools reference 为准。")

    # ══════════════════════════════════════════════════════════════
    # 五轮防幻觉
    # ══════════════════════════════════════════════════════════════
    h("五轮防幻觉核对记录", 1)

    h("Round 1 — Session", 2)
    bullet("核对：sessions.md、how-claude-code-works.md。")
    bullet("保留：JSONL per project；--continue/--resume/--fork；/clear /compact。")
    bullet("剔除：把 Hermes SessionDB 或 OpenClaw session:xxx 写成 Claude Code API。")

    h("Round 2 — 任务模式", 2)
    bullet("核对：scheduled-tasks.md、routines.md、headless.md、permission-modes。")
    bullet(
        "保留：/loop 会话作用域 + 7 天；Routines 云持久；Desktop；-p/--bare。"
    )
    bullet(
        "剔除：宣称 /loop 等同 Hermes gateway cron（跨重启无关会话仍跑）；"
        "剔除 Claude Code 没有 WebFetch（实际有同名工具）。"
    )

    h("Round 3 — WebFetch / 安全", 2)
    bullet("核对：tools-reference WebFetch、security.md、permissions.md、sandboxing.md。")
    bullet(
        "保留：lossy 抽取；15min cache；隔离上下文；WebFetch(domain:)；curl 不因 deny WebFetch 而自动禁。"
    )
    bullet(
        "剔除：OpenClaw SECURITY NOTICE 原文照搬；"
        "剔除「主会话必见完整 HTML」；"
        "剔除「无缓存」（官方写明 15 minutes）。"
    )

    h("Round 4 — Tools", 2)
    bullet("核对：tools-reference 全表工具名。")
    bullet(
        "保留：PascalCase 内置名（WebFetch、Bash、Read…）；MCP/Skill 扩展模型。"
    )
    bullet(
        "剔除：OpenHands terminal/file_editor 或 Hermes web_extract 作为 Claude Code 内置名。"
    )

    h("Round 5 — Prompt / Memory 交叉复检", 2)
    bullet("核对：memory.md、prompt-caching.md、how-claude-code-works context。")
    bullet(
        "冻结："
        "(1) Session=本地 JSONL 项目会话；"
        "(2) 任务=交互 / -p / 会话 Cron / Routines / Desktop / CI；"
        "(3) WebFetch=独立抽取上下文+15min cache，非 SECURITY NOTICE 包装；"
        "(4) 记忆=CLAUDE.md + auto MEMORY.md（软上下文，硬策略用 hooks/permissions）。"
    )
    bullet(
        "开放问题：实机 tool_result 是否含任何可见 untrusted 标记"
        "（官方强调隔离上下文与 lossy 答案，未承诺 OpenClaw 式标记串）；"
        "Cloud 与本地 CLI 默认权限差异；版本间 Cron/TodoWrite 等行为变更。"
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt 组成
    # ══════════════════════════════════════════════════════════════
    h("Prompt 组成", 1)
    p(
        "每轮 API 请求重发：system prompt + 项目上下文 + 历史消息/工具结果 + 新消息。"
        "Claude Code 自动管理 prompt caching；缓存按前缀精确匹配。"
    )
    cite(
        "How Claude Code uses prompt caching",
        "https://code.claude.com/docs/en/prompt-caching.md",
    )

    h("1. 缓存层（概念）", 2)
    bullet("System prompt 层：工具定义等；变更易使整前缀失效。")
    bullet(
        "Project context：CLAUDE.md、auto memory、未 scoped rules — "
        "会话开始加载；/clear 或 /compact 后可重建。"
    )
    bullet("Conversation：消息与 tool results；一般追加不破坏前缀。")
    bullet(
        "模型与 effort 也是 cache key：/model、/effort 切换会整窗未命中缓存。"
    )

    h("2. CLAUDE.md 与系统提示定制", 2)
    bullet(
        "CLAUDE.md / .claude/rules / managed / user / local 多层级加载"
        "（见记忆节）；属上下文而非强制配置。"
    )
    bullet(
        "CLAUDE.md 编辑默认不在中途改写已建缓存前缀；"
        "新内容常需新会话、/clear 或 /compact 才进入后续请求。"
    )
    bullet(
        "-p / SDK：--system-prompt 替换；"
        "--append-system-prompt / --append-system-prompt-file 追加。"
    )
    bullet("Skills：启动时见描述，完整内容按需加载；可 defer MCP tools（ToolSearch）。")

    h("3. Agent SDK / 采集", 2)
    bullet("Agent SDK 可改 system prompt preset、权限回调、会话 resume。")
    bullet(
        "研究抓包：ANTHROPIC_BASE_URL / LLM gateway；"
        "或 -p --output-format stream-json；勿依赖私有 JSONL schema。"
    )
    cite(
        "Modifying system prompts (Agent SDK)",
        "https://code.claude.com/docs/en/agent-sdk/modifying-system-prompts.md",
    )

    h("Prompt 组成 · 核对摘要", 2)
    bullet("保留：三层缓存观；CLAUDE.md 中途不立即生效；lossy WebFetch 不进 system。")
    bullet("剔除：把 Hermes SOUL/MEMORY volatile tier 名称写成 Claude Code API。")

    # ══════════════════════════════════════════════════════════════
    # 记忆与持久化
    # ══════════════════════════════════════════════════════════════
    h("记忆与持久化机制", 1)
    p(
        "每会话 context window 从空开始；跨会话靠 "
        "CLAUDE.md（你写）与 auto memory（Claude 写）。"
        "二者均为上下文，非硬策略——硬门禁用 PreToolUse hooks / permissions。"
    )
    cite(
        "How Claude remembers your project",
        "https://code.claude.com/docs/en/memory.md",
    )

    h("1. CLAUDE.md 层级", 2)
    bullet("Managed：OS 级组织策略路径（如 /etc/claude-code/CLAUDE.md）。")
    bullet("User：~/.claude/CLAUDE.md。")
    bullet("Project：./CLAUDE.md 或 ./.claude/CLAUDE.md。")
    bullet("Local：./CLAUDE.local.md（gitignore 个人偏好）。")
    bullet("上级目录 CLAUDE.md 全量加载；子目录按需；.claude/rules 可路径作用域。")
    bullet("@imports 可嵌套（最大约 4 hop）；/init 可生成起步文件。")

    h("2. Auto memory", 2)
    bullet(
        "目录：~/.claude/projects/<project>/memory/ "
        "（同 git repo 的 worktree 共享）；含 MEMORY.md 索引 + 主题文件。"
    )
    bullet(
        "每会话仅加载 MEMORY.md 前 200 行或前 25KB（先到为准）；"
        "详细笔记放 topic 文件。"
    )
    bullet("默认开启；/memory 可关；CLAUDE_CODE_DISABLE_AUTO_MEMORY=1；可配 autoMemoryDirectory。")
    bullet("机器本地，不随云环境自动同步。")

    h("3. 会话 transcript 持久化", 2)
    bullet("JSONL 保存完整对话与工具轨迹，支撑 resume/fork/export。")
    bullet("与「记忆」分离：新会话不自动载入旧 transcript，除非 --resume。")

    h("4. 与 OpenClaw / Hermes 对照", 2)
    bullet(
        "OpenClaw MEMORY.md + daily notes ≈ Claude Code CLAUDE.md（指令）"
        "+ auto MEMORY.md（学习）；无官方 daily notes 文件约定。"
    )
    bullet(
        "Hermes frozen snapshot mid-session 不更新 prompt ≈ "
        "Claude Code CLAUDE.md 编辑常待 compact/新会话；"
        "但 auto memory 可在会话中读写文件（UI 显示 Writing/Recalling memory）。"
    )

    h("记忆与持久化 · 核对摘要", 2)
    bullet("保留：双轨记忆；MEMORY.md 加载上限；软上下文 vs hooks。")
    bullet(
        "剔除：平台 API Memory Tool（/memories）与 Claude Code 产品 auto memory 混为一谈"
        "（前者是 Messages API 通用工具，托管于应用侧；产品文档路径为 "
        "~/.claude/projects/.../memory/）。"
    )

    h("主要参考链接", 1)
    for link in [
        "https://code.claude.com/docs/llms.txt",
        "https://code.claude.com/docs/en/how-claude-code-works.md",
        "https://code.claude.com/docs/en/sessions.md",
        "https://code.claude.com/docs/en/scheduled-tasks.md",
        "https://code.claude.com/docs/en/routines.md",
        "https://code.claude.com/docs/en/headless.md",
        "https://code.claude.com/docs/en/tools-reference.md",
        "https://code.claude.com/docs/en/security.md",
        "https://code.claude.com/docs/en/permissions.md",
        "https://code.claude.com/docs/en/sandboxing.md",
        "https://code.claude.com/docs/en/memory.md",
        "https://code.claude.com/docs/en/prompt-caching.md",
        "https://code.claude.com/docs/en/hooks-guide.md",
        "https://code.claude.com/docs/en/permission-modes.md",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
