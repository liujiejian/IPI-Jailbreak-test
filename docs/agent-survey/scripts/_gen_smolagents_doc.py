# -*- coding: utf-8 -*-
"""Generate smolagents research document — accuracy-focused with 5-round audit."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "Smolagents.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def add_code(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def cite(text: str, url: str) -> None:
        doc.add_paragraph(f"{text}（来源：{url}）")

    doc.add_heading("Smolagents", level=0)
    doc.add_paragraph(
        "本文档研究 Hugging Face smolagents 开源 Agent 框架机制，术语严格对齐官方资料。"
        "用户指定入口：https://github.com/huggingface/smolagents；"
        "技术细节以 HF Docs（https://huggingface.co/docs/smolagents/index）与源码为准。"
    )

    # ── Scope ──
    doc.add_heading("文档范围与术语分层（必读）", level=1)
    doc.add_paragraph("研究须区分以下层次，不可混用：")
    bullet(
        "Agent 运行时（MultiStepAgent 基类）：ReAct 循环 think → act → observe，"
        "默认 max_steps=20。"
    )
    bullet(
        "CodeAgent（默认推荐）：LLM 以 Python 代码片段表达动作，"
        "工具调用写成函数调用；依赖 LocalPythonExecutor 或远程沙箱执行。"
    )
    bullet(
        "ToolCallingAgent：LLM 以 JSON tool calling 表达动作，"
        "经 model.get_tool_call 解析后执行。"
    )
    bullet(
        "CLI 双入口：smolagent（通用 CodeAgent + 可选 toolbox）与 "
        "webagent（helium 浏览器自动化，vision_web_browser.py）。"
    )
    bullet(
        "多 Agent：managed_agents 层级；planning_interval 周期性 PlanningStep。"
    )
    bullet(
        "工具生态：@tool 装饰器、default_tools、MCP（ToolCollection.from_mcp）、"
        "LangChain/Hub Space 导入；Hub 上 push/pull agent 与 tool。"
    )
    cite(
        "README；Agents reference；intro_agents",
        "https://github.com/huggingface/smolagents",
    )
    cite(
        "官方文档索引",
        "https://huggingface.co/docs/smolagents/index",
    )

    # ── 1. LLM 请求采集 ──
    doc.add_heading("研究背景与 LLM 请求采集", level=1)
    doc.add_paragraph(
        "采集目标：每次 LLM 调用前的 messages、工具 schema、"
        "CodeAgent 生成的代码动作、ToolCallingAgent 的 tool_calls、"
        "执行观测（observations/errors）与多 Agent 委派轨迹。"
    )

    doc.add_heading("1. verbosity_level 与 AgentLogger", level=2)
    bullet(
        "MultiStepAgent 接受 verbosity_level（默认 LogLevel.INFO）；"
        "控制逐步日志详细程度。"
    )
    bullet(
        "agent.replay(detailed=False)：重放上次 run 的步骤；"
        "detailed=True 会指数级增长日志，仅调试使用。"
    )
    cite(
        "Agents reference > replay / verbosity_level",
        "https://huggingface.co/docs/smolagents/reference/agents",
    )

    doc.add_heading("2. AgentMemory 与 write_memory_to_messages", level=2)
    bullet(
        "AgentMemory：system_prompt（SystemPromptStep）+ steps 列表"
        "（TaskStep / ActionStep / PlanningStep）。"
    )
    bullet(
        "write_memory_to_messages()：将历史 llm_outputs、actions、observations/errors "
        "序列化为 LLM 输入 messages，并插入 PLAN、error 等关键词。"
    )
    bullet(
        "get_full_steps()：导出含 model input messages 的完整步骤字典；"
        "get_succinct_steps() 不含 model input。"
    )
    bullet(
        "研究侧可在每 step 后读取 agent.memory.steps[-1] 或 get_full_steps() 采样上下文。"
    )
    cite(
        "Memory reference；tutorials/memory",
        "https://huggingface.co/docs/smolagents/tutorials/memory",
    )

    doc.add_heading("3. agent.run(stream=True) 与 stream_outputs", level=2)
    bullet(
        "agent.run(stream=True)：返回生成器，每执行一步 yield 一步；"
        "须迭代消费才能看到逐步输出。"
    )
    bullet(
        "CodeAgent / ToolCallingAgent 另有 stream_outputs 参数（默认 False）："
        "控制代码/工具执行期是否流式输出。"
    )
    bullet(
        "README 示例使用 stream_outputs=True；与 run(stream=) 是不同层级的流式开关。"
    )
    cite(
        "Agents reference > run / CodeAgent stream_outputs",
        "https://huggingface.co/docs/smolagents/reference/agents",
    )
    cite(
        "README Quick demo",
        "https://github.com/huggingface/smolagents/blob/main/README.md",
    )

    doc.add_heading("4. OpenTelemetry 可观测性", level=2)
    bullet(
        "inspect_runs 教程：采用 OpenTelemetry 标准；"
        "SmolagentsInstrumentor（openinference-instrumentation-smolagents）"
        "可对接 Phoenix、MLflow、Langfuse。"
    )
    bullet(
        "MLflow：mlflow.smolagents.autolog() 一行开启 trace、span、"
        "输入输出与 token 用量。"
    )
    bullet(
        "多 Agent 场景下 trace 可展示 manager 调用 managed agent 的层级。"
    )
    cite(
        "tutorials/inspect_runs",
        "https://huggingface.co/docs/smolagents/tutorials/inspect_runs",
    )

    doc.add_heading("5. step_callbacks 与逐步执行", level=2)
    bullet(
        "step_callbacks：每步回调，可传入 list 或 dict[MemoryStep, Callable]；"
        "回调可读写 agent.memory（如 web browser 截图裁剪示例）。"
    )
    bullet(
        "agent.step(memory_step)：单步 ReAct，适合长耗时工具或逐步修改 memory。"
    )
    bullet(
        "return_full_result=True 时 run() 返回 RunResult 对象而非仅 final answer。"
    )
    cite(
        "tutorials/memory；Agents reference > step / return_full_result",
        "https://huggingface.co/docs/smolagents/tutorials/memory",
    )

    doc.add_heading("6. GradioUI 交互采集", level=2)
    bullet(
        "GradioUI(agent, file_upload_folder=..., reset_agent_memory=True)："
        "基于 gradio.ChatInterface 的 Web UI；"
        "reset_agent_memory 控制每次交互是否清空 memory。"
    )
    bullet(
        "stream_to_gradio：将 agent 消息流式投影为 Gradio ChatMessages。"
    )
    cite(
        "Agents reference > GradioUI / stream_to_gradio",
        "https://huggingface.co/docs/smolagents/reference/agents",
    )

    # ── 2. Prompt ──
    doc.add_heading("Prompt 设计组成", level=1)

    doc.add_heading("1. PromptTemplates 结构", level=2)
    bullet(
        "MultiStepAgent 使用 PromptTemplates：system_prompt、"
        "planning（plan / update_plan_pre_messages / update_plan_post_messages）、"
        "managed_agent（task / report）、final_answer（pre_messages / post_messages）。"
    )
    bullet(
        "instructions 参数：自定义指令插入 system prompt。"
    )
    bullet(
        "CodeAgent 另有 code_block_tags 控制代码块解析正则；"
        "use_structured_outputs_internally 可在 action step 使用结构化生成。"
    )
    cite(
        "Agents reference > PromptTemplates / CodeAgent",
        "https://huggingface.co/docs/smolagents/reference/agents",
    )

    doc.add_heading("2. CodeAgent 系统提示与代码动作格式", level=2)
    bullet(
        "CodeAgent 系统提示要求 LLM 以 Python 代码片段表达动作；"
        "工具作为函数调用；终止须调用 final_answer(...)。"
    )
    bullet(
        "extract_action(model_output, split_token)：从 LLM 输出解析代码动作。"
    )
    bullet(
        "additional_authorized_imports：扩展 LocalPythonExecutor 授权 import 列表；"
        "默认仅 BASE_BUILTIN_MODULES。"
    )
    cite(
        "README > How do Code agents work?；secure_code_execution",
        "https://github.com/huggingface/smolagents/blob/main/README.md",
    )

    doc.add_heading("3. ToolCallingAgent 工具 schema", level=2)
    bullet(
        "Tool 基类定义 name / description / inputs / output_type；"
        "to_dict() 生成 LLM 可见 schema。"
    )
    bullet(
        "execute_tool_call(tool_name, arguments)：执行工具或 managed agent；"
        "process_tool_calls 更新 ActionStep memory。"
    )
    bullet(
        "max_tool_threads：并行 tool call 线程池上限。"
    )
    cite(
        "reference/tools；Agents reference > ToolCallingAgent",
        "https://huggingface.co/docs/smolagents/reference/tools",
    )

    doc.add_heading("4. Planning 与 Managed Agent 提示", level=2)
    bullet(
        "planning_interval=N：每 N 步插入 PlanningStep，"
        "使用 planning 模板生成/更新计划。"
    )
    bullet(
        "managed_agents：子 Agent 须设 name + description；"
        "父 Agent 通过 managed_agent 模板委派任务并接收 report。"
    )
    bullet(
        "provide_run_summary：作为 managed agent 被调用时是否提供 run 摘要。"
    )
    cite(
        "Agents reference > planning_interval / managed_agents",
        "https://huggingface.co/docs/smolagents/reference/agents",
    )

    doc.add_heading("5. 典型执行轮次示意", level=2)
    doc.add_paragraph(
        "单 run 内：TaskStep(task) → [PlanningStep?] → "
        "循环 ActionStep（llm_output → 代码/tool call → observations/error）"
        "→ final_answer。"
    )
    add_code(
        """# 研究采集点（非固定 JSON schema）：
# - agent.memory.get_full_steps() 含 model_input_messages
# - ActionStep: llm_output, code_action, observations, error
# - write_memory_to_messages() 投影为 LLM chat messages
# - OTel trace 记录 manager → managed_agent 层级"""
    )
    cite(
        "intro_agents > multi-step loop；Memory reference",
        "https://huggingface.co/docs/smolagents/conceptual_guides/intro_agents",
    )

    # ── 3. Memory ──
    doc.add_heading("Memory 与持久化机制", level=1)
    doc.add_paragraph(
        "smolagents 的 AgentMemory 是单次 run 内的步骤轨迹，"
        "非跨会话向量库或语义长期记忆；框架未内置 checkpointer。"
    )

    doc.add_heading("1. AgentMemory 组成", level=2)
    bullet("system_prompt：SystemPromptStep，初始化后固定。")
    bullet(
        "steps：TaskStep（用户任务）+ ActionStep（每步动作/观测/错误/截图）"
        "+ PlanningStep（周期性计划）。"
    )
    bullet(
        "ActionStep 可含 observations_images（如浏览器截图 PIL Image 列表）。"
    )
    bullet("memory.reset()：清空 steps 保留 system_prompt。")
    cite(
        "Memory reference > AgentMemory",
        "https://huggingface.co/docs/smolagents/reference/agents",
    )

    doc.add_heading("2. run(reset=) 与会话延续", level=2)
    bullet(
        "agent.run(task, reset=True)：默认重置 memory 后开始新任务。"
    )
    bullet(
        "reset=False：延续上次 run 的 memory，实现同 Agent 实例多轮对话。"
    )
    bullet(
        "GradioUI reset_agent_memory=True 时每次 UI 交互清空 memory。"
    )
    cite(
        "Agents reference > run(reset)；GradioUI reset_agent_memory",
        "https://huggingface.co/docs/smolagents/reference/agents",
    )

    doc.add_heading("3. 动态修改 memory", level=2)
    bullet(
        "step_callbacks 可在每步后修改 agent.memory.steps（如删除旧截图省 token）。"
    )
    bullet(
        "逐步模式：手动 append TaskStep，循环 agent.step() 后 append ActionStep，"
        "可在步间替换 memory.steps[-1]。"
    )
    bullet(
        "可将另一 agent 的 memory.steps 赋给当前 agent（教程示例）。"
    )
    cite(
        "tutorials/memory > Dynamically change / Run agents one step at a time",
        "https://huggingface.co/docs/smolagents/tutorials/memory",
    )

    doc.add_heading("4. save / from_folder / push_to_hub", level=2)
    bullet(
        "agent.save(output_dir)：导出 tools/、managed_agents/、agent.json、"
        "prompt.yaml、app.py、requirements.txt。"
    )
    bullet(
        "from_folder / from_hub：从本地或 Hub 加载 agent 定义；"
        "from_hub 默认 trust_remote_code=False，须显式确认远程代码风险。"
    )
    bullet(
        "上述为 agent 配置与 prompt 快照，≠ 运行时 memory 自动持久化。"
    )
    cite(
        "Agents reference > save / from_hub",
        "https://huggingface.co/docs/smolagents/reference/agents",
    )

    doc.add_heading("5. 无内置语义/RAG 长期记忆", level=2)
    doc.add_paragraph(
        "所查官方文档与核心 API 均无 CrewAI Memory / LangGraph store / "
        "OpenClaw MEMORY.md 式跨会话语义记忆；"
        "长期状态须应用侧自行序列化 memory 或外接向量库。"
    )
    cite(
        "tutorials/memory；README",
        "https://huggingface.co/docs/smolagents/tutorials/memory",
    )

    # ── 4. Session ──
    doc.add_heading("Session 与会话模型", level=1)
    bullet(
        "会话单元 = 一个 MultiStepAgent 实例 + 其 AgentMemory；"
        "无框架级 thread_id / session key。"
    )
    bullet(
        "同实例 reset=False 的连续 run() 共享 memory.steps，"
        "等效于应用内多轮会话。"
    )
    bullet(
        "新 Agent() 或 memory.reset() / run(reset=True) 开启新会话。"
    )
    bullet(
        "UserInputTool：工具执行时阻塞式 input() 向终端用户提问；"
        "非异步 HITL middleware。"
    )
    bullet(
        "agent.interrupt()：中断当前 agent 执行。"
    )
    cite(
        "default_tools UserInputTool；Agents reference > interrupt",
        "https://github.com/huggingface/smolagents/blob/main/src/smolagents/default_tools.py",
    )

    # ── 5. Automation ──
    doc.add_heading("Automation 与定时任务", level=1)
    bullet(
        "框架层无 OpenClaw 式 Heartbeat、内置 cron 或定时调度器。"
    )
    bullet(
        "CLI smolagent / webagent 为一次性命令行执行；"
        "交互模式 smolagent 无 prompt 时启动配置向导。"
    )
    bullet(
        "push_to_hub 导出 Gradio Space（app.py），可部署为 HF Space 常驻 UI，"
        "但非 agent 自主定时任务。"
    )
    bullet(
        "逐步 agent.step() 循环可由应用侧调度器驱动长周期任务。"
    )
    cite(
        "README > CLI；Agents reference > push_to_hub",
        "https://github.com/huggingface/smolagents/blob/main/README.md",
    )

    # ── 6. External Web ──
    doc.add_heading("外部 Web 访问", level=1)
    doc.add_paragraph(
        "smolagents 无统一内置 web_fetch 或 EXTERNAL_UNTRUSTED_CONTENT 包裹；"
        "外部内容通过可选 Tool 或浏览器自动化进入 agent memory。"
    )

    doc.add_heading("1. 搜索与抓取 Tool", level=2)
    bullet(
        "WebSearchTool：duckduckgo / bing / exa 引擎；"
        "返回 markdown 标题+链接+摘要。"
    )
    bullet(
        "DuckDuckGoSearchTool / GoogleSearchTool（SerpAPI 或 Serper）/ "
        "ApiWebSearchTool（默认 Brave API）。"
    )
    bullet(
        "VisitWebpageTool：requests GET + markdownify 转 markdown，"
        "默认 max_output_length=40000 截断。"
    )
    bullet(
        "WikipediaSearchTool：wikipedia-api 拉取词条 summary 或全文。"
    )
    bullet(
        "add_base_tools=False（默认）：不自动挂载工具；"
        "add_base_tools=True 时 TOOL_MAPPING 含 python_interpreter、"
        "web_search（DuckDuckGo）、visit_webpage。"
    )
    cite(
        "default_tools.py；reference/tools",
        "https://github.com/huggingface/smolagents/blob/main/src/smolagents/default_tools.py",
    )

    doc.add_heading("2. webagent CLI 与 Helium 浏览器", level=2)
    bullet(
        "webagent：基于 vision_web_browser.py + helium + selenium，"
        "VLM 驱动浏览器点击/滚动/截图。"
    )
    bullet(
        "官方 web_browser 教程：自定义 @tool（go_back、close_popups、search_item_ctrl_f）"
        "+ step_callbacks 截图写入 ActionStep.observations_images。"
    )
    bullet(
        "页面 DOM/URL 进入 observations 文本与截图，"
        "无框架级 untrusted 边界标记。"
    )
    cite(
        "README > CLI webagent；examples/web_browser",
        "https://huggingface.co/docs/smolagents/examples/web_browser",
    )

    doc.add_heading("3. MCP 与第三方 Tool", level=2)
    bullet(
        "ToolCollection.from_mcp / MCPClient：连接 MCP server 暴露远程工具；"
        "工具返回原文进入 agent 上下文。"
    )
    bullet(
        "Tool.from_langchain、Tool.from_hub、Tool.from_space："
        "导入外部工具定义。"
    )
    cite(
        "reference/tools > from_mcp / from_hub",
        "https://huggingface.co/docs/smolagents/reference/tools",
    )

    doc.add_heading("4. 与 OpenClaw web_fetch 差异（研究对照）", level=2)
    bullet("OpenClaw：内置 web_fetch + EXTERNAL_UNTRUSTED_CONTENT 包裹。")
    bullet(
        "smolagents：VisitWebpageTool / WebSearchTool 等返回 markdown 字符串，"
        "经 ActionStep.observations 或 CodeAgent 变量注入 LLM messages；"
        "secure_code_execution 文档明确网页 prompt injection 可污染 agent memory。"
    )
    cite(
        "tutorials/secure_code_execution > prompt injection",
        "https://huggingface.co/docs/smolagents/tutorials/secure_code_execution",
    )

    # ── 7. Security ──
    doc.add_heading("Security 机制", level=1)
    doc.add_paragraph(
        "smolagents 核心风险在 CodeAgent 代码执行；"
        "LocalPythonExecutor 为 best-effort 缓解，非安全边界。"
    )

    doc.add_heading("1. LocalPythonExecutor 限制", level=2)
    bullet(
        "AST 解释执行：默认禁止未授权 import；子模块须显式授权（如 numpy.*）；"
        "限制循环迭代次数；未定义操作抛错。"
    )
    bullet(
        "README 与 SECURITY 明确：可被绕过，不可用于运行不可信代码；"
        "真隔离须远程沙箱。"
    )
    bullet(
        "PythonInterpreterTool（ToolCallingAgent 路径）复用同一 evaluate_python_code。"
    )
    cite(
        "tutorials/secure_code_execution > Local Python executor",
        "https://huggingface.co/docs/smolagents/tutorials/secure_code_execution",
    )

    doc.add_heading("2. 远程沙箱 executor_type", level=2)
    bullet(
        "CodeAgent(executor_type=\"blaxel\"|\"e2b\"|\"modal\"|\"docker\")："
        "仅将 LLM 生成的代码片段发往远程执行；"
        "模型调用仍在本地。"
    )
    bullet(
        "with CodeAgent(...) as agent 或 agent.cleanup() 释放沙箱资源。"
    )
    bullet(
        "方案 1（片段远程执行）不支持 managed_agents 多 Agent（凭证无法安全传入沙箱）。"
    )
    bullet(
        "方案 2：整段 agent 代码在 E2B/Docker 内运行，可支持 multi-agent，"
        "但可能须向沙箱传入 API key。"
    )
    cite(
        "tutorials/secure_code_execution > Sandbox approaches",
        "https://huggingface.co/docs/smolagents/tutorials/secure_code_execution",
    )

    doc.add_heading("3. final_answer_checks 与 Hub 信任", level=2)
    bullet(
        "final_answer_checks：接受 final answer 前的验证函数列表"
        "（参数：answer, memory, agent）。"
    )
    bullet(
        "from_hub(trust_remote_code=False 默认)：加载 Hub agent/tool 前须审查代码；"
        "文档警告等同 pip/npm 安装不可信包。"
    )
    cite(
        "Agents reference > final_answer_checks / from_hub",
        "https://huggingface.co/docs/smolagents/reference/agents",
    )

    doc.add_heading("4. IPI 研究含义", level=2)
    doc.add_paragraph(
        "注入面：VisitWebpageTool / WebSearchTool 返回、"
        "MCP tool 输出、managed_agent report、"
        "浏览器 observations 与截图 OCR 上下文、"
        "Hub 远程 tool 定义。"
        "内容经 write_memory_to_messages() 进入后续 LLM 轮次，"
        "无 EXTERNAL_UNTRUSTED_CONTENT 自动消毒。"
    )
    cite(
        "tutorials/secure_code_execution；default_tools VisitWebpageTool",
        "https://huggingface.co/docs/smolagents/tutorials/secure_code_execution",
    )

    doc.add_heading("5. 无框架级 Guardrail / PII 脱敏", level=2)
    doc.add_paragraph(
        "所查 Agents / tools / secure_code_execution 文档均无 "
        "HallucinationGuardrail、PII 中间件或 tool output 统一消毒层；"
        "安全依赖沙箱选型与应用侧 final_answer_checks。"
    )
    cite(
        "Agents reference",
        "https://huggingface.co/docs/smolagents/reference/agents",
    )

    # ── 8. Architecture ──
    doc.add_heading("架构总览", level=1)
    bullet(
        "核心代码 agents.py <1000 行（README 自述）；"
        "MultiStepAgent → CodeAgent / ToolCallingAgent。"
    )
    bullet(
        "ReAct 循环：memory → write_memory_to_messages → model → "
        "extract_action / process_tool_calls → executor / tools → observations → memory。"
    )
    bullet(
        "Model 抽象：InferenceClientModel、LiteLLMModel、OpenAIModel、"
        "TransformersModel、AzureOpenAIModel、AmazonBedrockModel 等。"
    )
    bullet(
        "多 Agent：CodeAgent 可管理 ToolCallingAgent（inspect_runs 示例）；"
        "managed agent 亦可为 CodeAgent。"
    )
    bullet(
        "实验性 API：官方文档注明 API 可能变更、结果因模型而异。"
    )
    cite(
        "README；intro_agents；reference/agents",
        "https://huggingface.co/docs/smolagents/index",
    )

    doc.add_heading(
        "与 OpenClaw / OpenHands / Hermes / LangGraph / CrewAI / AutoGen / OpenManus 的对照（研究映射）",
        level=2,
    )
    doc.add_paragraph("非 smolagents 官方术语：")
    bullet(
        "CodeAgent 代码动作 ≈ 可组合多工具单次 LLM 步；"
        "对比 ToolCallingAgent 类似 AutoGen/OpenManus 的 JSON tool_calls 路径。"
    )
    bullet(
        "AgentMemory.steps ≈ 单 run 轨迹；"
        "reset=False 类似会话延续；无 LangGraph checkpointer / CrewAI Memory 语义层。"
    )
    bullet(
        "managed_agents ≈ AutoGen AgentTool / CrewAI hierarchical process 简化版；"
        "planning_interval ≈ CrewAI planning=True 周期性计划。"
    )
    bullet(
        "webagent + Helium ≈ OpenManus BrowserUseTool / AutoGen MultimodalWebSurfer；"
        "VisitWebpageTool ≈ 轻量 HTTP 抓取，非完整浏览器 agent。"
    )
    bullet(
        "LocalPythonExecutor ≈ OpenManus PythonExecute（本地解释）；"
        "e2b/docker executor ≈ 可选沙箱层，须显式配置。"
    )
    bullet(
        "SmolagentsInstrumentor ≈ LangGraph LangSmith trace / CrewAI OTel 集成；"
        "非默认开启。"
    )

    # ── 9. Research ──
    doc.add_heading("研究建议", level=1)
    bullet(
        "采集：verbosity_level=2 + get_full_steps()；"
        "或 SmolagentsInstrumentor + Phoenix/MLflow。"
    )
    bullet(
        "对比 CodeAgent vs ToolCallingAgent 同任务下的步数与 IPI 易感面。"
    )
    bullet(
        "IPI：VisitWebpageTool markdown、WebSearchTool 摘要、"
        "webagent observations、MCP 返回、managed_agent report。"
    )
    bullet(
        "安全：local vs e2b/docker；验证 managed_agents 在方案 1 沙箱下的限制。"
    )
    bullet(
        "会话：reset=False 多轮 vs memory.reset()；"
        "无内置跨进程持久化须自研序列化。"
    )
    bullet(
        "流式：区分 run(stream=True) 逐步事件与 stream_outputs 执行期流式。"
    )

    # ── 10. Five-round audit ──
    doc.add_heading("自查：潜在幻觉与核实结论（第五轮）", level=1)
    doc.add_paragraph(
        "2026-07-12 第五轮复核：对照 GitHub README、HF Docs（agents/tools/memory/"
        "secure_code_execution/inspect_runs/intro_agents/web_browser）、"
        "default_tools.py、SECURITY.md 全部关键断言。"
        "结论：第四轮修正项仍成立；本轮未发现新的机制性幻觉。"
    )
    checks = [
        ("✓ 双 Agent 类型：CodeAgent（Python 代码动作）vs ToolCallingAgent（JSON）", "https://huggingface.co/docs/smolagents/reference/agents"),
        ("✓ MultiStepAgent ReAct：think → act → observe；max_steps 默认 20", "https://huggingface.co/docs/smolagents/conceptual_guides/intro_agents"),
        ("✓ AgentMemory = system_prompt + Task/Action/Planning steps；非跨会话向量库", "https://huggingface.co/docs/smolagents/tutorials/memory"),
        ("✓ run(reset=False) 延续 memory；reset=True 默认新任务", "https://huggingface.co/docs/smolagents/reference/agents"),
        ("✓ LocalPythonExecutor 为 AST 解释器；官方明确非安全边界", "https://huggingface.co/docs/smolagents/tutorials/secure_code_execution"),
        ("✓ 远程沙箱：executor_type=blaxel|e2b|modal|docker + cleanup", "https://huggingface.co/docs/smolagents/tutorials/secure_code_execution"),
        ("✓ 方案 1 片段沙箱不支持 managed_agents；方案 2 整系统进沙箱可支持", "https://huggingface.co/docs/smolagents/tutorials/secure_code_execution"),
        ("✓ add_base_tools 默认 False；True 时挂载 python_interpreter/web_search/visit_webpage", "https://github.com/huggingface/smolagents/blob/main/src/smolagents/default_tools.py"),
        ("✓ VisitWebpageTool：requests+markdownify；max_output_length 默认 40000", "https://github.com/huggingface/smolagents/blob/main/src/smolagents/default_tools.py"),
        ("✓ CLI：smolagent（通用）与 webagent（helium 浏览器）并存", "https://github.com/huggingface/smolagents/blob/main/README.md"),
        ("✓ OpenTelemetry：SmolagentsInstrumentor；MLflow autolog 支持", "https://huggingface.co/docs/smolagents/tutorials/inspect_runs"),
        ("✓ managed_agents + planning_interval 多 Agent/规划", "https://huggingface.co/docs/smolagents/reference/agents"),
        ("✓ MCP：ToolCollection.from_mcp；Hub from_hub 须 trust_remote_code", "https://huggingface.co/docs/smolagents/reference/tools"),
        ("✓ secure_code_execution 明确网页 prompt injection 风险", "https://huggingface.co/docs/smolagents/tutorials/secure_code_execution"),
        ("△ stream_outputs（执行期）vs run(stream=True)（逐步 yield）— 不同参数，均已分述", "https://huggingface.co/docs/smolagents/reference/agents"),
        ("△ Web 双路径：VisitWebpageTool/WebSearchTool 工具链 vs webagent CLI helium 浏览器", "https://huggingface.co/docs/smolagents/examples/web_browser"),
        ("△ save/push_to_hub 导出 agent 配置 — 非 runtime memory 自动持久化", "https://huggingface.co/docs/smolagents/reference/agents"),
        ("△ UserInputTool 阻塞 input() — 非异步 HITL middleware", "https://github.com/huggingface/smolagents/blob/main/src/smolagents/default_tools.py"),
        ("✗ 内置 web_fetch / EXTERNAL_UNTRUSTED_CONTENT — 所查文档均无", "https://huggingface.co/docs/smolagents/tutorials/secure_code_execution"),
        ("✗ OpenClaw 式 Heartbeat / 内置 cron — 框架无", "https://github.com/huggingface/smolagents/blob/main/README.md"),
        ("✗ 跨会话语义 Memory / checkpointer — 框架无；仅 AgentMemory 步骤列表", "https://huggingface.co/docs/smolagents/tutorials/memory"),
        ("✗ 框架级 HallucinationGuardrail / PII 脱敏 — 所查文档无", "https://huggingface.co/docs/smolagents/reference/agents"),
        ("✗ LocalPythonExecutor 可作生产安全沙箱 — README/SECURITY 明确否定", "https://github.com/huggingface/smolagents/blob/main/README.md"),
        ("— 第五轮：上述条目全部复检通过，无新增 ✗/△", "https://huggingface.co/docs/smolagents/index"),
    ]
    for text, url in checks:
        cite(text, url)

    doc.add_heading("防幻觉机制说明（五轮复核流程）", level=2)
    bullet("第一轮：确认 GitHub README 与 HF Docs 权威来源，建立 CodeAgent/ToolCallingAgent 分层。")
    bullet("第二轮：核对 AgentMemory（run 内 steps）≠ 跨会话语义/RAG 长期记忆。")
    bullet("第三轮：核对 Web 双路径（default_tools vs webagent/helium）及沙箱 executor_type。")
    bullet("第四轮：核对 stream_outputs vs run(stream=) 命名差异；方案 1/2 沙箱与 multi-agent 限制。")
    bullet("第五轮：全量重读核心文档与 default_tools.py，确认无 OpenClaw 式机制误标。")

    doc.add_heading("主要参考链接", level=1)
    for link in [
        "https://github.com/huggingface/smolagents",
        "https://github.com/huggingface/smolagents/blob/main/README.md",
        "https://github.com/huggingface/smolagents/blob/main/SECURITY.md",
        "https://github.com/huggingface/smolagents/blob/main/src/smolagents/default_tools.py",
        "https://github.com/huggingface/smolagents/blob/main/src/smolagents/vision_web_browser.py",
        "https://huggingface.co/docs/smolagents/index",
        "https://huggingface.co/docs/smolagents/conceptual_guides/intro_agents",
        "https://huggingface.co/docs/smolagents/reference/agents",
        "https://huggingface.co/docs/smolagents/reference/tools",
        "https://huggingface.co/docs/smolagents/tutorials/memory",
        "https://huggingface.co/docs/smolagents/tutorials/secure_code_execution",
        "https://huggingface.co/docs/smolagents/tutorials/inspect_runs",
        "https://huggingface.co/docs/smolagents/examples/web_browser",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
