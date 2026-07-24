# -*- coding: utf-8 -*-
"""Generate Cline research notes (English).

Same structure as Chinese _gen_cline_doc.py:
Session / task modes / web_fetch(security notice) / tools,
plus Prompt composition & memory/persistence. Five-round anti-hallucination vs docs.cline.bot.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "en" / "Cline.docx"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label} (Source: {url})")

    # ── Title ──
    h("Cline", 0)
    p(
        "This document follows the same structure as OpenHands.docx / Hermes.docx / "
        "Claude-Code.docx / Dify.docx / AutoGPT.docx: "
        "Session, task modes, web_fetch (and Security Notice comparison), Tools; "
        "plus Prompt composition, memory and persistence mechanisms, and five-round "
        "anti-hallucination verification at the end. "
        "Subject: open-source coding agent Cline (VS Code / JetBrains extension, CLI, Kanban, SDK shared agent core)."
    )
    p("Official documentation index: https://docs.cline.bot/llms.txt")
    p(
        "Positioning: AI coding agent in editor and terminal — read/write files, run commands, "
        "fetch web/browser, MCP extensions; user approval required per step by default "
        "(Auto Approve / YOLO can relax this)."
    )
    cite("Cline Overview", "https://docs.cline.bot/cline-overview")
    p(
        "Execution chain (research view): User → agentic loop (Plan/Act) → "
        "tools (read_files / editor / bash / fetch_web / MCP…) → "
        "results fed back into session → (optional) Checkpoints snapshot workspace."
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "Under the ClineCore / Hub system, a Session is a persistent entity with full "
        "conversation history, tool call records, and metadata, with a sessionId; "
        "can attach across clients; Spoke can continue execution after disconnect."
    )
    cite(
        "Hub-Spoke Architecture · Session Persistence",
        "https://docs.cline.bot/sdk/architecture/hub-spoke",
    )
    cite("ClineCore", "https://docs.cline.bot/sdk/clinecore")

    h("1. Storage location", 2)
    bullet(
        "Default: ~/.cline/data/sessions/ — sessions.db (SQLite index) + "
        "[session-id].json (authoritative session snapshot)."
    )
    bullet(
        "Config root: ~/.cline/ (global) and project .cline/; "
        "isolation via --config / CLINE_DATA_DIR / --data-dir."
    )
    bullet(
        "ClineCore.start() returns session.sessionId, manifestPath, messagesPath, etc."
    )
    cite("Config", "https://docs.cline.bot/getting-started/config")

    h("2. Hub-Spoke and multi-client", 2)
    bullet(
        "Hub: local singleton daemon (default 127.0.0.1:25463), coordinates sessions "
        "and approval routing; does not run the agent loop."
    )
    bullet("Spoke: worker process runs agent loop and tools; Client (CLI/IDE) can come and go.")
    bullet(
        "backendMode: auto | hub | remote | local; "
        "local = in-process + local storage, no shared session."
    )
    bullet(
        "Participant roles: creator / participant / observer; "
        "CLI can resume with --id <session-id>; cline history manages history."
    )
    cite("CLI Reference", "https://docs.cline.bot/cli/cli-reference")

    h("3. Task vs Session (product semantics)", 2)
    bullet(
        "UI often says \"Task\"; slash `/newtask` starts a new task after distilling context "
        "(new clean context window)."
    )
    bullet(
        "Research mapping: multi-turn messages under same sessionId ≈ Claude Code project session; "
        "Schedule/Kanban each trigger is closer to independent or attached new session (see task modes)."
    )

    h("4. Comparison with other agents (non-official terminology)", 2)
    bullet("Claude Code ~/.claude/projects/*.jsonl ≈ Cline ~/.cline/data/sessions/*.json.")
    bullet("Hermes SessionDB ≈ Cline sessions.db + JSON snapshot.")
    bullet("Dify conversation_id API field ≠ Cline sessionId (different protocol).")

    # ══════════════════════════════════════════════════════════════
    # Task modes
    # ══════════════════════════════════════════════════════════════
    h("Task Modes", 1)

    h("1. Plan & Act (dual mode)", 2)
    bullet(
        "Plan: can read codebase, search, discuss strategy; cannot modify files or execute commands."
    )
    bullet(
        "Act: retains full Plan conversation context then executes file edits/commands; can switch back and forth."
    )
    bullet("Can configure different models for Plan/Act; CLI `-p/--plan` starts Plan.")
    bullet("`/deep-planning`: deeper planning → implementation_plan.md → new implementation task.")
    cite("Plan & Act Mode", "https://docs.cline.bot/core-workflows/plan-and-act")

    h("2. Product entry points", 2)
    bullet("IDE: VS Code / JetBrains extension (and cross-editor ACP: Cursor, Windsurf, Zed, Neovim, etc.).")
    bullet("CLI interactive: `cline` / TUI (`-i`); `--zen` background hub session.")
    bullet(
        "Headless: `--json`, stdin pipe, stdout redirect → CI/scripts; "
        "`--auto-approve`, `--timeout`."
    )
    bullet("Kanban: multi-agent task board + per-card worktree.")
    cite("CLI Overview", "https://docs.cline.bot/usage/cli-overview")

    h("3. Auto Approve / YOLO", 2)
    bullet(
        "Per-tool-category approval: read/write project or all files, safe/all commands, "
        "Use the browser, Use MCP, notifications, etc."
    )
    bullet(
        "YOLO: auto-approve everything (including Plan→Act switch); officially marked dangerous; "
        "isolated environment recommended."
    )
    cite(
        "Auto Approve & YOLO Mode",
        "https://docs.cline.bot/features/auto-approve",
    )

    h("4. Scheduling / Teams / Subagents", 2)
    bullet(
        "Schedule: cron via hub; available in CLI/SDK/Kanban; "
        "docs state not applicable to VS Code / JetBrains extension."
    )
    bullet("Agent Teams: multi-agent collaboration; Subagents: parallel research without filling main context.")
    cite("Scheduling", "https://docs.cline.bot/cli/scheduling")
    cite("Subagents", "https://docs.cline.bot/features/subagents")

    h("Task Modes · verification summary", 2)
    bullet(
        "Keep: Plan/Act; IDE/CLI/Headless/Kanban; Auto Approve/YOLO; "
        "Schedule (limited surface); Teams/Subagents."
    )
    bullet("Drop: OpenClaw Heartbeat; writing Schedule as built-in extension capability.")

    # ══════════════════════════════════════════════════════════════
    # web_fetch
    # ══════════════════════════════════════════════════════════════
    h("web_fetch (and Security Notice)", 1)
    p(
        "Conclusion: official tools reference names the built-in external retrieval tool "
        "`fetch_web` (HTTP + HTML→markdown), categorized as External retrieval. "
        "Public docs do not document OpenClaw-style SECURITY NOTICE / EXTERNAL_UNTRUSTED_CONTENT wrapping, "
        "nor Claude Code WebFetch \"isolated small-model extraction + ~15min cache\" behavior. "
        "Separate Browser capability surface (Auto Approve \"Use the browser\"); "
        "Overview says \"use a browser\"."
    )

    h("1. Built-in names in docs (note naming drift)", 2)
    bullet(
        "tools-reference (aligned with ClineCore): "
        "`fetch_web` — HTTP requests with HTML-to-markdown conversion."
    )
    bullet(
        "sdk/tools table also lists `fetch_web_content` (Fetch web content); "
        "naming not fully consistent with above — record both at research freeze, do not treat one as absolute truth."
    )
    bullet(
        "Legacy extension/example XML-style names (read_file / execute_command, etc.) "
        "differ from current runtime names; docs require alignment with new names."
    )
    cite(
        "Tools Reference · Built-In",
        "https://docs.cline.bot/tools-reference/all-cline-tools",
    )
    cite("SDK Tools", "https://docs.cline.bot/sdk/tools")

    h("2. Browser vs fetch_web", 2)
    bullet(
        "Auto Approve has separate item \"Use the browser\": "
        "Browser tool for web fetching and searching — "
        "product semantics parallel to generic file/terminal approval."
    )
    bullet(
        "Public user docs do not give fixed internal function name list for browser tool "
        "(e.g. browser_action in historical community discussion); "
        "do not fabricate parameter protocols not appearing on docs.cline.bot."
    )
    bullet("MCP can add third-party browser/scraping tools on top.")

    h("3. Approval and enablement conditions", 2)
    bullet("Default design: user approval required before tool calls (unless Auto Approve / YOLO).")
    bullet(
        "SDK toolPolicies / requestToolApproval can fine-grain autoApprove per tool."
    )
    bullet(
        "Extension implementation side (non-Mintlify main doc terms): historical source has "
        "web_fetch API path for Cline account and clineWebToolsEnabled switch — "
        "mark as \"implementation detail / pending official doc solidification\", "
        "do not write as llms.txt formal contract."
    )

    h("4. Security Notice comparison", 2)
    bullet(
        "Docs check: no SECURITY NOTICE / untrusted_tool_result / "
        "EXTERNAL_UNTRUSTED_CONTENT default wrapping documented."
    )
    bullet(
        "Security-related public surface: "
        "approval gate, .clineignore, CLINE_COMMAND_PERMISSIONS, "
        "rules/hooks/plugins \"trusted sources only\", Enterprise MCP allowlist / YOLO control."
    )
    bullet(
        "Enterprise Prompt Storage can back up conversations to S3/R2 (compliance); "
        "not untrusted marking of web content."
    )

    h("Threat model mapping (research, non-official)", 2)
    bullet(
        "External web markdown/HTML → model context → subsequent editor/bash/MCP: "
        "Indirect Prompt Injection / Goal Hijacking paper surface."
    )
    bullet(
        "Browser interaction surface may be larger (navigation/forms); ablate separately from fetch_web alone."
    )
    bullet(
        "Compare: OpenClaw web_fetch+NOTICE; Claude Code WebFetch+isolated extraction; "
        "Hermes web_extract+untrusted tag; "
        "Cline ≈ fetch_web(+browser) + approval, no NOTICE at doc layer."
    )

    code(
        """Research sketch (non-official message binding):
User: Summarize https://example.com/docs
→ tool: fetch_web { url: "https://example.com/docs" }
← tool result: (HTML→markdown body; docs do not specify SECURITY NOTICE wrapper)
→ model may continue editor/bash/MCP
Note: concrete tool schema per implementation/SDK; do not write as OpenClaw EXTERNAL_UNTRUSTED_CONTENT."""
    )

    h("web_fetch · verification summary", 2)
    bullet(
        "Keep: fetch_web (and fetch_web_content alias in docs); "
        "browser approval category; no public SECURITY NOTICE; HTML→markdown."
    )
    bullet(
        "Drop: default 15min cache+isolated extraction; "
        "writing OpenClaw NOTICE text as Cline built-in."
    )

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "ClineCore built-in tools + MCP (.cline/mcp.json) + "
        "custom tools/plugins (docs: SDK/CLI/Kanban; extension side not fully applicable)."
    )

    h("1. ClineCore built-in (tools-reference)", 2)
    bullet("`bash` — execute shell.")
    bullet("`editor` — view and edit files.")
    bullet("`read_files` — batch read files.")
    bullet("`apply_patch` — unified diff patch.")
    bullet("`search` — ripgrep code search.")
    bullet("`fetch_web` — external HTTP + HTML→markdown.")
    bullet("`ask_question` — ask user questions.")
    p(
        "SDK Tools page also lists: search_codebase, run_commands, fetch_web_content, "
        "skills, submit_and_exit, etc. — same harness has naming evolution; cite source page when referencing."
    )

    h("2. MCP / Plugins / Skills", 2)
    bullet("MCP: loaded alongside built-in tools; Enterprise can allowlist.")
    bullet("Plugins: register tools, hooks, modify behavior (SDK/CLI/Kanban).")
    bullet("Skills: modular SKILL.md; trigger via slash `/skill-name`.")
    cite("MCP Overview", "https://docs.cline.bot/mcp/mcp-overview")
    cite("Skills", "https://docs.cline.bot/customization/skills")

    h("3. Access control", 2)
    bullet(".clineignore: restrict accessible paths.")
    bullet("CLINE_COMMAND_PERMISSIONS: allow/deny command patterns and redirects.")
    bullet("Sandbox: CLINE_SANDBOX / CLINE_SANDBOX_DATA_DIR.")

    h("Tools · verification summary", 2)
    bullet("Keep: ClineCore table + MCP + approval policy; names per official table with drift noted.")
    bullet("Drop: fabricating fixed PascalCase WebFetch as sole name.")

    # ══════════════════════════════════════════════════════════════
    # Five-round anti-hallucination
    # ══════════════════════════════════════════════════════════════
    h("Five-Round Anti-Hallucination Verification", 1)

    h("Round 1 — Session", 2)
    bullet("Checked: hub-spoke; clinecore; config; cli-reference --id/history.")
    bullet(
        "Keep: sessions.db + [id].json; sessionId; hub multi-client; local vs hub."
    )
    bullet("Drop: Claude Code JSONL path; Dify conversation_id field name.")

    h("Round 2 — Task modes", 2)
    bullet("Checked: plan-and-act; cli-overview; auto-approve; scheduling; kanban.")
    bullet("Keep: Plan/Act; Headless; YOLO; Schedule surface limitation.")
    bullet("Drop: schedule available in extension (contradicts doc Warning).")

    h("Round 3 — web_fetch / security", 2)
    bullet("Checked: all-cline-tools; sdk/tools; auto-approve browser; overview.")
    bullet(
        "Keep: fetch_web; HTML→markdown; browser category; no SECURITY NOTICE in docs."
    )
    bullet(
        "Drop: OpenClaw NOTICE; Claude Code 15min cache claim; "
        "undocumented browser_action parameter table."
    )

    h("Round 4 — Tools", 2)
    bullet("Checked: all-cline-tools; sdk/tools; mcp; plugins/skills; clineignore.")
    bullet("Keep: built-in seven-piece set (reference table) + MCP/plugins; legacy name deprecation note.")
    bullet("Drop: citing only sdk page or only reference page without noting drift.")

    h("Round 5 — Prompt / Memory cross-check", 2)
    bullet("Checked: cline-rules; memory-bank; using-commands; checkpoints; clinecore systemPrompt.")
    bullet(
        "Freeze: "
        "(1) Session = sessionId + ~/.cline/data/sessions; "
        "(2) Modes = Plan/Act + entry points (IDE/CLI/Headless/Kanban) + approval level; "
        "(3) External web = fetch_web(+browser/MCP), no doc SECURITY NOTICE; "
        "(4) Prompt = systemPrompt + Rules + Skills; "
        "(5) Cross-session project memory = Memory Bank methodology (rules-driven markdown), "
        "not Hermes automatic MEMORY.md loading."
    )
    bullet(
        "Open questions: whether fetch_web and fetch_web_content are same runtime implementation; "
        "official browser tool schema; whether extension web_fetch API remains primary path."
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt composition
    # ══════════════════════════════════════════════════════════════
    h("Prompt Composition", 1)
    p(
        "Runtime assembled by harness: configurable systemPrompt + auto-loaded Rules/Skills instructions "
        "+ user messages/@file context + tool results. Full system template not fully public."
    )

    h("1. systemPrompt (SDK/ClineCore)", 2)
    bullet("start/config can pass systemPrompt (example: helpful coding assistant).")
    bullet("CLI side also has system prompt override capability (see CLI Reference updates).")

    h("2. Rules (persistent instructions)", 2)
    bullet(
        "Workspace .clinerules/, global Rules directory, and auto-detect "
        ".cursorrules / .windsurfrules / AGENTS.md."
    )
    bullet("Workspace and global merged; workspace wins on conflict.")
    cite("Rules", "https://docs.cline.bot/customization/cline-rules")

    h("3. Skills and slash commands", 2)
    bullet("Skills module instructions enter capability instructions; `/skill` triggers.")
    bullet(
        "`/deep-planning`, `/newtask`, `/smol` (`/compact`), `/newrule` "
        "manage planning and context."
    )
    cite("Using Commands", "https://docs.cline.bot/core-workflows/using-commands")

    h("4. User context injection", 2)
    bullet("@ mentions and drag files into conversation (Adding Context).")
    bullet("CLI: `@./image.png`, pipe stdin as task context.")

    h("Prompt Composition · verification summary", 2)
    bullet("Keep: systemPrompt + Rules + Skills + @context + tool results.")
    bullet("Drop: CLAUDE.md auto system or Hermes SOUL.md as Cline default filename.")

    # ══════════════════════════════════════════════════════════════
    # Memory and persistence
    # ══════════════════════════════════════════════════════════════
    h("Memory and Persistence", 1)
    p(
        "Layers: "
        "(A) session-level JSON/SQLite persistence; (B) context compression slash; "
        "(C) Memory Bank project documentation methodology; (D) Checkpoints workspace shadow Git; "
        "(E) optional Enterprise Prompt Storage."
    )

    h("1. Session message persistence", 2)
    bullet("Full history + tool records written to session JSON; list/get/readMessages API.")
    bullet("Cross-process/client resume under Hub; unrelated to MEMORY.md file system.")

    h("2. Context window management", 2)
    bullet("`/smol`: compress conversation within same task.")
    bullet("`/newtask`: distilled new task, clean context.")
    bullet("CLI compaction flags see CLI Reference (agentic/basic/off, etc.; per current page).")

    h("3. Memory Bank (best practice, not built-in DB)", 2)
    bullet(
        "Methodology injected via Rules: project memory-bank/*.md "
        "(projectbrief, productContext, activeContext, systemPatterns, "
        "techContext, progress)."
    )
    bullet(
        "Instructions emphasize \"memory resets between sessions, must rely on Memory Bank files\" — "
        "requires user initialize/update; not automatic MEMORY.md loader."
    )
    cite("Memory Bank", "https://docs.cline.bot/best-practices/memory-bank")

    h("4. Checkpoints (code state, not semantic memory)", 2)
    bullet("Shadow Git: snapshot after each tool file edit/command; Restore Files / Task / Both.")
    bullet("Conversation can remain while code rolls back — serves experimentation, not cross-session knowledge base.")
    cite("Checkpoints", "https://docs.cline.bot/core-workflows/checkpoints")

    h("5. Other", 2)
    bullet("Skills/Rules/Hooks/cron spec files themselves are persistent config memory.")
    bullet("Enterprise Prompt Storage: conversation backup to S3/R2.")

    h("Memory and Persistence · verification summary", 2)
    bullet(
        "Keep: session JSON; /smol·/newtask; Memory Bank methodology; Checkpoints; "
        "optional enterprise backup."
    )
    bullet(
        "Drop: writing Memory Bank as out-of-box automatic built-in vector memory; "
        "Hermes ~/.hermes/memories/MEMORY.md path."
    )

    h("Primary Reference Links", 1)
    for link in [
        "https://docs.cline.bot/llms.txt",
        "https://docs.cline.bot/cline-overview",
        "https://docs.cline.bot/sdk/architecture/hub-spoke",
        "https://docs.cline.bot/sdk/clinecore",
        "https://docs.cline.bot/getting-started/config",
        "https://docs.cline.bot/core-workflows/plan-and-act",
        "https://docs.cline.bot/core-workflows/checkpoints",
        "https://docs.cline.bot/core-workflows/using-commands",
        "https://docs.cline.bot/features/auto-approve",
        "https://docs.cline.bot/tools-reference/all-cline-tools",
        "https://docs.cline.bot/sdk/tools",
        "https://docs.cline.bot/mcp/mcp-overview",
        "https://docs.cline.bot/customization/cline-rules",
        "https://docs.cline.bot/customization/skills",
        "https://docs.cline.bot/best-practices/memory-bank",
        "https://docs.cline.bot/usage/cli-overview",
        "https://docs.cline.bot/cli/cli-reference",
        "https://docs.cline.bot/cli/scheduling",
        "https://docs.cline.bot/customization/clineignore",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
