# -*- coding: utf-8 -*-
"""Generate MetaGPT Agent research document — accuracy-focused with 5-round audit."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "MetaGPT.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def add_code(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def cite(text: str, url: str) -> None:
        doc.add_paragraph(f"{text}（来源：{url}）")

    doc.add_heading("MetaGPT", level=0)
    doc.add_paragraph(
        "本文档研究 MetaGPT 多智能体框架机制，术语严格对齐官方资料。"
        "用户指定入口：https://docs.deepwisdom.ai/main/en/guide/get_started/introduction.html；"
        "技术细节以 docs.deepwisdom.ai 文档站与 GitHub geekan/MetaGPT 源码为准。"
        "README 另列商业产品 MGX（mgx.dev），与开源框架本体须区分。"
    )

    # ── Scope ──
    doc.add_heading("文档范围与术语分层（必读）", level=1)
    doc.add_paragraph(
        "introduction 页强调 Code = SOP(Team) 与软件公司多角色 SOP；"
        "研究须区分以下层次，不可混用："
    )
    bullet(
        "概念层：Agent = LLM + Observation + Thought + Action + Memory；"
        "MultiAgent = Agents + Environment + SOP + Communication + Economy。"
    )
    bullet(
        "框架原子层：Action（可含 _aask LLM 调用或非 LLM 逻辑）与 Role（Agent 抽象，"
        "含 memory、react_mode、_act/_think）。"
    )
    bullet(
        "多 Agent 运行时：Team + Environment；消息经 publish_message 广播，"
        "Role._watch 按 cause_by 订阅上游 Action。"
    )
    bullet(
        "Software Company 内置 SOP：CLI metagpt \"...\" / generate_repo() → "
        "ProductManager / Architect / ProjectManager / Engineer 等预置 Role。"
    )
    bullet(
        "DataInterpreter（DI）：独立 Role，@register_tool 工具注册表 + 代码执行，"
        "≠ Software Company SOP 路径。"
    )
    bullet(
        "RAG 模块（metagpt[rag]）：SimpleEngine 独立向量检索引擎，"
        "≠ Role.rc.memory 短期 Message 列表。"
    )
    bullet(
        "MGX / MGXEnv：README 商业产品与自然语言编程团队；"
        "Team 默认 use_mgx=True 可选用 MGXEnv，≠ 开源文档核心教程路径。"
    )
    cite(
        "introduction；concepts；README",
        "https://docs.deepwisdom.ai/main/en/guide/get_started/introduction.html",
    )
    cite(
        "agent_101；multi_agent_101；team.py",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/agent_101.html",
    )

    # ── 1. LLM 请求采集 ──
    doc.add_heading("研究背景与 LLM 请求采集", level=1)
    doc.add_paragraph(
        "采集目标：每次 LLM 调用前的 prompt、Action._aask 上下文、"
        "Message 链（content / instruct_content / cause_by）、"
        "工具输出、搜索/浏览结果与 cost_manager 费用轨迹。"
    )

    doc.add_heading("1. loguru 文件日志（metagpt/logs.py）", level=2)
    bullet(
        "define_log_level：stderr INFO + METAGPT_ROOT/logs/{date}.txt DEBUG。"
    )
    bullet(
        "logger.info/debug 贯穿 Role.run、Action.run、Team.run 等路径。"
    )
    bullet(
        "无内置 OpenTelemetry / Langfuse trace；观测依赖日志文件与 messages 导出。"
    )
    cite(
        "metagpt/logs.py",
        "https://github.com/geekan/MetaGPT/blob/main/metagpt/logs.py",
    )

    doc.add_heading("2. LLM 流式输出（log_llm_stream）", level=2)
    bullet(
        "create_llm_stream_queue + log_llm_stream：通过 ContextVar asyncio.Queue "
        "推送 token 流；未建队列时仅 partial(print) 到终端。"
    )
    bullet(
        "研究侧可 set_llm_stream_logfunc 自定义流式 hook。"
    )
    cite(
        "metagpt/logs.py > log_llm_stream / create_llm_stream_queue",
        "https://github.com/geekan/MetaGPT/blob/main/metagpt/logs.py",
    )

    doc.add_heading("3. 工具输出日志（log_tool_output）", level=2)
    bullet(
        "ToolLogItem（type/name/value）+ log_tool_output / log_tool_output_async；"
        "set_tool_output_logfunc 可替换默认 no-op。"
    )
    bullet(
        "DataInterpreter 执行 @register_tool 工具时可经此接口观测返回。"
    )
    cite(
        "metagpt/logs.py；create_and_use_tools",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/create_and_use_tools.html",
    )

    doc.add_heading("4. Action._aask 与 LLM 调用链", level=2)
    bullet(
        "Action 子类 run() 内 await self._aask(prompt, [system_text])；"
        "框架在 Action 基类封装 LLM API 调用。"
    )
    bullet(
        "Researcher CollectLinks：_aask 分解子问题、筛选 URL；"
        "WebBrowseAndSummarize 浏览后 _aask 摘要。"
    )
    bullet(
        "config2.yaml repair_llm_output: true 可在 JSON 输出无效时尝试修复。"
    )
    cite(
        "agent_101 > SimpleWriteCode；researcher > CollectLinks",
        "https://docs.deepwisdom.ai/main/en/guide/use_cases/agent/researcher.html",
    )
    cite(
        "config/config2.example.yaml",
        "https://github.com/geekan/MetaGPT/blob/main/config/config2.example.yaml",
    )

    doc.add_heading("5. Team 费用与轮次可观测", level=2)
    bullet(
        "Team.invest(investment) 设置 cost_manager.max_budget；"
        "run(n_round) 每轮 _check_balance，超预算抛 NoMoneyException。"
    )
    bullet(
        "env.is_idle 为 True 时提前结束；archive(auto_archive) 归档项目。"
    )
    cite(
        "metagpt/team.py > invest / run / _check_balance",
        "https://github.com/geekan/MetaGPT/blob/main/metagpt/team.py",
    )

    doc.add_heading("6. 序列化检查点（Team.serialize / deserialize）", level=2)
    bullet(
        "Team.serialize() 写 SERDESER_PATH/team/team.json + context 序列化；"
        "deserialize 恢复 Team 与 Context，用于断点续跑。"
    )
    bullet(
        "serialize_decorator 装饰 Team.run；Message instruct_content 经 pickle + schema mapping。"
    )
    bullet(
        "官方 in_depth serialization 页面 fetch 常为落地页；机制以 team.py / serialize.py 为准。"
    )
    cite(
        "metagpt/team.py；metagpt/utils/serialize.py",
        "https://github.com/geekan/MetaGPT/blob/main/metagpt/team.py",
    )

    # ── 2. Prompt ──
    doc.add_heading("Prompt 设计组成", level=1)

    doc.add_heading("1. Role 人设（name / profile / goal / constraints）", level=2)
    bullet(
        "预置 Role（ProductManager、Researcher 等）在 __init__ 设置 goal/constraints；"
        "自定义 Role 可覆盖 _act 逻辑。"
    )
    bullet(
        "Researcher 按 language（en-us / zh-cn）选择 research_system_text prompt。"
    )
    cite(
        "researcher > Researcher.__init__",
        "https://docs.deepwisdom.ai/main/en/guide/use_cases/agent/researcher.html",
    )

    doc.add_heading("2. Action PROMPT_TEMPLATE", level=2)
    bullet(
        "Action 类内 PROMPT_TEMPLATE 字符串 + run() 内 format(instruction=...) → _aask。"
    )
    bullet(
        "SimpleWriteCode：要求返回 ```python ... ``` 无其他文本。"
    )
    cite(
        "agent_101 > SimpleWriteCode",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/agent_101.html",
    )

    doc.add_heading("3. react_mode 与 Action 选择", level=2)
    bullet(
        "_set_react_mode(react_mode=\"by_order\")：按 set_actions 顺序执行；"
        "多 Action Role 的 self.rc.todo 随轮次切换。"
    )
    bullet(
        "文档亦提及 by_llm 等模式（Think and act 章节）；"
        "by_order 为 Researcher / RunnableCoder 示例默认。"
    )
    cite(
        "agent_101 > RunnableCoder；researcher",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/agent_101.html",
    )

    doc.add_heading("4. Message 结构化输出（instruct_content）", level=2)
    bullet(
        "Message：content、role、cause_by、sent_from、send_to、instruct_content（Pydantic BaseModel）。"
    )
    bullet(
        "Researcher 用 Report(topic, links/summaries/content) 在 Action 间传递结构化状态。"
    )
    cite(
        "agent_communication > Message；researcher",
        "https://docs.deepwisdom.ai/main/en/guide/in_depth_guides/agent_communication.html",
    )

    doc.add_heading("5. 按 Role 定制 LLM（config2 roles 段）", level=2)
    bullet(
        "config2.yaml roles 列表：按 ProductManager / Architect / Engineer 等 "
        "分别配置 api_type、model、base_url。"
    )
    bullet(
        "prompt_schema: json | markdown（configuration 页）。"
    )
    cite(
        "config/config2.example.yaml；configuration",
        "https://docs.deepwisdom.ai/main/en/guide/get_started/configuration.html",
    )

    doc.add_heading("6. Context 共享", level=2)
    bullet(
        "Context() 显式创建后传入 Role/Team；Role 与其 Action 隐式共享同一 Context（含 config、cost_manager）。"
    )
    cite(
        "agent_101 > ProductManager 示例",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/agent_101.html",
    )

    # ── 3. Memory ──
    doc.add_heading("Memory 与持久化", level=1)

    doc.add_heading("1. Role 短期记忆（rc.memory）", level=2)
    bullet(
        "Memory 存储 Role _observe 到的 Message 列表；get_memories(k) 按最近 k 条检索，k=0 返回全部。"
    )
    bullet(
        "use_memories 教程明确此为 short-term memory，检索基于 recency，非语义向量。"
    )
    bullet(
        "_act 中可 self.rc.memory.add(msg) 记住本 Role 上一轮输出。"
    )
    cite(
        "use_memories",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/use_memories.html",
    )

    doc.add_heading("2. RAG SimpleEngine（独立持久化）", level=2)
    bullet(
        "pip install metagpt[rag]；SimpleEngine.from_docs / from_objs / from_index。"
    )
    bullet(
        "检索：Faiss / BM25 / ChromaDB / ES 及混合；"
        "Post-retrieval：LLMRanker / Colbert / Cohere / BGE 等。"
    )
    bullet(
        "persist(persist_dir) + FAISSIndexConfig 加载；向量不必每次重建。"
    )
    cite(
        "rag_module",
        "https://docs.deepwisdom.ai/main/en/guide/in_depth_guides/rag_module.html",
    )

    doc.add_heading("3. role_zero 长期记忆（可选，默认关闭）", level=2)
    bullet(
        "config2 role_zero.enable_longterm_memory: false（默认）；"
        "longterm_memory_persist_path、memory_k、similarity_top_k。"
    )
    bullet(
        "与 use_memories 短期 Message 列表不同路径；须显式开启。"
    )
    cite(
        "config/config2.example.yaml > role_zero",
        "https://github.com/geekan/MetaGPT/blob/main/config/config2.example.yaml",
    )

    doc.add_heading("4. exp_pool 经验池（可选）", level=2)
    bullet(
        "exp_pool.enabled: false；persist_path、retrieval_type（bm25/chroma）、use_llm_ranker。"
    )
    bullet(
        "独立 Chroma/BM25 经验检索，≠ Role.rc.memory。"
    )
    cite(
        "config/config2.example.yaml > exp_pool",
        "https://github.com/geekan/MetaGPT/blob/main/config/config2.example.yaml",
    )

    doc.add_heading("5. Team / Context 序列化", level=2)
    bullet(
        "Team.serialize / deserialize：跨运行恢复多 Agent 状态与 context；"
        "SERDESER_PATH 默认存储目录。"
    )
    bullet(
        "无 OpenClaw 式跨会话语义 MEMORY.md 自动注入。"
    )
    cite(
        "metagpt/team.py",
        "https://github.com/geekan/MetaGPT/blob/main/metagpt/team.py",
    )

    # ── 4. Session ──
    doc.add_heading("Session 与会话边界", level=1)

    doc.add_heading("1. 单 Role 循环", level=2)
    bullet(
        "while msg: msg = await role.run(msg)：单 Role 多轮直至返回空或终止。"
    )
    bullet(
        "Context 贯穿单次 Python 进程内多次 run。"
    )
    cite(
        "agent_101 > ProductManager main",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/agent_101.html",
    )

    doc.add_heading("2. Team 多 Agent 会话", level=2)
    bullet(
        "team.hire([roles]) → team.invest(budget) → team.run(n_round=3, idea=...)。"
    )
    bullet(
        "run_project 发布 Message(content=idea) 到 Environment；"
        "各 Role 异步 observe → react → act → publish_message。"
    )
    cite(
        "team.py；multi_agent_101（SimpleCoder/Tester/Reviewer）",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/multi_agent_101.html",
    )

    doc.add_heading("3. 人机协作（is_human=True）", level=2)
    bullet(
        "SimpleReviewer(is_human=True)：轮到人类 Role 时进程暂停等待终端输入。"
    )
    bullet(
        "限制：自定义 _act 调用的 Action 须在 set_actions 列表内；"
        "当前仅终端 input，多行/结构化不便。"
    )
    cite(
        "human_engagement",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/human_engagement.html",
    )

    doc.add_heading("4. Software Company 工作区", level=2)
    bullet(
        "metagpt \"Create a 2048 game\" 在 ./workspace 生成仓库；"
        "ProjectRepo 可读目录结构。"
    )
    cite(
        "README > Usage",
        "https://github.com/geekan/MetaGPT/blob/main/README.md",
    )

    # ── 5. Automation ──
    doc.add_heading("Automation 与编排", level=1)

    doc.add_heading("1. Team.run 轮次驱动", level=2)
    bullet(
        "async Team.run(n_round, idea)：固定轮次循环 env.run()，"
        "非 cron / webhook 产品层调度。"
    )
    bullet(
        "investment 预算控制 LLM 总花费；无预算则 NoMoneyException。"
    )
    cite(
        "team.py",
        "https://github.com/geekan/MetaGPT/blob/main/metagpt/team.py",
    )

    doc.add_heading("2. SOP 内 Action 顺序（react_mode=by_order）", level=2)
    bullet(
        "Researcher：CollectLinks → WebBrowseAndSummarize → ConductResearch 固定流水线。"
    )
    bullet(
        "Software Company 内置各 Role 的 Action 链为预置 SOP，非用户画布配置。"
    )
    cite(
        "researcher；introduction",
        "https://docs.deepwisdom.ai/main/en/guide/use_cases/agent/researcher.html",
    )

    doc.add_heading("3. Environment 消息路由", level=2)
    bullet(
        "publish_message(Message(..., cause_by=ActionClass, send_to=...))；"
        "_watch({ActionA, ActionB}) 订阅 cause_by。"
    )
    bullet(
        "env.is_idle：所有 Role 无新消息时 True；agent_communication 示例 while not env.is_idle: await env.run()。"
    )
    cite(
        "agent_communication",
        "https://docs.deepwisdom.ai/main/en/guide/in_depth_guides/agent_communication.html",
    )

    doc.add_heading("4. 无内置 Heartbeat / Schedule", level=2)
    doc.add_paragraph(
        "所查文档与 team.py 均无 OpenClaw 式定时唤醒或外部 Trigger 产品；"
        "自动化须应用侧 cron 调用 CLI/API 或自写循环。"
    )

    # ── 6. External Web ──
    doc.add_heading("外部 Web 与工具", level=1)

    doc.add_heading("1. Researcher 搜索 + 浏览链", level=2)
    bullet(
        "CollectLinks：SearchEngine（serpapi/google/serper/ddg）收集 URL。"
    )
    bullet(
        "WebBrowseAndSummarize：WebBrowserEngine（playwright 默认 / selenium）"
        "抓取页面并 LLM 摘要。"
    )
    bullet(
        "ConductResearch：汇总 summaries 生成最终报告；write_report 输出文件。"
    )
    cite(
        "researcher",
        "https://docs.deepwisdom.ai/main/en/guide/use_cases/agent/researcher.html",
    )

    doc.add_heading("2. 搜索/浏览器配置（config2 search / browser）", level=2)
    bullet(
        "search.engine + api_key（及 google cse_id）；"
        "browser.engine + browser_type。"
    )
    bullet(
        "可选依赖：metagpt[search-google]、metagpt[search-ddg]、"
        "metagpt[playwright]、metagpt[selenium]。"
    )
    cite(
        "researcher > Dependencies；config2.example.yaml",
        "https://docs.deepwisdom.ai/main/en/guide/use_cases/agent/researcher.html",
    )

    doc.add_heading("3. DataInterpreter 工具注册", level=2)
    bullet(
        "@register_tool 将 metagpt/tools/libs 下函数/类注册；"
        "DataInterpreter(tools=[\"calculate_factorial\"]) 按 docstring 选型调用。"
    )
    bullet(
        "工具可执行代码、数学、文件等；非统一 web_fetch 抽象。"
    )
    cite(
        "create_and_use_tools",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/create_and_use_tools.html",
    )

    doc.add_heading("4. 其他用例（文档索引）", level=2)
    bullet(
        "Tool usage: web scraping、email summarization、text2image 等列于 Use Cases；"
        "各用例独立 Role/Action 实现。"
    )
    cite(
        "README Tutorial > Use Cases",
        "https://github.com/geekan/MetaGPT/blob/main/README.md",
    )

    doc.add_heading("5. IPI 研究含义", level=2)
    doc.add_paragraph(
        "注入面：SearchEngine 摘要、WebBrowserEngine 页面正文、"
        "RAG SimpleEngine 检索 chunk、DataInterpreter 工具返回、"
        "Environment 广播 Message content。"
        "外部网页内容进入 _aask prompt 或 instruct_content，"
        "无 EXTERNAL_UNTRUSTED_CONTENT 自动消毒层。"
    )

    # ── 7. Security ──
    doc.add_heading("Security 与安全机制", level=1)

    doc.add_heading("1. 预算熔断（cost_manager）", level=2)
    bullet(
        "investment 上限；超额 NoMoneyException 停止 Team.run。"
    )
    cite(
        "team.py",
        "https://github.com/geekan/MetaGPT/blob/main/metagpt/team.py",
    )

    doc.add_heading("2. 代码执行风险（DataInterpreter / SimpleRunCode）", level=2)
    bullet(
        "SimpleRunCode 示例 subprocess python3 -c；"
        "DataInterpreter 可执行生成代码，属高权限本地执行面。"
    )
    cite(
        "agent_101 > SimpleRunCode；create_and_use_tools",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/agent_101.html",
    )

    doc.add_heading("3. 无统一不可信内容标记", level=2)
    doc.add_paragraph(
        "所查 introduction、researcher、agent_communication、README 均无 "
        "EXTERNAL_UNTRUSTED_CONTENT 或等价 web 内容隔离协议。"
    )

    doc.add_heading("4. 人机审核（可选）", level=2)
    bullet(
        "is_human=True 将指定 Role 换为终端人工输入，可作质量闸门，非自动安全过滤。"
    )
    cite(
        "human_engagement",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/human_engagement.html",
    )

    # ── 8. Architecture ──
    doc.add_heading("架构总览", level=1)
    bullet(
        "定位：多智能体框架；核心哲学 Code = SOP(Team)；"
        "一行需求 → 软件公司全流程产出（PRD/设计/代码/文档）。"
    )
    bullet(
        "单 Agent：Role.run → observe → react（think + act）→ 返回 Message。"
    )
    bullet(
        "多 Agent：Team + Environment 消息总线；SOP 体现为 Action 链与 _watch 订阅图。"
    )
    bullet(
        "配置：~/.metagpt/config2.yaml（metagpt --init-config）；"
        "Python 3.9–3.11（<3.12）。"
    )
    bullet(
        "MGX（mgx.dev）：README 商业 AI agent 开发团队产品，与开源 CLI/库并列介绍。"
    )
    cite(
        "introduction；README",
        "https://docs.deepwisdom.ai/main/en/guide/get_started/introduction.html",
    )

    doc.add_heading(
        "与 OpenClaw / OpenHands / LangGraph / CrewAI / Dify / Qwen-Agent 的对照（研究映射）",
        level=2,
    )
    doc.add_paragraph("非 MetaGPT 官方术语：")
    bullet(
        "Role + Action ≈ CrewAI Agent + Task / LangGraph 节点函数；"
        "Environment.publish_message ≈ LangGraph 状态广播 / CrewAI 消息传递。"
    )
    bullet(
        "react_mode=by_order ≈ 固定边工作流；by_llm ≈ ReAct 式 Action 选择。"
    )
    bullet(
        "rc.memory（Message 列表）≈ 滑动窗口会话历史；"
        "SimpleEngine RAG ≈ Dify Knowledge Retrieval / Qwen-Agent Memory 类。"
    )
    bullet(
        "Team.invest + n_round ≈ LLM 预算 + 固定迭代上限；"
        "≠ Dify Schedule Trigger / LangGraph Cron。"
    )
    bullet(
        "Researcher 搜索浏览 ≈ Qwen-Agent web 工具 / OpenHands browse；"
        "无 OpenClaw 统一 web_fetch + 内容标记。"
    )
    bullet(
        "DataInterpreter @register_tool ≈ Smolagents Tool / Qwen-Agent @register_tool；"
        "≠ Software Company 多 Role SOP。"
    )

    # ── 9. Research ──
    doc.add_heading("研究建议", level=1)
    bullet("采集：logs/*.txt + Message 链（cause_by/instruct_content）+ cost_manager 费用。")
    bullet("对比 Software Company SOP vs 单 Researcher vs DataInterpreter 同任务 LLM 调用次数。")
    bullet("IPI：SearchEngine 摘要、浏览器正文、RAG chunk、Environment 广播 Message。")
    bullet("会话：Team.serialize 断点 vs 单进程 Context；区分 rc.memory 与 SimpleEngine。")
    bullet("区分 MGX 商业产品与 geekan/MetaGPT 开源框架文档。")
    bullet("人机：is_human Role 暂停点与消息格式约束。")

    # ── 10. Five-round audit ──
    doc.add_heading("自查：潜在幻觉与核实结论（第五轮）", level=1)
    doc.add_paragraph(
        "2026-07-12 第五轮复核：对照 introduction、concepts、agent_101、"
        "multi_agent_101、use_memories、rag_module、agent_communication、"
        "human_engagement、researcher、create_and_use_tools、configuration、"
        "README、team.py、logs.py、config2.example.yaml 全部关键断言。"
        "结论：第四轮修正项仍成立；本轮未发现新的机制性幻觉。"
    )
    checks = [
        ("✓ Code = SOP(Team) 与软件公司多角色哲学", "https://docs.deepwisdom.ai/main/en/guide/get_started/introduction.html"),
        ("✓ Agent = LLM + Observation + Thought + Action + Memory", "https://docs.deepwisdom.ai/main/en/guide/tutorials/concepts.html"),
        ("✓ Role + Action + react_mode（by_order 等）", "https://docs.deepwisdom.ai/main/en/guide/tutorials/agent_101.html"),
        ("✓ Team.hire / invest / run(n_round) + cost_manager 预算", "https://github.com/geekan/MetaGPT/blob/main/metagpt/team.py"),
        ("✓ Environment publish_message + _watch(cause_by) 订阅", "https://docs.deepwisdom.ai/main/en/guide/in_depth_guides/agent_communication.html"),
        ("✓ rc.memory = Message 短期列表；get_memories(k) 按 recency", "https://docs.deepwisdom.ai/main/en/guide/tutorials/use_memories.html"),
        ("✓ RAG SimpleEngine：Faiss/BM25/Chroma/ES + persist/load", "https://docs.deepwisdom.ai/main/en/guide/in_depth_guides/rag_module.html"),
        ("✓ Researcher：CollectLinks + WebBrowseAndSummarize + ConductResearch", "https://docs.deepwisdom.ai/main/en/guide/use_cases/agent/researcher.html"),
        ("✓ SearchEngine serpapi/google/serper/ddg；WebBrowserEngine playwright/selenium", "https://docs.deepwisdom.ai/main/en/guide/use_cases/agent/researcher.html"),
        ("✓ DataInterpreter + @register_tool 工具注册", "https://docs.deepwisdom.ai/main/en/guide/tutorials/create_and_use_tools.html"),
        ("✓ logs.py：loguru 文件 + log_llm_stream + log_tool_output", "https://github.com/geekan/MetaGPT/blob/main/metagpt/logs.py"),
        ("✓ Team.serialize/deserialize 断点恢复", "https://github.com/geekan/MetaGPT/blob/main/metagpt/team.py"),
        ("✓ is_human=True 人机 Role（终端 input）", "https://docs.deepwisdom.ai/main/en/guide/tutorials/human_engagement.html"),
        ("✓ config2.yaml + Python 3.9–3.11", "https://github.com/geekan/MetaGPT/blob/main/README.md"),
        ("△ MGX（mgx.dev）为商业产品 — 与开源框架本体分离", "https://github.com/geekan/MetaGPT/blob/main/README.md"),
        ("△ Software Company 图示标注 Gradually Implementing — 非全部 Role 已完备", "https://github.com/geekan/MetaGPT/blob/main/README.md"),
        ("△ serialization 文档页 fetch 常为落地页 — 机制以 team.py 为准", "https://github.com/geekan/MetaGPT/blob/main/metagpt/team.py"),
        ("△ role_zero 长期记忆默认 false — 与 use_memories 短期记忆不同路径", "https://github.com/geekan/MetaGPT/blob/main/config/config2.example.yaml"),
        ("△ Team 默认 use_mgx=True 可选 MGXEnv — 教程多用标准 Environment", "https://github.com/geekan/MetaGPT/blob/main/metagpt/team.py"),
        ("✗ EXTERNAL_UNTRUSTED_CONTENT / 统一 web 内容消毒 — 所查文档均无", "https://docs.deepwisdom.ai/main/en/guide/use_cases/agent/researcher.html"),
        ("✗ OpenClaw 式 Heartbeat / 定时唤醒 — 框架无", "https://github.com/geekan/MetaGPT/blob/main/metagpt/team.py"),
        ("✗ 跨会话语义 MEMORY.md 自动注入 — 无；靠 rc.memory / RAG / role_zero 可选", "https://docs.deepwisdom.ai/main/en/guide/tutorials/use_memories.html"),
        ("✗ tools.html 独立正文 — fetch 仅为落地页，工具细节在 create_and_use_tools + researcher", "https://docs.deepwisdom.ai/main/en/guide/get_started/tools.html"),
        ("— 第五轮：上述条目全部复检通过，无新增 ✗/△", "https://docs.deepwisdom.ai/main/en/guide/get_started/introduction.html"),
    ]
    for text, url in checks:
        cite(text, url)

    doc.add_heading("防幻觉机制说明（五轮复核流程）", level=2)
    bullet("第一轮：确认 introduction 营销叙事 vs concepts/agent_101 技术分层；区分 Software Company / DI / RAG。")
    bullet("第二轮：核对 Memory 三层（rc.memory 短期 / SimpleEngine RAG / role_zero 长期可选）不可混用。")
    bullet("第三轮：核对 Web 路径（Researcher SearchEngine+Browser / DI 工具）及 config2 search/browser 配置。")
    bullet("第四轮：核对 Team 预算与序列化；MGX 商业产品与开源分离；serialization 页落地页问题。")
    bullet("第五轮：全量重读 team.py、logs.py、researcher、agent_communication，确认无 OpenClaw 式机制误标。")

    doc.add_heading("主要参考链接", level=1)
    for link in [
        "https://docs.deepwisdom.ai/main/en/guide/get_started/introduction.html",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/concepts.html",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/agent_101.html",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/multi_agent_101.html",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/use_memories.html",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/create_and_use_tools.html",
        "https://docs.deepwisdom.ai/main/en/guide/tutorials/human_engagement.html",
        "https://docs.deepwisdom.ai/main/en/guide/in_depth_guides/rag_module.html",
        "https://docs.deepwisdom.ai/main/en/guide/in_depth_guides/agent_communication.html",
        "https://docs.deepwisdom.ai/main/en/guide/use_cases/agent/researcher.html",
        "https://docs.deepwisdom.ai/main/en/guide/get_started/configuration.html",
        "https://github.com/geekan/MetaGPT",
        "https://github.com/geekan/MetaGPT/blob/main/README.md",
        "https://github.com/geekan/MetaGPT/blob/main/metagpt/team.py",
        "https://github.com/geekan/MetaGPT/blob/main/metagpt/logs.py",
        "https://github.com/geekan/MetaGPT/blob/main/config/config2.example.yaml",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
