# -*- coding: utf-8 -*-
"""Generate Dify research notes (English).

Same structure as Chinese _gen_dify_doc.py: Session, task modes,
web_fetch/security notice, tools, Prompt, memory/persistence,
five-round anti-hallucination vs docs.dify.ai.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "en" / "Dify.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label} (Source: {url})")

    # ── Title ──
    h("Dify", 0)
    p(
        "This document follows the same structure as OpenHands.docx / Hermes.docx / "
        "Claude-Code.docx: Session, task modes, web_fetch (and Security Notice comparison), "
        "Tools; plus Prompt composition, memory and persistence mechanisms, and a "
        "five-round anti-hallucination review at the end. "
        "Subject: the open-source AI application platform Dify "
        "(Studio orchestration of Workflow / Chatflow / Agent apps, etc.)."
    )
    p("Official documentation index: https://docs.dify.ai/llms.txt")
    p(
        "Positioning: drag-and-drop agentic flows in Studio, published as Web App, API, "
        "or MCP Server. Under the hood, a unified Workflow engine; Chatflow / Agent are "
        "product forms built on top."
    )
    cite(
        "Key Concepts",
        "https://docs.dify.ai/en/use-dify/getting-started/key-concepts",
    )
    p(
        "Execution chain (research view, Agent/Chatflow): "
        "User → chat-messages (conversation) → LLM/Agent node → "
        "Tool / HTTP Request / Knowledge Retrieval → result re-injected into prompt → User."
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "Dify's session entity is Conversation, identified by the system variable "
        "sys.conversation_id. The same conversation aggregates multi-turn messages so the "
        "LLM continues on the same topic and context. Each Chatflow turn triggers one "
        "workflow run."
    )
    cite(
        "Key Concepts > Chatflow variables",
        "https://docs.dify.ai/en/use-dify/getting-started/key-concepts",
    )

    h("1. conversation_id and API continuation", 2)
    bullet(
        "Chat API: POST /v1/chat-messages; leave conversation_id empty for a new session, "
        "response returns a new ID; subsequent requests carry it to continue."
    )
    bullet(
        "When conversation_id is already set, new inputs are ignored; only query is "
        "processed (use conversation variables for dynamic state)."
    )
    bullet(
        "Service API and WebApp sessions are not shared: same user has separate history "
        "across the two channels."
    )
    bullet(
        "sys.dialogue_count: dialogue turn counter; can be used with If-Else branching."
    )
    cite(
        "Developing with APIs / Chat messages",
        "https://docs.dify.ai/en/self-host/use-dify/publish/developing-with-apis",
    )

    h("2. Workflow run vs Conversation", 2)
    bullet(
        "Workflow app: single-turn task; sys.workflow_run_id tracks each execution; "
        "no Chatflow-style multi-turn conversation (unless managed externally)."
    )
    bullet(
        "Chatflow: each turn = one Chatflow run + history/variables under the same "
        "conversation_id."
    )
    bullet("Logs/Traces: Monitor > Logs shows trigger source and run records.")

    h("3. Research mapping vs OpenClaw / Hermes / Claude Code (non-official terms)", 2)
    bullet(
        "OpenClaw Main session / Claude Code interactive session ≈ "
        "Dify Chatflow/Chatbot same conversation_id."
    )
    bullet(
        "OpenClaw Isolated / Automation new session ≈ "
        "Workflow single run, or each Schedule/Webhook Trigger firing an independent run."
    )
    bullet(
        "Cross-session knowledge: Knowledge Base / environment variables; "
        "not a MEMORY.md file system (see Memory section)."
    )

    # ══════════════════════════════════════════════════════════════
    # Task modes
    # ══════════════════════════════════════════════════════════════
    h("Task Modes", 1)
    p(
        "Application type and launch method determine whether multi-turn, persistent "
        "conversation, and unattended operation apply."
    )

    h("1. Chatflow (recommended for multi-turn)", 2)
    bullet(
        "Each user input triggers a full Chatflow; supports Memory, Conversation "
        "Variables, streaming Answer output."
    )
    bullet("Entry is fixed as User Input (not Trigger).")
    cite(
        "Workflow & Chatflow",
        "https://docs.dify.ai/en/cloud/use-dify/build/workflow-chatflow.md",
    )

    h("2. Workflow (single-turn / batch / triggers)", 2)
    bullet("Suited for one-shot tasks and batch; WebApp/API easy for batch runs.")
    bullet("Start: User Input (human/API trigger) or Trigger (automatic).")

    h("3. Trigger (Workflow only)", 2)
    bullet(
        "Schedule Trigger: cron / visual schedule; at most one schedule per workflow."
    )
    bullet("Integration Trigger: third-party event subscription (e.g. Slack).")
    bullet(
        "Webhook Trigger: external HTTP callback; variables extracted from "
        "query/header/body."
    )
    bullet(
        "Multiple Triggers allowed; Quick Settings can enable/disable published Triggers."
    )
    cite(
        "Trigger Overview; Schedule Trigger",
        "https://docs.dify.ai/en/use-dify/nodes/trigger/overview",
    )

    h("4. Basic application forms (legacy simplified UI)", 2)
    bullet("Chatbot: simple dialogue with model + prompt.")
    bullet("Agent (including New Agent): chat app with autonomous tool calling.")
    bullet("Text Generator: single-turn completion.")
    p(
        "Official note: underlying engine is the same workflow engine; UI is a simplified "
        "legacy form."
    )

    h("5. Publish and invocation surfaces", 2)
    bullet("Web App / Embed.")
    bullet("Service API (Bearer Key; Cloud base URL api.dify.ai/v1).")
    bullet("MCP Server: expose Dify App to Claude Desktop / Cursor, etc.")
    bullet("difyctl: terminal/CI/coding agent invocation.")

    h("6. Scheduling strategy comparison (research notes)", 2)
    p("OpenClaw: precise isolation → Cron; full session → Heartbeat.")
    p("Dify equivalents:")
    bullet(
        "Scheduled unattended → Workflow Schedule Trigger (new run, not Chatflow "
        "continuation)."
    )
    bullet("Event-driven → Webhook / Integration Trigger.")
    bullet("Dialogue memory needed → Chatflow + conversation_id (not Trigger).")

    # ══════════════════════════════════════════════════════════════
    # web_fetch / Security Notice
    # ══════════════════════════════════════════════════════════════
    h("web_fetch (and Security Notice)", 1)
    p(
        "[Conclusion] Dify has no built-in Agent tool named web_fetch, "
        "nor OpenClaw-style framework-level SECURITY NOTICE / EXTERNAL_UNTRUSTED_CONTENT "
        "wrapping. External web/API content mainly enters via HTTP Request node, "
        "Tool/Agent tools, and Knowledge \"Sync from website\" (Jina/Firecrawl)."
    )

    h("1. Path A: HTTP Request node (workflow-level fetch)", 2)
    bullet(
        "Supports GET/HEAD/POST/PUT/PATCH/DELETE; URL/Headers/Body/Auth can insert "
        "{{variable}}."
    )
    bullet(
        "Output split into Response Body, Status Code, Headers, Files, Size — "
        "downstream references structured variables; no automatic SECURITY NOTICE."
    )
    bullet(
        "Configurable timeout, retry (up to ~10), failure branch, ssl_verify."
    )
    cite(
        "HTTP Request Node",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/http-request",
    )

    h("2. Path B: Tool / Agent tool invocation", 2)
    bullet(
        "Agent node or Tool node calls Workspace tool plugins, Swagger, MCP, "
        "Workflow-as-Tool. E.g. search/crawl plugins (depends on install and Marketplace)."
    )
    bullet(
        "Tool return enters subsequent LLM as Observation / node output; "
        "official docs do not mandate automatic untrusted boundary text."
    )
    cite(
        "Agent Node; Tool Node; Dify Tools",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/agent.md",
    )

    h("3. Path C: Knowledge website import (delayed / index surface)", 2)
    p(
        "\"Sync from website\" uses Firecrawl or Jina Reader to parse public pages into "
        "Markdown then ingest into KB. This is knowledge-base build-time ingest, not "
        "same-turn Agent real-time web_fetch; injection surface is later Knowledge "
        "Retrieval → LLM Context."
    )
    cite(
        "Import Data from Website",
        "https://docs.dify.ai/en/cloud/use-dify/knowledge/create-knowledge/import-text-data/sync-from-website.md",
    )

    h("4. Security Notice comparison and SSRF (infrastructure)", 2)
    bullet("OpenClaw: web_fetch results wrapped with SECURITY NOTICE.")
    bullet("Hermes: <untrusted_tool_result>.")
    bullet(
        "Claude Code: WebFetch isolated context + lossy extraction + 15min cache."
    )
    bullet(
        "Dify: product node docs do not describe equivalent \"content marking\"; "
        "self-hosted default has ssrf_proxy (Squid) blocking/filtering sandbox service "
        "outbound requests for SSRF prevention; domains not in allowed_domains are blocked."
    )
    cite(
        "SSRF Proxy (Docker Issues / External KB)",
        "https://docs.dify.ai/en/self-host/deploy/troubleshooting/docker-issues",
    )
    p(
        "IPI research implication: assume HTTP/Tool response body can be treated as "
        "instructions by the model; add separation/sanitization at orchestration layer or "
        "rely on SSRF ACL — do not assume OpenClaw-style automatic wrapping. "
        "Knowledge-base web path is \"index-then-retrieve poisoning\" threat, distinct "
        "from real-time tool output."
    )

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "Tools are managed under Workspace Integrations > Tools, usable in "
        "Workflow/Chatflow Tool nodes, Agent nodes, and Agent-class apps."
    )
    cite(
        "Dify Tools",
        "https://docs.dify.ai/en/cloud/use-dify/workspace/tools.md",
    )

    h("1. Tool types", 2)
    bullet(
        "Tool Plugin: Marketplace / built-in (e.g. Current Time); some require authorization."
    )
    bullet("Swagger/OpenAPI: paste or URL import to generate tool UI.")
    bullet(
        "Workflow as Tool: Workflow starting from User Input only (Chatflow cannot be Tool)."
    )
    bullet(
        "MCP: HTTP transport only; Dynamic Client Registration / custom Header/timeout."
    )

    h("2. Key nodes in orchestration", 2)
    bullet(
        "Agent: Function Calling or ReAct; configure tools, Instructions, Max Iterations, "
        "Memory."
    )
    bullet("Tool: fixed call to a tool action (not autonomous tool selection).")
    bullet("HTTP Request: generic outbound HTTP.")
    bullet("Knowledge Retrieval: retrieve KB → context for LLM.")
    bullet("Code: Python/JS in sandbox (related to ssrf/sandbox components).")
    bullet("Variable Assigner: update Conversation Variables.")

    h("3. Agent strategies", 2)
    bullet("Function Calling: uses model native tools parameters.")
    bullet("ReAct: Thought → Action → Observation prompt loop.")
    bullet("More Agent Strategies installable from Marketplace.")

    h("4. Research simplified messages sketch (Chatflow + Agent)", 2)
    code(
        """{
  "messages": [
    {
      "role": "system",
      "content": "<Agent Instructions / LLM System prompt with {{vars}}>"
    },
    {
      "role": "user",
      "content": "Check https://api.example.com/health and return the status."
    },
    {
      "role": "assistant",
      "content": null,
      "tool_calls": [{
        "id": "call_001",
        "function": {
          "name": "<configured_http_or_plugin_tool>",
          "arguments": "{\\"url\\": \\"https://api.example.com/health\\"}"
        }
      }]
    },
    {
      "role": "tool",
      "tool_call_id": "call_001",
      "content": "{\\"status\\": 200, \\"body\\": \\"{\\\\\\"ok\\\\\\": true}\\"}"
    },
    {
      "role": "assistant",
      "content": "The API is healthy (status 200, ok=true)."
    }
  ]
}"""
    )
    p(
        "Note: actual path uses Dify workflow events and chat-messages API; "
        "tool names come from installed plugins/Swagger, not a fixed web_fetch."
    )

    # ══════════════════════════════════════════════════════════════
    # Five-round anti-hallucination
    # ══════════════════════════════════════════════════════════════
    h("Five-Round Anti-Hallucination Review", 1)

    h("Round 1 — Session", 2)
    bullet("Checked: key-concepts; chat-messages / developing-with-apis.")
    bullet("Kept: conversation_id; API/WebApp isolation; dialogue_count.")
    bullet(
        "Removed: writing Claude Code JSONL paths or Hermes SessionDB as Dify storage format."
    )

    h("Round 2 — Task modes", 2)
    bullet(
        "Checked: trigger/overview; schedule-trigger; workflow-chatflow; build/agent."
    )
    bullet(
        "Kept: Chatflow vs Workflow; Trigger Workflow-only; at most one Schedule per app."
    )
    bullet(
        "Removed: claiming Chatflow can use Schedule Trigger; "
        "removed Dify Heartbeat API."
    )

    h("Round 3 — web_fetch / security", 2)
    bullet(
        "Checked: nodes/http-request; sync-from-website; ssrf_proxy Docker docs."
    )
    bullet(
        "Kept: no built-in web_fetch name; HTTP Request structured output; "
        "Jina/Firecrawl for KB ingest; SSRF proxy."
    )
    bullet(
        "Removed: OpenClaw SECURITY NOTICE assumed by default; "
        "removed \"HTTP response must have 15min cache\" (Claude Code/OpenClaw feature, "
        "not Dify HTTP node doc item); "
        "removed treating website import as real-time Agent web_fetch."
    )

    h("Round 4 — Tools", 2)
    bullet("Checked: workspace/tools; nodes/tools; nodes/agent.")
    bullet("Kept: Plugin / Swagger / Workflow-as-Tool / MCP; Agent FC/ReAct.")
    bullet(
        "Removed: treating Claude Code `WebFetch` PascalCase tool as Dify built-in name."
    )

    h("Round 5 — Prompt / Memory cross-check", 2)
    bullet(
        "Checked: nodes/llm; key-concepts Conversation Variables; knowledge integrate."
    )
    bullet(
        "Frozen: "
        "(1) Session ≡ conversation_id; "
        "(2) Task modes = Chatflow / Workflow(+Triggers) / Agent·Chatbot; "
        "(3) External web = HTTP Request + Tools + KB crawl; no framework SECURITY NOTICE; "
        "(4) Memory = TokenBuffer Memory + Conversation Vars + Knowledge."
    )
    bullet(
        "Open questions: Cloud vs self-hosted SSRF policy differences; "
        "specific Marketplace \"web crawl\" plugin return format; "
        "need live capture of LLM messages to verify Memory window injection template."
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt composition
    # ══════════════════════════════════════════════════════════════
    h("Prompt Composition", 1)
    p(
        "LLM node uses Chat roles (System / User / Assistant) or Completion text by model "
        "type. Prompts reference workflow variables via {{variable}}; Jinja2 optional for "
        "loops/conditions."
    )
    cite(
        "LLM Node",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/llm",
    )

    h("1. System / user messages and variables", 2)
    bullet("System: role and constraints; User: {{user_input}}, etc.")
    bullet(
        "Variables substituted before model delivery; deep paths like "
        "{{api_response.data.items[0].id}} supported."
    )
    bullet(
        "Agent Instructions: natural language + Jinja2; Query can come from upstream node."
    )

    h("2. Context / RAG", 2)
    bullet(
        "Knowledge Retrieval output connects to LLM context; "
        "prompt references e.g. {{knowledge_retrieval.result}}; Dify can track citations."
    )

    h("3. Memory injection (Chatflow)", 2)
    bullet(
        "Memory enabled: prior turns merged into subsequent prompt as formatted "
        "user–assistant pairs; USER template editable; TokenBufferMemory; "
        "node-level, not across different conversations."
    )
    bullet(
        "Agent node also has Memory window (message count / cost tradeoff)."
    )

    h("4. Capturing research prompts", 2)
    bullet("Monitor Logs / workflow run details for node IO.")
    bullet(
        "Service API streaming events for text chunks; no documented \"full LLM mock\" path."
    )
    bullet(
        "Self-hosted can add proxy between API and model provider for packet capture."
    )

    h("Prompt Composition · Review summary", 2)
    bullet("Kept: roles + {{var}} + Jinja2 + Knowledge context + Memory toggle.")
    bullet(
        "Removed: Hermes SOUL.md / Claude Code CLAUDE.md as Dify default loaded files."
    )

    # ══════════════════════════════════════════════════════════════
    # Memory and persistence
    # ══════════════════════════════════════════════════════════════
    h("Memory and Persistence", 1)
    p(
        "Dify has no OpenClaw/Hermes-style dedicated MEMORY.md file, "
        "nor Claude Code auto memory directory. "
        "Cross-turn and cross-session capabilities are split across: dialogue Memory, "
        "Conversation Variables, Knowledge, Logs."
    )

    h("1. TokenBuffer Memory (node-level dialogue window)", 2)
    bullet(
        "Optional on LLM / Agent nodes; buffers recent user–assistant pairs."
    )
    bullet(
        "Larger window = fuller context, higher token cost; not written to other conversations."
    )

    h("2. Conversation Variables (Chatflow)", 2)
    bullet(
        "Session-level mutable state (todo list, accumulated cost, etc.); "
        "Variable Assigner updates."
    )
    bullet("API may have list/update endpoints for conversation variables.")
    cite(
        "Key Concepts > Conversation Variables",
        "https://docs.dify.ai/en/use-dify/getting-started/key-concepts",
    )

    h("3. Knowledge Bases (persistent external memory / RAG)", 2)
    bullet(
        "Documents/Notion/website crawl ingested; retrieval node injects into LLM."
    )
    bullet(
        "External knowledge API + SSRF whitelist (self-hosted) can connect custom retrieval."
    )
    bullet(
        "Threat surface: poison at index time → multi-session Retrieve "
        "(similar to RAG IPI papers, not real-time web_fetch)."
    )

    h("4. Environment Variables", 2)
    bullet(
        "App-level secrets/constants; can separate from secrets when sharing DSL; "
        "not casually changed at runtime."
    )

    h("5. Comparison with other agents", 2)
    bullet(
        "Hermes MEMORY.md frozen snapshot ≈ Dify Conversation Vars + Memory window "
        "(different mechanism: no md file snapshot layer)."
    )
    bullet(
        "Claude Code CLAUDE.md ≈ partially analogous to \"app prompt / DSL fixed copy\", "
        "but not a repo-root auto-load file system."
    )

    h("Memory and Persistence · Review summary", 2)
    bullet(
        "Kept: three main memory surfaces (buffer / conv vars / KB); "
        "API continuation via conversation_id."
    )
    bullet(
        "Removed: default MEMORY.md exists; "
        "removed Trigger runs automatically inheriting a user's full Chatflow memory "
        "(Trigger starts new Workflow run, no Chatflow conversation)."
    )

    h("Primary Reference Links", 1)
    for link in [
        "https://docs.dify.ai/llms.txt",
        "https://docs.dify.ai/en/use-dify/getting-started/key-concepts",
        "https://docs.dify.ai/en/cloud/use-dify/build/workflow-chatflow.md",
        "https://docs.dify.ai/en/use-dify/nodes/trigger/overview",
        "https://docs.dify.ai/en/use-dify/nodes/trigger/schedule-trigger",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/http-request",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/agent.md",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/llm",
        "https://docs.dify.ai/en/cloud/use-dify/nodes/tools.md",
        "https://docs.dify.ai/en/cloud/use-dify/workspace/tools.md",
        "https://docs.dify.ai/en/cloud/use-dify/knowledge/create-knowledge/import-text-data/sync-from-website.md",
        "https://docs.dify.ai/en/self-host/use-dify/publish/developing-with-apis",
        "https://docs.dify.ai/en/self-host/deploy/troubleshooting/docker-issues",
    ]:
        bullet(link)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
