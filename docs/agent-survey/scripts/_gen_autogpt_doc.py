# -*- coding: utf-8 -*-
"""Generate AutoGPT research notes (same structure as OpenHands / Hermes / Dify).

Paper track: Session / 任务模式 / web_fetch(security notice) / tools,
plus Prompt 组成 & 记忆持久化. 5-round anti-hallucination vs agpt.co/docs.
Primary object: AutoGPT Platform; Classic noted as legacy only.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "AutoGPT.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label}（来源：{url}）")

    # ── Title ──
    h("AutoGPT", 0)
    p(
        "本文档结构对齐 OpenHands.docx / Hermes.docx / Claude-Code.docx / Dify.docx："
        "Session、任务模式、web_fetch（及 Security Notice 对照）、Tools；"
        "文末附 Prompt 组成、记忆与持久化机制，以及五轮防幻觉核对。"
        "主对象为 AutoGPT Platform（Agent Builder + AutoPilot + Blocks）；"
        "AutoGPT Classic 仅作遗留对照（官方声明不再从安全角度维护）。"
    )
    p("官方文档索引：https://agpt.co/docs/llms.txt")
    p(
        "定位：低代码搭建 continuous agents；"
        "Frontend（Builder / Library / AutoPilot）+ Server（执行与 Marketplace）。"
        "Agent ≡ 由 Blocks 连成的自动化工作流；执行顺序完全由 data flow 决定。"
    )
    cite(
        "What is the AutoGPT Platform?",
        "https://agpt.co/docs/platform/what-is-autogpt-platform",
    )
    p(
        "执行链路（研究视角 · Platform Agent Task）："
        "Trigger/New Task → Input blocks yield → Action blocks"
        "（LLM / Send Web Request / Firecrawl / Stagehand / …）→ "
        "Output blocks → Task 结果。"
    )
    p(
        "执行链路（研究视角 · AutoPilot）："
        "User chat → AutoPilot（可直接跑 ~400 blocks / 建改 agent / MCP / "
        "one-shot 取页 / 多步 browser session）→ 回注入对话。"
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "官方文档对「Session」有两套产品语义，需分开谈："
        "(A) AutoPilot 的 chat / conversation；(B) Agent Library 的 Task 执行实例。"
        "公开文档未给出类似 Dify conversation_id 的稳定对外 ID 字段说明。"
    )

    h("1. AutoPilot chat session / conversation", 2)
    bullet(
        "入口：导航 Home → platform.agpt.co；自然语言完成跑 agent、搭 workflow、调研等。"
    )
    bullet(
        "文档与 changelog 称 chat sessions / conversations；"
        "可删除会话、导出 Markdown、带时间戳；"
        "断线重连会 replay 未看到的内容，历史不被清空。"
    )
    bullet(
        "长对话可触发 summarization（UI：「Summarizing earlier messages…」）。"
    )
    bullet(
        "会话内 browser/tool：多步浏览时，browser session 在同一次 conversation 内跨步骤保持。"
    )
    bullet(
        "changelog 提及 session-level dry-run 等标志；"
        "研究侧勿捏造「固定 JSONL 路径」或 Hermes SessionDB。"
    )
    cite("AutoPilot", "https://agpt.co/docs/platform/using-the-platform/autopilot")
    cite(
        "Changelog · browse + summarize",
        "https://agpt.co/docs/platform/changelog/changelog/february-26-march-4-2026",
    )

    h("2. Agent Task（工作流一次运行）", 2)
    bullet(
        "Library 中每次执行记为一条 Task：输入、输出、成本；"
        "可用任务 URL 分享结果视图。"
    )
    bullet(
        "手工：New Task → 填 Input → Start Task；"
        "Schedule：同一表单走 Schedule Task；"
        "Trigger 型 agent：New Task 变为 New Trigger，仅能由 webhook 启动。"
    )
    cite(
        "Agent Library",
        "https://agpt.co/docs/platform/using-the-platform/agent-library",
    )
    cite(
        "Scheduling & Triggers",
        "https://agpt.co/docs/platform/using-the-platform/scheduling-and-triggers",
    )

    h("3. Data-flow 执行模型（非会话协议）", 2)
    bullet("无独立 control-flow：块在「全部已连接且必填输入就绪」时运行。")
    bullet("Input 启动 → 沿 pin 传递 → Output 收集结果；无依赖的块可并行任意顺序。")
    cite(
        "Data Flow & Execution",
        "https://agpt.co/docs/platform/using-the-platform/data-flow-and-execution",
    )

    h("4. 与 OpenClaw / Hermes / Dify 的研究映射（非官方术语）", 2)
    bullet(
        "OpenClaw Main / Claude Code 交互会话 ≈ AutoPilot 同一 conversation"
        "（含跨步 browser session）。"
    )
    bullet(
        "Isolated / Automation 单次运行 ≈ Library 上的一次 Task "
        "（手动 / Schedule / Webhook）。"
    )
    bullet(
        "Dify conversation_id 级 API 续聊 ≈ AutoPilot 聊天历史 "
        "+（可选）Mem0 记忆块跨 Task；勿写成 Dify 同款 API 字段。"
    )

    # ══════════════════════════════════════════════════════════════
    # 任务模式
    # ══════════════════════════════════════════════════════════════
    h("任务模式", 1)

    h("1. AutoPilot（对话式平台助手）", 2)
    bullet("跑 Library agent、生成/编辑 agent、搜 Marketplace。")
    bullet(
        "直接执行单个 blocks（~400+）：研究、图像/视频、HTTP、跑代码"
        "（含委派 Claude Code）等，无需先搭完整图。"
    )
    bullet("管理日程、查看任务结果；changelog：Native scheduling from chat、MCP 自动发现等。")
    cite("AutoPilot", "https://agpt.co/docs/platform/using-the-platform/autopilot")

    h("2. Agent Builder（可视化工作流）", 2)
    bullet("Canvas 上拖拽 Input / Action / Output（及 Trigger）Blocks 并连线。")
    bullet("Input/Output 定义对外 schema；内部块运行时不对用户暴露。")
    bullet("保存即更新 Library（无单独 draft 状态文档说明）。")
    cite(
        "Agent Builder Guide",
        "https://agpt.co/docs/platform/using-the-platform/agent-builder-guide",
    )

    h("3. 手动 Task / Schedule / Webhook Trigger", 2)
    bullet("On-demand：New Task + Start Task。")
    bullet("Schedule：预填输入 + 频率/星期/时刻/时区；Scheduled 选项卡管理。")
    bullet(
        "Trigger：Builder 放入 trigger block → Library New Trigger → "
        "Webhook URL（如 backend.agpt.co/.../generic_webhook/...）；"
        "此类 agent 不可手动 New Task。"
    )

    h("4. Marketplace / 导入导出 / API / 自托管", 2)
    bullet("Marketplace 添加或分享 agent；Upload / Export 文件交换。")
    bullet("API：账号 API key；另有 OAuth/SSO 文档。")
    bullet("Cloud 与 Self-Host 双路径；platform 目录 Polyform Shield，仓库其余多为 MIT。")
    cite("API Introduction", "https://agpt.co/docs/platform/api-and-integrations/api-guide")

    h("5. AutoGPT Classic（遗留，非本文主对象）", 2)
    bullet(
        "LLM 循环决策 + 动作结果回灌 prompt 的 generalist agent；"
        "用户需授权动作。"
    )
    bullet(
        "官方 Maintenance Notice：不再从安全角度支持；"
        "依赖不更新、问题不修——论文威胁面分析勿默认 Classic 仍获补丁。"
    )
    cite(
        "Classic Introduction",
        "https://agpt.co/docs/classic/autogpt-classic/introduction",
    )

    h("任务模式 · 核对摘要", 2)
    bullet(
        "保留：AutoPilot ↔ Builder 等价能力面；Task/Schedule/Trigger；"
        "data-flow 执行；Classic = 遗留。"
    )
    bullet("剔除：OpenClaw Heartbeat、Dify「仅 Workflow 可 Schedule」套用到 AutoGPT。")

    # ══════════════════════════════════════════════════════════════
    # web_fetch
    # ══════════════════════════════════════════════════════════════
    h("web_fetch（及 Security Notice）", 1)
    p(
        "结论（相对 OpenClaw/Claude Code）："
        "AutoGPT Platform 文档中不存在名为 web_fetch / WebFetch 的内置工具，"
        "也未文档化 OpenClaw 式 SECURITY NOTICE / EXTERNAL_UNTRUSTED_CONTENT 包装。"
        "外部网页进入模型上下文的路径是：AutoPilot 浏览能力 + 各类抓取/HTTP Blocks；"
        "自定义发包块应使用后端 SSRF 防护 Requests 包装。"
    )

    h("1. AutoPilot 原生浏览（changelog 产品能力）", 2)
    bullet("One-shot：单步 fetch + extract 任意页面。")
    bullet(
        "Multi-step：完整 browser session，同 conversation 内跨步骤持久"
        "（登录、菜单导航、抽数等）。"
    )
    bullet(
        "文档未给出工具内部名、是否 lossy 小模型抽取、是否 15min cache——"
        "勿从 Claude Code WebFetch 或 OpenClaw 抄这些细节。"
    )
    cite(
        "Changelog · AutoPilot can browse",
        "https://agpt.co/docs/platform/changelog/changelog/february-26-march-4-2026",
    )

    h("2. 工作流 Blocks：抓取 / 内容 API（近似「读网页」）", 2)
    bullet(
        "Firecrawl Scrape：单 URL；markdown/html/screenshot 等；"
        "max_age 默认文档写 1 hour（页面缓存年龄参数，≠框架 SECURITY NOTICE）。"
    )
    bullet("Exa Contents：urls → 全文/highlights/summary；livecrawl 选项；context 字符串便于喂 LLM。")
    bullet(
        "Extract Website Content（Jina）：拉 HTML 并抽主文；"
        "raw_content 可选 raw 或 Jina Reader。"
    )
    bullet("Stagehand Act / Extract / Observe：Browserbase 上的 AI 浏览器自动化。")
    cite("Firecrawl Scrape", "https://agpt.co/docs/integrations/block-integrations/scrape")
    cite("Exa Contents", "https://agpt.co/docs/integrations/block-integrations/contents")
    cite("Jina Search / Extract", "https://agpt.co/docs/integrations/block-integrations/search-2")
    cite("Stagehand Blocks", "https://agpt.co/docs/integrations/block-integrations/blocks")

    h("3. HTTP / 通用出网", 2)
    bullet(
        "Send Web Request：任意 URL，方法 GET/POST/PUT/DELETE/PATCH；"
        "JSON/form/multipart；区分 client_error / server_error。"
    )
    bullet("Send Authenticated Web Request：按 host 注入凭据。")
    bullet("Read RSS Feed；Execute Code（E2B 沙箱，带 internet）。")
    bullet("AutoPilot 文档亦写明可「Make custom HTTP requests to any API」。")
    cite("Misc · Send Web Request", "https://agpt.co/docs/integrations/block-integrations/misc")

    h("4. SSRF 防护（块开发安全，非 content wrapper）", 2)
    bullet(
        "Build your own Blocks：外部 URL 必须用 "
        "backend.util.request 的 requests/Requests 包装。"
    )
    bullet(
        "特性：校验 URL/协议；解析 DNS；阻断私网（RFC 1918 等）；"
        "默认禁 redirect；可选 trusted_origins；非 200 可 raise。"
    )
    bullet(
        "这是请求层 SSRF 控制，不是把 tool output 包进 SECURITY NOTICE 文本。"
        "第三方 Firecrawl/Exa/Jina 通路是否共用同一包装，公开文档未逐一声明。"
    )
    cite(
        "Security Best Practices for SSRF Prevention",
        "https://agpt.co/docs/platform/building-blocks/new_blocks",
    )

    h("5. Classic 遗留浏览面（对照）", 2)
    bullet("WebSearchComponent、WebSeleniumComponent 等 Forge 内置组件提供搜索/浏览器命令。")
    bullet("因 Classic 安全停止维护，论文主线应放在 Platform。")

    h("威胁模型映射（研究，非官方）", 2)
    bullet(
        "外网文本 → pin 流入 AI Text Generator / AutoPilot 上下文 → "
        "后续块或对话动作：IPI / Goal Hijacking 论文面。"
    )
    bullet(
        "Schedule/Trigger 反复跑带 scrape→LLM 的图：持久化投毒可跨 Task 触发"
        "（若再写入 Mem0/Graphiti 则更像跨会话记忆污染）。"
    )
    bullet(
        "对比：OpenClaw web_fetch+SECURITY NOTICE；Claude Code WebFetch+隔离抽取；"
        "Hermes web_extract+untrusted_tool_result；"
        "Dify HTTP Request+ssrf_proxy——AutoGPT 更接近「Blocks/出网 + SSRF wrapper」，"
        "另加 AutoPilot 会话级 browser。"
    )

    code(
        """示例：Send Web Request → AI Text Generator（研究示意，非官方消息协议）
[Input:url] → [Send Web Request] → response body
                              ↓
                    [AI Text Generator]
                      prompt: summarize {{body}}
                      sys_prompt: ...
                              ↓
                         [Agent Output]
说明：真实执行是 data-flow pins，不是 OpenAI tool_calls 固定 schema；
AutoPilot 内部工具调用名以产品实现为准，勿写成 web_fetch。"""
    )

    h("web_fetch · 核对摘要", 2)
    bullet(
        "保留：无 web_fetch 名；无 SECURITY NOTICE 文档；"
        "AutoPilot browse；Firecrawl/Exa/Jina/Stagehand/Send Web Request；SSRF Requests。"
    )
    bullet(
        "剔除：默认 15min cache + SECURITY NOTICE；"
        "剔除把 Classic Selenium 当 Platform 默认。"
    )

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "Platform「工具」单位是 Block（及 AutoPilot 对 Block/MCP 的直接调用），"
        "不是 Claude Code 式固定 PascalCase 内置表。"
    )

    h("1. Blocks 作为能力原子", 2)
    bullet("三类：Input / Action / Output；另有 Trigger 特殊 Input。")
    bullet("集成覆盖 LLM、搜索、通讯、Notion、GitHub、Airtable、Exa、Firecrawl 等。")
    bullet("Agent Blocks：把完整 agent 图嵌进更大工作流。")
    cite("Agent Blocks Overview", "https://agpt.co/docs/platform/building-blocks/agent-blocks")
    cite("Integrations index", "https://agpt.co/docs/integrations/readme")

    h("2. AutoPilot 工具面", 2)
    bullet("直接跑任意 block；MCP：自然语言连接 Notion/Slack/Jira 等。")
    bullet("changelog：ask_question、技能自蒸馏 registry、MCP auto-discover 等。")

    h("3. 人机协同与其它", 2)
    bullet("Human In The Loop：暂停等待人工批准/可编辑后继续。")
    bullet("Credentials：OAuth / API key / 用户密码；块上 credential bar。")
    bullet("自定义块：Block SDK + SSRF 安全要求。")

    h("Tools · 核对摘要", 2)
    bullet("保留：Blocks ≈ tools；MCP；Agent-as-block；HITL。")
    bullet("剔除：伪造内置工具名 WebFetch / web_fetch。")

    # ══════════════════════════════════════════════════════════════
    # 五轮防幻觉
    # ══════════════════════════════════════════════════════════════
    h("五轮防幻觉核对记录", 1)

    h("Round 1 — Session", 2)
    bullet("核对：autopilot；agent-library；data-flow-and-execution；相关 changelog。")
    bullet(
        "保留：AutoPilot conversation + Library Task；"
        "断线 replay；会话内 browser 持久；无公开 conversation_id 契约。"
    )
    bullet("剔除：Claude Code ~/.claude/projects JSONL；Hermes state.db；Dify sys.conversation_id。")

    h("Round 2 — 任务模式", 2)
    bullet("核对：autopilot vs builder；scheduling-and-triggers；what-is-platform；classic intro。")
    bullet("保留：AutoPilot / Builder / Task / Schedule / Trigger / Marketplace / API；Classic legacy。")
    bullet("剔除：Heartbeat；把 Classic 写成当前推荐生产路径。")

    h("Round 3 — web_fetch / 安全", 2)
    bullet(
        "核对：february-26-march-4-2026 browse；scrape/contents/search-2/blocks/misc；"
        "new_blocks SSRF。"
    )
    bullet(
        "保留：无 web_fetch/SECURITY NOTICE；"
        "one-shot + multi-step browse；Send Web Request；第三方 scrape；SSRF wrapper。"
    )
    bullet(
        "剔除：OpenClaw NOTICE 文本默认存在；"
        "剔除 Claude Code「15min cache + isolated small-model extract」作为 AutoGPT 文档事实。"
    )

    h("Round 4 — Tools", 2)
    bullet("核对：agent-builder-guide；integrations；agent-blocks；MCP changelog。")
    bullet("保留：数百 Blocks；AutoPilot 直跑；MCP；自定义块 SDK。")
    bullet("剔除：固定工具白名单 WebFetch/Bash/Read。")

    h("Round 5 — Prompt / Memory 交叉复检", 2)
    bullet("核对：AI Text Generator；Mem0 Basic blocks；summarize changelog；Graphiti changelog。")
    bullet(
        "冻结："
        "(1) Session = AutoPilot chat + Task run；"
        "(2) 任务模式 = AutoPilot / Builder / Schedule / Trigger；"
        "(3) 外网 = browse + HTTP/scrape blocks，SSRFs 在请求层；"
        "(4) Prompt = 块级 prompt/sys_prompt + pin 注入；"
        "(5) 记忆 = 对话摘要 + Mem0(+Graphiti) + 文件工作区，非 MEMORY.md。"
    )
    bullet(
        "开放问题：AutoPilot browse 的内部 tool 名与是否包装 untrusted；"
        "Cloud 上第三方 scrape 与 SSRF wrapper 的边界；"
        "Mem0 user/run/agent 作用域默认值需实机确认。"
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt 组成
    # ══════════════════════════════════════════════════════════════
    h("Prompt 组成", 1)
    p(
        "Builder 路径：提示词落在 LLM 类 Blocks 的字段上，并由上游 pin 填变量；"
        "AutoPilot 路径：用户自然语言 + 平台侧系统策略（未公开完整 system 模板）。"
    )

    h("1. AI Text Generator（块级）", 2)
    bullet("prompt：主用户提示；可用 Prompt Values 的 {keys} / {{var}} 填充。")
    bullet("sys_prompt：系统侧额外上下文。")
    bullet("model / max_tokens / retry / ollama_host 等。")
    cite("LLM · AI Text Generator", "https://agpt.co/docs/integrations/block-integrations/llm")

    h("2. Data-flow 组装上下文", 2)
    bullet(
        "典型：Scrape/Exa/Send Web Request 的 markdown 或 body → "
        "连到 prompt / prompt_values → LLM → Output。"
    )
    bullet("列表可逐项迭代，对每个 URL 跑一遍抽文+生成。")

    h("3. AutoPilot 对话侧", 2)
    bullet("新会话 themed prompt categories：Learn / Create / Automate / Organize。")
    bullet("可导出 chat 为 Markdown；长上下文触发 summarizing indicator。")
    bullet("官方未公开完整 AutoPilot system prompt 或 tool schema 清单。")

    h("4. Classic（对照）", 2)
    bullet(
        "循环：目标 + ActionHistory 等组件摘要进入 prompt；"
        "命令由各 Component 的 CommandProvider 提供。"
    )

    h("Prompt 组成 · 核对摘要", 2)
    bullet("保留：prompt + sys_prompt + prompt_values；pin 注入外网文本。")
    bullet("剔除：CLAUDE.md / SOUL.md / Dify Jinja2 节点模板 当作 AutoGPT 默认。")

    # ══════════════════════════════════════════════════════════════
    # 记忆与持久化
    # ══════════════════════════════════════════════════════════════
    h("记忆与持久化机制", 1)
    p(
        "无 OpenClaw/Hermes 式仓库根 MEMORY.md 约定；"
        "跨轮/跨跑记忆分散在：AutoPilot 对话与摘要、Mem0 块、文件工作区、"
        "changelog 中的 Graphiti 长期记忆块、以及 Task/Library 元数据。"
    )

    h("1. AutoPilot 对话记忆与摘要", 2)
    bullet("多轮消息保留；过长时 summarize earlier messages（可见 UI）。")
    bullet("附件/生成文件可在对话内引用；changelog 称 persistent file workspace。")

    h("2. Mem0 记忆 Blocks（跨 workflow 执行）", 2)
    bullet("Add Memory：写入 Mem0；可按 user / 可选 run 或 agent 分段；支持 metadata。")
    bullet("Get All Memories / Get Latest Memory / Search Memory（语义检索）。")
    bullet("文档：Memories persist across workflow executions。")
    cite("Basic · Memory blocks", "https://agpt.co/docs/integrations/block-integrations/basic")

    h("3. Graphiti memory（changelog）", 2)
    bullet("v0.6.55：Long-term memory via Graphiti 以 block 形式提供 persistent agent knowledge。")
    bullet("细节以 Marketplace/块文档为准；勿与 Liang GraphRAG 论文环境混为一谈。")

    h("4. 文件与 Store Value", 2)
    bullet("File Store：URL/data URI/本地路径 → 临时目录供后续块使用。")
    bullet("Store Value：工作流内常量转发；非长期外部记忆。")

    h("5. Skills registry（AutoPilot changelog）", 2)
    bullet(
        "Self-distilled skills：复杂多步完成后可沉淀可复用 recipes——"
        "属于 procedural 记忆，不同于 Mem0 事实条目。"
    )

    h("6. Classic 对照", 2)
    bullet("ActionHistoryComponent、ContextComponent 等向 prompt 提供历史与文件上下文。")
    bullet("安全停止维护 ⇒ 论文主威胁面仍应以 Platform 为准。")

    h("记忆与持久化 · 核对摘要", 2)
    bullet(
        "保留：对话+摘要；Mem0；文件工作区；Graphiti block（changelog）；"
        "Task 历史不等于语义 memory。"
    )
    bullet(
        "剔除：默认 MEMORY.md；"
        "剔除把每次 Schedule 自动继承完整 AutoPilot 聊天 transcript"
        "（Schedule 跑的是 agent 图 + 预填输入，除非图内显式读 Mem0）。"
    )

    h("主要参考链接", 1)
    for link in [
        "https://agpt.co/docs/llms.txt",
        "https://agpt.co/docs/platform/what-is-autogpt-platform",
        "https://agpt.co/docs/platform/using-the-platform/autopilot",
        "https://agpt.co/docs/platform/using-the-platform/agent-builder-guide",
        "https://agpt.co/docs/platform/using-the-platform/agent-library",
        "https://agpt.co/docs/platform/using-the-platform/scheduling-and-triggers",
        "https://agpt.co/docs/platform/using-the-platform/data-flow-and-execution",
        "https://agpt.co/docs/platform/building-blocks/new_blocks",
        "https://agpt.co/docs/platform/building-blocks/agent-blocks",
        "https://agpt.co/docs/platform/changelog/changelog/february-26-march-4-2026",
        "https://agpt.co/docs/integrations/block-integrations/scrape",
        "https://agpt.co/docs/integrations/block-integrations/contents",
        "https://agpt.co/docs/integrations/block-integrations/search-2",
        "https://agpt.co/docs/integrations/block-integrations/blocks",
        "https://agpt.co/docs/integrations/block-integrations/misc",
        "https://agpt.co/docs/integrations/block-integrations/basic",
        "https://agpt.co/docs/integrations/block-integrations/llm",
        "https://agpt.co/docs/platform/api-and-integrations/api-guide",
        "https://agpt.co/docs/classic/autogpt-classic/introduction",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
