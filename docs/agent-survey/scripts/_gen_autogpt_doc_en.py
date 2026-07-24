# -*- coding: utf-8 -*-
"""Generate AutoGPT research notes (English).

Same structure as Chinese _gen_autogpt_doc.py: Session, task modes,
web_fetch/security notice, tools, Prompt, memory/persistence,
five-round anti-hallucination vs agpt.co/docs.
Primary object: AutoGPT Platform; Classic noted as legacy only.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "en" / "AutoGPT.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label} (Source: {url})")

    # ── Title ──
    h("AutoGPT", 0)
    p(
        "This document follows the same structure as OpenHands.docx / Hermes.docx / "
        "Claude-Code.docx / Dify.docx: Session, task modes, web_fetch (and Security "
        "Notice comparison), Tools; plus Prompt composition, memory and persistence "
        "mechanisms, and a five-round anti-hallucination review at the end. "
        "Primary subject: AutoGPT Platform (Agent Builder + AutoPilot + Blocks); "
        "AutoGPT Classic is legacy comparison only (officially no longer maintained "
        "from a security perspective)."
    )
    p("Official documentation index: https://agpt.co/docs/llms.txt")
    p(
        "Positioning: low-code continuous agents; "
        "Frontend (Builder / Library / AutoPilot) + Server (execution and Marketplace). "
        "Agent ≡ automation workflow of connected Blocks; execution order fully determined "
        "by data flow."
    )
    cite(
        "What is the AutoGPT Platform?",
        "https://agpt.co/docs/platform/what-is-autogpt-platform",
    )
    p(
        "Execution chain (research view · Platform Agent Task): "
        "Trigger/New Task → Input blocks yield → Action blocks "
        "(LLM / Send Web Request / Firecrawl / Stagehand / …) → "
        "Output blocks → Task result."
    )
    p(
        "Execution chain (research view · AutoPilot): "
        "User chat → AutoPilot (can directly run ~400 blocks / build-edit agent / MCP / "
        "one-shot page fetch / multi-step browser session) → re-injected into dialogue."
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "Official docs use two product semantics for \"Session\" — keep them separate: "
        "(A) AutoPilot chat / conversation; (B) Agent Library Task execution instance. "
        "Public docs do not describe a stable external ID field like Dify conversation_id."
    )

    h("1. AutoPilot chat session / conversation", 2)
    bullet(
        "Entry: navigate Home → platform.agpt.co; natural language to run agents, "
        "build workflows, research, etc."
    )
    bullet(
        "Docs and changelog refer to chat sessions / conversations; "
        "sessions can be deleted, exported as Markdown, with timestamps; "
        "reconnect replays unseen content; history is not cleared."
    )
    bullet(
        "Long conversations can trigger summarization (UI: \"Summarizing earlier messages…\")."
    )
    bullet(
        "In-session browser/tool: multi-step browsing keeps browser session across steps "
        "within the same conversation."
    )
    bullet(
        "Changelog mentions session-level dry-run flags, etc.; "
        "research side must not fabricate \"fixed JSONL paths\" or Hermes SessionDB."
    )
    cite("AutoPilot", "https://agpt.co/docs/platform/using-the-platform/autopilot")
    cite(
        "Changelog · browse + summarize",
        "https://agpt.co/docs/platform/changelog/changelog/february-26-march-4-2026",
    )

    h("2. Agent Task (one workflow run)", 2)
    bullet(
        "Each execution in Library is one Task: inputs, outputs, cost; "
        "task URL can share result view."
    )
    bullet(
        "Manual: New Task → fill Input → Start Task; "
        "Schedule: same form via Schedule Task; "
        "Trigger-type agent: New Task becomes New Trigger, webhook-only start."
    )
    cite(
        "Agent Library",
        "https://agpt.co/docs/platform/using-the-platform/agent-library",
    )
    cite(
        "Scheduling & Triggers",
        "https://agpt.co/docs/platform/using-the-platform/scheduling-and-triggers",
    )

    h("3. Data-flow execution model (not a session protocol)", 2)
    bullet(
        "No separate control-flow: blocks run when \"all connected and required inputs ready\"."
    )
    bullet(
        "Input starts → passes along pins → Output collects results; "
        "blocks without dependencies can run in parallel in any order."
    )
    cite(
        "Data Flow & Execution",
        "https://agpt.co/docs/platform/using-the-platform/data-flow-and-execution",
    )

    h("4. Research mapping vs OpenClaw / Hermes / Dify (non-official terms)", 2)
    bullet(
        "OpenClaw Main / Claude Code interactive session ≈ AutoPilot same conversation "
        "(including cross-step browser session)."
    )
    bullet(
        "Isolated / Automation single run ≈ one Task on Library "
        "(manual / Schedule / Webhook)."
    )
    bullet(
        "Dify conversation_id-level API continuation ≈ AutoPilot chat history "
        "+ (optional) Mem0 memory blocks across Tasks; do not write as Dify-style API field."
    )

    # ══════════════════════════════════════════════════════════════
    # Task modes
    # ══════════════════════════════════════════════════════════════
    h("Task Modes", 1)

    h("1. AutoPilot (conversational platform assistant)", 2)
    bullet("Run Library agents, generate/edit agents, search Marketplace.")
    bullet(
        "Directly execute individual blocks (~400+): research, image/video, HTTP, run code "
        "(including delegating Claude Code), etc., without building a full graph first."
    )
    bullet(
        "Manage schedules, view task results; changelog: Native scheduling from chat, "
        "MCP auto-discovery, etc."
    )
    cite("AutoPilot", "https://agpt.co/docs/platform/using-the-platform/autopilot")

    h("2. Agent Builder (visual workflow)", 2)
    bullet("Drag Input / Action / Output (and Trigger) Blocks on Canvas and connect.")
    bullet(
        "Input/Output define external schema; internal blocks not exposed to user at runtime."
    )
    bullet("Save updates Library (no separate draft state documented).")
    cite(
        "Agent Builder Guide",
        "https://agpt.co/docs/platform/using-the-platform/agent-builder-guide",
    )

    h("3. Manual Task / Schedule / Webhook Trigger", 2)
    bullet("On-demand: New Task + Start Task.")
    bullet(
        "Schedule: pre-filled inputs + frequency/day/time/timezone; manage in Scheduled tab."
    )
    bullet(
        "Trigger: place trigger block in Builder → Library New Trigger → "
        "Webhook URL (e.g. backend.agpt.co/.../generic_webhook/...); "
        "such agents cannot be manually New Task."
    )

    h("4. Marketplace / import-export / API / self-host", 2)
    bullet("Marketplace add or share agents; Upload / Export file exchange.")
    bullet("API: account API key; separate OAuth/SSO docs.")
    bullet(
        "Cloud and Self-Host paths; platform directory Polyform Shield, rest of repo mostly MIT."
    )
    cite("API Introduction", "https://agpt.co/docs/platform/api-and-integrations/api-guide")

    h("5. AutoGPT Classic (legacy, not primary subject)", 2)
    bullet(
        "LLM loop decision + action results fed back into prompt as generalist agent; "
        "user must authorize actions."
    )
    bullet(
        "Official Maintenance Notice: no longer supported from security perspective; "
        "dependencies not updated, issues not fixed — threat analysis papers should not "
        "assume Classic still receives patches."
    )
    cite(
        "Classic Introduction",
        "https://agpt.co/docs/classic/autogpt-classic/introduction",
    )

    h("Task Modes · Review summary", 2)
    bullet(
        "Kept: AutoPilot ↔ Builder equivalent capability surface; Task/Schedule/Trigger; "
        "data-flow execution; Classic = legacy."
    )
    bullet(
        "Removed: OpenClaw Heartbeat; applying Dify \"Schedule only on Workflow\" to AutoGPT."
    )

    # ══════════════════════════════════════════════════════════════
    # web_fetch
    # ══════════════════════════════════════════════════════════════
    h("web_fetch (and Security Notice)", 1)
    p(
        "Conclusion (vs OpenClaw/Claude Code): "
        "AutoGPT Platform docs have no built-in tool named web_fetch / WebFetch, "
        "nor documented OpenClaw-style SECURITY NOTICE / EXTERNAL_UNTRUSTED_CONTENT wrapping. "
        "External web enters model context via: AutoPilot browsing + various scrape/HTTP Blocks; "
        "custom outbound blocks should use backend SSRF-protected Requests wrapper."
    )

    h("1. AutoPilot native browsing (changelog product capability)", 2)
    bullet("One-shot: single-step fetch + extract any page.")
    bullet(
        "Multi-step: full browser session, persists across steps within same conversation "
        "(login, menu navigation, data extraction, etc.)."
    )
    bullet(
        "Docs do not give internal tool name, lossy small-model extraction, or 15min cache — "
        "do not copy these details from Claude Code WebFetch or OpenClaw."
    )
    cite(
        "Changelog · AutoPilot can browse",
        "https://agpt.co/docs/platform/changelog/changelog/february-26-march-4-2026",
    )

    h("2. Workflow Blocks: scrape / content API (approximate \"read web page\")", 2)
    bullet(
        "Firecrawl Scrape: single URL; markdown/html/screenshot, etc.; "
        "max_age default documented as 1 hour (page cache age parameter, ≠ framework SECURITY NOTICE)."
    )
    bullet(
        "Exa Contents: urls → full text/highlights/summary; livecrawl option; "
        "context string for LLM feeding."
    )
    bullet(
        "Extract Website Content (Jina): pull HTML and extract main text; "
        "raw_content optional raw or Jina Reader."
    )
    bullet(
        "Stagehand Act / Extract / Observe: AI browser automation on Browserbase."
    )
    cite("Firecrawl Scrape", "https://agpt.co/docs/integrations/block-integrations/scrape")
    cite("Exa Contents", "https://agpt.co/docs/integrations/block-integrations/contents")
    cite("Jina Search / Extract", "https://agpt.co/docs/integrations/block-integrations/search-2")
    cite("Stagehand Blocks", "https://agpt.co/docs/integrations/block-integrations/blocks")

    h("3. HTTP / generic outbound", 2)
    bullet(
        "Send Web Request: any URL, methods GET/POST/PUT/DELETE/PATCH; "
        "JSON/form/multipart; distinguishes client_error / server_error."
    )
    bullet("Send Authenticated Web Request: inject credentials by host.")
    bullet("Read RSS Feed; Execute Code (E2B sandbox, with internet).")
    bullet(
        "AutoPilot docs also state ability to \"Make custom HTTP requests to any API\"."
    )
    cite("Misc · Send Web Request", "https://agpt.co/docs/integrations/block-integrations/misc")

    h("4. SSRF protection (block development security, not content wrapper)", 2)
    bullet(
        "Build your own Blocks: external URLs must use "
        "backend.util.request requests/Requests wrapper."
    )
    bullet(
        "Features: validate URL/protocol; resolve DNS; block private networks (RFC 1918, etc.); "
        "redirects disabled by default; optional trusted_origins; non-200 can raise."
    )
    bullet(
        "This is request-layer SSRF control, not wrapping tool output in SECURITY NOTICE text. "
        "Whether third-party Firecrawl/Exa/Jina paths share the same wrapper is not "
        "individually stated in public docs."
    )
    cite(
        "Security Best Practices for SSRF Prevention",
        "https://agpt.co/docs/platform/building-blocks/new_blocks",
    )

    h("5. Classic legacy browsing surface (comparison)", 2)
    bullet(
        "WebSearchComponent, WebSeleniumComponent and other Forge built-in components "
        "provide search/browser commands."
    )
    bullet(
        "Classic security maintenance stopped — paper mainline should focus on Platform."
    )

    h("Threat model mapping (research, non-official)", 2)
    bullet(
        "External web text → pin flows into AI Text Generator / AutoPilot context → "
        "subsequent blocks or dialogue actions: IPI / Goal Hijacking paper surface."
    )
    bullet(
        "Schedule/Trigger repeatedly running scrape→LLM graphs: persistent poisoning can "
        "trigger across Tasks (if written to Mem0/Graphiti, more like cross-session memory pollution)."
    )
    bullet(
        "Comparison: OpenClaw web_fetch+SECURITY NOTICE; Claude Code WebFetch+isolated extraction; "
        "Hermes web_extract+untrusted_tool_result; "
        "Dify HTTP Request+ssrf_proxy — AutoGPT closer to \"Blocks/outbound + SSRF wrapper\", "
        "plus AutoPilot session-level browser."
    )

    code(
        """Example: Send Web Request → AI Text Generator (research sketch, not official message protocol)
[Input:url] → [Send Web Request] → response body
                              ↓
                    [AI Text Generator]
                      prompt: summarize {{body}}
                      sys_prompt: ...
                              ↓
                         [Agent Output]
Note: actual execution is data-flow pins, not fixed OpenAI tool_calls schema;
AutoPilot internal tool call names are product implementation — do not write as web_fetch."""
    )

    h("web_fetch · Review summary", 2)
    bullet(
        "Kept: no web_fetch name; no SECURITY NOTICE docs; "
        "AutoPilot browse; Firecrawl/Exa/Jina/Stagehand/Send Web Request; SSRF Requests."
    )
    bullet(
        "Removed: default 15min cache + SECURITY NOTICE; "
        "removed treating Classic Selenium as Platform default."
    )

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "Platform \"tools\" unit is Block (and AutoPilot direct Block/MCP invocation), "
        "not Claude Code-style fixed PascalCase built-in table."
    )

    h("1. Blocks as capability atoms", 2)
    bullet("Three types: Input / Action / Output; plus Trigger as special Input.")
    bullet(
        "Integrations cover LLM, search, comms, Notion, GitHub, Airtable, Exa, Firecrawl, etc."
    )
    bullet("Agent Blocks: embed full agent graph into larger workflow.")

    cite("Agent Blocks Overview", "https://agpt.co/docs/platform/building-blocks/agent-blocks")
    cite("Integrations index", "https://agpt.co/docs/integrations/readme")

    h("2. AutoPilot tool surface", 2)
    bullet(
        "Directly run any block; MCP: natural language connect Notion/Slack/Jira, etc."
    )
    bullet(
        "Changelog: ask_question, self-distilled skills registry, MCP auto-discover, etc."
    )

    h("3. Human-in-the-loop and others", 2)
    bullet(
        "Human In The Loop: pause for human approval/editable continue."
    )
    bullet("Credentials: OAuth / API key / user password; credential bar on blocks.")
    bullet("Custom blocks: Block SDK + SSRF security requirements.")

    h("Tools · Review summary", 2)
    bullet("Kept: Blocks ≈ tools; MCP; Agent-as-block; HITL.")
    bullet("Removed: fabricated built-in tool names WebFetch / web_fetch.")

    # ══════════════════════════════════════════════════════════════
    # Five-round anti-hallucination
    # ══════════════════════════════════════════════════════════════
    h("Five-Round Anti-Hallucination Review", 1)

    h("Round 1 — Session", 2)
    bullet(
        "Checked: autopilot; agent-library; data-flow-and-execution; related changelog."
    )
    bullet(
        "Kept: AutoPilot conversation + Library Task; "
        "reconnect replay; in-session browser persistence; no public conversation_id contract."
    )
    bullet(
        "Removed: Claude Code ~/.claude/projects JSONL; Hermes state.db; Dify sys.conversation_id."
    )

    h("Round 2 — Task modes", 2)
    bullet(
        "Checked: autopilot vs builder; scheduling-and-triggers; what-is-platform; classic intro."
    )
    bullet(
        "Kept: AutoPilot / Builder / Task / Schedule / Trigger / Marketplace / API; Classic legacy."
    )
    bullet("Removed: Heartbeat; writing Classic as current recommended production path.")

    h("Round 3 — web_fetch / security", 2)
    bullet(
        "Checked: february-26-march-4-2026 browse; scrape/contents/search-2/blocks/misc; "
        "new_blocks SSRF."
    )
    bullet(
        "Kept: no web_fetch/SECURITY NOTICE; "
        "one-shot + multi-step browse; Send Web Request; third-party scrape; SSRF wrapper."
    )
    bullet(
        "Removed: OpenClaw NOTICE text assumed by default; "
        "removed Claude Code \"15min cache + isolated small-model extract\" as AutoGPT doc fact."
    )

    h("Round 4 — Tools", 2)
    bullet("Checked: agent-builder-guide; integrations; agent-blocks; MCP changelog.")
    bullet("Kept: hundreds of Blocks; AutoPilot direct run; MCP; custom block SDK.")
    bullet("Removed: fixed tool whitelist WebFetch/Bash/Read.")

    h("Round 5 — Prompt / Memory cross-check", 2)
    bullet(
        "Checked: AI Text Generator; Mem0 Basic blocks; summarize changelog; Graphiti changelog."
    )
    bullet(
        "Frozen: "
        "(1) Session = AutoPilot chat + Task run; "
        "(2) Task modes = AutoPilot / Builder / Schedule / Trigger; "
        "(3) External web = browse + HTTP/scrape blocks, SSRF at request layer; "
        "(4) Prompt = block-level prompt/sys_prompt + pin injection; "
        "(5) Memory = dialogue summarization + Mem0(+Graphiti) + file workspace, not MEMORY.md."
    )
    bullet(
        "Open questions: AutoPilot browse internal tool name and untrusted wrapping; "
        "Cloud third-party scrape vs SSRF wrapper boundary; "
        "Mem0 user/run/agent scope defaults need live verification."
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt composition
    # ══════════════════════════════════════════════════════════════
    h("Prompt Composition", 1)
    p(
        "Builder path: prompts live on LLM-class Block fields, filled by upstream pins; "
        "AutoPilot path: user natural language + platform-side system policy "
        "(full system template not publicly disclosed)."
    )

    h("1. AI Text Generator (block-level)", 2)
    bullet(
        "prompt: main user prompt; fill via Prompt Values {keys} / {{var}}."
    )
    bullet("sys_prompt: system-side extra context.")
    bullet("model / max_tokens / retry / ollama_host, etc.")
    cite("LLM · AI Text Generator", "https://agpt.co/docs/integrations/block-integrations/llm")

    h("2. Data-flow context assembly", 2)
    bullet(
        "Typical: Scrape/Exa/Send Web Request markdown or body → "
        "connect to prompt / prompt_values → LLM → Output."
    )
    bullet("Lists can iterate per URL for extract+generate pipeline.")

    h("3. AutoPilot dialogue side", 2)
    bullet(
        "New session themed prompt categories: Learn / Create / Automate / Organize."
    )
    bullet("Export chat as Markdown; long context triggers summarizing indicator.")
    bullet(
        "Official docs do not disclose full AutoPilot system prompt or tool schema list."
    )

    h("4. Classic (comparison)", 2)
    bullet(
        "Loop: goal + ActionHistory and other component summaries enter prompt; "
        "commands provided by each Component's CommandProvider."
    )

    h("Prompt Composition · Review summary", 2)
    bullet("Kept: prompt + sys_prompt + prompt_values; pin injection of external web text.")
    bullet(
        "Removed: CLAUDE.md / SOUL.md / Dify Jinja2 node templates as AutoGPT defaults."
    )

    # ══════════════════════════════════════════════════════════════
    # Memory and persistence
    # ══════════════════════════════════════════════════════════════
    h("Memory and Persistence", 1)
    p(
        "No OpenClaw/Hermes-style repo-root MEMORY.md convention; "
        "cross-turn/cross-run memory is split across: AutoPilot dialogue and summarization, "
        "Mem0 blocks, file workspace, Graphiti long-term memory blocks in changelog, "
        "and Task/Library metadata."
    )

    h("1. AutoPilot dialogue memory and summarization", 2)
    bullet(
        "Multi-turn messages retained; when too long, summarize earlier messages (visible UI)."
    )
    bullet(
        "Attachments/generated files referenceable in dialogue; changelog calls it "
        "persistent file workspace."
    )

    h("2. Mem0 memory Blocks (across workflow executions)", 2)
    bullet(
        "Add Memory: write to Mem0; segment by user / optional run or agent; metadata supported."
    )
    bullet("Get All Memories / Get Latest Memory / Search Memory (semantic retrieval).")
    bullet("Docs: Memories persist across workflow executions.")
    cite("Basic · Memory blocks", "https://agpt.co/docs/integrations/block-integrations/basic")

    h("3. Graphiti memory (changelog)", 2)
    bullet(
        "v0.6.55: Long-term memory via Graphiti as block providing persistent agent knowledge."
    )
    bullet(
        "Details per Marketplace/block docs; do not conflate with Liang GraphRAG paper environment."
    )

    h("4. Files and Store Value", 2)
    bullet(
        "File Store: URL/data URI/local path → temp directory for downstream blocks."
    )
    bullet("Store Value: in-workflow constant forwarding; not long-term external memory.")

    h("5. Skills registry (AutoPilot changelog)", 2)
    bullet(
        "Self-distilled skills: after complex multi-step completion, reusable recipes can "
        "precipitate — procedural memory, distinct from Mem0 factual entries."
    )

    h("6. Classic comparison", 2)
    bullet(
        "ActionHistoryComponent, ContextComponent, etc. provide history and file context to prompt."
    )
    bullet(
        "Security maintenance stopped ⇒ paper main threat surface should remain Platform."
    )

    h("Memory and Persistence · Review summary", 2)
    bullet(
        "Kept: dialogue+summarization; Mem0; file workspace; Graphiti block (changelog); "
        "Task history ≠ semantic memory."
    )
    bullet(
        "Removed: default MEMORY.md; "
        "removed each Schedule automatically inheriting full AutoPilot chat transcript "
        "(Schedule runs agent graph + pre-filled inputs unless graph explicitly reads Mem0)."
    )

    h("Primary Reference Links", 1)
    for link in [
        "https://agpt.co/docs/llms.txt",
        "https://agpt.co/docs/platform/what-is-autogpt-platform",
        "https://agpt.co/docs/platform/using-the-platform/autopilot",
        "https://agpt.co/docs/platform/using-the-platform/agent-builder-guide",
        "https://agpt.co/docs/platform/using-the-platform/agent-library",
        "https://agpt.co/docs/platform/using-the-platform/scheduling-and-triggers",
        "https://agpt.co/docs/platform/using-the-platform/data-flow-and-execution",
        "https://agpt.co/docs/platform/building-blocks/new_blocks",
        "https://agpt.co/docs/platform/building-blocks/agent-blocks",
        "https://agpt.co/docs/platform/changelog/changelog/february-26-march-4-2026",
        "https://agpt.co/docs/integrations/block-integrations/scrape",
        "https://agpt.co/docs/integrations/block-integrations/contents",
        "https://agpt.co/docs/integrations/block-integrations/search-2",
        "https://agpt.co/docs/integrations/block-integrations/blocks",
        "https://agpt.co/docs/integrations/block-integrations/misc",
        "https://agpt.co/docs/integrations/block-integrations/basic",
        "https://agpt.co/docs/integrations/block-integrations/llm",
        "https://agpt.co/docs/platform/api-and-integrations/api-guide",
        "https://agpt.co/docs/classic/autogpt-classic/introduction",
    ]:
        bullet(link)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
