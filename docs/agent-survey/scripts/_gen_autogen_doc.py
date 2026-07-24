# -*- coding: utf-8 -*-
"""Generate AutoGen research document — accuracy-focused with 5-round audit."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "AutoGen.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def add_code(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def cite(text: str, url: str) -> None:
        doc.add_paragraph(f"{text}（来源：{url}）")

    doc.add_heading("AutoGen", level=0)
    doc.add_paragraph(
        "本文档研究 Microsoft AutoGen 生态中的 Agent 机制，术语严格对齐官方文档。"
        "用户指定入口：https://microsoft.github.io/autogen/stable/index.html"
    )

    # ── Scope ──
    doc.add_heading("文档范围与术语分层（必读）", level=1)
    doc.add_paragraph("官方文档实际分为五层，研究时不可混用：")
    bullet(
        "AutoGen Core（autogen-core）：事件驱动、Actor 模型、异步消息、"
        "可分布式 runtime；面向严肃多智能体系统。"
    )
    bullet(
        "AgentChat（autogen-agentchat）：基于 Core 的高层 API；"
        "预设 Agent（AssistantAgent 等）、Team（RoundRobin/Selector/Swarm/"
        "MagenticOne/GraphFlow）、Termination、HITL；初学者推荐入口。"
    )
    bullet(
        "Extensions（autogen-ext）：模型客户端、MCP Workbench、"
        "MultimodalWebSurfer、Docker 代码执行、ChromaDB/Redis Memory 等。"
    )
    bullet(
        "AutoGen Studio：低代码原型 UI（Team Builder / Playground / Gallery）；"
        "官方明确非生产就绪应用。"
    )
    bullet(
        "AutoGen 0.2（旧版）：与 0.4+ 架构不同；"
        "Migration Guide 指引从 0.2.x 迁移，勿与当前 AgentChat API 混谈。"
    )
    cite(
        "Index；AgentChat；Core；AutoGen Studio",
        "https://microsoft.github.io/autogen/stable/index.html",
    )

    # ── 1. LLM 请求采集 ──
    doc.add_heading("研究背景与 LLM 请求采集", level=1)
    doc.add_paragraph(
        "采集目标：每次 model 调用前的 messages、tools schema，"
        "tool_calls 与 FunctionExecutionResult、Memory 注入、Team 消息流。"
    )

    doc.add_heading("1. run_stream / Console 消息流", level=2)
    bullet(
        "agent.run_stream(task=) 与 team.run_stream(task=) 返回异步生成器；"
        "最后一项为 TaskResult。"
    )
    bullet(
        "消息类型（autogen_agentchat.messages）：TextMessage、"
        "ToolCallRequestEvent、ToolCallExecutionEvent、ToolCallSummaryMessage、"
        "MemoryQueryEvent、HandoffMessage、StructuredMessage、MultiModalMessage 等。"
    )
    bullet(
        "Console(stream, output_stats=True) 格式化打印并汇总 prompt/completion tokens。"
    )
    cite(
        "Agents > Streaming Messages；Teams > Observing a Team",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/agents.html",
    )

    doc.add_heading("2. Python Logging（TRACE / EVENT）", level=2)
    bullet(
        "AgentChat：logging.getLogger(TRACE_LOGGER_NAME) 与 "
        "logging.getLogger(EVENT_LOGGER_NAME) 分别启用 trace 与结构化 agent 间消息日志。"
    )
    bullet(
        "模型客户端调用与 runtime 事件日志见 Core Logging Guide（需额外配置）。"
    )
    cite(
        "Logging",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/logging.html",
    )

    doc.add_heading("3. flow.stream_events 通道投影（Flow 类运行时）", level=2)
    bullet(
        "Consuming Streams 文档描述 flow.stream_events / stream_turn 的 "
        "llm、tools、flow、messages 通道投影；"
        "须消费完毕后再读 stream.result。"
    )
    bullet(
        "frame.channel / frame.type / frame.content / frame.event 承载元数据。"
    )
    bullet(
        "注：此为 AgentChat 流式运行时契约文档；"
        "与 CrewAI 同名 API 不同产品，勿混用术语。"
    )
    cite(
        "Consuming Streams",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/consuming-streams.html",
    )

    doc.add_heading("4. 直接访问 model_context", level=2)
    bullet(
        "AssistantAgent 内部维护 model_context（ChatCompletionContext）；"
        "Memory.update_context 将检索结果写入 SystemMessage。"
    )
    bullet(
        "可 await agent._model_context.get_messages() 检查注入后完整 LLM 消息列表（文档示例）。"
    )
    bullet(
        "BufferedChatCompletionContext(buffer_size=N) 可限制送入模型的历史长度。"
    )
    cite(
        "Memory and RAG；Agents > model_context",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/memory.html",
    )

    doc.add_heading("5. 第三方可观测性", level=2)
    bullet(
        "Observability 目录列 Langfuse、Datadog、MLflow、OpenLIT、"
        "Arize Phoenix、Opik 等 OpenTelemetry 集成；属可选部署配置。"
    )
    cite(
        "Observability overview",
        "https://microsoft.github.io/autogen/stable/user-guide/observability/overview.html",
    )

    # ── 2. Prompt ──
    doc.add_heading("Prompt 设计组成", level=1)

    doc.add_heading("1. AssistantAgent 层", level=2)
    bullet(
        "system_message：系统指令；description：agent 文本描述（Team 选择器可用）。"
    )
    bullet(
        "run(task=) 传入新消息（非完整历史）；"
        "agent 为有状态对象，run 会将消息追加到内部状态。"
    )
    bullet(
        "reflect_on_tool_use=True：工具返回非自然语言时，额外让模型总结 tool output。"
    )
    bullet(
        "output_content_type=Pydantic 模型：结构化输出（StructuredMessage）。"
    )
    bullet(
        "handoffs=[Handoff(target=...)]：生成 transfer_to_* 工具供模型切换 agent。"
    )
    cite(
        "Agents",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/agents.html",
    )

    doc.add_heading("2. Team 层（共享上下文）", level=2)
    bullet(
        "RoundRobinGroupChat：参与者轮询发言，每次广播给全体（共享 context）。"
    )
    bullet(
        "SelectorGroupChat：每轮由 ChatCompletion 模型选择下一发言者。"
    )
    bullet(
        "Swarm：HandoffMessage 驱动 agent 切换。"
    )
    bullet(
        "MagenticOneGroupChat：Orchestrator 规划 + 子 agent 执行（Task/Progress Ledger）。"
    )
    bullet(
        "GraphFlow：DiGraph 有向图控制执行顺序（顺序/并行/条件/循环）；标注为 experimental。"
    )
    cite(
        "Teams；GraphFlow；Magentic-One",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/teams.html",
    )

    doc.add_heading("3. Memory / RAG 注入", level=2)
    bullet(
        "Memory 协议：add / query / update_context / clear / close。"
    )
    bullet(
        "ListMemory：按时间顺序将全部记忆追加为 SystemMessage "
        "（\"Relevant memory content (in chronological order):\"）。"
    )
    bullet(
        "ChromaDBVectorMemory / RedisMemory：向量检索后 update_context；"
        "RAG 质量依赖 chunking 与 embedding 选择。"
    )
    bullet(
        "每次 run 前触发 MemoryQueryEvent，可见检索到的 MemoryContent 列表。"
    )
    cite(
        "Memory and RAG",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/memory.html",
    )

    doc.add_heading("4. Core 层（自建 Agent）", level=2)
    bullet(
        "Core Tools 文档：RoutedAgent + model_client.create(messages, tools=) "
        "自行组装 SystemMessage + UserMessage + AssistantMessage + "
        "FunctionExecutionResultMessage 多轮循环。"
    )
    bullet(
        "官方注明：Core API 最小化，预组装 agent 请用 AgentChat。"
    )
    cite(
        "Core > Tools > Tool-Equipped Agent",
        "https://microsoft.github.io/autogen/stable/user-guide/core-user-guide/components/tools.html",
    )

    doc.add_heading("5. 典型 LLM 请求示意（含 Memory 注入）", level=2)
    add_code(
        """[
  UserMessage(content="What is the weather in New York?", source="user"),
  SystemMessage(content="\\nRelevant memory content (in chronological order):\\n1. The weather should be in metric units\\n"),
  AssistantMessage(content=[FunctionCall(name="get_weather", arguments='{"city":"New York","units":"metric"}')]),
  FunctionExecutionResultMessage(content=[FunctionExecutionResult(content="23 °C and Sunny.", ...)])
]"""
    )
    cite(
        "Memory and RAG > ListMemory 示例 _model_context.get_messages()",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/memory.html",
    )

    # ── 3. Memory ──
    doc.add_heading("Memory 与持久化机制", level=1)

    doc.add_heading("1. Memory 协议（语义/RAG 上下文）", level=2)
    bullet("ListMemory：简单列表，全量或按 query 检索。")
    bullet("ChromaDBVectorMemory：ChromaDB + embedding（默认 SentenceTransformer 或 OpenAI）。")
    bullet("RedisMemory：Redis 向量库。")
    bullet(
        "update_context 在 AssistantAgent 每步前调用，"
        "将检索结果格式化为 SystemMessage 写入 model_context。"
    )
    cite(
        "Memory and RAG",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/memory.html",
    )

    doc.add_heading("2. save_state / load_state（执行状态）", level=2)
    bullet(
        "AssistantAgent.save_state() 返回含 llm_messages 的 AssistantAgentState。"
    )
    bullet(
        "Team.save_state() 保存全部参与者状态 + message_thread + current_turn 等。"
    )
    bullet(
        "状态为可 JSON 序列化 dict，可写文件/数据库；"
        "team.reset() 清空状态，load_state 后恢复对话历史。"
    )
    bullet(
        "自定义 Agent 应覆写 save_state/load_state；"
        "AssistantAgent 默认仅保存 model_context。"
    )
    cite(
        "Managing State",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/state.html",
    )

    doc.add_heading("3. Component 序列化（配置持久化）", level=2)
    bullet(
        "dump_component() / load_component()：Agent、Team、TerminationCondition "
        "可导出 JSON 声明式配置。"
    )
    bullet(
        "警告：ONLY LOAD COMPONENTS FROM TRUSTED SOURCES；"
        "序列化函数可能含可执行代码。"
    )
    bullet("tools 序列化尚未支持（Serialize Components 原文）；selector_func 不可序列化。")
    cite(
        "Serializing Components",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/serialize-components.html",
    )

    doc.add_heading("4. 与 LangGraph checkpointer / CrewAI checkpoint 的对照", level=2)
    doc.add_paragraph(
        "AutoGen 无内置 event-driven checkpoint/resume/fork 机制（所查 AgentChat 文档）；"
        "持久化依赖 save_state/load_state 或应用侧存储。"
        "研究映射：≈ LangGraph checkpointer 的应用层自管；≠ CrewAI CheckpointConfig 自动写入。"
    )

    # ── 4. Session ──
    doc.add_heading("Session 与会话模型", level=1)
    bullet(
        "Agent 有状态：on_messages / run 仅传入自上次调用以来的新消息；"
        "内部维护完整 model_context。"
    )
    bullet(
        "Team 有状态：run 后 termination condition 自动 reset，"
        "但 agent 内部状态与 message_thread 保留（除非 team.reset()）。"
    )
    bullet(
        "max_turns（RoundRobin/Selector/Swarm）：限制单轮 run 的 agent 发言次数；"
        "resume 后 turn count 从 0 重计，conversation history 保留。"
    )
    bullet(
        "HandoffTermination + HandoffMessage：agent 将控制权交给 user 或另一 agent；"
        "Swarm resume 时 task 须为 HandoffMessage。"
    )
    bullet(
        "Consuming Streams 文档：Conversational Flow 支持 session_id 多轮 stream_turn。"
    )
    cite(
        "Agents（stateful）；Human-in-the-Loop；Termination",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/human-in-the-loop.html",
    )

    doc.add_heading("AutoGen Studio 会话", level=2)
    bullet(
        "Playground：交互式测试、实时消息流、UserProxyAgent 交互、暂停/停止。"
    )
    bullet(
        "官方：非生产应用；无完整鉴权/越狱防护；"
        "生产应使用 AutoGen 框架自行实现安全特性。"
    )
    cite(
        "AutoGen Studio",
        "https://microsoft.github.io/autogen/stable/user-guide/autogenstudio-user-guide/index.html",
    )

    # ── 5. Automation ──
    doc.add_heading("Automation 与定时任务", level=1)
    bullet(
        "开源框架层无 OpenClaw 式 Heartbeat 或内置 cron 文档。"
    )
    bullet(
        "应用模式：外部调度器调用 team.run()；"
        "或 FastAPI WebSocket + UserProxyAgent 自定义 input_func 构建 Web 服务。"
    )
    bullet(
        "Human-in-the-Loop 文档引用 AgentChat FastAPI / ChainLit 示例。"
    )
    bullet(
        "AutoGen Studio Deployment：导出 Python 代码、Docker 容器运行 team。"
    )
    cite(
        "Human-in-the-Loop；AutoGen Studio > Deployment",
        "https://microsoft.github.io/autogen/stable/user-guide/autogenstudio-user-guide/index.html",
    )

    # ── 6. External Web ──
    doc.add_heading("外部 Web 访问", level=1)
    doc.add_paragraph(
        "AutoGen 无统一内置 web_fetch；外部内容通过 Extensions 工具或专用 Agent 进入。"
    )

    doc.add_heading("1. MCP Workbench（mcp-server-fetch）", level=2)
    bullet(
        "McpWorkbench(StdioServerParams(command=\"uvx\", args=[\"mcp-server-fetch\"]))："
        "通过 MCP 拉取 URL 内容；reflect_on_tool_use=True 可让模型总结页面。"
    )
    cite(
        "Agents > MCP Workbench",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/agents.html",
    )

    doc.add_heading("2. autogen_ext.tools.http", level=2)
    bullet("Core/Extensions 提供 HttpTool 用于 REST API 请求（非专用网页抓取 agent）。")
    cite(
        "Core > Tools > Built-in Tools",
        "https://microsoft.github.io/autogen/stable/user-guide/core-user-guide/components/tools.html",
    )

    doc.add_heading("3. MultimodalWebSurfer（浏览器 Agent）", level=2)
    bullet(
        "Playwright Chromium：截图 + accessibility tree + set-of-marks prompting；"
        "支持 visit_url、click、scroll、web search 等动作。"
    )
    bullet(
        "返回含页面截图、metadata、动作描述、inner text 的复合响应。"
    )
    bullet(
        "安全警告：网页可 prompt injection；可能自动点 cookie；"
        "须人工监督；Windows 需 WindowsProactorEventLoopPolicy。"
    )
    cite(
        "MultimodalWebSurfer API；Magentic-One",
        "https://microsoft.github.io/autogen/stable/reference/python/autogen_ext.agents.web_surfer.html",
    )

    doc.add_heading("4. 自定义 FunctionTool", level=2)
    bullet(
        "Agents 教程用 mock web_search 函数演示 tool calling；"
        "schema 由 FunctionTool 从签名与 docstring 自动生成。"
    )
    doc.add_paragraph("与 OpenClaw web_fetch 差异（研究对照）：")
    bullet("OpenClaw：内置 web_fetch + EXTERNAL_UNTRUSTED_CONTENT 包裹。")
    bullet(
        "AutoGen：MCP fetch / WebSurfer / 自定义 tool；"
        "tool 结果进入 FunctionExecutionResultMessage，无框架级 untrusted 边界。"
    )

    # ── 7. Security ──
    doc.add_heading("Security 机制", level=1)

    doc.add_heading("1. 代码执行隔离", level=2)
    bullet(
        "DockerCommandLineCodeExecutor + PythonCodeExecutionTool："
        "Docker 容器内执行模型生成代码。"
    )
    bullet(
        "MagenticOne approval_func：代码执行前 ApprovalRequest/ApprovalResponse 人工批准。"
    )
    bullet(
        "Magentic-One 建议：容器隔离、虚拟环境、监控日志、限制网络访问。"
    )
    cite(
        "Core > PythonCodeExecutionTool；Magentic-One",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/magentic-one.html",
    )

    doc.add_heading("2. Human-in-the-Loop", level=2)
    bullet(
        "UserProxyAgent + input_func：run 期间阻塞等待用户输入；"
        "会阻碍 team 保存/恢复，仅适合短交互（批准按钮等）。"
    )
    bullet(
        "推荐异步模式：max_turns 或 HandoffTermination 暂停 → "
        "应用保存 state → 用户反馈后再次 run()。"
    )
    bullet(
        "ExternalTermination：UI「Stop」按钮等外部控制终止。"
    )
    cite(
        "Human-in-the-Loop；Termination > ExternalTermination",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/human-in-the-loop.html",
    )

    doc.add_heading("3. MCP / Web 安全", level=2)
    bullet(
        "MCP Workbench 连接不可信 server 时，tool metadata 可影响 LLM 行为（与 CrewAI MCP 文档同类风险）。"
    )
    bullet(
        "MultimodalWebSurfer：官方明确易受网页 prompt injection；"
        "勿在无监督环境运行。"
    )
    bullet(
        "Serialize Components：仅加载可信来源配置，防代码执行型反序列化攻击。"
    )

    doc.add_heading("4. Task/Agent 层 guardrail", level=2)
    doc.add_paragraph(
        "所查 AgentChat 文档无 CrewAI 式 HallucinationGuardrail 或 LangChain PIIMiddleware；"
        "安全控制依赖：termination、handoff、approval_func、"
        "应用侧对 tool 返回的过滤，以及 Core/Extensions 自选组件。"
    )

    doc.add_heading("5. IPI 研究含义", level=2)
    doc.add_paragraph(
        "注入面：FunctionExecutionResult / ToolCallSummaryMessage 内容、"
        "Memory.update_context 写入的 SystemMessage、"
        "MultimodalWebSurfer 返回的 page inner text 与截图描述、"
        "MCP tool metadata 与 result、Team 广播的跨 agent 消息。"
        "不能假设 EXTERNAL_UNTRUSTED_CONTENT 自动包裹。"
    )

    # ── 8. Architecture ──
    doc.add_heading("架构总览", level=1)
    bullet(
        "推荐路径：简单任务用单 AssistantAgent；"
        "复杂协作用 Team；需确定性流程用 GraphFlow；"
        "开放域 Web/文件任务用 MagenticOne。"
    )
    bullet(
        "AssistantAgent 工具循环：默认 max_tool_iterations=1；"
        "可增至 10；同轮多 tool call 默认并行（可 parallel_tool_calls=False）。"
    )
    bullet(
        "Core：SingleThreadedAgentRuntime / GrpcWorkerAgentRuntime 分布式；"
        "RoutedAgent + message_handler 异步消息。"
    )
    bullet(
        "Workbench：共享状态的工具集合（如 McpWorkbench）；"
        "与单 FunctionTool 相对。"
    )
    cite(
        "Index；Agents > Tool Iterations；Core",
        "https://microsoft.github.io/autogen/stable/index.html",
    )

    doc.add_heading("与 OpenClaw / OpenHands / Hermes / LangGraph / CrewAI 的对照（研究映射）", level=2)
    doc.add_paragraph("非 AutoGen 官方术语：")
    bullet(
        "AgentChat Team ≈ CrewAI Crew / LangGraph 多 node 子图；"
        "GraphFlow ≈ CrewAI Flow / LangGraph StateGraph 确定性边。"
    )
    bullet(
        "save_state/load_state ≈ LangGraph checkpointer 手动序列化；"
        "≠ CrewAI 自动 CheckpointConfig 事件写入。"
    )
    bullet(
        "Memory.update_context → SystemMessage ≈ CrewAI memory recall 注入 / "
        "LangGraph store→system_msg 应用逻辑。"
    )
    bullet(
        "MultimodalWebSurfer ≈ MUZZLE Web Agent / OpenHands browser 能力；"
        "无 OpenClaw 单一 web_fetch 抽象。"
    )
    bullet(
        "UserProxyAgent 阻塞输入 ≈ LangGraph interrupt/HITL；"
        "HandoffMessage ≈ CrewAI Flow @human_feedback 路由。"
    )

    # ── 9. Research ──
    doc.add_heading("研究建议", level=1)
    bullet(
        "采集：run_stream + Console；TRACE/EVENT logger；"
        "或直接 _model_context.get_messages() 在 Memory 注入后采样。"
    )
    bullet("Memory vs save_state：RAG 注入与对话历史恢复分开测试。")
    bullet("Session：team.load_state 跨请求 vs agent.run 单会话追加。")
    bullet(
        "IPI：MCP fetch 返回、WebSurfer inner text、"
        "ListMemory/ChromaDB 投毒、Team 广播消息。"
    )
    bullet(
        "安全：approval_func vs UserProxyAgent 阻塞；"
        "Docker 代码执行 vs LocalCommandLineCodeExecutor。"
    )

    # ── 10. Five-round audit ──
    doc.add_heading("自查：潜在幻觉与核实结论（第五轮）", level=1)
    doc.add_paragraph(
        "2026-07-12 第五轮复核：对照 Index / AgentChat Agents / Teams / "
        "Memory / State / Serialize / HITL / Termination / GraphFlow / "
        "Magentic-One / MultimodalWebSurfer / Core Tools / Logging / "
        "AutoGen Studio 全部关键断言。"
        "结论：第四轮修正项仍成立；本轮未发现新的机制性幻觉。"
    )
    checks = [
        ("✓ 五层架构：Core / AgentChat / Extensions / Studio / 0.2 分离", "https://microsoft.github.io/autogen/stable/index.html"),
        ("✓ AssistantAgent 有状态；run 传新消息非完整历史", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/agents.html"),
        ("✓ v0.4 工具由同一 agent 在 run 内直接执行（非 0.2 分离 executor）", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/agents.html"),
        ("✓ Memory 协议：add/query/update_context；ListMemory 写 SystemMessage", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/memory.html"),
        ("✓ save_state/load_state：Agent 与 Team 可 JSON 持久化", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/state.html"),
        ("✓ MCP Workbench + mcp-server-fetch 拉取 URL", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/agents.html"),
        ("✓ MultimodalWebSurfer：Playwright + 网页 prompt injection 警告", "https://microsoft.github.io/autogen/stable/reference/python/autogen_ext.agents.web_surfer.html"),
        ("✓ UserProxyAgent 阻塞；不推荐长时阻塞（影响 save/resume）", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/human-in-the-loop.html"),
        ("✓ GraphFlow 标注 experimental", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/graph-flow.html"),
        ("✓ AutoGen Studio 非生产就绪（官方 Caution）", "https://microsoft.github.io/autogen/stable/user-guide/autogenstudio-user-guide/index.html"),
        ("△ consuming-streams — flow.stream_events 文档与 AgentChat run_stream 并存，采集时需区分 API", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/consuming-streams.html"),
        ("△ Flow stream_events — 属 AgentChat 流式契约文档，与 CrewAI 同名 API 不同产品", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/consuming-streams.html"),
        ("△ max_tool_iterations 默认 1 — 多步 tool 链需显式提高", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/agents.html"),
        ("△ tools 不可 dump_component 序列化", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/serialize-components.html"),
        ("✗ 内置 web_fetch / EXTERNAL_UNTRUSTED_CONTENT — 所查页面均无", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/agents.html"),
        ("✗ OpenClaw 式 Heartbeat — 所查页面无等价机制", "https://microsoft.github.io/autogen/stable/index.html"),
        ("✗ 自动 checkpoint/resume/fork — 无 CrewAI/LangGraph 式内置；仅 save_state", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/state.html"),
        ("✗ 框架级 HallucinationGuardrail / PIIMiddleware — AgentChat 文档未提供", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/agents.html"),
        ("✗ AssistantAgent 适合生产 — 官方称 kitchen sink，建议理解后自实现", "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/agents.html"),
        ("— 第五轮：上述条目全部复检通过，无新增 ✗/△", "https://microsoft.github.io/autogen/stable/index.html"),
    ]
    for text, url in checks:
        cite(text, url)

    doc.add_heading("防幻觉机制说明（五轮复核流程）", level=2)
    bullet("第一轮：按 Index 建 Core/AgentChat/Extensions/Studio 四层，逐节标注来源。")
    bullet("第二轮：核对 Memory（语义注入）与 save_state（执行状态）是否混谈。")
    bullet("第三轮：核对 Web 路径（MCP fetch / WebSurfer / mock tool）与 untrusted 边界。")
    bullet("第四轮：核对 HITL 双路径、Team 预设、GraphFlow experimental 标注。")
    bullet("第五轮：全量重读核心页，确认 0.2 vs 0.4 API 未混入正文。")

    doc.add_heading("主要参考链接", level=1)
    for link in [
        "https://microsoft.github.io/autogen/stable/index.html",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/index.html",
        "https://microsoft.github.io/autogen/stable/user-guide/core-user-guide/index.html",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/agents.html",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/teams.html",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/memory.html",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/state.html",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/serialize-components.html",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/human-in-the-loop.html",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/tutorial/termination.html",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/graph-flow.html",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/magentic-one.html",
        "https://microsoft.github.io/autogen/stable/user-guide/core-user-guide/components/tools.html",
        "https://microsoft.github.io/autogen/stable/user-guide/agentchat-user-guide/logging.html",
        "https://microsoft.github.io/autogen/stable/user-guide/autogenstudio-user-guide/index.html",
        "https://microsoft.github.io/autogen/stable/reference/python/autogen_ext.agents.web_surfer.html",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
