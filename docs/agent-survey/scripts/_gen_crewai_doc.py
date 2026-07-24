# -*- coding: utf-8 -*-
"""Generate CrewAI research notes (same structure as OpenHands / Codex / Cline).

Paper track: Session / 任务模式 / web_fetch(security notice) / tools,
plus Prompt 组成 & 记忆持久化. 5-round anti-hallucination vs docs.crewai.com
(v1.15.2).
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "zh" / "CrewAI.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label}（来源：{url}）")

    # ── Title ──
    h("CrewAI", 0)
    p(
        "本文档结构对齐 OpenHands.docx / Hermes.docx / Claude-Code.docx / "
        "Dify.docx / AutoGPT.docx / Cline.docx / Codex-CLI.docx："
        "Session、任务模式、web_fetch（及 Security Notice 对照）、Tools；"
        "文末附 Prompt 组成、记忆与持久化机制，以及五轮防幻觉核对。"
        "对象为开源多智能体框架 CrewAI（Crews + Flows）；"
        "CrewAI AMP 企业层仅作部署/触发对照。"
    )
    p("官方文档索引：https://docs.crewai.com/llms.txt")
    p(
        "定位：用 Flows 管状态与控制流，用 Crews（角色化 Agent + Task）做协作求解。"
        "官方建议：生产应用从 Flow 起步，复杂自主任务再委托 Crew。"
    )
    cite(
        "Introduction",
        "https://docs.crewai.com/v1.15.2/en/introduction",
    )
    p(
        "执行链路（研究视角 · Crew）："
        "kickoff → Agent(LLM) → Tool/MCP/App 调用 → 结果回灌 → Task 完成 → "
        "（可选）Memory remember / Checkpoint。"
    )
    p(
        "执行链路（研究视角 · Conversational Flow）："
        "handle_turn(message, session_id) → 同 session 消息历史 + 本轮 Flow 运行 → "
        "append_assistant_message。"
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "CrewAI「会话」因入口不同而语义分裂："
        "(A) Conversational Flow 的 session_id/message history；"
        "(B) Crew/Flow 一次 kickoff 的运行态 + Checkpoint 血缘；"
        "(C) AMP API 的 kickoff_id / resume。勿把三者当成同一字段。"
    )

    h("1. Conversational Flows（真正多轮 chat session）", 2)
    bullet(
        "每个用户输入 = 一次新的 flow run，但共享同一 session_id。"
    )
    bullet(
        "`handle_turn(message, session_id=...)` → 内部 "
        "`kickoff(inputs={\"id\": session_id})` → `state.id`；"
        "消息写入 `state.messages`。"
    )
    bullet(
        "`stream_turn` / `flow.chat()` REPL / `ChatSession`（SSE·WebSocket）。"
    )
    bullet(
        "`Flow.kickoff()` 不接受 `user_message=` / `session_id=` 关键字；"
        "对话面走 handle_turn。"
    )
    cite(
        "Conversational Flows",
        "https://docs.crewai.com/v1.15.2/en/guides/flows/conversational-flows",
    )

    h("2. Crew / Flow 运行 + Checkpoint", 2)
    bullet(
        "普通 `crew.kickoff()` / Flow 运行偏单次（或批）工作流，不是聊天 session。"
    )
    bullet(
        "Checkpoint：事件驱动快照（默认 task_completed）；"
        "存完整配置、任务进度、中间输出、memory/knowledge 等；"
        "可 resume / fork（新 lineage）。"
    )
    bullet(
        "存储：JsonProvider（如 `./.checkpoints/`）或 SqliteProvider；"
        "`max_checkpoints` 可裁剪。"
    )
    bullet(
        "Flow Persistence：`state.id` 上持久化；fork 可新 id 承状态。"
    )
    cite(
        "Checkpointing",
        "https://docs.crewai.com/v1.15.2/en/concepts/checkpointing",
    )

    h("3. AMP API（企业）", 2)
    bullet("POST /kickoff 启动部署后的 crew；GET /status/{kickoff_id}；POST /resume（人工反馈续跑）。")
    bullet("与本地 Conversational session_id 不同命名空间。")
    cite(
        "AMP kickoff / resume",
        "https://docs.crewai.com/v1.15.2/en/api-reference/kickoff",
    )

    h("4. 研究映射（非官方术语）", 2)
    bullet(
        "Codex/Cline 交互 session ≈ CrewAI Conversational Flow session_id。"
    )
    bullet(
        "OpenClaw Isolated / Automation 单次 ≈ crew.kickoff / AMP kickoff_id。"
    )
    bullet(
        "Checkpoint resume ≈ 失败恢复，不等同聊天续聊 API。"
    )

    # ══════════════════════════════════════════════════════════════
    # 任务模式
    # ══════════════════════════════════════════════════════════════
    h("任务模式", 1)

    h("1. Flows（骨干）", 2)
    bullet("@start / @listen 事件驱动；条件 or/and、Router、循环。")
    bullet("可嵌 Crew 或直接 Agent；内置 Flow memory 方法（remember/recall/extract）。")
    bullet("HITL：`@human_feedback`、ask() 等（步骤审批 ≠ 下一句聊天）。")
    cite("Flows", "https://docs.crewai.com/v1.15.2/en/concepts/flows")

    h("2. Crews + Processes", 2)
    bullet("Crew = Agents + Tasks；`kickoff` / 异步 / for_each / streaming。")
    bullet("Process.sequential：任务列表顺序，前序输出作后序上下文。")
    bullet(
        "Process.hierarchical：经理 Agent/LLM 委派与验收；需 manager_llm 或 manager_agent。"
    )
    bullet("`planning=True` 可启用规划能力。")
    cite("Crews", "https://docs.crewai.com/v1.15.2/en/concepts/crews")
    cite("Processes", "https://docs.crewai.com/v1.15.2/en/concepts/processes")

    h("3. 本地 CLI / AMP Automations", 2)
    bullet("crewai CLI：项目脚手架、run、replay 等。")
    bullet(
        "AMP：Automations、Triggers（Gmail/Slack/Webhook…）、"
        "Crew Studio、Traces、Webhook Streaming。"
    )

    h("任务模式 · 核对摘要", 2)
    bullet("保留：Flow 为生产骨架；Crew sequential/hierarchical；Conversational handle_turn；Checkpoint。")
    bullet("剔除：把「仅有 Agent 单聊」写成框架全部能力。")

    # ══════════════════════════════════════════════════════════════
    # web_fetch
    # ══════════════════════════════════════════════════════════════
    h("web_fetch（及 Security Notice）", 1)
    p(
        "结论：CrewAI 框架层没有名为 `web_fetch` / `WebFetch` 的内置工具，"
        "也没有文档化的 OpenClaw 式 SECURITY NOTICE / EXTERNAL_UNTRUSTED_CONTENT 包装。"
        "网页进入上下文依赖 crewai_tools / MCP / Apps 等可安装或可连接能力；"
        "MCP 文档单独警告 tool metadata prompt injection。"
    )

    h("1. 典型网页工具（非内置唯一名）", 2)
    bullet(
        "`ScrapeWebsiteTool`：HTTP 拉页并解析 HTML；"
        "可固定 website_url 或运行时任意 URL。"
    )
    bullet(
        "Firecrawl / Scrapfly / Selenium / Spider / Stagehand / Browserbase / "
        "Hyperbrowser / Bright Data / Oxylabs / You.com Contents 等 scraping 套件。"
    )
    bullet(
        "搜索侧：`SerperDevTool`、Tavily、Exa、Brave、WebsiteSearchTool（网站内容 RAG）等。"
    )
    bullet("ApifyActorsTool、MultiOnTool：更偏爬虫/浏览器自动化平台。")
    cite(
        "Web Scraping Overview",
        "https://docs.crewai.com/v1.15.2/en/tools/web-scraping/overview",
    )
    cite(
        "ScrapeWebsiteTool",
        "https://docs.crewai.com/v1.15.2/en/tools/web-scraping/scrapewebsitetool",
    )
    cite(
        "SerperDevTool",
        "https://docs.crewai.com/v1.15.2/en/tools/search-research/serperdevtool",
    )

    h("2. Knowledge URL 源 ≠ 实时 web_fetch", 2)
    bullet(
        "Knowledge：文档/文件/URL 语义检索（RAG）注入「知道什么」；"
        "与 agent 调用 ScrapeWebsiteTool 的实时抓取不同威胁时序。"
    )

    h("3. Security Notice 对照", 2)
    bullet(
        "工具文档不说明把 scrape/search 结果包进 SECURITY NOTICE。"
    )
    bullet(
        "MCP Security：仅信任的 MCP；"
        "恶意 tool metadata 可在「列出工具」阶段即污染 LLM；"
        "含「hijack reasoning / prompt injection」表述。"
    )
    bullet(
        "Memory Privacy Note：记忆内容会送给分析用 LLM——"
        "敏感数据宜本地 LLM；这不是网页 untrusted wrapper。"
    )
    cite(
        "MCP Security Considerations",
        "https://docs.crewai.com/v1.15.2/en/mcp/security",
    )

    h("威胁模型映射（研究，非官方）", 2)
    bullet(
        "外网 scrape/search → tool result → Agent LLM → "
        "后续 Tool/App/委派：间接注入面。"
    )
    bullet(
        "恶意 MCP metadata：无需实际调用工具即可注入——额外攻击面。"
    )
    bullet(
        "若记忆开启且写入含网页派生内容：跨 kickoff/session 的持久化污染"
        "（配合 scopes/source 分析）。"
    )
    bullet(
        "对比：Codex web_search+treat untrusted；Cline fetch_web；"
        "OpenClaw web_fetch+NOTICE；CrewAI = 可选工具生态 + MCP 警告。"
    )

    code(
        """典型（官方示例思路）：
from crewai_tools import SerperDevTool, ScrapeWebsiteTool
agent = Agent(..., tools=[SerperDevTool(), ScrapeWebsiteTool()])
# 工具返回正文进入后续 LLM 轮次；文档无 SECURITY NOTICE 外壳约定
"""
    )

    h("web_fetch · 核对摘要", 2)
    bullet(
        "保留：无框架级 web_fetch；Scrape/Search/Browser 工具族；"
        "无 NOTICE；MCP metadata 风险。"
    )
    bullet("剔除：默认内置 web_fetch；把 Knowledge URL ingest 当成 live fetch。")

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "五类扩展：**Tools / MCP / Apps / Skills / Knowledge**。"
        "前三者在运行时都解析为 BaseTool 统一列表；后两者改 prompt/上下文。"
    )
    cite(
        "Agent Capabilities",
        "https://docs.crewai.com/v1.15.2/en/concepts/agent-capabilities",
    )

    h("1. Action：Tools · MCP · Apps", 2)
    bullet("Tools：`pip install 'crewai[tools]'`；自建 BaseTool；可选缓存/异步/Pydantic 输出。")
    bullet("MCP：远程工具服务器；仅连接受信任方。")
    bullet("Apps：平台集成（Gmail 等）经 CrewAI 平台 token。")

    h("2. Context：Skills · Knowledge", 2)
    bullet("Skills：文件系统技能包，注入如何思考；非 callable。")
    bullet("Knowledge：RAG 事实；与 Memory 统一存储不同。")

    h("3. Collaboration / 其它", 2)
    bullet("coworker 委托工具、代码执行选项（Agent 参数）、reasoning 等见 Agents 文档。")

    h("Tools · 核对摘要", 2)
    bullet("保留：五能力分层；Tools≠Skills≠Knowledge。")
    bullet("剔除：只有一个固定 web_fetch 内置名。")

    # ══════════════════════════════════════════════════════════════
    # 五轮防幻觉
    # ══════════════════════════════════════════════════════════════
    h("五轮防幻觉核对记录", 1)

    h("Round 1 — Session", 2)
    bullet("核对：conversational-flows；checkpointing；AMP kickoff/resume。")
    bullet(
        "保留：session_id+handle_turn；kickoff 运行态；checkpoint lineage；kickoff_id。"
    )
    bullet("剔除：把 kickoff_id 写成 chat session_id。")

    h("Round 2 — 任务模式", 2)
    bullet("核对：introduction；flows；crews；processes；planning。")
    bullet("保留：Flow+Crew；sequential/hierarchical；Conversational；AMP triggers。")
    bullet("剔除：宣称无 Flow、仅单 Agent。")

    h("Round 3 — web_fetch / 安全", 2)
    bullet("核对：web-scraping overview；scrapewebsitetool；mcp/security。")
    bullet("保留：工具生态抓网；无 SECURITY NOTICE；MCP metadata PI。")
    bullet("剔除：OpenClaw NOTICE；Codex web_search 默认模式抄到 CrewAI。")

    h("Round 4 — Tools", 2)
    bullet("核对：agent-capabilities；tools；skills；knowledge。")
    bullet("保留：五类；Tools/MCP/Apps 统一为 BaseTool。")
    bullet("剔除：Skills 当成可调用 scrape 工具。")

    h("Round 5 — Prompt / Memory 交叉复检", 2)
    bullet("核对：agents；memory；knowledge；skills。")
    bullet(
        "冻结："
        "(1) Session = conversational session_id 或 kickoff/checkpoint；"
        "(2) 任务 = Flow / Crew(process) / AMP；"
        "(3) 外网 = 可选 scrape/search 工具 + MCP 风险；"
        "(4) Prompt = role/goal/backstory + Task + Skills/Knowledge；"
        "(5) Memory = 统一 Memory 类（非 Checkpoint）。"
    )
    bullet(
        "开放问题：具体 tool 结果模板是否含未文档化的安全前缀；"
        "默认记忆存储路径与跨进程隔离；"
        "AMP 与本地 memory 文件是否共享。"
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt 组成
    # ══════════════════════════════════════════════════════════════
    h("Prompt 组成", 1)
    p(
        "Agent：role / goal / backstory（及可选系统模板参数）+"
        "Task description/expected_output +"
        "Skills 指令 + Knowledge 检索片段 +"
        "工具 schema 与 tool 返回。"
    )

    h("1. Agent 人格字段", 2)
    bullet("role、goal、backstory 为批判性参数；影响工具选择与风格。")
    bullet("可 JSONC/YAML/代码定义；支持 reasoning、动态日期等高级选项。")
    cite("Agents", "https://docs.crewai.com/v1.15.2/en/concepts/agents")

    h("2. Task 与协作上下文", 2)
    bullet("Task.description / expected_output；context= 前序任务输出。")
    bullet("Hierarchical 时经理规划与委派进入执行上下文。")

    h("3. Skills / Knowledge", 2)
    bullet("Skills：领域程序注入「如何想」。")
    bullet("Knowledge：语义检索「知道什么」。")

    h("4. Context window", 2)
    bullet(
        "`respect_context_window`（默认 True）可自动处理过长上下文；"
        "大资料优先 RAG/Knowledge 而非塞进 prompt。"
    )

    h("Prompt 组成 · 核对摘要", 2)
    bullet("保留：role/goal/backstory + task + skills/knowledge + tools。")
    bullet("剔除：AGENTS.md / CLAUDE.md 作为 CrewAI 默认装载（除非用户自建约定）。")

    # ══════════════════════════════════════════════════════════════
    # 记忆与持久化
    # ══════════════════════════════════════════════════════════════
    h("记忆与持久化机制", 1)
    p(
        "三条主线必须分开："
        "① 统一 Memory 类（语义记忆）；"
        "② Checkpointing（执行快照）；"
        "③ Conversational state.messages / Flow state persistence。"
    )

    h("1. Unified Memory", 2)
    bullet(
        "单一 `Memory` API 取代旧 short/long/entity/external 分家；"
        "保存时 LLM 推断 scope/categories/importance；"
        "recall 综合语义+时效+重要性。"
    )
    bullet(
        "用法：standalone；`Crew(memory=True)` 或传入 Memory；"
        "Agent 级 scoped view；Flow 内 `remember/recall/extract_memories`。"
    )
    bullet(
        "默认 embedder 常为 OpenAI text-embedding-3-large（可改）；"
        "分析 LLM 默认 gpt-4o-mini 可改（含 Ollama）。"
    )
    bullet("Scopes / Slices / private+source 标签；RecallFlow 深检索。")
    cite("Memory", "https://docs.crewai.com/v1.15.2/en/concepts/memory")

    h("2. Checkpointing（非语义记忆）", 2)
    bullet("恢复任务进度与中间产物；可重新水合 memory/knowledge 配置状态。")
    bullet("职责：故障恢复 / fork；不是「用户偏好知识库」API。")

    h("3. Conversational / Flow 状态", 2)
    bullet("ConversationState.messages 跨 turn 的同一 session_id。")
    bullet("Flow Persistence 按 state.id 快照业务状态。")

    h("4. Knowledge（RAG）", 2)
    bullet("文档/URL 等源的检索增强；与 Memory 统一类并存但目的不同。")

    h("记忆与持久化 · 核对摘要", 2)
    bullet(
        "保留：Memory vs Checkpoint vs chat messages；scopes；分析 LLM 隐私注记。"
    )
    bullet(
        "剔除：把 Checkpoint 目录当成 MEMORY.md；"
        "剔除默认存在 OpenClaw 式 MEMORY.md。"
    )

    h("主要参考链接", 1)
    for link in [
        "https://docs.crewai.com/llms.txt",
        "https://docs.crewai.com/v1.15.2/en/introduction",
        "https://docs.crewai.com/v1.15.2/en/concepts/flows",
        "https://docs.crewai.com/v1.15.2/en/concepts/crews",
        "https://docs.crewai.com/v1.15.2/en/concepts/processes",
        "https://docs.crewai.com/v1.15.2/en/concepts/agents",
        "https://docs.crewai.com/v1.15.2/en/concepts/tools",
        "https://docs.crewai.com/v1.15.2/en/concepts/agent-capabilities",
        "https://docs.crewai.com/v1.15.2/en/concepts/memory",
        "https://docs.crewai.com/v1.15.2/en/concepts/knowledge",
        "https://docs.crewai.com/v1.15.2/en/concepts/checkpointing",
        "https://docs.crewai.com/v1.15.2/en/concepts/skills",
        "https://docs.crewai.com/v1.15.2/en/guides/flows/conversational-flows",
        "https://docs.crewai.com/v1.15.2/en/tools/web-scraping/overview",
        "https://docs.crewai.com/v1.15.2/en/tools/web-scraping/scrapewebsitetool",
        "https://docs.crewai.com/v1.15.2/en/tools/search-research/serperdevtool",
        "https://docs.crewai.com/v1.15.2/en/mcp/security",
        "https://docs.crewai.com/v1.15.2/en/api-reference/kickoff",
    ]:
        bullet(link)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
