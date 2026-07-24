# -*- coding: utf-8 -*-
"""Generate Codex CLI research notes (English).

Same structure as Chinese _gen_codex_cli_doc.py: Session, task modes,
web_fetch/security notice, tools, Prompt, memory/persistence,
five-round anti-hallucination vs developers.openai.com/codex.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "en" / "Codex-CLI.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def code(text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def cite(label: str, url: str) -> None:
        doc.add_paragraph(f"{label} (Source: {url})")

    # ── Title ──
    h("Codex CLI", 0)
    p(
        "This document follows the same structure as OpenHands.docx / Hermes.docx / "
        "Claude-Code.docx / Dify.docx / AutoGPT.docx / Cline.docx: Session, task modes, "
        "web_fetch (and Security Notice comparison), Tools; plus Prompt composition, memory "
        "and persistence mechanisms, and a five-round anti-hallucination review at the end. "
        "Primary subject: local terminal agent Codex CLI (shares config layer with IDE "
        "extension / Desktop); Codex Cloud is comparison only, not primary threat surface."
    )
    p("Official documentation index: https://developers.openai.com/codex/llms.txt")
    p(
        "Positioning: local coding agent — read/modify files and execute commands within OS "
        "sandbox; optional web search / MCP / Skills / Subagents; "
        "network off by default for commands in workspace-write sandbox."
    )
    cite(
        "Agent approvals & security",
        "https://developers.openai.com/codex/agent-approvals-security",
    )
    p(
        "Execution chain (research view · CLI): User → TUI/`codex exec` turn → "
        "tools (shell / web_search / MCP / skill scripts…) → "
        "results into transcript → (optional) Memories / AGENTS.md feedback."
    )

    # ══════════════════════════════════════════════════════════════
    # Session
    # ══════════════════════════════════════════════════════════════
    h("Session", 1)
    p(
        "Interactive CLI persists resumable session / conversation units; "
        "non-interactive `codex exec` JSON stream exposes thread_id. "
        "Local state root is CODEX_HOME (default ~/.codex)."
    )

    h("1. Storage and continuation", 2)
    bullet(
        "By default writes local session transcript to CODEX_HOME "
        "(doc example: ~/.codex/history.jsonl)."
    )
    bullet(
        "[history] persistence = \"none\" can disable; "
        "history.max_bytes truncates oldest entries and compacts."
    )
    bullet(
        "`codex resume`: continue by ID or recent session; "
        "`--last` defaults to current working directory, `--all` cross-directory."
    )
    bullet(
        "`codex archive` / `unarchive` / `delete`; "
        "TUI has `/resume` `/archive` `/delete` `/fork` `/rename`."
    )
    bullet(
        "`codex exec --ephemeral`: do not persist session rollout files to disk."
    )
    cite(
        "History persistence",
        "https://developers.openai.com/codex/config-file/config-advanced",
    )
    cite(
        "CLI reference · resume / exec",
        "https://developers.openai.com/codex/cli/reference",
    )

    h("2. Task vs Session (product semantics)", 2)
    bullet(
        "`/new`: start new task within same CLI session (clears task context, stays in process)."
    )
    bullet("`/fork`: fork current conversation into new task.")
    bullet("`/compact`: compress visible dialogue to free context window.")
    bullet(
        "JSONL events: thread.started / turn.* / item.* "
        "(including web searches, MCP, command_execution)."
    )

    h("3. Comparison with other agents (non-official terms)", 2)
    bullet("Claude Code project JSONL ≈ Codex history.jsonl + resume.")
    bullet("Cline sessions.db/[id].json ≈ Codex session ID + history persistence.")
    bullet("Do not write Dify conversation_id API field as Codex contract.")

    # ══════════════════════════════════════════════════════════════
    # Task modes
    # ══════════════════════════════════════════════════════════════
    h("Task Modes", 1)

    h("1. Interactive TUI (`codex`)", 2)
    bullet(
        "Launch interactive UI without subcommand; can attach images; slash commands mid-session."
    )
    bullet(
        "Web search cached by default; `--search` switches to live. "
        "Common local dev: `--sandbox workspace-write --ask-for-approval on-request`."
    )
    cite("CLI reference · interactive", "https://developers.openai.com/codex/cli/reference")

    h("2. Non-interactive (`codex exec`)", 2)
    bullet(
        "Scripts/CI: progress on stderr, final message on stdout; `--json` yields JSONL event stream."
    )
    bullet(
        "Default read-only sandbox; explicit `--sandbox workspace-write` needed for disk writes."
    )
    bullet(
        "`--ignore-user-config` / `--ignore-rules` for controlled automation; "
        "`--full-auto` deprecated compatibility path."
    )
    cite(
        "Non-interactive mode",
        "https://developers.openai.com/codex/non-interactive-mode",
    )

    h("3. Sandbox × Approval combinations", 2)
    bullet(
        "Two layers: **Sandbox** (technical boundary: write scope, network) + "
        "**Approval policy** (when to stop and ask human)."
    )
    bullet(
        "Common: Auto (workspace-write + on-request); "
        "read-only; untrusted approval; "
        "`--yolo` / danger-full-access (no sandbox no approval, officially not recommended)."
    )
    bullet(
        "`approvals_reviewer = auto_review`: qualified approval requests can pass reviewer agent first."
    )
    bullet(
        "Version-controlled dirs recommend Auto; non-version dirs tend read-only; "
        "can switch mid-session via `/permissions`."
    )
    cite(
        "Agent approvals & security",
        "https://developers.openai.com/codex/agent-approvals-security",
    )
    cite("Sandbox", "https://developers.openai.com/codex/sandboxing")

    h("4. Automation / multi-agent / cloud (boundaries)", 2)
    bullet(
        "Scheduled tasks / Automations: docs mainly Desktop \"Scheduled\"; can configure skills."
    )
    bullet("Subagents / multi_agent feature: collaborative tools can be enabled.")
    bullet(
        "Cloud: isolated container + two phases (setup can use network for deps, agent phase "
        "offline by default) — different from local CLI sandbox; do not mix in papers."
    )

    h("Task Modes · Review summary", 2)
    bullet("Kept: TUI / exec; sandbox×approval; yolo; resume; exec default read-only.")
    bullet(
        "Removed: writing Cloud default internet as CLI workspace-write default network on."
    )

    # ══════════════════════════════════════════════════════════════
    # web_fetch
    # ══════════════════════════════════════════════════════════════
    h("web_fetch (and Security Notice)", 1)
    p(
        "Conclusion: Codex CLI docs expose external web via first-party **web search tool** "
        "(config key `web_search`), not a tool named `web_fetch` / `WebFetch`. "
        "Official docs explicitly require treating all web results as untrusted input, "
        "and discuss cached mode reducing prompt injection from arbitrary live content, "
        "but public docs do not provide OpenClaw-style SECURITY NOTICE / "
        "EXTERNAL_UNTRUSTED_CONTENT wrapper text."
    )

    h("1. web_search modes", 2)
    bullet(
        "`cached` (default): OpenAI-maintained web index/cache, not instant arbitrary page pull."
    )
    bullet("`indexed`: external access gated through search index.")
    bullet("`live` (=`--search`): fetch newer live data.")
    bullet("`disabled`: turn off the tool.")
    bullet(
        "Using `--yolo` or other full-access sandbox, web search defaults to live."
    )
    cite("Web search", "https://developers.openai.com/codex/web-search")
    cite(
        "Config basics · Web search mode",
        "https://developers.openai.com/codex/config-file/config-basic",
    )

    h("2. Untrusted wording vs SECURITY NOTICE", 2)
    bullet(
        "Doc intent: \"Treat all web results as untrusted input\"; "
        "cached \"reduces but does not eliminate\" PI; "
        "when enabling network or live search, caution — "
        "\"agent may fetch and follow untrusted instructions\"."
    )
    bullet(
        "Differs from Claude Code \"isolated small-model extraction + ~15min cache\" description; "
        "differs from OpenClaw fixed NOTICE wrapper — do not cross-copy."
    )
    bullet(
        "Also: `memories.disable_on_external_context` can exclude tasks with MCP / web search "
        "from memory generation (legacy alias: no_memories_if_mcp_or_web_search)."
    )

    h("3. Separated from \"command outbound network\"", 2)
    bullet(
        "Default workspace-write: **spawned commands have no network** "
        "(needs `[sandbox_workspace_write] network_access = true`)."
    )
    bullet(
        "Can use web search tool without opening full command network "
        "(approvals page explicitly references platform web_search guide)."
    )
    bullet(
        "Optional features.network_proxy: domain allowlist / local block / DNS check "
        "for command network when enabled."
    )

    h("Threat model mapping (research, non-official)", 2)
    bullet(
        "cached/live web results → model context → shell/MCP/file edits: "
        "indirect injection / goal hijacking surface; official docs name PI."
    )
    bullet(
        "live + command network_access=true expands executable outbound surface; "
        "yolo removes both sandbox and approval."
    )
    bullet(
        "Comparison table: OpenClaw web_fetch+NOTICE; Claude WebFetch; "
        "Cline fetch_web; Codex = web_search(cached|live) + \"treat as untrusted\"."
    )

    code(
        """Config example (official docs):
web_search = \"cached\"   # default
# web_search = \"live\"    # same as --search
# web_search = \"disabled\"

# Command outbound network (separate from web_search)
[sandbox_workspace_write]
network_access = true
"""
    )

    h("web_fetch · Review summary", 2)
    bullet(
        "Kept: tool semantics=web_search; four modes; default cached; untrusted warning; "
        "no SECURITY NOTICE docs; command network off by default."
    )
    bullet("Removed: built-in tool name web_fetch; OpenClaw NOTICE assumed by default.")

    # ══════════════════════════════════════════════════════════════
    # Tools
    # ══════════════════════════════════════════════════════════════
    h("Tools", 1)
    p(
        "Capabilities from: sandboxed shell/exec, web_search, MCP, Skills scripts, "
        "(optional) multi_agent/subagent collaborative tools; config can disable namespaces."
    )

    h("1. Built-in execution and retrieval", 2)
    bullet("`shell_tool` feature (default true): default shell tool.")
    bullet(
        "`unified_exec`: unified PTY-backed exec (Windows default policy differs)."
    )
    bullet("web_search feature / top-level `web_search = ...` controls retrieval tool.")
    bullet(
        "JSON item types include command_execution, file changes, web searches, "
        "MCP tool calls, plan updates."
    )

    h("2. MCP / Skills / Rules", 2)
    bullet(
        "MCP: STDIO / Streamable HTTP; config in config.toml; "
        "Desktop/CLI/IDE shared; `/mcp` lists tools."
    )
    bullet(
        "Skills: SKILL.md + scripts/references; `/skills`; skill scripts can require approval."
    )
    bullet(
        "Rules (experimental): `prefix_rule` controls allow/prompt/forbidden for "
        "out-of-sandbox commands."
    )
    cite("MCP", "https://developers.openai.com/codex/extend/mcp")
    cite("Rules", "https://developers.openai.com/codex/agent-configuration/rules")

    h("3. Approval surface extends beyond shell", 2)
    bullet("Side-effecting app/connector/MCP calls also require approval.")
    bullet("MCP calls marked destructive always require approval when annotated.")
    bullet(
        "Granular approval can separately control sandbox / rules / MCP / "
        "request_permissions / skill_approval."
    )

    h("Tools · Review summary", 2)
    bullet("Kept: shell/exec + web_search + MCP + Skills + Rules.")
    bullet("Removed: fabricated fixed PascalCase WebFetch as sole name.")

    # ══════════════════════════════════════════════════════════════
    # Five-round anti-hallucination
    # ══════════════════════════════════════════════════════════════
    h("Five-Round Anti-Hallucination Review", 1)

    h("Round 1 — Session", 2)
    bullet(
        "Checked: config-advanced History; cli-reference resume/archive; exec JSON thread_id."
    )
    bullet("Kept: ~/.codex/history.jsonl; resume; /new·/fork·/compact; ephemeral.")
    bullet("Removed: Claude ~/.claude/projects paths; Cline sessions.db structure.")

    h("Round 2 — Task modes", 2)
    bullet("Checked: cli-reference; non-interactive; approvals; sandboxing.")
    bullet("Kept: TUI vs exec; sandbox×approval table; yolo; Cloud two-phase comparison.")
    bullet(
        "Removed: asserting Desktop Scheduled as CLI sole scheduling engine without note."
    )

    h("Round 3 — web_fetch / security", 2)
    bullet("Checked: web-search; config-basic web_search; approvals network section.")
    bullet(
        "Kept: web_search four modes; cached default; treat as untrusted; "
        "command network off by default; no SECURITY NOTICE wrapper docs."
    )
    bullet(
        "Removed: tool name web_fetch; OpenClaw NOTICE; Claude 15min isolated extract copied verbatim."
    )

    h("Round 4 — Tools", 2)
    bullet("Checked: features table shell_tool/web_search; mcp; rules; skills.")
    bullet("Kept: shell + web_search + MCP + Skills; destructive MCP approval.")
    bullet("Removed: discussing only curl without official web_search tool.")

    h("Round 5 — Prompt / Memory cross-check", 2)
    bullet("Checked: agents-md; memories; customization overview; /compact·/memories.")
    bullet(
        "Frozen: "
        "(1) Session = history + resume/session id; "
        "(2) Modes = TUI/exec + sandbox/approval; "
        "(3) External retrieval = web_search(cached|live|…) + untrusted guidance; "
        "(4) Prompt = AGENTS.md chain + Skills + (optional) Memories; "
        "(5) Memories off by default, path ~/.codex/memories/."
    )
    bullet(
        "Open questions: web_search result message wrapping format into model; "
        "whether live equals arbitrary URL fetch; "
        "history.jsonl vs per-session rollout file layout needs live verification."
    )

    # ══════════════════════════════════════════════════════════════
    # Prompt composition
    # ══════════════════════════════════════════════════════════════
    h("Prompt Composition", 1)
    p(
        "On startup loads instruction chain + user message/attachments + tool results; "
        "full system template not fully public. Configurable compact prompt override, etc."
    )

    h("1. AGENTS.md instruction chain", 2)
    bullet("Global: ~/.codex/AGENTS.override.md preferred, else AGENTS.md.")
    bullet(
        "Project: walk from Git root to cwd, at most one "
        "AGENTS.override.md / AGENTS.md / fallback filename per level."
    )
    bullet(
        "Concatenated top-down; closer to cwd later and overrides; "
        "limited by project_doc_max_bytes (default 32KiB)."
    )
    cite(
        "Custom instructions with AGENTS.md",
        "https://developers.openai.com/codex/agent-configuration/agents-md",
    )

    h("2. Skills / MCP instructions / user layer", 2)
    bullet("Skills metadata visible to agent; SKILL.md loaded when selected.")
    bullet("MCP initialization instructions field as server-wide guidance.")
    bullet("personality, /personality; debug can inspect instruction discovery.")

    h("3. Memories injection (optional)", 2)
    bullet("`memories.use_memories`: whether to inject existing memories into future sessions.")
    bullet(
        "Separate from mandatory team rules: rules should go in AGENTS.md, not memories alone."
    )

    h("4. Context management slash commands", 2)
    bullet(
        "`/compact`; compact_prompt / experimental_compact_prompt_file can override summary prompt."
    )
    bullet("`/status`: model, approval, writable roots, context remaining.")

    h("Prompt Composition · Review summary", 2)
    bullet("Kept: AGENTS discovery order; Skills; optional Memories; compact.")
    bullet("Removed: CLAUDE.md or Hermes SOUL.md as Codex default filenames.")

    # ══════════════════════════════════════════════════════════════
    # Memory and persistence
    # ══════════════════════════════════════════════════════════════
    h("Memory and Persistence", 1)

    h("1. Session transcripts (history)", 2)
    bullet("~/.codex/history.jsonl (can disable/cap); resume/archive/delete/fork.")
    bullet("exec JSON exposes thread/turn event stream.")

    h("2. Local Memories (experimental, off by default)", 2)
    bullet("[features] memories = true; files in ~/.codex/memories/.")
    bullet(
        "Background generation from idle long-enough qualified tasks; "
        "skips active/too-short sessions; field redaction."
    )
    bullet(
        "`/memories`: whether current task uses / generates; not same as toggling global switch."
    )
    bullet(
        "disable_on_external_context can avoid writing MCP/web_search tasks to memory — "
        "especially relevant for PI persistence research."
    )
    cite("Memories", "https://developers.openai.com/codex/customization/memories")

    h("3. AGENTS.md (persistent rules layer)", 2)
    bullet(
        "Cross-session, Git-committable team guidance; official recommends division from Memories."
    )

    h("4. Other local state", 2)
    bullet("config.toml / profiles / auth / logs / caches under CODEX_HOME.")
    bullet(
        "Trusted projects only load `.codex/` project layer (config/hooks/rules)."
    )
    bullet(
        "Chronicle: Desktop-only, resume recent context from screen — not CLI core path."
    )

    h("Memory and Persistence · Review summary", 2)
    bullet(
        "Kept: history; optional memories directory; AGENTS; "
        "external_context exclusion; trust boundary for loading .codex/."
    )
    bullet(
        "Removed: Memories on by default; "
        "removed Hermes fixed MEMORY.md path equivalent to ~/.codex/memories."
    )

    h("Primary Reference Links", 1)
    for link in [
        "https://developers.openai.com/codex/llms.txt",
        "https://developers.openai.com/codex/cli/reference",
        "https://developers.openai.com/codex/non-interactive-mode",
        "https://developers.openai.com/codex/agent-approvals-security",
        "https://developers.openai.com/codex/sandboxing",
        "https://developers.openai.com/codex/web-search",
        "https://developers.openai.com/codex/config-file/config-basic",
        "https://developers.openai.com/codex/config-file/config-advanced",
        "https://developers.openai.com/codex/agent-configuration/agents-md",
        "https://developers.openai.com/codex/agent-configuration/rules",
        "https://developers.openai.com/codex/extend/mcp",
        "https://developers.openai.com/codex/customization/overview",
        "https://developers.openai.com/codex/customization/memories",
        "https://developers.openai.com/codex/build-skills",
        "https://developers.openai.com/codex/cli/slash-commands",
    ]:
        bullet(link)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
